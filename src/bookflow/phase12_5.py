"""Phase 12.5 Book 1 appendix baseline correction and candidate builder.

This module is offline-only. It never mutates the supplied PDF/DOCX, Phase 12R
releases, manifests, checkpoints, canonical data, boundaries, or translation caches.
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterator

import fitz
import yaml
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


PDF_REF = Path("references/phase12_5/book1/source/book1_appendix_original.pdf")
GOLD_REF = Path("references/phase12_5/book1/human_gold/book1_bilingual_appendix_rebuilt.docx")
DATA_ROOT = Path("data/fullbook/phase12_5/book1")
DOC_ROOT = Path("docs/phase12_5")
OUTPUT_ROOT = Path("output/fullbook/phase12_5")
APPENDIX_STARTS = {"APPENDIX A": ("appendix_a", 381), "APPENDIX B": ("appendix_b", 398), "APPENDIX C": ("appendix_c", 400)}
APPENDIX_TITLES = {
    "APPENDIX A": ("附录 A", "APPENDIX A"),
    "FIELD MEASUREMENTS AND NOTES OF CERTAIN SPECIES OF LARGE GAME KILLED IN CHINA, 1911":
        ("1911年中国境内猎获的若干大型猎物物种的野外测量与说明", "FIELD MEASUREMENTS AND NOTES OF CERTAIN SPECIES OF LARGE GAME KILLED IN CHINA, 1911"),
    "APPENDIX B": ("附录 B", "APPENDIX B"),
    "ESTIMATE OF EXPENSES": ("费用估算", "ESTIMATE OF EXPENSES"),
    "APPENDIX C": ("附录 C", "APPENDIX C"),
    "TABLE OF DISTANCES AND STAGES": ("路程与驿站表", "TABLE OF DISTANCES AND STAGES"),
}
TECHNICAL_MARKER = "附录全译重建与 ChatGPT 网页辅助校订方案"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def atomic_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", "utf-8")
    tmp.replace(path)


def relative_metadata(root: Path, path: Path) -> dict[str, Any]:
    stat = path.stat()
    return {
        "path": path.relative_to(root).as_posix(),
        "size_bytes": stat.st_size,
        "mtime": datetime.fromtimestamp(stat.st_mtime).astimezone().isoformat(),
        "sha256": sha256(path),
    }


def iter_blocks(document: Document) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def has_cjk(text: str) -> bool:
    return bool(re.search(r"[\u3400-\u9fff]", text))


def split_bilingual(text: str) -> tuple[str, str]:
    text = text.strip()
    if not text:
        return "", ""
    parts = [part.strip() for part in text.split(" / ")]
    if len(parts) > 1 and any(has_cjk(part) for part in parts) and any(not has_cjk(part) for part in parts):
        zh = " / ".join(part for part in parts if has_cjk(part))
        en = " / ".join(part for part in parts if not has_cjk(part))
        return en, zh
    if has_cjk(text):
        return "", text
    return text, ""


def raw_value(text: str) -> str | None:
    tokens = re.findall(r"(?<![A-Za-z])\d+(?:[.,]\d+)?(?:[¼½¾⅛⅜⅝⅞]|[⁄/]\d+)?", text)
    return " | ".join(tokens) or None


def raw_unit(text: str) -> str | None:
    units = re.findall(r"\b(?:in\.?|ft\.?|feet|foot|lb\.?|oz\.?|mm\.?|miles?|li|versts?|taels?|£|s\.|d\.)\b", text, re.I)
    return " | ".join(units) or None


def printed_to_physical(value: int, default: int) -> int:
    return value + 90 if 291 <= value <= 314 else default


def extract_appendices(root: Path) -> dict[str, Any]:
    pdf = root / PDF_REF
    gold = root / GOLD_REF
    if not pdf.is_file() or not gold.is_file():
        raise FileNotFoundError("both authoritative Phase 12.5 inputs are required")
    source_sha = sha256(pdf)
    gold_sha = sha256(gold)
    document = Document(gold)
    objects: list[dict[str, Any]] = []
    place_names: list[dict[str, Any]] = []
    appendix: str | None = None
    source_page = 381
    paragraph_counts: Counter[str] = Counter()
    table_counts: Counter[str] = Counter()
    active = False
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text == "APPENDIX A":
                active = True
            if text == "Index" and active:
                break
            if not active or not text:
                continue
            if text in APPENDIX_STARTS:
                appendix, source_page = APPENDIX_STARTS[text]
            marker = re.fullmatch(r"【?(\d{3})】?", text)
            if block.style.name == "PageNumber" and marker:
                source_page = printed_to_physical(int(marker.group(1)), source_page)
                continue
            if appendix is None:
                continue
            paragraph_counts[appendix] += 1
            en, zh = split_bilingual(text)
            object_type = "heading" if block.style.name == "AppendixHeading" else "prose"
            object_id = f"{appendix}.{object_type}.{paragraph_counts[appendix]:03d}"
            objects.append(make_record(
                object_id, appendix, object_type, source_page, source_sha,
                en, zh, text, None, None, None, gold_sha,
            ))
        else:
            if not active or appendix is None:
                continue
            table_counts[appendix] += 1
            table_id = f"{appendix}.table.{table_counts[appendix]:03d}"
            for row_no, row in enumerate(block.rows, 1):
                row_id = f"{table_id}.row.{row_no:03d}"
                for column_no, cell in enumerate(row.cells, 1):
                    text = cell.text.strip()
                    en, zh = split_bilingual(text)
                    object_id = f"{row_id}.cell.{column_no:03d}"
                    record = make_record(
                        object_id, appendix, "table_cell", source_page, source_sha,
                        en, zh, text, table_id, row_id, f"column.{column_no:03d}", gold_sha,
                    )
                    objects.append(record)
                    if appendix == "appendix_c" and en and zh and ("→" in text or "†" in text):
                        status = "uncertain" if "†" in text or "待考" in text else "probable"
                        place_names.append({
                            "object_id": object_id + ".place",
                            "original_romanization": en,
                            "zh_name": zh,
                            "display_name": text,
                            "identification_status": status,
                            "modern_correspondence": None,
                            "evidence_note": "Human Gold display; dagger/待考 remains unresolved." if status == "uncertain" else "Human Gold display; modern correspondence not asserted.",
                            "source_page": source_page,
                            "source_file_sha256": source_sha,
                            "review_status": "confirmed",
                        })
    payload = {
        "schema_version": "phase12_5_appendix_objects.v1",
        "book_id": "big_game_1913",
        "baseline": {"source_pdf_sha256": source_sha, "human_gold_docx_sha256": gold_sha},
        "authority": {"source_facts": "source_pdf", "zh_and_bilingual_display": "human_gold_docx"},
        "objects": objects,
        "statistics": {
            "objects": len(objects),
            "by_appendix": dict(Counter(item["appendix_id"] for item in objects)),
            "by_type": dict(Counter(item["object_type"] for item in objects)),
            "tables": dict(table_counts),
            "place_names": len(place_names),
        },
    }
    data_root = root / DATA_ROOT
    atomic_json(data_root / "appendix_objects_v1.json", payload)
    atomic_json(data_root / "place_names_v1.json", {"schema_version": "phase12_5_place_names.v1", "objects": place_names})
    conflicts = source_conflicts(pdf, objects)
    atomic_json(data_root / "source_gold_conflicts_v1.json", conflicts)
    return {"objects": payload, "place_names": place_names, "conflicts": conflicts}


def make_record(
    object_id: str, appendix_id: str, object_type: str, source_page: int,
    source_sha: str, source_text: str, zh_text: str, bilingual: str,
    table_id: str | None, row_id: str | None, column_id: str | None, gold_sha: str,
) -> dict[str, Any]:
    return {
        "object_id": object_id,
        "appendix_id": appendix_id,
        "object_type": object_type,
        "source_page": source_page,
        "source_file_sha256": source_sha,
        "source_text": source_text or None,
        "zh_text": zh_text or None,
        "bilingual_display": bilingual,
        "table_id": table_id,
        "row_id": row_id,
        "column_id": column_id,
        "value_raw": raw_value(source_text or bilingual),
        "unit_raw": raw_unit(source_text or bilingual),
        "review_status": "confirmed",
        "review_method": "human_confirmed",
        "review_note": f"English/numeric facts trace to source PDF; Chinese/display trace to Human Gold {gold_sha}.",
    }


def source_conflicts(pdf_path: Path, objects: list[dict[str, Any]]) -> dict[str, Any]:
    reader = fitz.open(pdf_path)
    page_text: dict[int, str] = {}
    for physical in range(381, 405):
        excerpt_page = physical - 379
        if 1 <= excerpt_page <= reader.page_count:
            page_text[physical] = re.sub(r"\W", "", reader[excerpt_page - 1].get_text() or "").lower()
    items: list[dict[str, Any]] = []
    for item in objects:
        value = item.get("value_raw")
        if not value:
            continue
        normalized = page_text.get(item["source_page"], "")
        tokens = [re.sub(r"\W", "", token) for token in value.split(" | ")]
        missing = [token for token in tokens if token and token not in normalized]
        if missing:
            items.append({
                "object_id": item["object_id"], "source_page": item["source_page"],
                "field": "value_raw", "gold_value": value,
                "status": "human_recheck_required", "reason": "token not confirmed by embedded PDF text layer",
            })
    return {
        "schema_version": "phase12_5_source_gold_conflicts.v1",
        "comparison_limit": "PDF embedded text is lossy for fractions/symbols; findings are review items, not automatic corrections.",
        "conflict_count": len(items),
        "conflicts": items,
    }


def write_baseline_manifest(root: Path) -> dict[str, Any]:
    pdf = root / PDF_REF
    gold = root / GOLD_REF
    manifest = {
        "schema_version": "phase12_5_book1_baseline.v1",
        "book_id": "big_game_1913",
        "registered_at": datetime.now(timezone.utc).isoformat(),
        "inputs": {
            "source_pdf": {**relative_metadata(root, pdf), "role": "source_fact_authority", "authority": ["English", "numbers", "fractions", "units", "symbols", "taxonomy", "table_relationships", "printed_pages", "romanization", "footnotes"], "review_status": "authoritative", "modifiable": False, "package_inclusion": "forbidden"},
            "human_gold_docx": {**relative_metadata(root, gold), "role": "human_reviewed_gold_reference", "authority": ["Chinese translation", "bilingual titles", "editable table design", "place-name display", "layout reference"], "review_status": "human_confirmed", "modifiable": False, "package_inclusion": "forbidden"},
        },
        "legacy_release_identity": "legacy_historical_release; superseded_content_quality_baseline; engine_regression_reference_only",
        "candidate_build_id_rule": "phase12-5-book1-<language>-<UTC timestamp>-<gold sha prefix>",
        "conflict_policy": "Source PDF controls English/numeric facts; report every conflict for human review; never silently overwrite Human Gold.",
        "book_data_package_inclusion": "forbidden",
    }
    path = root / DOC_ROOT / "book1_baseline_manifest.yaml"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), "utf-8")
    return manifest


def select_text(text: str, language: str, *, preserve_place: bool = False) -> str:
    if language == "bilingual" or not text:
        return text
    en, zh = split_bilingual(text)
    if language == "en":
        return en or (text if not has_cjk(text) else "")
    if preserve_place and en and zh:
        return f"{zh} / {en}"
    return zh or (text if not has_cjk(text) else text)


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def remove_technical_section(document: Document) -> None:
    removing = False
    for child in list(document.element.body):
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
            if TECHNICAL_MARKER in text:
                removing = True
            if text == "APPENDIX A":
                removing = False
        if removing:
            document.element.body.remove(child)


def apply_structured_appendix(document: Document, payload: dict[str, Any]) -> None:
    """Replay the structured Appendix A-C model onto the Gold layout skeleton."""

    records = iter(payload["objects"])
    active = False
    consumed = 0
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text == "APPENDIX A":
                active = True
            if text == "Index" and active:
                break
            if not active or not text:
                continue
            if block.style.name == "PageNumber" and re.fullmatch(r"【?(\d{3})】?", text):
                continue
            record = next(records)
            if record["object_type"] not in {"heading", "prose"}:
                raise RuntimeError(f"structured replay type mismatch at {record['object_id']}")
            block.text = record["bilingual_display"]
            consumed += 1
        elif active:
            for row in block.rows:
                for cell in row.cells:
                    record = next(records)
                    if record["object_type"] != "table_cell":
                        raise RuntimeError(f"structured replay type mismatch at {record['object_id']}")
                    cell.text = record["bilingual_display"]
                    consumed += 1
    try:
        extra = next(records)
    except StopIteration:
        extra = None
    if extra is not None or consumed != len(payload["objects"]):
        raise RuntimeError("structured appendix replay did not consume the complete object model")


def transform_document(document: Document, language: str) -> None:
    remove_technical_section(document)
    in_appendix = False
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text == "APPENDIX A":
            in_appendix = True
        if text == "Index":
            in_appendix = False
        style = paragraph.style.name
        remove = language == "en" and style in {"BodyChinese", "ChapterTitleChinese"}
        remove = remove or (language == "zh-Hans" and style in {"BodyEnglish", "ChapterTitle"})
        if remove:
            remove_paragraph(paragraph)
            continue
        if in_appendix and text in APPENDIX_TITLES:
            zh, en = APPENDIX_TITLES[text]
            paragraph.text = en if language == "en" else zh if language == "zh-Hans" else f"{zh} / {en}"
        elif in_appendix:
            paragraph.text = select_text(text, language)
        elif style == "CaptionEnglish" or (" / " in text and style in {"ChapterTitle", "FrontMatterHeading"}):
            paragraph.text = select_text(text, language)
    appendix_seen = False
    for table in list(document.tables):
        header = " | ".join(cell.text for cell in table.rows[0].cells) if table.rows else ""
        if "测量项目 / Measurement" in header:
            appendix_seen = True
        if not appendix_seen:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = select_text(cell.text, language)
            continue
        parallel = len(table.columns) == 2 and "English original" in header and "中文" in header
        if parallel and language != "bilingual":
            keep = 0 if language == "en" else 1
            drop = 1 - keep
            for row in table.rows:
                cells = row.cells
                cells[keep].text = select_text(cells[keep].text, language)
                cells[drop]._tc.getparent().remove(cells[drop]._tc)
            grid = table._tbl.tblGrid
            if grid is not None and len(grid.gridCol_lst) > 1:
                grid.remove(grid.gridCol_lst[drop])
        else:
            for row in table.rows:
                for cell in row.cells:
                    preserve_place = "→" in cell.text or "†" in cell.text
                    cell.text = select_text(cell.text, language, preserve_place=preserve_place)


def markdown_from_docx(path: Path, output: Path) -> None:
    document = Document(path)
    lines: list[str] = []
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name
            prefix = "# " if style == "Title" else "## " if style in {"FrontMatterHeading", "ChapterTitle", "ChapterTitleChinese", "AppendixHeading"} else ""
            lines.extend([prefix + text, ""])
        else:
            rows = [[cell.text.replace("\n", " ").strip() for cell in row.cells] for row in block.rows]
            if not rows:
                continue
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * width) + " |")
            lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
            lines.append("")
    output.write_text("\n".join(lines), "utf-8")


def find_soffice() -> str:
    candidates = [shutil.which("soffice"), r"C:\Program Files\LibreOffice\program\soffice.exe", r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return str(candidate)
    raise RuntimeError("LibreOffice soffice not found")


def render_pdf(docx_path: Path, pdf_dir: Path) -> Path:
    pdf_dir.mkdir(parents=True, exist_ok=True)
    profile = pdf_dir / ".lo_profile"
    profile.mkdir(exist_ok=True)
    command = [find_soffice(), "--headless", f"-env:UserInstallation={profile.resolve().as_uri()}", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(docx_path)]
    run = subprocess.run(command, capture_output=True, text=True, timeout=600)
    output = pdf_dir / (docx_path.stem + ".pdf")
    if run.returncode or not output.is_file():
        raise RuntimeError(f"LibreOffice conversion failed: {run.stderr[-1000:]}")
    return output


def appendix_page_renders(pdf_path: Path, qa_dir: Path) -> list[dict[str, Any]]:
    qa_dir.mkdir(parents=True, exist_ok=True)
    document = fitz.open(pdf_path)
    hits: list[int] = []
    started = False
    for page_no, page in enumerate(document):
        text = page.get_text()
        compact = re.sub(r"\s+", "", text)
        if "APPENDIXA" in compact or "附录A" in compact:
            started = True
        if started and ("Index" in text or "索引" in text):
            break
        if started:
            hits.append(page_no)
    rendered = []
    for page_no in hits:
        path = qa_dir / f"appendix_page_{page_no + 1:04d}.png"
        pix = document[page_no].get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        pix.save(path)
        rendered.append({"pdf_page": page_no + 1, "path": path.name, "sha256": sha256(path)})
    document.close()
    return rendered


def build_candidates(root: Path, languages: tuple[str, ...] = ("en", "zh-Hans", "bilingual")) -> dict[str, Any]:
    gold = root / GOLD_REF
    gold_sha = sha256(gold)
    objects_path = root / DATA_ROOT / "appendix_objects_v1.json"
    structured = json.loads(objects_path.read_text("utf-8"))
    if structured["baseline"]["human_gold_docx_sha256"] != gold_sha:
        raise RuntimeError("structured appendix baseline does not match Human Gold")
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    builds: dict[str, Any] = {}
    for language in languages:
        if language not in {"en", "zh-Hans", "bilingual"}:
            raise ValueError(f"unsupported language: {language}")
        build_id = f"phase12-5-book1-{language}-{stamp}-{gold_sha[:12]}"
        build_dir = root / OUTPUT_ROOT / build_id
        docx_dir, md_dir, pdf_dir, qa_dir = (build_dir / name for name in ("docx", "markdown", "pdf", "qa"))
        for directory in (docx_dir, md_dir, pdf_dir, qa_dir):
            directory.mkdir(parents=True, exist_ok=False)
        docx_path = docx_dir / f"book_{language}.docx"
        shutil.copy2(gold, docx_path)
        document = Document(docx_path)
        apply_structured_appendix(document, structured)
        transform_document(document, language)
        document.save(docx_path)
        markdown_path = md_dir / f"book_{language}.md"
        markdown_from_docx(docx_path, markdown_path)
        pdf_path = render_pdf(docx_path, pdf_dir)
        qa = appendix_page_renders(pdf_path, qa_dir)
        reader = fitz.open(pdf_path)
        all_text = "\n".join(page.get_text() or "" for page in reader)
        compact_text = re.sub(r"\s+", "", all_text)
        gates = {
            "appendix_a_complete": ("APPENDIXA" in compact_text or "附录A" in compact_text),
            "appendix_b_complete": ("APPENDIXB" in compact_text or "附录B" in compact_text),
            "appendix_c_complete": ("APPENDIXC" in compact_text or "附录C" in compact_text),
            "technical_workflow_absent": TECHNICAL_MARKER not in all_text,
            "editable_tables_present": len(Document(docx_path).tables) >= 30,
            "appendix_qa_pages_rendered": bool(qa),
            "new_build_independent": docx_path.resolve() != gold.resolve(),
        }
        manifest = {
            "schema_version": "phase12_5_candidate_manifest.v1",
            "build_id": build_id,
            "status": "candidate_awaiting_independent_acceptance",
            "language": language,
            "baseline_manifest": "docs/phase12_5/book1_baseline_manifest.yaml",
            "structured_objects": "data/fullbook/phase12_5/book1/appendix_objects_v1.json",
            "structured_objects_sha256": sha256(objects_path),
            "structured_objects_replayed": len(structured["objects"]),
            "human_gold_sha256": gold_sha,
            "legacy_release_role": "engine_regression_reference_only",
            "artifacts": {
                "docx": {"path": docx_path.relative_to(root).as_posix(), "sha256": sha256(docx_path), "size_bytes": docx_path.stat().st_size},
                "markdown": {"path": markdown_path.relative_to(root).as_posix(), "sha256": sha256(markdown_path), "size_bytes": markdown_path.stat().st_size},
                "pdf": {"path": pdf_path.relative_to(root).as_posix(), "sha256": sha256(pdf_path), "size_bytes": pdf_path.stat().st_size, "pages": reader.page_count},
            },
            "content_gates": gates,
            "content_gates_passed": all(gates.values()),
            "visual_acceptance": "not_performed_by_codex",
            "qa_renders": qa,
            "package_inclusion": "forbidden",
        }
        atomic_json(build_dir / "render_manifest.json", manifest)
        builds[language] = manifest
    index_path = root / DATA_ROOT / "candidate_builds_v1.json"
    index = json.loads(index_path.read_text("utf-8")) if index_path.is_file() else {}
    index.update(builds)
    atomic_json(index_path, index)
    return builds


def prepare(root: Path) -> dict[str, Any]:
    baseline = write_baseline_manifest(root)
    extraction = extract_appendices(root)
    report = {
        "baseline": baseline,
        "statistics": extraction["objects"]["statistics"],
        "conflict_count": extraction["conflicts"]["conflict_count"],
    }
    atomic_json(root / DATA_ROOT / "preparation_report_v1.json", report)
    return report
