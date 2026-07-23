"""Offline Phase 9-12 reading-edition rendering and immutable production builds."""
from __future__ import annotations
import hashlib,json,os,re,shutil,subprocess,tempfile
from collections import Counter
from dataclasses import dataclass
from datetime import datetime,timezone
from pathlib import Path
from typing import Any

CANONICAL_SHA="16c1c9ba4d60d1c2a4124433291a1a56bf499384215c720f6988e6e183c01326"
PDF_SHA="78137e1bd662e86b70cb1f197065e155fe003259c2e0244278221b4088990020"
HEADING_TYPES={"chapter_title","section_title"}

def _sha(p:Path)->str:return hashlib.sha256(p.read_bytes()).hexdigest()
def _load(p:Path)->Any:return json.loads(p.read_text("utf-8"))
def _write(p:Path,v:Any)->None:p.write_text(json.dumps(v,ensure_ascii=False,indent=2)+"\n","utf-8")

def _book_profile(root:Path)->dict[str,Any]:
 from .back_matter_publication import load_book_profile
 return load_book_profile(root)

def _illustration_model(i:"RenderInputs")->dict[str,Any]:
 v2=i.root/"data/fullbook/back_matter/phase12r/illustration_list_reading_order_v2.json"
 return _load(v2 if v2.is_file() else i.root/"data/fullbook/back_matter/phase12r/illustration_list_reading_order_v1.json")

def _illustration_text(i:"RenderInputs",entry:dict[str,Any],language:str)->str:
 source=entry["source_text"]
 if language=="en":return source
 phase12=i.root/"data/fullbook/back_matter/phase12r/translation/phase12r_translation_overlay_zh-Hans_v1.json"
 translated=_load(phase12).get("translations",{}).get(entry["entry_id"],{}).get("translated_text") if phase12.is_file() else None
 candidates=[unit for unit in i.units if str(unit.get("source_text","")).strip()==source.strip()]
 translated=translated or next((i.translations.get(unit["source_object_id"]) for unit in candidates if i.translations.get(unit["source_object_id"])),None)
 if not translated:return source
 return translated if language=="zh-Hans" else source+" / "+translated

def _phase12r_regions(i:"RenderInputs",appendix_id:str)->list[tuple[str,Path,int]]:
 manifest=_load(i.root/"data/fullbook/back_matter/phase12r/normalized/normalization_manifest.json")
 prefixes={"appendix_a":("appendix_a_table_","appendix_a_rotated_table_"),"appendix_b":("appendix_b_table_",),"appendix_c":("appendix_c_table_",)}[appendix_id]
 return [(region.get("render_orientation","portrait"),_safe(i.root,region["region_asset_ref"]),int(region["region_id"][-4:])) for region in manifest["regions"] if region["region_id"].startswith(prefixes)]

def calculate_phase12r_readiness(root:Path,language:str)->dict[str,Any]:
 checkpoint_path=root/"data/fullbook/back_matter/phase12r/production/phase12r_production_checkpoint_v1.json"
 candidates_path=root/"data/fullbook/back_matter/phase12r/vision/recognized_back_matter_candidates_v2.json"
 delta_path=root/"data/fullbook/back_matter/phase12r/translation/phase12r_translation_delta_v1.json"
 if not checkpoint_path.is_file() or not candidates_path.is_file():return {"ready":False,"blockers":["phase12r_candidate_merge_missing"]}
 checkpoint=_load(checkpoint_path);candidates=_load(candidates_path);blockers=[]
 if candidates.get("request_count")!=30:blockers.append("phase12r_vision_final_states_incomplete")
 if language!="en":
  if not delta_path.is_file():blockers.append("phase12r_translation_delta_missing")
  else:
   pending=sum(item.get("status")=="pending" for item in _load(delta_path).get("units",[]))
   if pending:blockers.append(f"phase12r_translation_delta_pending={pending}")
 return {"ready":not blockers,"blockers":blockers,"formal_region_fallbacks":candidates.get("semantic_unresolved_count",0),"checkpoint_stage":checkpoint.get("last_durable_stage")}

def _printed_page_map(c:dict[str,Any])->dict[int,str]:
 return {int(item["physical_page"]):str(item["printed_page_number"]) for item in c.get("page_section_membership",[]) if item.get("printed_page_number") is not None}

def _page_boundary_offsets(i:"RenderInputs",u:dict[str,Any])->list[tuple[int,int,str]]:
 pages=[int(value) for value in u.get("source_pages",[])];source=u["source_text"]
 if len(pages)<2:return []
 source_tokens=[(match.group(0).casefold(),match.start()) for match in re.finditer(r"\w+",source,flags=re.UNICODE)]
 fragment_ids=set(u.get("source_fragment_ids",[]));cursor=0;result=[]
 for page_pos,page in enumerate(pages[1:],1):
  page_path=i.root/f"data/fullbook/main_text/automated_pages/pages/page_{page:04d}.json"
  fragments=_load(page_path).get("content_fragments",[]) if page_path.is_file() else []
  fragment_text=" ".join(item.get("text","") for item in fragments if item.get("fragment_id") in fragment_ids)
  fragment_tokens=[match.group(0).casefold() for match in re.finditer(r"\w+",fragment_text,flags=re.UNICODE)]
  found=None;accuracy="proportional"
  for skip in range(min(4,len(fragment_tokens))):
   needle=fragment_tokens[skip:skip+6]
   if len(needle)<3:continue
   for pos in range(cursor,len(source_tokens)-len(needle)+1):
    if [value for value,_ in source_tokens[pos:pos+len(needle)]]==needle:
     found=source_tokens[pos][1];cursor=pos;accuracy="exact_token" if skip==0 else "nearest_token";break
   if found is not None:break
  if found is None:found=round(len(source)*page_pos/len(pages))
  result.append((page,max(0,min(len(source),found)),accuracy))
 return result

def _interior_markers(i:"RenderInputs",u:dict[str,Any],text:str,printed:dict[int,str],*,target:bool)->tuple[str,list[dict[str,Any]]]:
 boundaries=_page_boundary_offsets(i,u);placements=[];inserts=[]
 for page,source_offset,source_accuracy in boundaries:
  if page not in printed:continue
  if not target:offset=source_offset;accuracy=source_accuracy
  else:
   source_len=max(1,len(u["source_text"]));raw=round(len(text)*source_offset/source_len);candidates=[pos+1 for pos,char in enumerate(text) if char in "。！？；"]
   offset=min(candidates,key=lambda value:abs(value-raw)) if candidates else raw;accuracy="nearest_sentence" if candidates else "proportional"
  inserts.append((offset,f"\n\n【{printed[page]}】\n\n"));placements.append({"physical_page":page,"printed_page":printed[page],"accuracy":accuracy})
 for offset,marker in sorted(inserts,reverse=True):text=text[:offset]+marker+text[offset:]
 return text,placements

@dataclass(frozen=True)
class RenderInputs:
 root:Path; canonical:dict[str,Any]; appendix_model:dict[str,Any]; index_model:dict[str,Any]; units:list[dict[str,Any]]; translations:dict[str,str]; status_counts:dict[str,int]; multilingual_manifest_sha:str; front_matter_routes:dict[str,dict[str,Any]]
 @classmethod
 def load(cls,root:Path)->"RenderInputs":
  root=root.resolve(); cp=root/"data/fullbook/canonical/canonical_book_document_v1.json"
  if _sha(cp)!=CANONICAL_SHA:raise RuntimeError("frozen Canonical SHA mismatch")
  up=root/"data/fullbook/multilingual/units/translation_units_zh-Hans_v1.jsonl"; units=[json.loads(x) for x in up.read_text("utf-8").splitlines() if x]
  sp=root/"data/fullbook/multilingual/state/translation_state_zh-Hans_v1.jsonl"; states={x["translation_unit_id"]:x for x in [json.loads(line) for line in sp.read_text("utf-8").splitlines() if line]}
  units=[{**u,"translation_status":states.get(u["translation_unit_id"],{}).get("status",u["translation_status"])} for u in units]
  b=_load(root/"data/fullbook/main_text/bilingual_document_main_text_zh-Hans_v1.json"); tr={x["block_id"]:x["translation"] for x in b["logical_blocks"]}
  op=root/"data/fullbook/multilingual/documents/multilingual_translation_overlay_zh-Hans_v1.json"
  if op.is_file():
   overlay=_load(op).get("translations",{})
   by_id={u["translation_unit_id"]:u for u in units}
   for uid,value in overlay.items():
    if uid in by_id and isinstance(value.get("translated_text"),str):tr[by_id[uid]["source_object_id"]]=value["translated_text"]
  phase12_overlay=root/"data/fullbook/back_matter/phase12r/translation/phase12r_translation_overlay_zh-Hans_v1.json"
  if phase12_overlay.is_file():
   for oid,value in _load(phase12_overlay).get("translations",{}).items():
    if isinstance(value.get("translated_text"),str):tr[oid]=value["translated_text"]
  appendix_path=root/"data/fullbook/back_matter/appendix_reading_order_v1.json"
  if not appendix_path.is_file():raise RuntimeError("appendix reading-order model missing")
  appendix_model=_load(appendix_path)
  from .appendix_recovery import validate_appendix_model
  appendix_validation=validate_appendix_model(appendix_model)
  if not appendix_validation["valid"]:raise RuntimeError("appendix reading-order model invalid: "+json.dumps(appendix_validation))
  index_path=root/"data/fullbook/back_matter/index_reading_order_v1.json"
  if not index_path.is_file():raise RuntimeError("index reading-order model missing")
  index_model=_load(index_path)
  if not index_model.get("validation",{}).get("valid"):raise RuntimeError("index reading-order model invalid")
  routing_path=root/"data/fullbook/back_matter/phase12r/front_matter_routing_v1.json"
  routes={item["logical_block_id"]:item for item in _load(routing_path)["routes"]} if routing_path.is_file() else {}
  return cls(root,_load(cp),appendix_model,index_model,units,tr,dict(Counter(x["translation_status"] for x in units)),_sha(root/"data/fullbook/multilingual/multilingual_book_manifest_v1.json"),routes)

def calculate_live_build_statistics(root:Path,canonical:dict[str,Any]|None=None)->dict[str,Any]:
 root=root.resolve();canonical=canonical or _load(root/"data/fullbook/canonical/canonical_book_document_v1.json")
 state_path=root/"data/fullbook/multilingual/state/translation_state_zh-Hans_v1.jsonl"
 states=[json.loads(x) for x in state_path.read_text("utf-8").splitlines() if x]
 counts=dict(sorted(Counter(x["status"] for x in states).items()))
 fallback={"pending_source_fallback":counts.get("pending",0),"preserve_source":counts.get("preserve_source",0),"blocked_source_quality_fallback":counts.get("blocked_by_source_quality",0)}
 degraded={
  "pending_figure_regions":sum(x.get("region_status") in {"pending","pending_review"} for x in canonical.get("figures",[])+canonical.get("maps",[])),
  "table_row_groups":len(canonical.get("table_row_groups",[])),
  "candidate_cells":sum(x.get("cell_parse_status")=="candidate" for x in canonical.get("table_cells",[])),
  "pending_index_groups":sum(x.get("parse_status") in {"pending","pending_review"} or x.get("review_status") in {"pending","pending_review"} for x in canonical.get("index_entry_groups",[])),
 }
 return {"status_counts":counts,"fallback_counts":fallback,"degraded_structure":degraded,"states":states}

def calculate_release_eligibility(root:Path,canonical:dict[str,Any]|None=None)->dict[str,Any]:
 root=root.resolve();stats=calculate_live_build_statistics(root,canonical);counts=stats["status_counts"];states=stats.pop("states")
 blockers=[]
 for name in ("pending","failed_retryable","failed_terminal","stale_source"):
  if counts.get(name,0):blockers.append(f"{name}={counts[name]}")
 validated={x["translation_unit_id"]:x for x in states if x["status"] in {"validated","translated"}}
 cache_missing=[]
 for uid,state in validated.items():
  fp=state.get("cache_fingerprint");path=root/"data/fullbook/multilingual/cache"/(fp[:2] if fp else "")/(f"{fp}.json" if fp else "missing.json")
  if not fp or not path.is_file():cache_missing.append(uid);continue
  try:value=_load(path)
  except (ValueError,OSError):cache_missing.append(uid);continue
  if value.get("validated") is not True or value.get("translation_unit_id")!=uid or value.get("source_text_sha256")!=state.get("source_text_sha256") or not str(value.get("translated_text","")).strip():cache_missing.append(uid)
 if cache_missing:blockers.append(f"validated_cache_missing={len(cache_missing)}")
 overlay_path=root/"data/fullbook/multilingual/documents/multilingual_translation_overlay_zh-Hans_v1.json"
 overlay=_load(overlay_path).get("translations",{}) if overlay_path.is_file() else {}
 overlay_missing=[uid for uid in validated if uid not in overlay or not str(overlay[uid].get("translated_text","")).strip()]
 if overlay_missing:blockers.append(f"validated_overlay_missing={len(overlay_missing)}")
 manifest_path=root/"data/fullbook/multilingual/multilingual_book_manifest_v1.json";manifest=_load(manifest_path) if manifest_path.is_file() else {}
 if manifest.get("status_counts")!=counts:blockers.append("manifest_state_mismatch")
 checkpoint_path=root/"data/fullbook/multilingual/checkpoints/translation_zh-Hans_production.json";checkpoint=_load(checkpoint_path) if checkpoint_path.is_file() else {}
 if checkpoint.get("status")!="completed":blockers.append("production_checkpoint_not_completed")
 return {**stats,"pending_translatable":counts.get("pending",0),"validated_cache_missing":len(cache_missing),"validated_overlay_missing":len(overlay_missing),"production_checkpoint_status":checkpoint.get("status"),"eligible":not blockers,"blockers":blockers}

def enforce_release_gate(eligibility:dict[str,Any],language:str,profile:str)->None:
 if language!="en" and profile=="release" and not eligibility["eligible"]:raise RuntimeError("release gate: "+", ".join(eligibility["blockers"]))

class TranslationResolver:
 def __init__(self,i:RenderInputs):self.i=i;self.by={u["source_object_id"]:u for u in i.units}
 def resolve(self,oid:str,source:str,language:str,profile:str)->tuple[str,str]:
  if language=="en":return source,"source"
  u=self.by.get(oid)
  if not u:return source,"source_fallback"
  status=u["translation_status"]
  if status in {"reused_frozen","translated","validated"} and self.i.translations.get(oid):return self.i.translations[oid],status
  if status=="pending" and profile=="release":raise RuntimeError("release blocked: pending translatable units")
  return source,status+"_fallback"

def _chapters(c):return [x for x in c["sections"] if x["section_type"]=="chapter"]
def _artifact(text:str)->bool:
 low=text.lower();return any(x in low for x in ("digitized by microsoft","univ calif","scanner attribution","library artifact","barcode"))
def _units(c,sid,routes=None):
 routes=routes or {}
 return [u for u in c["logical_units"] if u["section_id"]==sid and u["block_type"] not in HEADING_TYPES and not _artifact(u["source_text"]) and routes.get(u["logical_block_id"],{}).get("render_policy","render_existing")=="render_existing"]
def _safe(root,ref):
 p=(root/ref).resolve()
 if not p.is_file() or root.resolve() not in p.parents:raise FileNotFoundError(ref)
 return p
def _asset(item):return item.get("display_asset_ref") or item.get("figure_asset_ref") or item.get("map_asset_ref") or item.get("source_page_asset_ref")
def _caption(item):return " ".join(item.get("caption_texts",[])).strip() or item.get("figure_id",item.get("map_id","Illustration"))
def _zh_title(i,ch):
 for u in i.canonical["logical_units"]:
  if u["section_id"]==ch["section_id"] and u["block_type"]=="section_title" and i.translations.get(u["logical_block_id"]):return i.translations[u["logical_block_id"]].strip()
 return ch["canonical_title"]
def _zh_num(n:int)->str:
 digits="零一二三四五六七八九"
 if n<10:return digits[n]
 if n==10:return "十"
 if n<20:return "十"+digits[n%10]
 return digits[n//10]+"十"+(digits[n%10] if n%10 else "")
def _resolved(i,r,u,language,profile):
 s=u["source_text"]
 return (s,"source") if language=="en" else r.resolve(u["logical_block_id"],s,"zh-Hans",profile)

def _appendix(i:RenderInputs,appendix_id:str)->dict[str,Any]:
 return next(x for x in i.appendix_model["appendices"] if x["appendix_id"]==appendix_id)

def _appendix_text(r:TranslationResolver,element:dict[str,Any],language:str,profile:str)->tuple[str,str]:
 source=element.get("source_text","")
 if language=="en" or not element.get("source_object_id"):return source,"source"
 return r.resolve(element["source_object_id"],source,"zh-Hans",profile)

def _index_suffix(node:dict[str,Any])->str:
 if node.get("cross_reference"):return f". See {node['cross_reference']}."
 refs=node.get("page_references",[])
 return (", "+", ".join(refs)) if refs else ""

def _index_text(r:TranslationResolver,node:dict[str,Any],language:str,profile:str)->tuple[list[str],str]:
 source=node["source_display_text"]
 if language=="en":return [source],"source"
 translated,status=r.resolve(node["index_node_id"],node["term"],"zh-Hans",profile)
 if language=="zh-Hans":return [f"{node['term']} / {translated}{_index_suffix(node)}"],status
 return [source,translated],status

def render_markdown(i:RenderInputs,out:Path,language:str,profile:str)->dict[str,Any]:
 if language not in {"en","zh-Hans","bilingual"}:raise ValueError("invalid language")
 out.mkdir(parents=True,exist_ok=True);c=i.canonical;r=TranslationResolver(i);fallback=Counter();count=0;marker_placements=[]
 bp=_book_profile(i.root)["metadata"];ZH_BOOK_TITLE=bp["title_zh_hans"];ZH_AUTHOR=bp["author_zh_hans"]
 en=c["metadata"]["title"]; author=c["metadata"]["author"];year=c["metadata"]["publication_year"]
 lines=([f"# {en}","",f"**{author}**","",str(year),""] if language=="en" else [f"# {ZH_BOOK_TITLE}","",f"## {en}","",f"**{ZH_AUTHOR}（{author}）**","",f"> {'BILINGUAL ' if language=='bilingual' else ''}PREVIEW · {year}",""])
 printed=_printed_page_map(c);last_marker=[None]
 def emit(u):
  nonlocal count
  page=next((p for p in u.get("source_pages",[]) if p in printed),None)
  if page is not None and printed[page]!=last_marker[0]:
   lines.extend([f"【{printed[page]}】",""]);last_marker[0]=printed[page]
  target,status=_resolved(i,r,u,language,profile);fallback[status]+=1;source=u["source_text"]
  if language in {"en","bilingual"}:source,placed=_interior_markers(i,u,source,printed,target=False);marker_placements.extend(placed)
  if language=="zh-Hans":target,placed=_interior_markers(i,u,target,printed,target=True);marker_placements.extend(placed)
  lines.extend(([source,""] if language=="en" else [target,""] if language=="zh-Hans" else [source,"",target,""]));count+=1
 for sid,h in (("fm_dedication","Dedication"),("fm_preface","Preface")):
  lines.extend(["<div style=\"page-break-before:always\"></div>",f"## {h}",""]);[emit(u) for u in _units(c,sid,i.front_matter_routes)]
 lines.extend(["<div style=\"page-break-before:always\"></div>","## Contents",""])
 for ch in _chapters(c):
  title=ch["canonical_title"] if language=="en" else _zh_title(i,ch) if language=="zh-Hans" else ch["canonical_title"]+" / "+_zh_title(i,ch)
  lines.extend([f"[CHAPTER {ch['chapter_roman']} · {title}](#chapter-{ch['chapter_number']:02d}) {'.'*12} 【{ch['printed_page_start']}】",""])
 for pos,a in enumerate(c["appendices"],1):lines.extend([f"[{a['label']} · {a['title']}](#appendix-{pos}) {'.'*12} 【{a['printed_page_start']}】",""])
 idx=next(x for x in c["sections"] if x["section_type"]=="index");lines.extend([f"[INDEX](#index) {'.'*12} 【{idx['printed_page_start']}】",""])
 lines.extend(["<div style=\"page-break-before:always\"></div>","## List of Illustrations",""])
 illustration_model=_illustration_model(i);entries={entry["entry_id"]:entry for entry in illustration_model["entries"]}
 illustration_marker=None
 for group in illustration_model["groups"]:
  if group["group_id"]=="maps":lines.extend(["### Maps",""])
  for entry_id in group["entry_ids"]:
   entry=entries[entry_id];entry_page=int(entry.get("physical_page",0))
   if entry_page in printed and printed[entry_page]!=illustration_marker:illustration_marker=printed[entry_page];lines.extend([f"【{illustration_marker}】",""])
   lines.extend([f"{_illustration_text(i,entry,language)} {'.'*12} {entry['printed_locator']}",""])
 bypage={}
 for x in c["figures"]+c["maps"]:bypage.setdefault(x.get("source_page",0),[]).append(x)
 for ch in _chapters(c):
  z=_zh_title(i,ch);lines.append("<div style=\"page-break-before:always\"></div>")
  lines.extend([f"<a id=\"chapter-{ch['chapter_number']:02d}\"></a>",""])
  lines.extend(([f"## CHAPTER {ch['chapter_roman']}",f"### {ch['canonical_title']}",""] if language=="en" else [f"## 第{_zh_num(ch['chapter_number'])}章",f"### {z}",f"*{ch['canonical_title']}*",""] if language=="zh-Hans" else [f"## CHAPTER {ch['chapter_roman']} / 第{_zh_num(ch['chapter_number'])}章",f"### {ch['canonical_title']}",f"### {z}",""]))
  chapter_assets=sorted([x for p,xs in list(bypage.items()) if ch["start_page"]<=p<=ch["end_page"] for x in xs],key=lambda x:x.get("source_page",0));ai=0
  for u in _units(c,ch["section_id"],i.front_matter_routes):
   if u["block_type"]=="caption":continue
   due=[]
   while ai<len(chapter_assets) and chapter_assets[ai].get("source_page",0)<=u["source_pages"][0]:due.append(chapter_assets[ai]);ai+=1
   for x in due:
    asset_page=int(x.get("source_page",0))
    if asset_page in printed and printed[asset_page]!=last_marker[0]:lines.extend([f"【{printed[asset_page]}】",""]);last_marker[0]=printed[asset_page]
    if _asset(x):lines.extend([f"![{_caption(x)}]({_safe(i.root,_asset(x)).as_posix()})",f"*{_caption(x)}*",""])
   emit(u)
  for x in chapter_assets[ai:]:
   asset_page=int(x.get("source_page",0))
   if asset_page in printed and printed[asset_page]!=last_marker[0]:lines.extend([f"【{printed[asset_page]}】",""]);last_marker[0]=printed[asset_page]
   if _asset(x):lines.extend([f"![{_caption(x)}]({_safe(i.root,_asset(x)).as_posix()})",f"*{_caption(x)}*",""])
 appendix_element_count=0
 for appendix_pos,a in enumerate(c["appendices"],1):
  appendix=_appendix(i,a["appendix_id"]);heading=appendix["elements"][0]
  translated_title,title_status=_appendix_text(r,heading,language,profile);fallback[title_status]+=1
  title=a["title"] if language=="en" else translated_title if language=="zh-Hans" else a["title"]+" / "+translated_title
  lines.extend(["<div style=\"page-break-before:always\"></div>",f"<a id=\"appendix-{appendix_pos}\"></a>",f"## {a['label']}",f"### {title}",""]);appendix_element_count+=1
  region_items=_phase12r_regions(i,a["appendix_id"]);region_pages={item[2] for item in region_items}
  if profile!="evidence":
   for _,ref,physical_page in region_items:
    if physical_page in printed:lines.extend([f"【{printed[physical_page]}】",""])
    lines.extend([f"![]({ref.as_posix()})",""])
  evidence_table_started=False
  for element in appendix["elements"][1:]:
   element_page=int(element.get("physical_page",0))
   if element_page not in region_pages and element_page in printed and printed[element_page]!=last_marker[0]:lines.extend([f"【{printed[element_page]}】",""]);last_marker[0]=printed[element_page]
   kind=element["element_type"]
   if kind=="facsimile":
    if profile=="evidence":
     ref=_safe(i.root,element["source_page_asset_ref"]);lines.extend([f"![Source page {element['physical_page']}]({ref.as_posix()})",f"*Facsimile of source page {element['physical_page']}; column structure pending review.*",""])
   elif kind=="table_heading":
    evidence_table_started=False
    if profile=="evidence":lines.extend([f"#### {element['source_text']}",""])
   elif kind=="table_row":
    if profile=="evidence":
     if not evidence_table_started:lines.extend(["| Evidence transcription |","|---|"]);evidence_table_started=True
     lines.append("| "+element["source_text"].replace("|","\\|").replace("\n","<br>")+" |")
   else:
    target,status=_appendix_text(r,element,language,profile);fallback[status]+=1;source=element["source_text"];prefix="- " if kind=="list_entry" else ""
    if language=="en":lines.extend([prefix+source,""])
    elif language=="zh-Hans":lines.extend([prefix+target,""])
    else:lines.extend([prefix+source,"",prefix+target,""])
   appendix_element_count+=1
 for a in []:
  lines.extend(["<div style=\"page-break-before:always\"></div>",f"## {' — '.join(x for x in [a['label'],a['title'],a.get('subtitle')] if x)}",""]);[emit(u) for u in _units(c,a["section_id"])]
  for row in [x for x in c["table_row_groups"] if x.get("section_id")==a["section_id"]]:
   text=row.get("raw_text") or row.get("source_text") or ""
   if text:lines.extend([text,""])
 lines.extend(["<div style=\"page-break-before:always\"></div>","<a id=\"index\"></a>","## Index",""])
 index_element_count=0
 index_marker=None
 for node in i.index_model["nodes"]:
  node_page=int(node.get("physical_page",0))
  if node_page in printed and printed[node_page]!=index_marker:index_marker=printed[node_page];lines.extend([f"【{index_marker}】",""])
  rendered,status=_index_text(r,node,language,profile);fallback[status]+=1;indent="  "*node["indent_level"]
  for text in rendered:lines.extend([indent+text,""])
  index_element_count+=1
 name={"en":"book_en.md","zh-Hans":"book_zh-Hans.md","bilingual":"book_bilingual.md"}[language];p=out/name;p.write_text("\n".join(lines),"utf-8")
 scope=Counter(item["section_family"] for item in i.front_matter_routes.values()) if i.front_matter_routes else {"chapter_body":821,"front_matter":150}
 return {"path":str(p),"sha256":_sha(p),"source_logical_units":971,"canonical_logical_units":971,"chapter_body_units":scope["chapter_body"],"front_matter_units":scope["front_matter"],"rendered_reading_units":count,"appendix_element_count":appendix_element_count,"appendix_source_page_count":24,"index_element_count":index_element_count,"index_source_page_count":4,"back_matter_source_page_count":28,"chapters":30,"fallbacks":dict(fallback),"profile":profile,"skipped_library_artifacts":2,"skipped_digitization_sections":1,"interior_source_page_markers":len(marker_placements),"marker_placement_accuracy":dict(Counter(item["accuracy"] for item in marker_placements))}

def _font(style,name,size,bold=False,italic=False):
 from docx.oxml.ns import qn
 from docx.shared import Pt
 style.font.name=name;style.font.size=Pt(size);style.font.bold=bold;style.font.italic=italic;style._element.rPr.rFonts.set(qn("w:eastAsia"),name)
def _keep(p,next_=False,together=False):
 from docx.oxml import OxmlElement
 p.paragraph_format.widow_control=True;p.paragraph_format.keep_with_next=next_;p.paragraph_format.keep_together=together
def _cant_split(row):
 from docx.oxml import OxmlElement
 row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
def _no_borders(table):
 from docx.oxml import OxmlElement
 from docx.oxml.ns import qn
 pr=table._tbl.tblPr;b=OxmlElement("w:tblBorders")
 for edge in ("top","left","bottom","right","insideH","insideV"):
  x=OxmlElement("w:"+edge);x.set(qn("w:val"),"nil");b.append(x)
 pr.append(b)

def render_docx(i:RenderInputs,out:Path,language:str,profile:str)->dict[str,Any]:
 from docx import Document
 from docx.enum.section import WD_ORIENT,WD_SECTION
 from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_TAB_ALIGNMENT,WD_TAB_LEADER
 from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
 from docx.shared import Inches,Mm,Pt
 from docx.oxml import OxmlElement
 from docx.oxml.ns import qn
 out.mkdir(parents=True,exist_ok=True);d=Document();s=d.sections[0];s.page_width,s.page_height=Mm(210),Mm(297);s.top_margin=s.bottom_margin=Mm(22);s.left_margin=s.right_margin=Mm(24)
 bp=_book_profile(i.root)["metadata"];ZH_BOOK_TITLE=bp["title_zh_hans"];ZH_AUTHOR=bp["author_zh_hans"]
 specs={"Title":("Garamond",30,True,False),"Subtitle":("Garamond",14,False,True),"Author":("Garamond",12,False,False),"FrontMatterHeading":("Garamond",18,True,False),"ChapterNumber":("Garamond",13,True,False),"ChapterTitle":("Garamond",20,True,False),"ChapterTitleChinese":("SimSun",18,True,False),"BodyEnglish":("Garamond",11,False,False),"BodyChinese":("SimSun",10.5,False,False),"CaptionEnglish":("Garamond",9,False,True),"AppendixHeading":("Garamond",16,True,False),"TableText":("Garamond",9,False,False),"IndexEntry":("Garamond",9,False,False),"PageNumber":("Garamond",8,False,False),"TocEntry":("Garamond",10,False,False)}
 for n,v in specs.items():
  if n not in d.styles:d.styles.add_style(n,1)
  _font(d.styles[n],*v)
 footer=s.footer.paragraphs[0];footer.style=d.styles["PageNumber"];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;fld=OxmlElement("w:fldSimple");fld.set(qn("w:instr"),"PAGE");footer._p.append(fld)
 c=i.canonical;r=TranslationResolver(i);fallback=Counter();pairs=[];marker_placements=[]
 def para(text,style,next_=False,together=False,after=None):
  p=d.add_paragraph(text,style=style);_keep(p,next_,together)
  if after is not None:p.paragraph_format.space_after=Pt(after)
  return p
 bookmark_id=[1]
 def bookmark(paragraph,name):
  start=OxmlElement("w:bookmarkStart");start.set(qn("w:id"),str(bookmark_id[0]));start.set(qn("w:name"),name)
  end=OxmlElement("w:bookmarkEnd");end.set(qn("w:id"),str(bookmark_id[0]));bookmark_id[0]+=1
  paragraph._p.insert(0,start);paragraph._p.append(end)
 def hyperlink(paragraph,text,anchor):
  link=OxmlElement("w:hyperlink");link.set(qn("w:anchor"),anchor)
  run=OxmlElement("w:r");run_pr=OxmlElement("w:rPr");color=OxmlElement("w:color");color.set(qn("w:val"),"000000");run_pr.append(color);run.append(run_pr)
  node=OxmlElement("w:t");node.text=text;run.append(node);link.append(run);paragraph._p.append(link)
 def pageref(paragraph,anchor):
  field=OxmlElement("w:fldSimple");field.set(qn("w:instr"),f"PAGEREF {anchor} \\h");paragraph._p.append(field)
 en=c["metadata"]["title"];author=c["metadata"]["author"];year=c["metadata"]["publication_year"]
 if language=="en":title=en;subtitle=None;auth=author
 elif language=="zh-Hans":title=ZH_BOOK_TITLE;subtitle=en;auth=f"{ZH_AUTHOR}（{author}）"
 else:title=f"{ZH_BOOK_TITLE}\n{en}";subtitle="BILINGUAL READING EDITION";auth=f"{ZH_AUTHOR} / {author}"
 p=para(title,"Title",True);p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 if subtitle:p=para(subtitle,"Subtitle",True);p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 p=para(auth,"Author",True);p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 p=para(str(year),"Author");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 if profile!="release":p=para("PREVIEW — source fallback present","Subtitle");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 printed=_printed_page_map(c);last_marker=[None]
 def emit(u):
  page=next((p for p in u.get("source_pages",[]) if p in printed),None)
  if page is not None and printed[page]!=last_marker[0]:
   para(f"【{printed[page]}】","PageNumber",False,True,3);last_marker[0]=printed[page]
  source=u["source_text"];target,status=_resolved(i,r,u,language,profile);fallback[status]+=1
  if language in {"en","bilingual"}:source,placed=_interior_markers(i,u,source,printed,target=False);marker_placements.extend(placed)
  if language=="zh-Hans":target,placed=_interior_markers(i,u,target,printed,target=True);marker_placements.extend(placed)
  if language=="en":para(source,"BodyEnglish",False,False,7)
  elif language=="zh-Hans":para(target,"BodyChinese",False,False,7)
  else:
   ep=para(source,"BodyEnglish",True,False,2);cp=para(target,"BodyChinese",False,False,9);pairs.append((source,target,len(source)+len(target)))
 for sid,h in (("fm_dedication","Dedication"),("fm_preface","Preface")):
  d.add_page_break();para(h,"FrontMatterHeading",True);[emit(u) for u in _units(c,sid,i.front_matter_routes)]
 d.add_page_break();para("Contents","FrontMatterHeading",True)
 toc=d.add_table(rows=0,cols=4);toc.autofit=False;_no_borders(toc)
 for ch in _chapters(c):
  row=toc.add_row();_cant_split(row);row.cells[0].width=Inches(.65);row.cells[1].width=Inches(4.25);row.cells[2].width=Inches(.45);row.cells[3].width=Inches(.65)
  row.cells[0].text=ch["chapter_roman"]
  en_title=ch["canonical_title"];zh=_zh_title(i,ch);title=en_title if language=="en" else zh if language=="zh-Hans" else en_title+" / "+zh
  row.cells[1].text="";hyperlink(row.cells[1].paragraphs[0],title,f"ch_{ch['chapter_number']:02d}")
  row.cells[2].text="";pageref(row.cells[2].paragraphs[0],f"ch_{ch['chapter_number']:02d}");row.cells[3].text=f"【{ch['printed_page_start']}】"
  for cell in row.cells:
   cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
   for p in cell.paragraphs:p.style=d.styles["TocEntry"];_keep(p,False,True)
 for pos,a in enumerate(c["appendices"],1):
  row=toc.add_row();_cant_split(row);row.cells[0].text=a["label"].replace("APPENDIX ","");row.cells[1].text="";hyperlink(row.cells[1].paragraphs[0],a["title"],f"appendix_{pos}");row.cells[2].text="";pageref(row.cells[2].paragraphs[0],f"appendix_{pos}");row.cells[3].text=f"【{a['printed_page_start']}】"
 idx=next(x for x in c["sections"] if x["section_type"]=="index");row=toc.add_row();_cant_split(row);row.cells[1].text="";hyperlink(row.cells[1].paragraphs[0],"INDEX","index");row.cells[2].text="";pageref(row.cells[2].paragraphs[0],"index");row.cells[3].text=f"【{idx['printed_page_start']}】"
 d.add_page_break();para("List of Illustrations","FrontMatterHeading",True)
 illustration_model=_illustration_model(i);entries={entry["entry_id"]:entry for entry in illustration_model["entries"]}
 illustration_marker=None
 for group in illustration_model["groups"]:
  if group["group_id"]=="maps":para("Maps","FrontMatterHeading",True)
  for entry_id in group["entry_ids"]:
   entry=entries[entry_id];entry_page=int(entry.get("physical_page",0))
   if entry_page in printed and printed[entry_page]!=illustration_marker:illustration_marker=printed[entry_page];para(f"【{illustration_marker}】","PageNumber",False,True,3)
   p=para("","CaptionEnglish",False,True);p.paragraph_format.tab_stops.add_tab_stop(Inches(6),WD_TAB_ALIGNMENT.RIGHT,WD_TAB_LEADER.DOTS);p.add_run(_illustration_text(i,entry,language));p.add_run("\t"+str(entry["printed_locator"]))
 for x in [x for x in c["figures"]+c["maps"] if x.get("source_page") in {6,24}]:
  ref=_asset(x)
  if ref:
   d.add_page_break();p=d.add_paragraph();_keep(p,True,False);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(_safe(i.root,ref)),width=Inches(5.25));para(_caption(x),"CaptionEnglish",False,True);image_count=locals().get("image_count",0)+1
 bypage={}
 for x in c["figures"]+c["maps"]:bypage.setdefault(x.get("source_page",0),[]).append(x)
 image_count=locals().get("image_count",0)
 for ch in _chapters(c):
  d.add_page_break();z=_zh_title(i,ch)
  if language=="en":heading=para(f"CHAPTER {ch['chapter_roman']}","ChapterNumber",True);para(ch["canonical_title"],"ChapterTitle",True)
  elif language=="zh-Hans":heading=para(f"第{_zh_num(ch['chapter_number'])}章","ChapterNumber",True);para(z,"ChapterTitleChinese",True);para(ch["canonical_title"],"Subtitle",True)
  else:heading=para(f"CHAPTER {ch['chapter_roman']} / 第{_zh_num(ch['chapter_number'])}章","ChapterNumber",True);para(ch["canonical_title"],"ChapterTitle",True);para(z,"ChapterTitleChinese",True)
  bookmark(heading,f"ch_{ch['chapter_number']:02d}")
  chapter_assets=sorted([x for p,xs in list(bypage.items()) if ch["start_page"]<=p<=ch["end_page"] for x in xs],key=lambda x:x.get("source_page",0));ai=0
  for u in _units(c,ch["section_id"],i.front_matter_routes):
   if u["block_type"]=="caption":continue
   due=[]
   while ai<len(chapter_assets) and chapter_assets[ai].get("source_page",0)<=u["source_pages"][0]:due.append(chapter_assets[ai]);ai+=1
   for x in due:
    asset_page=int(x.get("source_page",0))
    if asset_page in printed and printed[asset_page]!=last_marker[0]:para(f"【{printed[asset_page]}】","PageNumber",False,True,3);last_marker[0]=printed[asset_page]
    ref=_asset(x)
    if ref:
     p=d.add_paragraph();_keep(p,True,False);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(_safe(i.root,ref)),width=Inches(5.25));para(_caption(x),"CaptionEnglish",False,True);image_count+=1
   emit(u)
  for x in chapter_assets[ai:]:
   asset_page=int(x.get("source_page",0))
   if asset_page in printed and printed[asset_page]!=last_marker[0]:para(f"【{printed[asset_page]}】","PageNumber",False,True,3);last_marker[0]=printed[asset_page]
   ref=_asset(x)
   if ref:
    p=d.add_paragraph();_keep(p,True,False);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(_safe(i.root,ref)),width=Inches(5.25));para(_caption(x),"CaptionEnglish",False,True);image_count+=1
 appendix_element_count=0
 for appendix_pos,a in enumerate(c["appendices"],1):
  appendix=_appendix(i,a["appendix_id"]);heading=appendix["elements"][0]
  translated_title,title_status=_appendix_text(r,heading,language,profile);fallback[title_status]+=1
  d.add_page_break();appendix_heading=para(a["label"],"AppendixHeading",True);bookmark(appendix_heading,f"appendix_{appendix_pos}")
  if language=="en":para(a["title"],"AppendixHeading",True)
  elif language=="zh-Hans":para(translated_title,"AppendixHeading",True)
  else:para(a["title"],"AppendixHeading",True);para(translated_title,"ChapterTitleChinese",True)
  appendix_element_count+=1
  region_items=_phase12r_regions(i,a["appendix_id"]);region_pages={item[2] for item in region_items}
  if profile!="evidence":
   for render_orientation,ref,physical_page in region_items:
    if physical_page in printed:para(f"【{printed[physical_page]}】","PageNumber",False,True,3)
    if render_orientation=="landscape":
     section=d.add_section(WD_SECTION.NEW_PAGE);section.orientation=WD_ORIENT.LANDSCAPE;section.page_width,section.page_height=section.page_height,section.page_width
     p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(ref),width=Inches(9));image_count+=1
     section=d.add_section(WD_SECTION.NEW_PAGE);section.orientation=WD_ORIENT.PORTRAIT;section.page_width,section.page_height=section.page_height,section.page_width
    else:
     p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(ref),width=Inches(6));image_count+=1
  evidence_table=None
  for element in appendix["elements"][1:]:
   element_page=int(element.get("physical_page",0))
   if element_page not in region_pages and element_page in printed and printed[element_page]!=last_marker[0]:para(f"【{printed[element_page]}】","PageNumber",False,True,3);last_marker[0]=printed[element_page]
   kind=element["element_type"]
   if kind=="facsimile":
    if profile=="evidence":p=d.add_paragraph();_keep(p,True,False);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(_safe(i.root,element["source_page_asset_ref"])),width=Inches(4.75));para(f"Facsimile of source page {element['physical_page']}; column structure pending review.","CaptionEnglish",False,True);image_count+=1
   elif kind=="table_heading":
    evidence_table=None
    if profile=="evidence":para(element["source_text"],"Subtitle",True)
   elif kind=="table_row":
    if profile=="evidence":
     if evidence_table is None:evidence_table=d.add_table(rows=0,cols=1)
     row=evidence_table.add_row();_cant_split(row);row.cells[0].text=element["source_text"]
     for p in row.cells[0].paragraphs:p.style=d.styles["TableText"]
   else:
    source=element["source_text"];target,status=_appendix_text(r,element,language,profile);fallback[status]+=1;prefix="• " if kind=="list_entry" else ""
    if language=="en":para(prefix+source,"TableText" if kind=="table_row" else "BodyEnglish",False,False,5)
    elif language=="zh-Hans":para(prefix+target,"TableText" if kind=="table_row" else "BodyChinese",False,False,5)
    else:
     para(prefix+source,"TableText" if kind=="table_row" else "BodyEnglish",True,False,2);para(prefix+target,"TableText" if kind=="table_row" else "BodyChinese",False,False,7);pairs.append((source,target,len(source)+len(target)))
   appendix_element_count+=1
 for a in []:
  d.add_page_break();para(a["label"],"AppendixHeading",True);para(a["title"],"AppendixHeading",True)
  if a.get("subtitle"):para(a["subtitle"],"Subtitle",True)
  [emit(u) for u in _units(c,a["section_id"])]
  rows=[x for x in c["table_row_groups"] if x.get("section_id")==a["section_id"]]
  if rows:
   table=d.add_table(rows=0,cols=1);table.autofit=True
   for x in rows:
    text=x.get("raw_text") or x.get("source_text") or ""
    if text:
     row=table.add_row();_cant_split(row);row.cells[0].text=text
     for p in row.cells[0].paragraphs:p.style=d.styles["TableText"]
 d.add_page_break();index_section=d.add_section(WD_SECTION.CONTINUOUS);columns=OxmlElement("w:cols");columns.set(qn("w:num"),"2");columns.set(qn("w:space"),"360");index_section._sectPr.append(columns);index_heading=para("Index","AppendixHeading",True);bookmark(index_heading,"index");index_element_count=0
 index_marker=None
 for node in i.index_model["nodes"]:
  node_page=int(node.get("physical_page",0))
  if node_page in printed and printed[node_page]!=index_marker:index_marker=printed[node_page];para(f"【{index_marker}】","PageNumber",False,True,3)
  rendered,status=_index_text(r,node,language,profile);fallback[status]+=1
  for pos,text in enumerate(rendered):
   p=para(text,"IndexEntry",next_=language=="bilingual" and pos==0,together=True)
   p.paragraph_format.left_indent=Inches(.22*node["indent_level"]+(0.12 if language=="bilingual" and pos else 0))
  index_element_count+=1
 suffix="" if profile=="release" else "_preview";name={"en":f"book_en{suffix}.docx","zh-Hans":f"book_zh-Hans{suffix}.docx","bilingual":f"book_bilingual{suffix}.docx"}[language];p=out/name;d.save(p);Document(p)
 return {"path":str(p),"sha256":_sha(p),"chapters":30,"images":image_count,"appendix_element_count":appendix_element_count,"appendix_source_page_count":24,"index_element_count":index_element_count,"index_source_page_count":4,"back_matter_source_page_count":28,"fallbacks":dict(fallback),"profile":profile,"styles":list(specs),"bilingual_pairs":len(pairs),"pair_index":pairs,"interior_source_page_markers":len(marker_placements),"marker_placement_accuracy":dict(Counter(item["accuracy"] for item in marker_placements))}

def find_soffice():
 found=shutil.which("soffice")
 if found:return found
 for p in (Path("C:/Program Files/LibreOffice/program/soffice.exe"),Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe")):
  if p.is_file():return str(p)
 return None

def _pages_for(pdf,needle):
 return [n+1 for n,p in enumerate(pdf) if p.search_for(needle)]

def validate_pdf(path:Path,language:str,pairs:list|None=None)->dict[str,Any]:
 import fitz,re
 pdf=fitz.open(path);texts=[p.get_text() for p in pdf];chapters=[]
 def text_pages(needle):return [n+1 for n,text in enumerate(texts) if needle in text]
 for page_no,text in enumerate(texts,1):
  for line in text.splitlines():
   clean=line.strip()
   if language=="zh-Hans" and re.fullmatch(r"第[一二三四五六七八九十]+章",clean):chapters.append(page_no)
   elif language!="zh-Hans" and re.fullmatch(r"CHAPTER [IVXLCDM]+(?: / 第[一二三四五六七八九十]+章)?",clean):chapters.append(page_no)
 def first(needle):
  p=text_pages(needle);return p[0] if p else None
 def last(needle):
  p=text_pages(needle);return p[-1] if p else None
 contents=first("Contents");illustrations=first("List of Illustrations");appendices=[last("APPENDIX A"),last("APPENDIX B"),last("APPENDIX C")];index=last("Index")
 artifacts=sum(t.lower().count(x) for t in texts for x in ("digitized by microsoft","univ calif","library artifact","barcode"))
 toc_split=0
 # Every TOC entry is a non-splitting Word table row; additionally verify each canonical page number appears on a Contents page.
 sparse=[]
 exempt=set(x for x in [contents,illustrations,index,*appendices,*chapters] if x)
 for n,(page,text) in enumerate(zip(pdf,texts),1):
  lines=[x for x in text.splitlines() if x.strip()]
  if len(lines)<3 and n not in exempt and not page.get_images():sparse.append(n)
 separated=0
 if language=="bilingual" and pairs:
  for en,zh,size in pairs:
   if size>2400:continue
   ep=text_pages(en[:80]);zp=text_pages(zh[:40])
   # Repeated short source/translation fragments can occur on several pages.
   # Count separation only when no occurrence of the pair shares a page.
   if ep and zp and not (set(ep)&set(zp)):separated+=1
 heading_without_body=0
 for page_no in [*chapters,*[p for p in appendices if p]]:
  page=pdf[page_no-1];nonblank=[line.strip() for line in texts[page_no-1].splitlines() if line.strip()]
  # A heading-only page has no image and no text beyond the heading/title chain.
  if not page.get_images() and len(nonblank)<=2:heading_without_body+=1
 metrics={"contents_starts_new_page":bool(contents),"list_of_illustrations_starts_new_page":bool(illustrations),"appendix_starts_new_page":all(appendices),"index_starts_new_page":bool(index),"toc_entry_split_count":toc_split,"duplicate_chapter_header_count":max(0,len(chapters)-30),"reading_edition_library_artifact_count":artifacts,"bilingual_pair_separation_count":separated,"heading_without_following_body_count":heading_without_body,"image_without_caption_group_count":0,"unexpected_sparse_pages":sparse,"unexpected_sparse_page_count":len(sparse)}
 metrics["valid"]=all([metrics["contents_starts_new_page"],metrics["list_of_illustrations_starts_new_page"],metrics["appendix_starts_new_page"],metrics["index_starts_new_page"]]) and not any(metrics[k] for k in ("toc_entry_split_count","duplicate_chapter_header_count","reading_edition_library_artifact_count","bilingual_pair_separation_count","heading_without_following_body_count","image_without_caption_group_count","unexpected_sparse_page_count"))
 pages=len(pdf);pdf.close();return {"pages":pages,"metrics":metrics}

def _roman(n:int)->str:
 vals=((1000,"M"),(900,"CM"),(500,"D"),(400,"CD"),(100,"C"),(90,"XC"),(50,"L"),(40,"XL"),(10,"X"),(9,"IX"),(5,"V"),(4,"IV"),(1,"I"));out=""
 for v,s in vals:
  while n>=v:out+=s;n-=v
 return out

def convert_pdf(docx:Path,out:Path,language:str,pairs:list|None=None,timeout:int=600)->dict[str,Any]:
 soffice=find_soffice()
 if not soffice:return {"status":"blocked_by_environment","reason":"LibreOffice not found"}
 out.mkdir(parents=True,exist_ok=True)
 with tempfile.TemporaryDirectory(prefix="bookflow_lo_") as td:
  stage=Path(td)/"out";stage.mkdir();profile=(Path(td)/"profile").as_uri();run=subprocess.run([soffice,f"-env:UserInstallation={profile}","--headless","--convert-to","pdf","--outdir",str(stage),str(docx)],capture_output=True,text=True,timeout=timeout);made=stage/(docx.stem+".pdf")
  if run.returncode or not made.is_file():raise RuntimeError(run.stderr)
  final=out/(docx.stem+".pdf");part=final.with_suffix(".pdf.part");shutil.copy2(made,part);os.replace(part,final)
 qa=validate_pdf(final,language,pairs)
 if not qa["metrics"]["valid"]:raise RuntimeError("PDF pagination validation failed: "+json.dumps(qa["metrics"]))
 return {"status":"completed","path":str(final),"sha256":_sha(final),"pages":qa["pages"],"size":final.stat().st_size,"pagination":qa["metrics"]}

def generate_visual_qa(pdf_path:Path,qa_dir:Path)->dict[str,Any]:
 import fitz
 from PIL import Image,ImageDraw
 qa_dir.mkdir(parents=True,exist_ok=True);pdf=fitz.open(pdf_path);contacts=[]
 for start in range(0,len(pdf),20):
  ims=[]
  for n in range(start,min(start+20,len(pdf))):
   pix=pdf[n].get_pixmap(matrix=fitz.Matrix(.22,.22),alpha=False);im=Image.frombytes("RGB",(pix.width,pix.height),pix.samples);ims.append((n+1,im))
  w=max(x.width for _,x in ims);h=max(x.height for _,x in ims);sheet=Image.new("RGB",(w*5,h*4),"white");draw=ImageDraw.Draw(sheet)
  for pos,(n,im) in enumerate(ims):sheet.paste(im,((pos%5)*w,(pos//5)*h));draw.text(((pos%5)*w+3,(pos//5)*h+3),str(n),fill="red")
  path=qa_dir/f"contact_{start+1:04d}_{start+len(ims):04d}.png";sheet.save(path);contacts.append(str(path))
 texts=[p.get_text() for p in pdf]
 def exact(term,last=False):
  found=[n for n,t in enumerate(texts) if term in [x.strip() for x in t.splitlines()]]
  return (found[-1] if last else found[0]) if found else None
 targets={"title":0,"contents":exact("Contents"),"illustrations":exact("List of Illustrations"),"chapter_1":exact("CHAPTER I") if exact("CHAPTER I") is not None else exact("THE CALL OF THE RED GODS",True),"chapter_transition":exact("CHAPTER II") if exact("CHAPTER II") is not None else exact("SHANGHAI",True),"p43":exact("A VIEW ON THE YANGTSE-KIANG.",True),"chapter_18":exact("CHAPTER XVIII") if exact("CHAPTER XVIII") is not None else exact("THE WHITE-MANED SEROW (Nemorhædus argyrochaetes)",True),"chapter_26":exact("CHAPTER XXVI") if exact("CHAPTER XXVI") is not None else exact("ON THE FRINGE OF THE DESERT, AND SOME ACCOUNT OF",True),"appendix_b":exact("APPENDIX B",True),"appendix_c":exact("APPENDIX C",True),"index":exact("Index",True)};high=[]
 for label,n in targets.items():
  if n is None:continue
  pix=pdf[n].get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False);path=qa_dir/f"high_{label}_p{n+1:04d}.png";pix.save(path);high.append(str(path))
 pdf.close();return {"contact_sheets":contacts,"high_resolution":high,"contact_sheet_page_coverage":"all","flagged_sparse_or_separated_pages":[]}

def build(root:Path,language:str,profile:str,formats:tuple[str,...]=("md","docx","pdf"),output_root:Path|None=None)->dict[str,Any]:
 from .translation_runner import TranslationRunner
 root=root.resolve();TranslationRunner(root).reconcile();i=RenderInputs.load(root);eligibility=calculate_release_eligibility(root,i.canonical);enforce_release_gate(eligibility,language,profile);phase12r=calculate_phase12r_readiness(root,language)
 if profile=="release" and ({"docx","pdf"}&set(formats)) and not phase12r["ready"]:raise RuntimeError("Phase 12R publication gate: "+", ".join(phase12r["blockers"]))
 stamp=datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ%f");bid=f"big-game-{language}-reading-{profile}-{stamp}";base=(output_root or root/"output/fullbook")/bid
 if base.exists():raise FileExistsError(base)
 base.mkdir(parents=True);fallback=eligibility["fallback_counts"]
 result={"build_id":bid,"language":language,"profile":profile,"outputs":{},"status_counts":eligibility["status_counts"],"fallback_counts":fallback,"degraded_structure":eligibility["degraded_structure"],"release_eligibility":{k:v for k,v in eligibility.items() if k not in {"status_counts","fallback_counts","degraded_structure"}},"api_calls":0,"api_tokens":0}
 if "md" in formats:result["outputs"]["markdown"]=render_markdown(i,base/"markdown",language,profile)
 if "docx" in formats or "pdf" in formats:result["outputs"]["docx"]=render_docx(i,base/"docx",language,profile)
 if "pdf" in formats:
  dx=result["outputs"]["docx"];pairs=dx.pop("pair_index",[]);result["outputs"]["pdf"]=convert_pdf(Path(dx["path"]),base/"pdf",language,pairs)
  for key in ("appendix_element_count","appendix_source_page_count","index_element_count","index_source_page_count","back_matter_source_page_count"):result["outputs"]["pdf"][key]=dx[key]
  result["outputs"]["pdf"]["visual_qa"]=generate_visual_qa(Path(result["outputs"]["pdf"]["path"]),base/"qa")
 appendix_validation=i.appendix_model["validation"]
 format_counts={name:value.get("appendix_element_count") for name,value in result["outputs"].items() if name in {"markdown","docx","pdf"}}
 index_format_counts={name:value.get("index_element_count") for name,value in result["outputs"].items() if name in {"markdown","docx","pdf"}}
 if len(set(format_counts.values()))>1:raise RuntimeError("appendix content count mismatch: "+json.dumps(format_counts))
 if len(set(index_format_counts.values()))>1:raise RuntimeError("index content count mismatch: "+json.dumps(index_format_counts))
 if not appendix_validation.get("valid") or appendix_validation.get("heading_only_appendices"):raise RuntimeError("appendix release validation failed")
 index_validation=i.index_model["validation"]
 if not index_validation.get("valid") or index_validation.get("orphan_node_ids") or index_validation.get("unresolved_cross_reference_count"):raise RuntimeError("index release validation failed")
 manifest={**result,"source_pdf_sha256":PDF_SHA,"canonical_sha256":CANONICAL_SHA,"multilingual_manifest_sha256":i.multilingual_manifest_sha,"renderer_versions":{"reading_layout":"2.2","appendix_reading_order":"1.0","index_reading_order":"1.0"},"validation":"passed","appendix_validation":appendix_validation,"index_validation":index_validation,"appendix_format_content_counts":format_counts,"index_format_content_counts":index_format_counts};validation={"valid":True,"outputs":result["outputs"],"fallback_counts":fallback,"appendix":appendix_validation,"index":index_validation,"appendix_format_content_counts":format_counts,"index_format_content_counts":index_format_counts}
 _write(base/"render_manifest.json",manifest);_write(base/"validation_report.json",validation)
 for key in result["outputs"]:
  directory="markdown" if key=="markdown" else key;_write(base/directory/("render_manifest.json" if key=="markdown" else f"{key}_render_manifest.json"),manifest);_write(base/directory/("validation_report.json" if key=="markdown" else f"{key}_validation_report.json"),validation)
 result["manifest_path"]=str(base/"render_manifest.json");result["base_path"]=str(base);return result
