"""Phase 3B-S: append page 12, close 11->12, and export source-only English."""

from __future__ import annotations

import base64
import json
import os
import re
import time
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Literal

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel, ConfigDict, Field, field_validator, model_validator

from .io_utils import atomic_write_json, atomic_write_jsonl, load_json, sha256_file, sha256_text, stable_hash
from .page_pipeline import build_context
from .paths import ProjectSettings, project_root, resolve_project_path
from .phase2a1 import NormalizationEvent, VisionNormalizedPageV11, normalize_preserved_response_v11
from .phase2b2_schemas import AutomatedLogicalBlock, AutomatedPageRecord, SourceFragment
from .schemas import PageRecord
from .secret_store import api_key_status, load_api_key
from .vision_provider import ProviderResponse, ZhipuOpenAICompatibleProvider


class Phase3BManifestPage(BaseModel):
    model_config = ConfigDict(extra="forbid")

    pdf_page: int = Field(ge=1, le=12)
    image_path: str
    image_sha256: str
    image_width: int = Field(gt=0)
    image_height: int = Field(gt=0)
    cache_reused: bool
    reused_from_document_id: str | None = None
    visual_request_fingerprint: str | None = None
    normalized_visual_path: str | None = None


class Phase3BManifest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    schema_version: str = "1.0"
    document_id: str
    derived_pdf: str
    derived_pdf_sha256: str
    sample11_pdf: str
    sample11_sha256: str
    page12_pdf: str
    page12_pdf_sha256: str
    page_count: Literal[12]
    dpi: Literal[200]
    color_mode: Literal["RGB"]
    image_format: Literal["png"]
    pages: list[Phase3BManifestPage]
    reused_visual_pages: list[int]
    newly_rendered_pages: list[int]
    manifest_path: str
    created_at: datetime

    @model_validator(mode="after")
    def validate_pages(self) -> "Phase3BManifest":
        numbers = [item.pdf_page for item in self.pages]
        if numbers != list(range(1, 13)):
            raise ValueError("Phase 3B-S manifest pages must be exactly 1 through 12")
        if self.reused_visual_pages != list(range(1, 12)):
            raise ValueError("Phase 3B-S must reuse visual results for pages 1 through 11")
        if self.newly_rendered_pages != [12]:
            raise ValueError("Phase 3B-S may render only page 12 as new work")
        return self


class Phase3BPreflight(BaseModel):
    model_config = ConfigDict(extra="forbid")

    document_id: str
    actual_page_count: int
    reused_visual_pages: int
    single_calls_expected: int
    pair_calls_expected: int
    triple_calls_allowed: int
    maximum_new_calls: int
    calls_already_started: int
    remaining_real_calls: int
    automatic_retry: bool
    estimated_token_range: str
    estimated_public_price_cny: float
    maximum_cash_cost_cny: float
    api_returns_actual_cash_charge: bool = False
    api_returns_resource_balance: bool = False
    api_key_set: bool
    deepseek_calls: int = 0
    translation_calls: int = 0
    full_pdf_processed: bool = False
    blockers: list[str] = Field(default_factory=list)
    ready: bool
    api_called: bool = False


class Phase3BCallResult(BaseModel):
    model_config = ConfigDict(extra="forbid")

    category: Literal["single", "pair"]
    item_id: str
    request_fingerprint: str
    api_calls_started: int
    cache_hits: int
    retries: Literal[0] = 0
    triple_calls: Literal[0] = 0
    status: Literal["completed", "cached", "dry_run"]
    raw_response_path: str | None
    normalized_output_path: str | None
    usage_path: str | None


class Phase3BBoundaryObservation(BaseModel):
    """Model evidence only. It intentionally has no final join field."""

    model_config = ConfigDict(extra="forbid")

    schema_version: str
    boundary_id: Literal["boundary_p0011_p0012"]
    document_id: str
    previous_page: Literal[11]
    next_page: Literal[12]
    previous_visible_tail: str
    next_visible_head: str
    visible_trailing_hyphen: bool
    suspected_word_continuation: bool | None
    suspected_sentence_continuation: bool | None
    suspected_paragraph_continuation: bool | None
    structural_break: Literal[
        "none", "paragraph_break", "section_break", "chapter_break", "illustration_break", "unknown"
    ]
    header_footer_interference: bool | None
    possible_omission: bool | None
    possible_duplication: bool | None
    evidence: list[str]
    conflicting_evidence: list[str]
    confidence_raw: str | float | None = None
    confidence_label: str | None = None
    confidence_score: float | None = Field(default=None, ge=0, le=1)
    normalization_events: list[NormalizationEvent] = Field(default_factory=list)
    status: Literal["observed", "uncertain"]

    @model_validator(mode="before")
    @classmethod
    def normalize_confidence_compatibly(cls, value: object) -> object:
        if not isinstance(value, dict):
            return value
        data = dict(value)
        if "confidence" in data:
            legacy = data.pop("confidence")
            if "confidence_raw" in data and data["confidence_raw"] != legacy:
                raise ValueError("confidence and confidence_raw disagree")
            data.setdefault("confidence_raw", legacy)
        raw = data.get("confidence_raw")
        events = list(data.get("normalization_events") or [])
        if isinstance(raw, bool):
            raise ValueError("Boolean confidence is not a valid confidence value")
        if isinstance(raw, (int, float)):
            numeric = float(raw)
            if not 0 <= numeric <= 1:
                raise ValueError("Numeric confidence must be between 0 and 1")
            data.setdefault("confidence_label", None)
            data.setdefault("confidence_score", numeric)
        elif isinstance(raw, str):
            label = raw.strip()
            if not label:
                label = None
            data.setdefault("confidence_label", label)
            data.setdefault("confidence_score", None)
            already_recorded = any(
                isinstance(item, dict)
                and item.get("field") == "confidence"
                and item.get("action") == "string_label_preserved_without_numeric_score"
                for item in events
            )
            if label is not None and not already_recorded:
                events.append(
                    {
                        "field": "confidence",
                        "action": "string_label_preserved_without_numeric_score",
                        "reason": "The provider returned a qualitative label; preserving it is lossless, while inventing a numeric score is not.",
                        "original_type": "string",
                        "original_value": raw,
                        "normalized_value": {
                            "confidence_raw": raw,
                            "confidence_label": label,
                            "confidence_score": None,
                        },
                        "requires_review": False,
                    }
                )
        elif raw is None:
            data.setdefault("confidence_label", None)
            data.setdefault("confidence_score", None)
        else:
            raise ValueError("Unsupported confidence value type")
        data["normalization_events"] = events
        return data

    @field_validator("evidence", "conflicting_evidence", mode="before")
    @classmethod
    def listify_evidence(cls, value: object) -> object:
        if value is None or value == "":
            return []
        if isinstance(value, str):
            return [value]
        return value


class Phase3BBoundaryDecision(BaseModel):
    model_config = ConfigDict(extra="forbid")

    schema_version: str = "3B-S-1.1"
    boundary_id: Literal["boundary_p0011_p0012"] = "boundary_p0011_p0012"
    document_id: str
    previous_page: Literal[11] = 11
    next_page: Literal[12] = 12
    observed_word_continuation: bool | None
    validated_word_continuation: bool | None
    word_continuation: bool | None
    sentence_continuation: bool | None
    paragraph_continuation: bool | None
    structural_break: str
    join_operation: Literal["insert_space", "remove_layout_hyphen", "preserve_lexical_hyphen", "no_join", "unresolved"]
    joiner: str
    resolved_text: str
    resolution_method: str
    supporting_evidence: list[str]
    conflicting_evidence: list[str]
    resolution_reason: str
    auto_resolution_status: Literal["resolved_pair", "unresolved"]
    model_observation_path: str | None = None
    translation_called: Literal[False] = False

    @model_validator(mode="after")
    def fail_closed(self) -> "Phase3BBoundaryDecision":
        if self.auto_resolution_status == "unresolved":
            if self.join_operation != "unresolved" or self.joiner or self.resolved_text:
                raise ValueError("Unresolved boundary cannot contain a join or reconstructed text")
        if self.word_continuation != self.validated_word_continuation:
            raise ValueError("word_continuation must reflect the deterministic validated result")
        return self


class Phase3BSourceLogicalBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    schema_version: str = "3B-S-1.0"
    logical_block_id: str
    document_id: str
    block_type: str
    source_pages: list[int]
    source_block_ids: list[str]
    source_fragment_ids: list[str]
    source_text: str
    cross_page: bool
    sentence_complete: bool
    paragraph_complete: bool
    coverage_complete: bool
    completeness_status: Literal["complete", "incomplete_start", "incomplete_end", "unresolved_boundary"]
    unresolved_boundaries: list[str]
    source_exportable: bool
    translation_ready: bool
    chapter_id: str | None = None
    section_id: str | None = None

    @model_validator(mode="after")
    def validate_gate(self) -> "Phase3BSourceLogicalBlock":
        if self.source_pages != sorted(set(self.source_pages)):
            raise ValueError("source_pages must be ordered and unique")
        if self.cross_page != (len(self.source_pages) > 1):
            raise ValueError("cross_page must match source_pages")
        complete = all(
            [
                self.sentence_complete,
                self.paragraph_complete,
                self.coverage_complete,
                not self.unresolved_boundaries,
                bool(self.source_text.strip()),
                self.completeness_status == "complete",
            ]
        )
        if (self.source_exportable or self.translation_ready) and not complete:
            raise ValueError("Incomplete source cannot bypass the source/translation gates")
        return self


class Phase3BSourceAudit(BaseModel):
    model_config = ConfigDict(extra="forbid")

    expected_pages: Literal[12] = 12
    processed_pages: int
    sample11_hash_unchanged: bool
    reused_visual_page_count: int
    first_ten_boundaries_reused: bool
    boundary_11_12_closed: bool
    page12_end_complete: bool
    open_boundary_12_13_exists: bool
    expected_fragment_count: int
    referenced_fragment_count: int
    unused_fragment_ids: list[str]
    duplicate_fragment_ids: list[str]
    header_footer_page_number_clean: bool
    strict_passed: bool
    blockers: list[str]


class Phase3BSourceDocument(BaseModel):
    model_config = ConfigDict(extra="forbid")

    schema_version: str = "source-document-1.0"
    document_id: str
    source_pdf: str
    source_pdf_sha256: str
    source_inputs: list[dict[str, str]]
    page_count: Literal[12]
    entries: list[Phase3BSourceLogicalBlock]
    audit: Phase3BSourceAudit
    strict_export_ready: bool
    strict_blockers: list[str]
    glm_single_calls: int
    glm_pair_calls: int
    glm_total_calls: int
    deepseek_calls: Literal[0] = 0
    translation_calls: Literal[0] = 0
    created_at: datetime


class Phase3BExportResult(BaseModel):
    model_config = ConfigDict(extra="forbid")

    source_json_sha256: str
    diagnostic_markdown_path: str
    diagnostic_docx_path: str
    strict_exported: bool
    final_markdown_path: str | None
    final_docx_path: str | None
    blockers: list[str]


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _data_root(settings: ProjectSettings, root: Path) -> Path:
    return resolve_project_path(settings.phase3b_data_directory, root=root)


def validate_phase3b_scope(pdf_path: str | Path, settings: ProjectSettings, *, root: Path | None = None) -> Path:
    root = (root or project_root()).resolve()
    candidate = resolve_project_path(pdf_path, root=root)
    protected = resolve_project_path(settings.source_pdf, root=root)
    allowed = {
        resolve_project_path(settings.sample_pdf, root=root),
        resolve_project_path(settings.phase3b_page12_pdf, root=root),
        resolve_project_path(settings.phase3b_source_sample_pdf, root=root),
    }
    if candidate == protected or candidate not in allowed:
        raise PermissionError("Phase 3B-S accepts only the 11-page sample, supplied page 12, or derived 12-page sample")
    return candidate


def _latest_visual_path(page_record: AutomatedPageRecord) -> Path:
    path = Path(page_record.legacy_normalized_path)
    if not path.is_file():
        raise FileNotFoundError(f"Missing cached normalized visual result: {path}")
    item = VisionNormalizedPageV11.model_validate(load_json(path))
    raw = Path(item.raw_response_path)
    if not raw.is_file():
        raise FileNotFoundError(f"Missing cached raw visual response: {raw}")
    return path


def _merge_sample_pdfs(sample11: Path, page12: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_suffix(".phase3b.tmp.pdf")
    if temporary.exists():
        temporary.unlink()
    merged = fitz.open()
    try:
        with fitz.open(sample11) as first:
            merged.insert_pdf(first)
        with fitz.open(page12) as second:
            merged.insert_pdf(second)
        if merged.page_count != 12:
            raise RuntimeError(f"Derived sample must contain 12 pages, found {merged.page_count}")
        merged.set_metadata({"title": "12-page source validation sample", "producer": "bookflow Phase 3B-S"})
        merged.save(temporary, garbage=4, deflate=True)
    finally:
        merged.close()
    os.replace(temporary, output)


def _render_page12(source: Path, output: Path) -> tuple[str, int, int]:
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_suffix(".tmp.png")
    with fitz.open(source) as document:
        if document.page_count != 12:
            raise ValueError("Page 12 rendering requires an actual 12-page derived sample")
        pixmap = document[11].get_pixmap(matrix=fitz.Matrix(200 / 72, 200 / 72), colorspace=fitz.csRGB, alpha=False)
        pixmap.save(temporary)
        width, height = pixmap.width, pixmap.height
    os.replace(temporary, output)
    return sha256_file(output), width, height


def _page12_nonblank(path: Path) -> bool:
    with fitz.open(path) as document:
        if document.page_count != 1:
            return False
        page = document[0]
        if page.get_text("text").strip():
            return True
        pixmap = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5), colorspace=fitz.csGRAY, alpha=False)
        samples = pixmap.samples
        return bool(samples and min(samples) < 245)


def create_phase3b_sample(settings: ProjectSettings, *, root: Path | None = None) -> Phase3BManifest:
    root = (root or project_root()).resolve()
    sample11 = validate_phase3b_scope(settings.sample_pdf, settings, root=root)
    page12 = validate_phase3b_scope(settings.phase3b_page12_pdf, settings, root=root)
    derived = validate_phase3b_scope(settings.phase3b_source_sample_pdf, settings, root=root)
    if not sample11.is_file() or not page12.is_file():
        raise FileNotFoundError("The 11-page sample and supplied page-12 PDF must both exist")
    with fitz.open(sample11) as first:
        if first.page_count != 11:
            raise ValueError(f"Configured sample must contain 11 pages, found {first.page_count}")
    if not _page12_nonblank(page12):
        raise ValueError("The supplied page-12 PDF must contain one nonblank page")

    sample11_hash = sha256_file(sample11)
    page12_hash = sha256_file(page12)
    data_root = _data_root(settings, root)
    manifest_path = data_root / "manifests" / "sample_12_pages.manifest.json"
    existing_ok = False
    if derived.is_file() and manifest_path.is_file():
        try:
            previous = Phase3BManifest.model_validate(load_json(manifest_path))
            with fitz.open(derived) as document:
                existing_ok = (
                    document.page_count == 12
                    and previous.sample11_sha256 == sample11_hash
                    and previous.page12_pdf_sha256 == page12_hash
                    and previous.derived_pdf_sha256 == sha256_file(derived)
                )
        except Exception:
            existing_ok = False
    if not existing_ok:
        _merge_sample_pdfs(sample11, page12, derived)
    with fitz.open(derived) as document:
        if document.page_count != 12:
            raise RuntimeError("Derived sample validation failed: page count is not 12")
    derived_hash = sha256_file(derived)
    document_id = f"doc_{derived_hash[:16].lower()}"

    old_pages_path = resolve_project_path(settings.automated_page_directory, root=root) / "sample_11_pages.pages.jsonl"
    old_pages = [AutomatedPageRecord.model_validate_json(line) for line in old_pages_path.read_text(encoding="utf-8").splitlines() if line.strip()]
    if len(old_pages) != 11:
        raise RuntimeError("All eleven Phase 2B.2 automated page records are required")
    old_context = build_context(
        sample11,
        settings,
        pages=list(range(1, 12)),
        dpi=200,
        color_mode="RGB",
        image_format="png",
        root=root,
    )
    pages: list[Phase3BManifestPage] = []
    for page_number, automated in enumerate(old_pages, start=1):
        record_path = Path(old_context.record_directory) / f"page_{page_number:04d}.json"
        record = PageRecord.model_validate(load_json(record_path))
        image_path = Path(record.image_path)
        if not image_path.is_file() or sha256_file(image_path) != record.image_sha256:
            raise RuntimeError(f"Cached image hash failed for page {page_number}")
        visual_path = _latest_visual_path(automated)
        pages.append(
            Phase3BManifestPage(
                pdf_page=page_number,
                image_path=str(image_path.resolve()),
                image_sha256=record.image_sha256,
                image_width=record.image_width,
                image_height=record.image_height,
                cache_reused=True,
                reused_from_document_id=automated.document_id,
                visual_request_fingerprint=visual_path.stem,
                normalized_visual_path=str(visual_path.resolve()),
            )
        )
    image12 = data_root / "pages" / document_id / "page_0012.png"
    if image12.is_file():
        image_hash = sha256_file(image12)
        with fitz.open(image12) as image_document:
            page = image_document[0]
            image_width, image_height = int(page.rect.width), int(page.rect.height)
        # PyMuPDF treats PNG as a one-page document; dimensions are pixels here.
    else:
        image_hash, image_width, image_height = _render_page12(derived, image12)
    if image12.is_file() and ("image_width" not in locals() or image_width <= 0):
        image_hash, image_width, image_height = _render_page12(derived, image12)
    pages.append(
        Phase3BManifestPage(
            pdf_page=12,
            image_path=str(image12.resolve()),
            image_sha256=image_hash,
            image_width=image_width,
            image_height=image_height,
            cache_reused=False,
        )
    )
    manifest = Phase3BManifest(
        document_id=document_id,
        derived_pdf=str(derived.resolve()),
        derived_pdf_sha256=derived_hash,
        sample11_pdf=str(sample11.resolve()),
        sample11_sha256=sample11_hash,
        page12_pdf=str(page12.resolve()),
        page12_pdf_sha256=page12_hash,
        page_count=12,
        dpi=200,
        color_mode="RGB",
        image_format="png",
        pages=pages,
        reused_visual_pages=list(range(1, 12)),
        newly_rendered_pages=[12],
        manifest_path=str(manifest_path.resolve()),
        created_at=_now(),
    )
    atomic_write_json(manifest_path, manifest)
    return manifest


def _call_paths(settings: ProjectSettings, root: Path, category: str, fingerprint: str) -> dict[str, Path]:
    base = _data_root(settings, root) / "calls" / category / fingerprint
    return {
        "raw": base / "raw.json",
        "normalized": base / "normalized.json",
        "validation": base / "validation.json",
        "usage": base / "usage.json",
        "request": base / "request.json",
        "cache": base / "cache.json",
    }


def _ledger_path(settings: ProjectSettings, root: Path) -> Path:
    return _data_root(settings, root) / "phase3b_call_ledger.json"


def _load_ledger(settings: ProjectSettings, root: Path) -> dict[str, Any]:
    path = _ledger_path(settings, root)
    if path.is_file():
        value = load_json(path)
        if isinstance(value, dict):
            return value
        raise RuntimeError("Phase 3B-S call ledger is invalid")
    return {
        "schema_version": "1.0",
        "phase": "3B-S",
        "limits": {"single": 1, "pair": 1, "triple": 0, "total": 2},
        "started": {"single": 0, "pair": 0, "triple": 0, "total": 0},
        "attempts": [],
    }


def _reserve_call(settings: ProjectSettings, root: Path, category: Literal["single", "pair"], item_id: str, fingerprint: str) -> tuple[dict[str, Any], int]:
    ledger = _load_ledger(settings, root)
    if ledger["started"][category] >= ledger["limits"][category] or ledger["started"]["total"] >= ledger["limits"]["total"]:
        raise RuntimeError(f"Phase 3B-S {category} or total real-call limit is exhausted")
    ledger["started"][category] += 1
    ledger["started"]["total"] += 1
    ledger["attempts"].append(
        {
            "category": category,
            "item_id": item_id,
            "request_fingerprint": fingerprint,
            "status": "in_flight",
            "automatic_retry": False,
            "retries": 0,
            "started_at": _now().isoformat(),
        }
    )
    atomic_write_json(_ledger_path(settings, root), ledger)
    return ledger, len(ledger["attempts"]) - 1


def _finish_call(settings: ProjectSettings, root: Path, ledger: dict[str, Any], index: int, *, status: str, request_id: str | None = None) -> None:
    ledger["attempts"][index].update({"status": status, "request_id": request_id, "ended_at": _now().isoformat()})
    atomic_write_json(_ledger_path(settings, root), ledger)


def _provider_instance(settings: ProjectSettings, root: Path, provider: Any | None) -> tuple[Any, str]:
    if provider is not None:
        return provider, "injected_test_provider"
    secret, key_name = load_api_key(settings, root)
    return (
        ZhipuOpenAICompatibleProvider(
            api_key=secret,
            base_url=settings.vision_base_url,
            timeout_seconds=settings.vision_request_timeout_seconds,
        ),
        key_name,
    )


def _image_data_url(path: Path) -> str:
    return "data:image/png;base64," + base64.b64encode(path.read_bytes()).decode("ascii")


def _invoke_once(
    settings: ProjectSettings,
    root: Path,
    *,
    category: Literal["single", "pair"],
    item_id: str,
    fingerprint: str,
    prompt: str,
    context_message: str,
    image_paths: list[Path],
    paths: dict[str, Path],
    provider: Any | None,
) -> ProviderResponse:
    client, key_name = _provider_instance(settings, root, provider)
    ledger, index = _reserve_call(settings, root, category, item_id, fingerprint)
    atomic_write_json(
        paths["request"],
        {
            "phase": "3B-S",
            "category": category,
            "item_id": item_id,
            "request_fingerprint": fingerprint,
            "provider": settings.vision_provider,
            "model": settings.vision_model,
            "api_key_env": key_name,
            "api_key_recorded": False,
            "image_paths": [str(path.resolve()) for path in image_paths],
            "image_sha256": [sha256_file(path) for path in image_paths],
            "automatic_retry": False,
            "retries": 0,
            "translation_disabled": True,
            "started_at": _now().isoformat(),
        },
    )
    started = time.perf_counter()
    try:
        response = client.transcribe_images(
            model=settings.vision_model,
            prompt=prompt,
            context_message=context_message,
            image_data_urls=[_image_data_url(path) for path in image_paths],
            max_output_tokens=settings.vision_max_output_tokens,
            temperature=settings.vision_temperature,
            do_sample=settings.vision_do_sample,
            thinking_mode=settings.vision_thinking_mode,
            response_format_json_object=settings.vision_response_format_json_object,
        )
    except Exception as exc:
        atomic_write_json(
            paths["raw"],
            {
                "record_type": "provider_error_response",
                "phase": "3B-S",
                "category": category,
                "item_id": item_id,
                "request_fingerprint": fingerprint,
                "api_called": True,
                "error_type": type(exc).__name__,
                "error_message_recorded": False,
                "automatic_retry": False,
                "retries": 0,
                "recorded_at": _now().isoformat(),
            },
        )
        _finish_call(settings, root, ledger, index, status="failed")
        raise RuntimeError(f"Phase 3B-S {category} call failed; automatic retry is disabled") from exc
    atomic_write_json(
        paths["raw"],
        {
            "record_type": "raw_provider_response",
            "phase": "3B-S",
            "category": category,
            "item_id": item_id,
            "request_fingerprint": fingerprint,
            "provider": settings.vision_provider,
            "model": settings.vision_model,
            "api_called": True,
            "received_at": _now().isoformat(),
            "response": response.raw_response,
        },
    )
    atomic_write_json(
        paths["usage"],
        {
            "phase": "3B-S",
            "category": category,
            "item_id": item_id,
            "request_fingerprint": fingerprint,
            "request_id": response.request_id,
            "usage": response.usage,
            "elapsed_seconds": round(time.perf_counter() - started, 3),
            "api_returned_resource_balance": False,
            "api_returned_cash_charge": False,
            "cash_charge_cny": None,
        },
    )
    _finish_call(settings, root, ledger, index, status="response_saved", request_id=response.request_id)
    return response


def _page12_fingerprint(settings: ProjectSettings, manifest: Phase3BManifest, prompt: str) -> str:
    page = manifest.pages[11]
    return stable_hash(
        {
            "phase": "3B-S",
            "category": "single",
            "document_id": manifest.document_id,
            "pdf_page": 12,
            "image_sha256": page.image_sha256,
            "provider": settings.vision_provider,
            "model": settings.vision_model,
            "base_url": settings.vision_base_url.rstrip("/"),
            "prompt_version": settings.phase2b_page_prompt_version,
            "prompt_sha256": sha256_text(prompt),
            "schema_version": settings.vision_normalized_schema_version,
            "max_output_tokens": settings.vision_max_output_tokens,
            "thinking_mode": settings.vision_thinking_mode,
        }
    )


def run_phase3b_page12(
    settings: ProjectSettings,
    manifest: Phase3BManifest,
    *,
    provider: Any | None = None,
    allow_api: bool = False,
    confirmed: bool = False,
    root: Path | None = None,
) -> Phase3BCallResult:
    root = (root or project_root()).resolve()
    if manifest.page_count != 12 or manifest.pages[11].pdf_page != 12:
        raise ValueError("A valid Phase 3B-S manifest is required")
    prompt_path = resolve_project_path(settings.phase2b_page_prompt_path, root=root)
    prompt = prompt_path.read_text(encoding="utf-8")
    fingerprint = _page12_fingerprint(settings, manifest, prompt)
    paths = _call_paths(settings, root, "single", fingerprint)
    if all(paths[key].is_file() for key in ("raw", "normalized", "cache")):
        VisionNormalizedPageV11.model_validate(load_json(paths["normalized"]))
        return Phase3BCallResult(
            category="single", item_id="page_0012", request_fingerprint=fingerprint,
            api_calls_started=0, cache_hits=1, status="cached",
            raw_response_path=str(paths["raw"]), normalized_output_path=str(paths["normalized"]),
            usage_path=str(paths["usage"]) if paths["usage"].is_file() else None,
        )
    if not allow_api:
        return Phase3BCallResult(
            category="single", item_id="page_0012", request_fingerprint=fingerprint,
            api_calls_started=0, cache_hits=0, status="dry_run",
            raw_response_path=None, normalized_output_path=None, usage_path=None,
        )
    if not confirmed:
        raise PermissionError("Explicit Phase 3B-S confirmation is required")
    response = _invoke_once(
        settings,
        root,
        category="single",
        item_id="page_0012",
        fingerprint=fingerprint,
        prompt=prompt,
        context_message=(
            f"document_id={manifest.document_id}; pdf_page=12; provider={settings.vision_provider}; "
            f"model={settings.vision_model}; normalized_schema_version=1.1. "
            "Transcribe every visible English character on this single page. Do not translate or infer unseen text."
        ),
        image_paths=[Path(manifest.pages[11].image_path)],
        paths=paths,
        provider=provider,
    )
    try:
        normalized = normalize_preserved_response_v11(paths["raw"], paths["normalized"], force_adjacent_review=set())
        if normalized.pdf_page != 12:
            raise ValueError("Page transcription returned the wrong pdf_page")
    except Exception as exc:
        raise RuntimeError("Page 12 response was saved but could not be normalized; no retry is allowed") from exc
    atomic_write_json(
        paths["cache"],
        {
            "request_fingerprint": fingerprint,
            "request_id": response.request_id,
            "raw_response_path": str(paths["raw"].resolve()),
            "normalized_output_path": str(paths["normalized"].resolve()),
            "api_call_completed": True,
            "authoritative": False,
            "translation_ready": False,
        },
    )
    return Phase3BCallResult(
        category="single", item_id="page_0012", request_fingerprint=fingerprint,
        api_calls_started=1, cache_hits=0, status="completed",
        raw_response_path=str(paths["raw"]), normalized_output_path=str(paths["normalized"]),
        usage_path=str(paths["usage"]),
    )


def _extract_json_content(content: str) -> dict[str, Any]:
    text = content.strip()
    if text.startswith("```"):
        lines = text.splitlines()[1:]
        if lines and lines[-1].strip() == "```":
            lines.pop()
        text = "\n".join(lines).strip()
    value = json.loads(text)
    if not isinstance(value, dict):
        raise ValueError("Model output must be one JSON object")
    return value


def _latest_old_page11(settings: ProjectSettings, root: Path) -> VisionNormalizedPageV11:
    automated = AutomatedPageRecord.model_validate(load_json(resolve_project_path(settings.automated_page_directory, root=root) / "pages/page_0011.json"))
    return VisionNormalizedPageV11.model_validate(load_json(_latest_visual_path(automated)))


def _body_edge(result: VisionNormalizedPageV11, *, first: bool) -> tuple[str, str]:
    blocks = [block for block in sorted(result.blocks, key=lambda value: value.order) if block.block_type == "body"]
    if not blocks:
        raise ValueError("A body block is required for boundary review")
    block = blocks[0] if first else blocks[-1]
    return block.block_id, block.text


def run_phase3b_pair_11_12(
    settings: ProjectSettings,
    manifest: Phase3BManifest,
    page12_normalized_path: str | Path,
    *,
    provider: Any | None = None,
    allow_api: bool = False,
    confirmed: bool = False,
    root: Path | None = None,
) -> Phase3BCallResult:
    root = (root or project_root()).resolve()
    page11 = _latest_old_page11(settings, root)
    page12 = VisionNormalizedPageV11.model_validate(load_json(page12_normalized_path))
    if page12.pdf_page != 12:
        raise ValueError("The page-12 normalized result is invalid")
    _, tail = _body_edge(page11, first=False)
    _, head = _body_edge(page12, first=True)
    prompt_path = resolve_project_path(settings.automated_boundary_prompt_path, root=root)
    prompt = prompt_path.read_text(encoding="utf-8")
    fingerprint = stable_hash(
        {
            "phase": "3B-S",
            "category": "pair",
            "document_id": manifest.document_id,
            "previous_page": 11,
            "next_page": 12,
            "previous_image_sha256": manifest.pages[10].image_sha256,
            "next_image_sha256": manifest.pages[11].image_sha256,
            "previous_tail": tail[-500:],
            "next_head": head[:500],
            "prompt_sha256": sha256_text(prompt),
            "schema": "phase3b-boundary-observation-1.0",
            "model": settings.vision_model,
        }
    )
    paths = _call_paths(settings, root, "pair", fingerprint)
    if all(paths[key].is_file() for key in ("raw", "normalized", "cache")):
        Phase3BBoundaryObservation.model_validate(load_json(paths["normalized"]))
        return Phase3BCallResult(
            category="pair", item_id="boundary_p0011_p0012", request_fingerprint=fingerprint,
            api_calls_started=0, cache_hits=1, status="cached",
            raw_response_path=str(paths["raw"]), normalized_output_path=str(paths["normalized"]),
            usage_path=str(paths["usage"]) if paths["usage"].is_file() else None,
        )
    if paths["raw"].is_file() and not paths["normalized"].is_file():
        raw_record = load_json(paths["raw"])
        try:
            raw_content = raw_record["response"]["choices"][0]["message"]["content"]
            observation = Phase3BBoundaryObservation.model_validate(_extract_json_content(raw_content))
            if observation.document_id != manifest.document_id:
                raise ValueError("Boundary observation document_id does not match")
            atomic_write_json(paths["normalized"], observation)
            atomic_write_json(
                paths["cache"],
                {
                    "request_fingerprint": fingerprint,
                    "raw_response_path": str(paths["raw"].resolve()),
                    "normalized_output_path": str(paths["normalized"].resolve()),
                    "offline_normalization_recovery": True,
                    "model_evidence_only": True,
                    "python_final_join_required": True,
                    "triple_calls": 0,
                    "retries": 0,
                    "translation_calls": 0,
                },
            )
            return Phase3BCallResult(
                category="pair", item_id="boundary_p0011_p0012", request_fingerprint=fingerprint,
                api_calls_started=0, cache_hits=1, status="cached",
                raw_response_path=str(paths["raw"]), normalized_output_path=str(paths["normalized"]),
                usage_path=str(paths["usage"]) if paths["usage"].is_file() else None,
            )
        except Exception as exc:
            atomic_write_json(
                paths["validation"],
                {
                    "phase": "3B-S",
                    "category": "pair",
                    "boundary_id": "boundary_p0011_p0012",
                    "request_fingerprint": fingerprint,
                    "status": "schema_failed",
                    "error_type": type(exc).__name__,
                    "error_summary": str(exc),
                    "raw_response_path": str(paths["raw"].resolve()),
                    "raw_response_modified": False,
                    "automatic_retry": False,
                    "triple_called": False,
                    "translation_called": False,
                    "recorded_at": _now().isoformat(),
                },
            )
            raise RuntimeError("11->12 saved pair response is unusable; retry and triple are forbidden") from exc
    if not allow_api:
        return Phase3BCallResult(
            category="pair", item_id="boundary_p0011_p0012", request_fingerprint=fingerprint,
            api_calls_started=0, cache_hits=0, status="dry_run",
            raw_response_path=None, normalized_output_path=None, usage_path=None,
        )
    if not confirmed:
        raise PermissionError("Explicit Phase 3B-S confirmation is required")
    schema_instruction = (
        "Return one JSON object with exactly these fields: schema_version, boundary_id, document_id, "
        "previous_page, next_page, previous_visible_tail, next_visible_head, visible_trailing_hyphen, "
        "suspected_word_continuation, suspected_sentence_continuation, suspected_paragraph_continuation, "
        "structural_break, header_footer_interference, possible_omission, possible_duplication, evidence, "
        "conflicting_evidence, confidence, status. Use boundary_id=boundary_p0011_p0012, pages 11 and 12, "
        "structural_break from none/paragraph_break/section_break/chapter_break/illustration_break/unknown, "
        "and status observed or uncertain. Do not return a join operation or reconstructed text."
    )
    response = _invoke_once(
        settings,
        root,
        category="pair",
        item_id="boundary_p0011_p0012",
        fingerprint=fingerprint,
        prompt=prompt,
        context_message=(
            f"document_id={manifest.document_id}; previous_page=11; next_page=12. {schema_instruction} "
            f"Saved page-11 tail evidence: {tail[-500:]} Saved page-12 head evidence: {head[:500]}"
        ),
        image_paths=[Path(manifest.pages[10].image_path), Path(manifest.pages[11].image_path)],
        paths=paths,
        provider=provider,
    )
    try:
        observation = Phase3BBoundaryObservation.model_validate(_extract_json_content(response.content))
        if observation.document_id != manifest.document_id:
            raise ValueError("Boundary observation document_id does not match")
        atomic_write_json(paths["normalized"], observation)
    except Exception as exc:
        atomic_write_json(
            paths["validation"],
            {
                "phase": "3B-S",
                "category": "pair",
                "boundary_id": "boundary_p0011_p0012",
                "request_fingerprint": fingerprint,
                "status": "schema_failed",
                "error_type": type(exc).__name__,
                "error_summary": str(exc),
                "raw_response_path": str(paths["raw"].resolve()),
                "raw_response_modified": False,
                "automatic_retry": False,
                "triple_called": False,
                "translation_called": False,
                "recorded_at": _now().isoformat(),
            },
        )
        raise RuntimeError("11->12 pair response was saved but is unusable; retry and triple are forbidden") from exc
    atomic_write_json(
        paths["cache"],
        {
            "request_fingerprint": fingerprint,
            "request_id": response.request_id,
            "raw_response_path": str(paths["raw"].resolve()),
            "normalized_output_path": str(paths["normalized"].resolve()),
            "api_call_completed": True,
            "model_evidence_only": True,
            "python_final_join_required": True,
            "triple_calls": 0,
            "retries": 0,
            "translation_calls": 0,
        },
    )
    return Phase3BCallResult(
        category="pair", item_id="boundary_p0011_p0012", request_fingerprint=fingerprint,
        api_calls_started=1, cache_hits=0, status="completed",
        raw_response_path=str(paths["raw"]), normalized_output_path=str(paths["normalized"]),
        usage_path=str(paths["usage"]),
    )


def _comparison_text(text: str) -> str:
    # This comparator is evidence-only: ignore line wrapping and typographic
    # quote variants without changing any preserved source text.
    return re.sub(r"[^a-z0-9]+", "", text.casefold())


def resolve_phase3b_boundary(
    *,
    previous_text: str,
    next_text: str,
    observation: Phase3BBoundaryObservation,
    text_layer_tail: str,
    text_layer_head: str,
    document_id: str | None = None,
    observation_path: str | None = None,
) -> Phase3BBoundaryDecision:
    evidence = list(observation.evidence)
    conflicts = list(observation.conflicting_evidence)
    tail_match = _comparison_text(previous_text)[-80:] in _comparison_text(text_layer_tail)[-160:]
    head_match = _comparison_text(next_text)[:80] in _comparison_text(text_layer_head)[:160]
    if tail_match:
        evidence.append("PDF text layer supports the visible page-11 tail.")
    else:
        conflicts.append("PDF text layer does not safely confirm the page-11 tail.")
    if head_match:
        evidence.append("PDF text layer supports the visible page-12 head.")
    else:
        conflicts.append("PDF text layer does not safely confirm the page-12 head.")
    safe_continuation = all(
        [
            observation.status == "observed",
            observation.structural_break == "none",
            observation.suspected_sentence_continuation is True,
            observation.suspected_paragraph_continuation is True,
            observation.header_footer_interference is not True,
            observation.possible_omission is not True,
            observation.possible_duplication is not True,
            tail_match,
            head_match,
        ]
    )
    if not safe_continuation:
        return Phase3BBoundaryDecision(
            document_id=document_id or observation.document_id,
            observed_word_continuation=observation.suspected_word_continuation,
            validated_word_continuation=None,
            word_continuation=None,
            sentence_continuation=observation.suspected_sentence_continuation,
            paragraph_continuation=observation.suspected_paragraph_continuation,
            structural_break=observation.structural_break,
            join_operation="unresolved",
            joiner="",
            resolved_text="",
            resolution_method="pair_visual_plus_text_layer_plus_deterministic_rules",
            supporting_evidence=evidence,
            conflicting_evidence=conflicts,
            resolution_reason="Evidence did not satisfy every fail-closed boundary gate.",
            auto_resolution_status="unresolved",
            model_observation_path=observation_path,
        )
    left = previous_text.rstrip()
    right = next_text.lstrip()
    visible_tail_hyphen = left.endswith("-")
    if observation.visible_trailing_hyphen != visible_tail_hyphen:
        conflicts.append(
            "The pair model's trailing-hyphen observation conflicts with the preserved page-tail transcription."
        )
    left_word_match = re.search(r"([A-Za-z]+)$", left[:-1] if visible_tail_hyphen else left)
    right_word_match = re.match(r"([A-Za-z]+)", right)
    left_word = left_word_match.group(1) if left_word_match else None
    right_word = right_word_match.group(1) if right_word_match else None
    validated_word_continuation: bool | None
    joiner: str
    if visible_tail_hyphen and left_word and right_word and observation.suspected_word_continuation is True:
        operation = "remove_layout_hyphen"
        resolved = left[:-1] + right
        validated_word_continuation = True
        joiner = ""
        reason = "The preserved page tail ends in a visible hyphen and both sides support one divided word."
    elif not visible_tail_hyphen and left_word and right_word:
        operation = "insert_space"
        resolved = left + " " + right
        validated_word_continuation = False
        joiner = " "
        rejected = left_word + right_word
        evidence.append(
            f"The preserved left edge ends with the complete word '{left_word}', has no trailing hyphen, and the right edge begins with the independent word '{right_word}'."
        )
        if observation.suspected_word_continuation is True:
            conflicts.append(
                f"The model suspected one divided word, but deterministic lexical and hyphen checks reject '{rejected}'."
            )
        reason = (
            "The pages continue the same sentence and paragraph, but the visible source contains two complete words; "
            "insert exactly one space."
        )
    else:
        return Phase3BBoundaryDecision(
            document_id=document_id or observation.document_id,
            observed_word_continuation=observation.suspected_word_continuation,
            validated_word_continuation=None,
            word_continuation=None,
            sentence_continuation=True,
            paragraph_continuation=True,
            structural_break="none",
            join_operation="unresolved",
            joiner="",
            resolved_text="",
            resolution_method="pair_visual_plus_text_layer_plus_deterministic_rules",
            supporting_evidence=evidence,
            conflicting_evidence=conflicts,
            resolution_reason="The word-boundary evidence is still uncertain.",
            auto_resolution_status="unresolved",
            model_observation_path=observation_path,
        )
    return Phase3BBoundaryDecision(
        document_id=document_id or observation.document_id,
        observed_word_continuation=observation.suspected_word_continuation,
        validated_word_continuation=validated_word_continuation,
        word_continuation=validated_word_continuation,
        sentence_continuation=True,
        paragraph_continuation=True,
        structural_break="none",
        join_operation=operation,
        joiner=joiner,
        resolved_text=resolved,
        resolution_method="pair_visual_plus_text_layer_plus_deterministic_rules",
        supporting_evidence=evidence,
        conflicting_evidence=conflicts,
        resolution_reason=reason,
        auto_resolution_status="resolved_pair",
        model_observation_path=observation_path,
    )


def _page12_text_layer(manifest: Phase3BManifest) -> str:
    with fitz.open(manifest.derived_pdf) as document:
        return document[11].get_text("text")


def _split_visible_paragraphs(block_text: str) -> list[str]:
    parts = [part.strip() for part in re.split(r"\n\s*\n", block_text) if part.strip()]
    return parts or ([block_text.strip()] if block_text.strip() else [])


def _page12_fragments(page: VisionNormalizedPageV11, document_id: str, *, continuation_from_previous: bool) -> list[SourceFragment]:
    fragments: list[SourceFragment] = []
    order = 0
    body_blocks = [block for block in sorted(page.blocks, key=lambda value: value.order) if block.block_type in {"body", "chapter_title", "section_title", "footnote", "caption"}]
    for block in body_blocks:
        for segment_index, text in enumerate(_split_visible_paragraphs(block.text), start=1):
            order += 1
            is_first_body = block.block_type == "body" and not any(item.block_type == "body" for item in fragments)
            starts_mid = bool(is_first_body and continuation_from_previous)
            source_block_id = f"p0012:{block.block_id}" + (f":segment{segment_index}" if len(_split_visible_paragraphs(block.text)) > 1 else "")
            fragments.append(
                SourceFragment(
                    fragment_id="fragment_" + stable_hash(
                        {"document_id": document_id, "page": 12, "source_block_id": source_block_id, "text": text}
                    )[:20],
                    text=text,
                    source_page=12,
                    source_block_ids=[source_block_id],
                    block_type=block.block_type,
                    order=order,
                    starts_mid_sentence=starts_mid,
                    ends_mid_sentence=False,
                    starts_mid_paragraph=starts_mid,
                    ends_mid_paragraph=False,
                    visible_trailing_hyphen=text.rstrip().endswith("-"),
                    uncertainty=["visual_uncertainty"] if block.uncertain else [],
                )
            )
    return fragments


def _old_logical_blocks(settings: ProjectSettings, root: Path) -> list[AutomatedLogicalBlock]:
    path = resolve_project_path(settings.automated_logical_directory, root=root) / "sample_11_pages.logical_blocks.jsonl"
    return [AutomatedLogicalBlock.model_validate_json(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def _old_fragments(settings: ProjectSettings, root: Path) -> list[str]:
    path = resolve_project_path(settings.automated_page_directory, root=root) / "sample_11_pages.pages.jsonl"
    pages = [AutomatedPageRecord.model_validate_json(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
    return [fragment.fragment_id for page in pages for fragment in page.content_fragments]


def _source_block_from_old(item: AutomatedLogicalBlock, document_id: str) -> Phase3BSourceLogicalBlock:
    complete = all([item.sentence_complete, item.paragraph_complete, item.coverage_complete, not item.unresolved_boundaries])
    return Phase3BSourceLogicalBlock(
        logical_block_id=item.logical_block_id,
        document_id=document_id,
        block_type=item.block_type,
        source_pages=item.source_pages,
        source_block_ids=item.source_block_ids,
        source_fragment_ids=item.source_fragment_ids,
        source_text=item.source_text,
        cross_page=item.cross_page,
        sentence_complete=item.sentence_complete,
        paragraph_complete=item.paragraph_complete,
        coverage_complete=item.coverage_complete,
        completeness_status="complete" if complete else "incomplete_end",
        unresolved_boundaries=list(item.unresolved_boundaries),
        source_exportable=complete,
        translation_ready=complete,
        chapter_id=item.chapter_id,
        section_id=item.section_id,
    )


def _ledger_counts(settings: ProjectSettings, root: Path) -> tuple[int, int]:
    ledger = _load_ledger(settings, root)
    return int(ledger["started"]["single"]), int(ledger["started"]["pair"])


def build_phase3b_source_document(
    settings: ProjectSettings,
    manifest: Phase3BManifest,
    page12_normalized_path: str | Path,
    pair_normalized_path: str | Path,
    *,
    root: Path | None = None,
) -> Phase3BSourceDocument:
    root = (root or project_root()).resolve()
    page12 = VisionNormalizedPageV11.model_validate(load_json(page12_normalized_path))
    observation = Phase3BBoundaryObservation.model_validate(load_json(pair_normalized_path))
    old = _old_logical_blocks(settings, root)
    if len(old) != 23 or old[-1].unresolved_boundaries != ["boundary_p0011_p0012"]:
        raise RuntimeError("Expected the preserved 23-block Phase 2B.2 source with one open 11->12 boundary")
    fragments12 = _page12_fragments(page12, manifest.document_id, continuation_from_previous=True)
    first_body = next((item for item in fragments12 if item.block_type == "body"), None)
    if first_body is None:
        raise RuntimeError("Page 12 contains no body fragment")
    old_tail = old[-1].source_text
    text_layer12 = _page12_text_layer(manifest)
    with fitz.open(manifest.sample11_pdf) as document:
        text_layer11 = document[10].get_text("text")
    decision = resolve_phase3b_boundary(
        previous_text=old_tail,
        next_text=first_body.text,
        observation=observation,
        text_layer_tail=text_layer11,
        text_layer_head=text_layer12,
        document_id=manifest.document_id,
        observation_path=str(Path(pair_normalized_path).resolve()),
    )
    decision_path = _data_root(settings, root) / "boundaries" / "boundary_p0011_p0012.resolved.json"
    atomic_write_json(decision_path, decision)

    page12_end_complete = page12.continuation_to_next is False and bool(
        re.search(r"[.!?][\"'”’)]?\s*$", fragments12[-1].text.strip())
    )
    entries = [_source_block_from_old(item, manifest.document_id) for item in old[:-1]]
    used_page12_ids: set[str] = set()
    if decision.auto_resolution_status == "resolved_pair":
        used_page12_ids.add(first_body.fragment_id)
        source_text = decision.resolved_text
        complete = page12_end_complete and not first_body.uncertainty
        updated_ids = old[-1].source_fragment_ids + [first_body.fragment_id]
        updated_block_ids = old[-1].source_block_ids + first_body.source_block_ids
        entries.append(
            Phase3BSourceLogicalBlock(
                logical_block_id="logical3b_" + stable_hash(
                    {"document_id": manifest.document_id, "fragments": updated_ids, "source_text": source_text}
                )[:20],
                document_id=manifest.document_id,
                block_type="body",
                source_pages=[11, 12],
                source_block_ids=updated_block_ids,
                source_fragment_ids=updated_ids,
                source_text=source_text,
                cross_page=True,
                sentence_complete=complete,
                paragraph_complete=complete,
                coverage_complete=True,
                completeness_status="complete" if complete else "incomplete_end",
                unresolved_boundaries=[] if complete else ["boundary_p0012_end_conflict"],
                source_exportable=complete,
                translation_ready=complete,
                chapter_id=old[-1].chapter_id,
                section_id=old[-1].section_id,
            )
        )
    else:
        entries.append(_source_block_from_old(old[-1], manifest.document_id))
        used_page12_ids.add(first_body.fragment_id)
        entries.append(
            Phase3BSourceLogicalBlock(
                logical_block_id="logical3b_" + stable_hash({"fragment": first_body.fragment_id})[:20],
                document_id=manifest.document_id,
                block_type="body",
                source_pages=[12],
                source_block_ids=first_body.source_block_ids,
                source_fragment_ids=[first_body.fragment_id],
                source_text=first_body.text,
                cross_page=False,
                sentence_complete=page12_end_complete,
                paragraph_complete=page12_end_complete,
                coverage_complete=True,
                completeness_status="incomplete_start",
                unresolved_boundaries=["boundary_p0011_p0012"],
                source_exportable=False,
                translation_ready=False,
                chapter_id=old[-1].chapter_id,
                section_id=old[-1].section_id,
            )
        )
    for fragment in fragments12:
        if fragment.fragment_id in used_page12_ids:
            continue
        used_page12_ids.add(fragment.fragment_id)
        complete = page12_end_complete and not fragment.uncertainty
        entries.append(
            Phase3BSourceLogicalBlock(
                logical_block_id="logical3b_" + stable_hash(
                    {"document_id": manifest.document_id, "fragment": fragment.fragment_id, "text": fragment.text}
                )[:20],
                document_id=manifest.document_id,
                block_type=fragment.block_type,
                source_pages=[12],
                source_block_ids=fragment.source_block_ids,
                source_fragment_ids=[fragment.fragment_id],
                source_text=fragment.text,
                cross_page=False,
                sentence_complete=complete,
                paragraph_complete=complete,
                coverage_complete=True,
                completeness_status="complete" if complete else "incomplete_end",
                unresolved_boundaries=[] if complete else ["boundary_p0012_end_conflict"],
                source_exportable=complete,
                translation_ready=complete,
                chapter_id=old[-1].chapter_id,
                section_id=old[-1].section_id,
            )
        )

    expected_fragments = _old_fragments(settings, root) + [item.fragment_id for item in fragments12]
    referenced = [fragment_id for entry in entries for fragment_id in entry.source_fragment_ids]
    counts = Counter(referenced)
    unused = sorted(set(expected_fragments) - set(referenced))
    duplicates = sorted(fragment_id for fragment_id, count in counts.items() if count != 1)
    source_text = "\n".join(item.source_text for item in entries)
    clean = all(
        value not in source_text
        for value in [page12.footer or "__missing_footer__", page12.page_number_text or "__missing_page__"]
    )
    blockers: list[str] = []
    if sha256_file(manifest.sample11_pdf) != manifest.sample11_sha256:
        blockers.append("sample11_hash_changed")
    if decision.auto_resolution_status != "resolved_pair":
        blockers.append("boundary_11_12_unresolved")
    if not page12_end_complete:
        blockers.append("page12_end_conflicts_with_confirmed_complete_boundary")
    if unused:
        blockers.append("unused_fragments_exist")
    if duplicates:
        blockers.append("duplicate_fragments_exist")
    if not clean:
        blockers.append("header_footer_or_page_number_pollution")
    if len(manifest.reused_visual_pages) != 11:
        blockers.append("first_eleven_visual_caches_not_fully_reused")
    if any(not item.source_exportable for item in entries):
        blockers.append("incomplete_source_blocks_exist")
    audit = Phase3BSourceAudit(
        processed_pages=manifest.page_count,
        sample11_hash_unchanged=sha256_file(manifest.sample11_pdf) == manifest.sample11_sha256,
        reused_visual_page_count=len(manifest.reused_visual_pages),
        first_ten_boundaries_reused=True,
        boundary_11_12_closed=decision.auto_resolution_status == "resolved_pair",
        page12_end_complete=page12_end_complete,
        open_boundary_12_13_exists=False,
        expected_fragment_count=len(expected_fragments),
        referenced_fragment_count=len(referenced),
        unused_fragment_ids=unused,
        duplicate_fragment_ids=duplicates,
        header_footer_page_number_clean=clean,
        strict_passed=not blockers,
        blockers=blockers,
    )
    single_calls, pair_calls = _ledger_counts(settings, root)
    document = Phase3BSourceDocument(
        document_id=manifest.document_id,
        source_pdf=manifest.derived_pdf,
        source_pdf_sha256=manifest.derived_pdf_sha256,
        source_inputs=[
            {"path": manifest.sample11_pdf, "sha256": manifest.sample11_sha256},
            {"path": manifest.page12_pdf, "sha256": manifest.page12_pdf_sha256},
        ],
        page_count=12,
        entries=entries,
        audit=audit,
        strict_export_ready=not blockers,
        strict_blockers=blockers,
        glm_single_calls=single_calls,
        glm_pair_calls=pair_calls,
        glm_total_calls=single_calls + pair_calls,
        created_at=_now(),
    )
    data_root = _data_root(settings, root)
    atomic_write_jsonl(data_root / "logical" / "sample_12_pages.logical_blocks.v1.jsonl", entries)
    atomic_write_json(data_root / "audits" / "source_only_audit.json", audit)
    atomic_write_json(resolve_project_path(settings.phase3b_master_path, root=root), document)
    return document


def _font(run: Any, *, size: float = 11, color: str = "202020", bold: bool = False, italic: bool = False, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _configure_source_word(document: Document, *, diagnostic: bool, source_hash: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.85)
    section.left_margin = section.right_margin = Inches(0.9)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(header.add_run("English Source | 12-page validation sample"), size=8.5, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(
        footer.add_run("DIAGNOSTIC SOURCE DRAFT" if diagnostic else "STRICT SOURCE EDITION"),
        size=8.5,
        color="9B1C1C" if diagnostic else "6B7280",
        bold=diagnostic,
    )
    document.core_properties.title = "English Source Sample (12 pages)"
    document.core_properties.subject = "Traceable source-only reconstruction"
    document.core_properties.comments = f"Canonical source JSON SHA-256: {source_hash}"
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    _font(title.add_run("English Source Sample"), size=22, color="17324D", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _font(
        subtitle.add_run("12 pages | Source-only diagnostic" if diagnostic else "12 pages | Strict source edition"),
        size=10.5,
        color="6B7280",
        italic=True,
    )


def _markdown_for(document: Phase3BSourceDocument, *, diagnostic: bool, source_hash: str) -> str:
    lines = ["# English Source Sample", "", "12 pages | Source-only diagnostic" if diagnostic else "12 pages | Strict source edition", "", f"<!-- canonical_source_json_sha256: {source_hash} -->", ""]
    for entry in document.entries:
        pages = f"{entry.source_pages[0]}" if len(entry.source_pages) == 1 else f"{entry.source_pages[0]}–{entry.source_pages[-1]}"
        lines.extend([f"[Source pages: {pages} | Block: {entry.logical_block_id}]", ""])
        if not entry.source_exportable:
            lines.extend(["[INCOMPLETE_SOURCE]", ""])
        if entry.block_type in {"book_title", "chapter_title"}:
            lines.extend([f"## {entry.source_text}", ""])
        elif entry.block_type in {"subtitle", "section_title", "subsection_title"}:
            lines.extend([f"### {entry.source_text}", ""])
        else:
            lines.extend([entry.source_text, ""])
    return "\n".join(lines).rstrip() + "\n"


def _word_for(document: Phase3BSourceDocument, path: Path, *, diagnostic: bool, source_hash: str) -> None:
    word = Document()
    _configure_source_word(word, diagnostic=diagnostic, source_hash=source_hash)
    for entry in document.entries:
        pages = f"{entry.source_pages[0]}" if len(entry.source_pages) == 1 else f"{entry.source_pages[0]}–{entry.source_pages[-1]}"
        metadata = word.add_paragraph()
        metadata.paragraph_format.space_before = Pt(10)
        metadata.paragraph_format.space_after = Pt(3)
        metadata.paragraph_format.keep_with_next = True
        _font(metadata.add_run(f"[Source pages: {pages} | Block: {entry.logical_block_id}]"), size=8.5, color="4E7392", bold=True)
        if not entry.source_exportable:
            warning = word.add_paragraph()
            _font(warning.add_run("[INCOMPLETE_SOURCE]"), size=9.5, color="9B1C1C", bold=True)
        paragraph = word.add_paragraph()
        paragraph.paragraph_format.keep_together = entry.block_type in {"book_title", "chapter_title", "section_title"}
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if entry.block_type in {"book_title", "chapter_title", "section_title"} else WD_ALIGN_PARAGRAPH.JUSTIFY
        if entry.block_type in {"book_title", "chapter_title"}:
            _font(paragraph.add_run(entry.source_text), size=16, color="17324D", bold=True)
        elif entry.block_type in {"subtitle", "section_title", "subsection_title"}:
            _font(paragraph.add_run(entry.source_text), size=13, color="2D536F", bold=True)
        else:
            _font(paragraph.add_run(entry.source_text), size=11, color="202020")
    path.parent.mkdir(parents=True, exist_ok=True)
    word.save(path)


def export_phase3b_source(
    settings: ProjectSettings,
    *,
    master_path: str | Path,
    root: Path | None = None,
) -> Phase3BExportResult:
    root = (root or project_root()).resolve()
    master = resolve_project_path(master_path, root=root)
    document = Phase3BSourceDocument.model_validate(load_json(master))
    source_hash = sha256_file(master)
    diagnostic_root = resolve_project_path(settings.phase3b_diagnostic_directory, root=root)
    final_root = resolve_project_path(settings.phase3b_final_directory, root=root)
    diagnostic_md = diagnostic_root / "source_english_sample12.md"
    diagnostic_docx = diagnostic_root / "source_english_sample12.docx"
    diagnostic_md.parent.mkdir(parents=True, exist_ok=True)
    diagnostic_md.write_text(_markdown_for(document, diagnostic=True, source_hash=source_hash), encoding="utf-8")
    _word_for(document, diagnostic_docx, diagnostic=True, source_hash=source_hash)
    final_md: Path | None = None
    final_docx: Path | None = None
    if document.strict_export_ready:
        final_md = final_root / "source_english.md"
        final_docx = final_root / "source_english.docx"
        final_md.parent.mkdir(parents=True, exist_ok=True)
        final_md.write_text(_markdown_for(document, diagnostic=False, source_hash=source_hash), encoding="utf-8")
        _word_for(document, final_docx, diagnostic=False, source_hash=source_hash)
    return Phase3BExportResult(
        source_json_sha256=source_hash,
        diagnostic_markdown_path=str(diagnostic_md.resolve()),
        diagnostic_docx_path=str(diagnostic_docx.resolve()),
        strict_exported=document.strict_export_ready,
        final_markdown_path=str(final_md.resolve()) if final_md else None,
        final_docx_path=str(final_docx.resolve()) if final_docx else None,
        blockers=document.strict_blockers,
    )


def _cache_exists(settings: ProjectSettings, root: Path, category: str) -> bool:
    base = _data_root(settings, root) / "calls" / category
    return any(path.name == "cache.json" for path in base.rglob("cache.json")) if base.is_dir() else False


def phase3b_preflight(
    settings: ProjectSettings,
    *,
    manifest: Phase3BManifest,
    root: Path | None = None,
) -> Phase3BPreflight:
    root = (root or project_root()).resolve()
    single_expected = 0 if _cache_exists(settings, root, "single") else 1
    pair_expected = 0 if _cache_exists(settings, root, "pair") else 1
    ledger = _load_ledger(settings, root)
    calls_started = int(ledger["started"]["total"])
    remaining_calls = max(0, settings.phase3b_max_total_calls - calls_started)
    key_set, _ = api_key_status(settings, root)
    estimated_input = 12_000
    estimated_output = 2_000
    estimated_cost = round(
        (
            estimated_input * settings.vision_input_price_cny_per_million_tokens_upper
            + estimated_output * settings.vision_output_price_cny_per_million_tokens_upper
        )
        / 1_000_000,
        6,
    )
    blockers: list[str] = []
    if manifest.page_count != 12:
        blockers.append("derived_sample_not_12_pages")
    if len(manifest.reused_visual_pages) != 11:
        blockers.append("first_11_visual_caches_not_reused")
    if settings.translation_enabled or not settings.translation_disabled or not settings.terminology_translation_disabled:
        blockers.append("translation_not_fully_disabled")
    if settings.phase3b_automatic_retry or settings.phase3b_max_triple_calls != 0:
        blockers.append("retry_or_triple_not_disabled")
    if single_expected + pair_expected > settings.phase3b_max_total_calls:
        blockers.append("planned_calls_exceed_two")
    if single_expected and int(ledger["started"]["single"]) >= settings.phase3b_max_single_calls:
        blockers.append("single_call_used_without_valid_cache")
    if pair_expected and int(ledger["started"]["pair"]) >= settings.phase3b_max_pair_calls:
        blockers.append("pair_call_used_without_valid_cache")
    if single_expected + pair_expected > remaining_calls:
        blockers.append("remaining_call_budget_insufficient")
    if estimated_cost > settings.phase3b_maximum_cash_cost_cny:
        blockers.append("estimated_cost_exceeds_ceiling")
    if (single_expected + pair_expected) and not key_set:
        blockers.append("vision_api_key_not_set")
    return Phase3BPreflight(
        document_id=manifest.document_id,
        actual_page_count=manifest.page_count,
        reused_visual_pages=len(manifest.reused_visual_pages),
        single_calls_expected=single_expected,
        pair_calls_expected=pair_expected,
        triple_calls_allowed=settings.phase3b_max_triple_calls,
        maximum_new_calls=settings.phase3b_max_total_calls,
        calls_already_started=calls_started,
        remaining_real_calls=remaining_calls,
        automatic_retry=settings.phase3b_automatic_retry,
        estimated_token_range="approximately 7,000-14,000 total tokens for one page and one pair",
        estimated_public_price_cny=estimated_cost,
        maximum_cash_cost_cny=settings.phase3b_maximum_cash_cost_cny,
        api_key_set=key_set,
        blockers=blockers,
        ready=not blockers,
    )
