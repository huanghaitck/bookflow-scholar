"""Offline page normalization and publication objects for back matter.

The module is deliberately provider-neutral.  It prepares deterministic page and
region assets, records coordinate transforms, and renders only reviewed layout
objects.  Provider calls belong to a later, explicitly authorized task.
"""
from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Iterable

from PIL import Image, ImageChops, ImageFilter, ImageOps


LAYOUT_OBJECT_TYPES = frozenset(
    {
        "prose",
        "heading",
        "list_entry",
        "figure",
        "caption",
        "table",
        "rotated_table",
        "multi_page_table",
        "illustration_list",
        "index",
        "multi_column_text",
        "footnote",
        "facsimile_region",
    }
)
RENDER_PROFILES = frozenset({"reading", "evidence"})
ORIENTATIONS = frozenset({0, 90, 180, 270})


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", "utf-8")


def load_book_profile(root: Path, path: str | Path | None = None) -> dict[str, Any]:
    profile_path = root / (path or "config/books/big_game_1913.phase12r.json")
    profile = json.loads(profile_path.read_text("utf-8"))
    required = {"book_id", "source_page_asset_pattern", "normalization", "render_profiles"}
    missing = sorted(required - set(profile))
    if missing:
        raise ValueError(f"book profile missing fields: {', '.join(missing)}")
    return profile


def _projection_score(image: Image.Image) -> float:
    """Prefer rotations with text-like horizontal runs; used only without hints."""
    thumb = ImageOps.grayscale(image)
    thumb.thumbnail((320, 320))
    pixels = thumb.point(lambda value: 255 if value < 175 else 0)
    horizontal = list(pixels.resize((1, pixels.height)).getdata())
    vertical = list(pixels.resize((pixels.width, 1)).getdata())
    return float(sum(value * value for value in horizontal) - sum(value * value for value in vertical))


def detect_orientation(
    image: Image.Image,
    *,
    scorer: Callable[[Image.Image, int], float] | None = None,
) -> tuple[int, dict[str, float]]:
    """Return clockwise rotation in {0,90,180,270} and auditable scores.

    A caller may supply an OCR or vision-backed scorer later.  The offline
    fallback detects horizontal versus vertical layout and intentionally treats
    0/180 and 90/270 ties conservatively.
    """
    scores: dict[str, float] = {}
    for degrees in sorted(ORIENTATIONS):
        candidate = image.rotate(-degrees, expand=True)
        score = scorer(candidate, degrees) if scorer else _projection_score(candidate)
        if not scorer and degrees in {180, 270}:
            score -= 0.001
        scores[str(degrees)] = round(float(score), 6)
    selected = max((int(key) for key in scores), key=lambda value: scores[str(value)])
    return selected, scores


def _content_bbox(image: Image.Image, margin: int = 12) -> tuple[int, int, int, int]:
    gray = ImageOps.grayscale(image).filter(ImageFilter.MedianFilter(3))
    background = Image.new("L", gray.size, 245)
    difference = ImageChops.difference(gray, background).point(lambda value: 255 if value > 34 else 0)
    bbox = difference.getbbox() or (0, 0, image.width, image.height)
    return (
        max(0, bbox[0] - margin),
        max(0, bbox[1] - margin),
        min(image.width, bbox[2] + margin),
        min(image.height, bbox[3] + margin),
    )


def _rotation_matrix(width: int, height: int, degrees: int) -> tuple[list[list[float]], tuple[int, int]]:
    if degrees == 0:
        return [[1, 0, 0], [0, 1, 0], [0, 0, 1]], (width, height)
    if degrees == 90:
        return [[0, -1, height], [1, 0, 0], [0, 0, 1]], (height, width)
    if degrees == 180:
        return [[-1, 0, width], [0, -1, height], [0, 0, 1]], (width, height)
    return [[0, 1, 0], [-1, 0, width], [0, 0, 1]], (height, width)


def _crop_matrix(bbox: tuple[int, int, int, int]) -> list[list[float]]:
    return [[1, 0, -bbox[0]], [0, 1, -bbox[1]], [0, 0, 1]]


def _deskew_matrix(
    input_size: tuple[int, int],
    output_size: tuple[int, int],
    degrees_clockwise: float,
) -> list[list[float]]:
    from math import cos, radians, sin

    angle = radians(degrees_clockwise)
    cosine, sine = cos(angle), sin(angle)
    input_x, input_y = input_size[0] / 2, input_size[1] / 2
    output_x, output_y = output_size[0] / 2, output_size[1] / 2
    return [
        [cosine, sine, output_x - cosine * input_x - sine * input_y],
        [-sine, cosine, output_y + sine * input_x - cosine * input_y],
        [0, 0, 1],
    ]


def _matmul(left: list[list[float]], right: list[list[float]]) -> list[list[float]]:
    return [
        [sum(left[row][offset] * right[offset][column] for offset in range(3)) for column in range(3)]
        for row in range(3)
    ]


@dataclass(frozen=True)
class NormalizationResult:
    manifest: dict[str, Any]
    normalized_asset: Path


def normalize_page(
    source: Path,
    destination: Path,
    *,
    physical_page: int,
    orientation_hint: int | None = None,
    crop_bbox_after_rotation: tuple[int, int, int, int] | None = None,
    deskew_degrees: float = 0.0,
    artifact_regions: dict[str, list[list[int]]] | None = None,
) -> NormalizationResult:
    """Normalize one page in the required rotate -> crop -> deskew order."""
    if orientation_hint is not None and orientation_hint not in ORIENTATIONS:
        raise ValueError("orientation hint must be 0, 90, 180, or 270")
    with Image.open(source) as opened:
        raw = ImageOps.exif_transpose(opened).convert("RGB")
        detected, scores = detect_orientation(raw)
        orientation = orientation_hint if orientation_hint is not None else detected
        rotation_matrix, rotated_size = _rotation_matrix(raw.width, raw.height, orientation)
        rotated = raw.rotate(-orientation, expand=True, resample=Image.Resampling.BICUBIC)
        crop_bbox = crop_bbox_after_rotation or _content_bbox(rotated)
        if not (0 <= crop_bbox[0] < crop_bbox[2] <= rotated.width and 0 <= crop_bbox[1] < crop_bbox[3] <= rotated.height):
            raise ValueError("crop bbox is outside the rotated image")
        cropped = rotated.crop(crop_bbox)
        normalized = cropped.rotate(
            -deskew_degrees,
            expand=True,
            fillcolor=(255, 255, 255),
            resample=Image.Resampling.BICUBIC,
        )
        destination.parent.mkdir(parents=True, exist_ok=True)
        normalized.save(destination, format="PNG", optimize=True)
    rotate_then_crop = _matmul(_crop_matrix(crop_bbox), rotation_matrix)
    deskew_matrix = _deskew_matrix(cropped.size, normalized.size, deskew_degrees)
    raw_to_normalized = _matmul(deskew_matrix, rotate_then_crop)
    manifest = {
        "schema_version": "page-normalization-1.0",
        "physical_page": physical_page,
        "source_asset_ref": source.as_posix(),
        "source_sha256": _sha256(source),
        "normalized_asset_ref": destination.as_posix(),
        "normalized_sha256": _sha256(destination),
        "orientation_degrees_clockwise": orientation,
        "orientation_detection_scores": scores,
        "orientation_source": "book_profile" if orientation_hint is not None else "offline_projection",
        "operation_order": ["rotate", "crop", "deskew"],
        "rotated_size": list(rotated_size),
        "crop_bbox_after_rotation": list(crop_bbox),
        "deskew_degrees_clockwise": deskew_degrees,
        "raw_to_normalized_transform": raw_to_normalized,
        "transform_chain": {
            "rotation": rotation_matrix,
            "crop": _crop_matrix(crop_bbox),
            "deskew": deskew_matrix,
        },
        "artifact_regions_raw": artifact_regions or {},
    }
    return NormalizationResult(manifest=manifest, normalized_asset=destination)


def crop_normalized_region(
    normalized: Path,
    destination: Path,
    bbox: tuple[int, int, int, int],
    *,
    region_id: str,
) -> dict[str, Any]:
    with Image.open(normalized) as opened:
        if not (0 <= bbox[0] < bbox[2] <= opened.width and 0 <= bbox[1] < bbox[3] <= opened.height):
            raise ValueError("region bbox is outside normalized page")
        region = opened.crop(bbox)
        destination.parent.mkdir(parents=True, exist_ok=True)
        region.save(destination, format="PNG", optimize=True)
    return {
        "region_id": region_id,
        "normalized_page_asset_ref": normalized.as_posix(),
        "bbox_normalized": list(bbox),
        "region_asset_ref": destination.as_posix(),
        "region_sha256": _sha256(destination),
    }


def normalize_profile_pages(root: Path, profile: dict[str, Any]) -> dict[str, Any]:
    settings = profile["normalization"]
    pages = [int(value) for value in settings["pages"]]
    orientation_hints = {int(key): int(value) for key, value in settings.get("orientation_hints", {}).items()}
    output_root = root / settings["output_root"]
    manifests: list[dict[str, Any]] = []
    for page in pages:
        source = root / profile["source_page_asset_pattern"].format(page=page)
        destination = output_root / f"page_{page:04d}.normalized.png"
        result = normalize_page(
            source,
            destination,
            physical_page=page,
            orientation_hint=orientation_hints.get(page),
            deskew_degrees=float(settings.get("deskew_hints", {}).get(str(page), 0.0)),
            artifact_regions={
                **settings.get("artifact_region_defaults", {}),
                **settings.get("artifact_regions", {}).get(str(page), {}),
            },
        )
        manifests.append(result.manifest)
    regions: list[dict[str, Any]] = []
    for region in settings.get("regions", []):
        page = int(region["physical_page"])
        normalized = output_root / f"page_{page:04d}.normalized.png"
        destination = output_root / "regions" / f"{region['region_id']}.png"
        with Image.open(normalized) as opened:
            if region.get("bbox_fraction"):
                fraction = region["bbox_fraction"]
                bbox = (
                    round(opened.width * fraction[0]),
                    round(opened.height * fraction[1]),
                    round(opened.width * fraction[2]),
                    round(opened.height * fraction[3]),
                )
            else:
                bbox = tuple(region["bbox_normalized"])
        region_manifest = crop_normalized_region(normalized, destination, bbox, region_id=region["region_id"])
        region_manifest["render_orientation"] = region.get("render_orientation", "portrait")
        regions.append(region_manifest)
    value = {
        "schema_version": "page-normalization-batch-1.0",
        "book_id": profile["book_id"],
        "operation_order": ["rotate", "crop", "deskew"],
        "pages": manifests,
        "regions": regions,
    }
    _write_json(output_root / "normalization_manifest.json", value)
    return value


def validate_layout_object(value: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if value.get("object_type") not in LAYOUT_OBJECT_TYPES:
        errors.append("invalid object_type")
    for key in ("object_id", "source_pages", "bbox", "orientation", "reading_order", "confidence", "provenance", "render_policy"):
        if key not in value:
            errors.append(f"missing {key}")
    if value.get("orientation") not in ORIENTATIONS:
        errors.append("invalid orientation")
    if not isinstance(value.get("source_pages"), list) or not value.get("source_pages"):
        errors.append("source_pages must be non-empty")
    if not isinstance(value.get("confidence"), (int, float)) or not 0 <= value.get("confidence", -1) <= 1:
        errors.append("confidence must be between 0 and 1")
    policy = value.get("render_policy", {})
    if set(policy) != RENDER_PROFILES:
        errors.append("render_policy must define reading and evidence")
    return errors


def validate_layout_model(model: dict[str, Any]) -> dict[str, Any]:
    errors = {
        item.get("object_id", f"object_{index}"): validate_layout_object(item)
        for index, item in enumerate(model.get("objects", []))
    }
    errors = {key: value for key, value in errors.items() if value}
    return {"valid": not errors, "object_count": len(model.get("objects", [])), "errors": errors}


def objects_for_profile(model: dict[str, Any], profile: str) -> list[dict[str, Any]]:
    if profile not in RENDER_PROFILES:
        raise ValueError(f"unknown render profile: {profile}")
    selected = []
    for item in model.get("objects", []):
        action = item["render_policy"][profile]
        if action != "omit":
            selected.append(item)
    return sorted(selected, key=lambda item: (item["reading_order"], item["object_id"]))


def validate_quality_gates(
    *,
    normalization: dict[str, Any],
    layout_model: dict[str, Any],
    illustration_model: dict[str, Any],
) -> dict[str, Any]:
    page_manifests = {item["physical_page"]: item for item in normalization["pages"]}
    objects = layout_model["objects"]
    illustration_entries = illustration_model["entries"]
    formal_387 = [
        item
        for item in objects
        if 387 in item["source_pages"] and item["object_type"] in {"table", "rotated_table"}
    ]
    tables = [
        item
        for item in objects
        if item["object_type"] in {"table", "rotated_table", "multi_page_table"}
    ]
    checks = {
        "orientation_387": page_manifests[387]["orientation_degrees_clockwise"] == 90,
        "rotate_before_crop": all(item["operation_order"][:2] == ["rotate", "crop"] for item in page_manifests.values()),
        "illustration_pages_covered": {item["physical_page"] for item in illustration_entries} == {21, 22},
        "illustration_entries_have_printed_locator": all(item.get("printed_locator") for item in illustration_entries),
        "printed_not_physical": all(str(item["printed_locator"]) != str(item["physical_page"]) for item in illustration_entries),
        "page_387_no_formal_rows": all(not item.get("rows") for item in formal_387),
        "numeric_unresolved_explicit": all(
            item.get("numeric_status") in {"resolved", "unresolved", "not_applicable"}
            for item in objects
            if item["object_type"] in {"table", "rotated_table", "multi_page_table"}
        ),
        "table_counts_accountable": all(
            isinstance(item.get("schema_candidate", {}).get("column_count"), int)
            and isinstance(item.get("schema_candidate", {}).get("header_row_count"), int)
            and item.get("schema_candidate", {}).get("row_count_status") in {"resolved", "unresolved"}
            and (
                item.get("schema_candidate", {}).get("row_count_status") == "unresolved"
                or isinstance(item.get("schema_candidate", {}).get("row_count"), int)
            )
            for item in tables
        ),
        "headers_footers_not_reading": all(
            item["render_policy"]["reading"] == "omit"
            for item in objects
            if item.get("provenance", {}).get("artifact_kind") in {"header", "footer", "watermark", "page_number"}
        ),
        "reading_technical_labels_zero": not any(
            item.get("technical_label_visible", False) and item["render_policy"]["reading"] != "omit"
            for item in objects
        ),
        "reading_full_page_facsimiles_zero": not any(
            item["object_type"] == "facsimile_region"
            and item.get("facsimile_scope") == "full_page"
            and item["render_policy"]["reading"] != "omit"
            for item in objects
        ),
        "landscape_tables_not_clipped": all(
            item.get("publication", {}).get("overflow_policy") in {"landscape_section", "safe_column_split"}
            for item in tables
            if item["object_type"] == "rotated_table" or item.get("publication", {}).get("wide")
        ),
    }
    return {"valid": all(checks.values()), "checks": checks}


def count_translation_delta(
    illustration_model: dict[str, Any],
    units: Iterable[dict[str, Any]],
    translated_source_object_ids: set[str],
) -> dict[str, int]:
    units_by_text: dict[str, list[dict[str, Any]]] = {}
    for unit in units:
        units_by_text.setdefault(str(unit.get("source_text", "")).strip(), []).append(unit)
    reusable = 0
    for entry in illustration_model["entries"]:
        matches = units_by_text.get(entry["source_text"].strip(), [])
        if any(unit.get("source_object_id") in translated_source_object_ids for unit in matches):
            reusable += 1
    planned = sum(entry.get("translation_policy") == "translate" for entry in illustration_model["entries"])
    return {"planned": planned, "reusable": reusable, "delta": planned - reusable}


def _object_counts(objects: Iterable[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for item in objects:
        counts[item["object_type"]] = counts.get(item["object_type"], 0) + 1
    return dict(sorted(counts.items()))


def build_render_plan_manifest(layout_model: dict[str, Any]) -> dict[str, Any]:
    profiles: dict[str, Any] = {}
    for profile in sorted(RENDER_PROFILES):
        selected = objects_for_profile(layout_model, profile)
        counts = _object_counts(selected)
        profiles[profile] = {
            "object_ids": [item["object_id"] for item in selected],
            "object_counts": counts,
            "format_object_counts": {"markdown": counts, "docx": counts, "pdf": counts},
            "technical_label_count": sum(bool(item.get("technical_label_visible")) for item in selected),
            "full_page_facsimile_count": sum(
                item["object_type"] == "facsimile_region" and item.get("facsimile_scope") == "full_page"
                for item in selected
            ),
        }
    return {"schema_version": "back-matter-render-plan-1.0", "profiles": profiles}


def render_illustration_list_markdown(
    illustration_model: dict[str, Any],
    *,
    language: str,
    translations: dict[str, str] | None = None,
) -> list[str]:
    translations = translations or {}
    entries = {item["entry_id"]: item for item in illustration_model["entries"]}
    lines = ["## List of Illustrations", ""]
    for group in illustration_model["groups"]:
        if group["group_id"] == "maps":
            lines.extend(["### Maps", ""])
        for entry_id in group["entry_ids"]:
            entry = entries[entry_id]
            translated = translations.get(entry_id)
            if language == "en" or not translated:
                label = entry["source_text"]
            elif language == "zh-Hans":
                label = translated
            else:
                label = f"{entry['source_text']} / {translated}"
            lines.append(f"{label} {'.' * 12} {entry['printed_locator']}")
        lines.append("")
    return lines


def _markdown_table(item: dict[str, Any]) -> list[str]:
    schema = item["schema_candidate"]
    columns = schema["columns"]
    rows = item.get("rows", [])
    if not rows:
        return []
    lines = ["| " + " | ".join(columns) + " |", "| " + " | ".join("---" for _ in columns) + " |"]
    for row in rows:
        cells = row.get("cells", [])
        if len(cells) != len(columns):
            raise ValueError(f"{item['object_id']} row width does not match schema")
        lines.append("| " + " | ".join(str(cell).replace("|", "\\|") for cell in cells) + " |")
    return lines + [""]


def render_layout_markdown(
    layout_model: dict[str, Any],
    illustration_model: dict[str, Any],
    *,
    profile: str,
    language: str = "en",
    translations: dict[str, str] | None = None,
    region_assets: dict[str, str] | None = None,
) -> tuple[str, dict[str, Any]]:
    """Render a back-matter fragment; no release directory is created."""
    region_assets = region_assets or {}
    selected = objects_for_profile(layout_model, profile)
    lines: list[str] = []
    rendered_ids: list[str] = []
    for item in selected:
        kind = item["object_type"]
        if kind == "illustration_list":
            lines.extend(render_illustration_list_markdown(illustration_model, language=language, translations=translations))
        elif kind in {"table", "rotated_table", "multi_page_table"}:
            lines.extend(_markdown_table(item))
        elif kind == "facsimile_region":
            region_ids = item["provenance"].get("region_ids") or [item["provenance"].get("region_id")]
            for region_id in region_ids:
                if region_id in region_assets:
                    lines.extend([f"![{item['object_id']}]({region_assets[region_id]})", ""])
        elif kind == "index":
            lines.extend(["## Index", "", '<div class="index-two-column">', "", "</div>", ""])
        elif kind == "heading":
            lines.extend([f"## {item['provenance'].get('section', item['object_id'])}", ""])
        if profile == "evidence":
            lines.extend([f"<!-- source-object: {item['object_id']} -->", ""])
        rendered_ids.append(item["object_id"])
    counts = _object_counts(selected)
    return "\n".join(lines), {
        "profile": profile,
        "format": "markdown",
        "object_ids": rendered_ids,
        "object_counts": counts,
        "technical_label_count": sum(bool(item.get("technical_label_visible")) for item in selected),
        "full_page_facsimile_count": sum(
            item["object_type"] == "facsimile_region" and item.get("facsimile_scope") == "full_page"
            for item in selected
        ),
    }


def render_layout_docx(
    layout_model: dict[str, Any],
    illustration_model: dict[str, Any],
    destination: Path,
    *,
    profile: str,
    language: str = "en",
    translations: dict[str, str] | None = None,
    region_assets: dict[str, Path] | None = None,
) -> dict[str, Any]:
    """Render native Word tables, region fallbacks, and a two-column index."""
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    translations = translations or {}
    region_assets = region_assets or {}
    selected = objects_for_profile(layout_model, profile)
    document = Document()
    for item in selected:
        kind = item["object_type"]
        if kind == "illustration_list":
            document.add_heading("List of Illustrations", level=1)
            entries = {entry["entry_id"]: entry for entry in illustration_model["entries"]}
            for group in illustration_model["groups"]:
                if group["group_id"] == "maps":
                    document.add_heading("Maps", level=2)
                for entry_id in group["entry_ids"]:
                    entry = entries[entry_id]
                    translated = translations.get(entry_id)
                    label = entry["source_text"] if language == "en" or not translated else translated if language == "zh-Hans" else f"{entry['source_text']} / {translated}"
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.tab_stops.add_tab_stop(
                        Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
                    )
                    paragraph.add_run(label)
                    paragraph.add_run("\t" + str(entry["printed_locator"]))
        elif kind in {"table", "rotated_table", "multi_page_table"} and item.get("rows"):
            schema = item["schema_candidate"]
            if item.get("publication", {}).get("wide"):
                section = document.add_section(WD_SECTION.NEW_PAGE)
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = section.page_height, section.page_width
            table = document.add_table(rows=1, cols=schema["column_count"])
            for index, value in enumerate(schema["columns"]):
                table.rows[0].cells[index].text = value
            for value in item["rows"]:
                cells = value.get("cells", [])
                if len(cells) != schema["column_count"]:
                    raise ValueError(f"{item['object_id']} row width does not match schema")
                row = table.add_row()
                for index, cell in enumerate(cells):
                    row.cells[index].text = str(cell)
        elif kind == "facsimile_region":
            region_ids = item["provenance"].get("region_ids") or [item["provenance"].get("region_id")]
            for region_id in region_ids:
                asset = region_assets.get(region_id)
                if asset and asset.is_file():
                    document.add_picture(str(asset), width=Inches(6.2))
        elif kind == "index":
            document.add_heading("Index", level=1)
            section = document.add_section(WD_SECTION.CONTINUOUS)
            columns = OxmlElement("w:cols")
            columns.set(qn("w:num"), "2")
            columns.set(qn("w:space"), "360")
            section._sectPr.append(columns)
        elif kind == "heading":
            document.add_heading(item["provenance"].get("section", item["object_id"]), level=1)
        if profile == "evidence":
            document.add_paragraph(f"Source object: {item['object_id']}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return {
        "profile": profile,
        "format": "docx",
        "path": str(destination),
        "sha256": _sha256(destination),
        "object_ids": [item["object_id"] for item in selected],
        "object_counts": _object_counts(selected),
        "technical_label_count": sum(bool(item.get("technical_label_visible")) for item in selected),
        "full_page_facsimile_count": sum(
            item["object_type"] == "facsimile_region" and item.get("facsimile_scope") == "full_page"
            for item in selected
        ),
    }
