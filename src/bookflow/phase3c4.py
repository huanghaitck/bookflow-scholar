"""Phase 3C+4: translate the complete 12-page sample and export candidates."""

from __future__ import annotations

import json
import math
import subprocess
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Literal
from urllib.parse import quote

import fitz
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel, ConfigDict, Field, model_validator

from .io_utils import atomic_write_json, atomic_write_text, load_json, sha256_file, sha256_text, stable_hash
from .paths import ProjectSettings, project_root, resolve_project_path
from .phase3b_source import Phase3BSourceDocument, Phase3BSourceLogicalBlock
from .secret_store import load_translation_api_key, translation_api_key_status
from .translation_provider import DeepSeekOpenAICompatibleProvider


TRANSLATABLE_TYPES = {
    "book_title", "subtitle", "chapter_title", "section_title", "subsection_title",
    "body", "footnote", "caption", "table_title", "table_cell", "epigraph",
    "other_translatable",
}
TITLE_TYPES = {"book_title", "subtitle", "chapter_title", "section_title", "subsection_title"}


def _now() -> datetime:
    return datetime.now(timezone.utc)


class Phase3C4TranslationUnit(BaseModel):
    model_config = ConfigDict(extra="forbid")

    target_block_id: str
    block_type: str
    source_text: str
    source_pages: list[int]
    source_fragment_ids: list[str]
    completeness_status: str
    chapter_id: str | None
    section_id: str | None
    chapter_title_block_id: str | None
    section_title_block_id: str | None
    chapter_title_context: str | None
    section_title_context: str | None
    context_before_block_ids: list[str]
    context_after_block_ids: list[str]
    context_before_text: str | None
    context_after_text: str | None
    translate_target_only: Literal[True] = True
    translation_ready: Literal[True] = True
    source_language: Literal["en"] = "en"
    target_language: Literal["zh-Hans"] = "zh-Hans"

    def provider_payload(self) -> dict[str, Any]:
        return self.model_dump(mode="json")


class Phase3C4ModelPayload(BaseModel):
    model_config = ConfigDict(extra="forbid")

    target_block_id: str
    block_type: str
    translation: str
    untranslated_source_terms: list[str] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


class Phase3C4TranslationResult(Phase3C4ModelPayload):
    model_config = ConfigDict(extra="forbid")

    schema_version: str = "1.0"
    source_text_sha256: str
    context_sha256: str
    request_fingerprint: str
    prompt_version: str
    prompt_sha256: str
    language_profile_sha256: str
    raw_response_path: str
    request_id: str | None
    usage: dict[str, Any] | None
    status: Literal["translated"] = "translated"


class Phase3C4Preflight(BaseModel):
    model_config = ConfigDict(extra="forbid")

    target_block_count: int
    cached_block_count: int
    estimated_new_content_calls: int
    maximum_content_calls: int
    estimated_model_list_calls: Literal[0] = 0
    maximum_model_list_calls: int
    total_source_characters: int
    total_context_characters: int
    estimated_input_tokens: int
    estimated_output_tokens: int
    maximum_output_tokens: int
    estimated_cost_lower_cny: float
    estimated_cost_upper_cny: float
    maximum_cash_cost_cny: float
    api_key_set: bool
    automatic_retry: Literal[False] = False
    api_returns_actual_cash_charge: Literal[False] = False
    api_returns_resource_balance: Literal[False] = False
    blockers: list[str]
    ready_for_real_call: bool
    api_called: Literal[False] = False


class Phase3C4BatchResult(BaseModel):
    model_config = ConfigDict(extra="forbid")

    results: list[Phase3C4TranslationResult]
    api_calls: int
    cache_hits: int
    failed: int
    retries: Literal[0] = 0
    model_list_calls: Literal[0] = 0
    glm_calls: Literal[0] = 0
    preflight: Phase3C4Preflight


class Phase3C4BilingualBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    block_id: str
    block_type: str
    chapter_id: str | None
    section_id: str | None
    source_pages: list[int]
    source_fragment_ids: list[str]
    source_block_ids: list[str]
    source_text: str
    completeness_status: str
    translation_ready: Literal[True] = True
    translation: str
    untranslated_source_terms: list[str]
    warnings: list[str]
    translation_status: Literal["translated"] = "translated"
    raw_response_reference: str
    cache_fingerprint: str
    usage_reference: str | None


class Phase3C4StructuralEntry(BaseModel):
    model_config = ConfigDict(extra="forbid")

    block_id: str
    block_type: str
    source_text: str
    translation: str


class Phase3C4AlignmentAudit(BaseModel):
    model_config = ConfigDict(extra="forbid")

    source_block_count: int
    translated_block_count: int
    missing_translation_ids: list[str]
    duplicate_translation_ids: list[str]
    extra_translation_ids: list[str]
    source_text_modified: bool
    untranslated_source_term_count: int
    strict_passed: bool


class Phase3C4BilingualDocument(BaseModel):
    model_config = ConfigDict(extra="forbid")

    schema_version: str = "bilingual-document-1.0"
    document_id: str
    source_document_path: str
    source_document_sha256: str
    page_count: Literal[12]
    source_language: Literal["en"] = "en"
    target_language: Literal["zh-Hans"] = "zh-Hans"
    logical_blocks: list[Phase3C4BilingualBlock]
    chapters: list[Phase3C4StructuralEntry]
    sections: list[Phase3C4StructuralEntry]
    audit: Phase3C4AlignmentAudit
    created_at: datetime


class Phase3C4ExportResult(BaseModel):
    model_config = ConfigDict(extra="forbid")

    source_json_sha256: str
    source_markdown_path: str
    source_docx_path: str
    bilingual_markdown_path: str
    bilingual_docx_path: str


class LibreOfficeSmokeResult(BaseModel):
    model_config = ConfigDict(extra="forbid")

    libreoffice_found: bool
    executable_path: str | None
    completed: bool
    fatal_error: bool
    source_pdf_path: str | None
    bilingual_pdf_path: str | None
    source_pdf_page_count: int | None
    bilingual_pdf_page_count: int | None
    warnings: list[str] = Field(default_factory=list)


def is_phase3c4_translatable_type(block_type: str) -> bool:
    return block_type in TRANSLATABLE_TYPES


def _source_document(settings: ProjectSettings, root: Path) -> tuple[Path, Phase3BSourceDocument]:
    path = resolve_project_path(settings.phase3c4_source_document_path, root=root)
    document = Phase3BSourceDocument.model_validate(load_json(path))
    if not document.strict_export_ready or not document.audit.strict_passed:
        raise ValueError("Phase 3C+4 requires the strict 12-page source audit to pass")
    if len(document.entries) != settings.phase3c4_expected_block_count:
        raise ValueError("Phase 3C+4 source block count is not exactly 24")
    return path, document


def _title(blocks: dict[str, Phase3BSourceLogicalBlock], block_id: str | None, target: str) -> tuple[str | None, str | None]:
    if not block_id or block_id == target:
        return None, None
    block = blocks.get(block_id)
    if not block or block.block_type not in TITLE_TYPES:
        return None, None
    return block.logical_block_id, block.source_text


def build_phase3c4_units(settings: ProjectSettings, *, root: Path | None = None) -> list[Phase3C4TranslationUnit]:
    root = (root or project_root()).resolve()
    _, source = _source_document(settings, root)
    eligible = [
        block for block in source.entries
        if is_phase3c4_translatable_type(block.block_type)
        and block.translation_ready and not block.unresolved_boundaries and block.source_text.strip()
    ]
    if len(eligible) != settings.phase3c4_expected_block_count:
        raise ValueError("Exactly 24 complete translatable blocks are required")
    by_id = {block.logical_block_id: block for block in source.entries}
    body_by_chapter: dict[str | None, list[Phase3BSourceLogicalBlock]] = {}
    for block in eligible:
        if block.block_type == "body":
            body_by_chapter.setdefault(block.chapter_id, []).append(block)
    units: list[Phase3C4TranslationUnit] = []
    for block in eligible:
        before = after = None
        if block.block_type == "body":
            sequence = body_by_chapter.get(block.chapter_id, [])
            index = sequence.index(block)
            before = sequence[index - 1] if index else None
            after = sequence[index + 1] if index + 1 < len(sequence) else None
        chapter_id, chapter_text = _title(by_id, block.chapter_id, block.logical_block_id)
        section_id, section_text = _title(by_id, block.section_id, block.logical_block_id)
        units.append(Phase3C4TranslationUnit(
            target_block_id=block.logical_block_id,
            block_type=block.block_type,
            source_text=block.source_text,
            source_pages=block.source_pages,
            source_fragment_ids=block.source_fragment_ids,
            completeness_status=block.completeness_status,
            chapter_id=block.chapter_id,
            section_id=block.section_id,
            chapter_title_block_id=chapter_id,
            section_title_block_id=section_id,
            chapter_title_context=chapter_text,
            section_title_context=section_text,
            context_before_block_ids=[before.logical_block_id] if before else [],
            context_after_block_ids=[after.logical_block_id] if after else [],
            context_before_text=before.source_text if before else None,
            context_after_text=after.source_text if after else None,
        ))
    return units


def _profile_text(settings: ProjectSettings, root: Path) -> str:
    path = resolve_project_path(settings.translation_language_profile_path, root=root)
    profile = yaml.safe_load(path.read_text(encoding="utf-8"))
    return yaml.safe_dump(profile, allow_unicode=True, sort_keys=True)


def phase3c4_request_fingerprint(
    unit: Phase3C4TranslationUnit,
    settings: ProjectSettings,
    *, prompt_sha256: str,
    profile_sha256: str,
) -> str:
    return stable_hash({
        "phase": "3C+4",
        "unit": unit.model_dump(mode="json"),
        "prompt_version": settings.phase3c4_prompt_version,
        "prompt_sha256": prompt_sha256,
        "profile_version": settings.translation_language_profile_version,
        "profile_sha256": profile_sha256,
        "provider": settings.translation_provider,
        "base_url": settings.translation_base_url,
        "model": settings.translation_model,
        "thinking_mode": settings.translation_thinking_mode,
        "temperature": settings.translation_temperature,
        "max_output_tokens": settings.translation_max_output_tokens,
        "response_format": "json_object",
    })


def _data_root(settings: ProjectSettings, root: Path) -> Path:
    return resolve_project_path(settings.phase3c4_data_directory, root=root)


def _call_paths(settings: ProjectSettings, root: Path, fingerprint: str) -> dict[str, Path]:
    folder = _data_root(settings, root) / "calls" / fingerprint
    return {name: folder / f"{name}.json" for name in ("request", "raw", "normalized", "usage", "cache")}


def _load_cached(path: Path, fingerprint: str) -> Phase3C4TranslationResult | None:
    if not path.is_file():
        return None
    result = Phase3C4TranslationResult.model_validate(load_json(path))
    return result if result.request_fingerprint == fingerprint else None


def normalize_phase3c4_translation(
    *, content: str, unit: Phase3C4TranslationUnit, settings: ProjectSettings,
    fingerprint: str, prompt_sha256: str, profile_sha256: str,
    raw_path: Path, usage: dict[str, Any] | None, request_id: str | None,
) -> Phase3C4TranslationResult:
    try:
        payload = Phase3C4ModelPayload.model_validate_json(content)
    except Exception as error:
        raise ValueError(f"Invalid Phase 3C+4 JSON output: {error}") from error
    # target_block_id is program control metadata; the authoritative value
    # comes from the translation unit, not the model echo.  A mismatch
    # (e.g. a single garbled hex character) is repaired rather than
    # discarding an otherwise correct translation.
    _returned_id = payload.target_block_id
    _id_mismatch = _returned_id != unit.target_block_id
    if _id_mismatch:
        payload = payload.model_copy(update={"target_block_id": unit.target_block_id})
    if payload.block_type != unit.block_type:
        raise ValueError("block_type does not match the requested unit")
    translation = payload.translation.strip()
    if not translation:
        raise ValueError("translation is empty")
    for label, context in (("context_before", unit.context_before_text), ("context_after", unit.context_after_text)):
        if context and translation == context.strip():
            raise ValueError(f"translation copied {label} instead of translating source_text")
    # Only enforce strict copy/length checks on prose-like body blocks, not
    # short labels, captions, dates, abbreviations or table-of-contents entries.
    _source_stripped = unit.source_text.strip()
    _english_word_count = len(_source_stripped.split())
    _is_prose_body = (
        unit.block_type == "body"
        and len(_source_stripped) > 50
        and _english_word_count >= 8
    )
    if (
        translation == _source_stripped
        and unit.block_type == "body"
        and len(_source_stripped) > 30
        and _english_word_count >= 5
    ):
        raise ValueError("translation copied the English source_text")
    if _is_prose_body and len(translation) < max(8, int(len(unit.source_text) * 0.12)):
        raise ValueError("translation is implausibly short for a complete body block")
    terms = [item.strip() for item in payload.untranslated_source_terms if item.strip()]
    if len(terms) != len(set(terms)):
        raise ValueError("untranslated_source_terms contains duplicates")
    warnings = list(payload.warnings)
    missing_preserved_terms: list[str] = []
    for term in terms:
        if term not in translation:
            missing_preserved_terms.append(term)
    if missing_preserved_terms:
        translation = translation.rstrip() + "（原文保留：" + "；".join(missing_preserved_terms) + "）"
        warnings.append(
            "Provider listed source term(s) without preserving them in the translation; "
            "the exact terms were deterministically appended during normalization."
        )
    context_hash = stable_hash({
        "before": unit.context_before_text,
        "after": unit.context_after_text,
        "chapter": unit.chapter_title_context,
        "section": unit.section_title_context,
    })
    normalized_payload = payload.model_dump()
    normalized_payload["translation"] = translation
    normalized_payload["untranslated_source_terms"] = terms
    normalized_payload["warnings"] = warnings
    return Phase3C4TranslationResult(
        **normalized_payload,
        source_text_sha256=sha256_text(unit.source_text),
        context_sha256=context_hash,
        request_fingerprint=fingerprint,
        prompt_version=settings.phase3c4_prompt_version,
        prompt_sha256=prompt_sha256,
        language_profile_sha256=profile_sha256,
        raw_response_path=str(raw_path), request_id=request_id, usage=usage,
    )


def normalize_saved_phase3c4_response(
    settings: ProjectSettings, *, target_block_id: str, root: Path | None = None,
) -> Phase3C4TranslationResult:
    """Normalize an already-saved provider response without making a network call."""

    root = (root or project_root()).resolve()
    unit = next(
        (item for item in build_phase3c4_units(settings, root=root) if item.target_block_id == target_block_id),
        None,
    )
    if unit is None:
        raise ValueError(f"Unknown Phase 3C+4 target block: {target_block_id}")
    prompt = resolve_project_path(settings.phase3c4_prompt_path, root=root).read_text(encoding="utf-8")
    profile = _profile_text(settings, root)
    prompt_hash, profile_hash = sha256_text(prompt), sha256_text(profile)
    fingerprint = phase3c4_request_fingerprint(
        unit, settings, prompt_sha256=prompt_hash, profile_sha256=profile_hash
    )
    paths = _call_paths(settings, root, fingerprint)
    if not paths["raw"].is_file():
        raise FileNotFoundError(f"Saved raw Phase 3C+4 response not found: {paths['raw']}")
    raw = load_json(paths["raw"])
    choices = raw.get("choices") or []
    message = choices[0].get("message") if choices and isinstance(choices[0], dict) else None
    content = message.get("content") if isinstance(message, dict) else None
    if not isinstance(content, str) or not content.strip():
        raise ValueError("Saved raw Phase 3C+4 response has no JSON content")
    usage = load_json(paths["usage"]) if paths["usage"].is_file() else raw.get("usage")
    result = normalize_phase3c4_translation(
        content=content, unit=unit, settings=settings, fingerprint=fingerprint,
        prompt_sha256=prompt_hash, profile_sha256=profile_hash, raw_path=paths["raw"],
        usage=usage if isinstance(usage, dict) else None,
        request_id=str(raw.get("id")) if raw.get("id") else None,
    )
    atomic_write_json(paths["normalized"], result)
    atomic_write_json(paths["cache"], result)
    ledger_path, ledger = _ledger(settings, root)
    for entry in reversed(ledger.get("entries", [])):
        if entry.get("fingerprint") == fingerprint and entry.get("status") == "failed":
            entry["status"] = "completed_offline_normalization"
            entry["completed_at"] = _now().isoformat()
            entry["normalization_only"] = True
            break
    atomic_write_json(ledger_path, ledger)
    return result


def phase3c4_preflight(
    settings: ProjectSettings, *, root: Path | None = None,
    key_status_resolver: Callable[[ProjectSettings, Path], tuple[bool, str | None]] = translation_api_key_status,
) -> Phase3C4Preflight:
    root = (root or project_root()).resolve()
    units = build_phase3c4_units(settings, root=root)
    prompt = resolve_project_path(settings.phase3c4_prompt_path, root=root).read_text(encoding="utf-8")
    profile = _profile_text(settings, root)
    prompt_hash, profile_hash = sha256_text(prompt), sha256_text(profile)
    cached = 0
    for unit in units:
        fingerprint = phase3c4_request_fingerprint(unit, settings, prompt_sha256=prompt_hash, profile_sha256=profile_hash)
        if _load_cached(_call_paths(settings, root, fingerprint)["cache"], fingerprint):
            cached += 1
    source_chars = sum(len(unit.source_text) for unit in units)
    context_chars = sum(
        len(value or "") for unit in units for value in (
            unit.context_before_text, unit.context_after_text,
            unit.chapter_title_context, unit.section_title_context,
        )
    )
    serialized_chars = sum(len(json.dumps(unit.provider_payload(), ensure_ascii=False)) for unit in units)
    input_tokens = math.ceil((len(prompt) * len(units) + len(profile) * len(units) + serialized_chars) / 4)
    output_tokens = math.ceil(source_chars * 0.65)
    max_output = len(units) * settings.translation_max_output_tokens
    lower = (input_tokens * settings.translation_input_cache_hit_price_cny_per_million_tokens + output_tokens * settings.translation_output_price_cny_per_million_tokens) / 1_000_000
    upper = (input_tokens * settings.translation_input_cache_miss_price_cny_per_million_tokens + max_output * settings.translation_output_price_cny_per_million_tokens) / 1_000_000
    key_set, _ = key_status_resolver(settings, root)
    new_calls = len(units) - cached
    blockers: list[str] = []
    if len(units) != settings.phase3c4_expected_block_count:
        blockers.append("target block count is not 24")
    if new_calls > settings.phase3c4_maximum_content_calls:
        blockers.append("estimated calls exceed the 24-call hard limit")
    if upper > settings.phase3c4_maximum_cash_cost_cny:
        blockers.append("upper cost estimate exceeds the 2.00 CNY hard limit")
    if new_calls and not key_set:
        blockers.append("translation API key is not set")
    if settings.phase3c4_automatic_retry or settings.translation_automatic_retry:
        blockers.append("automatic retry must remain disabled")
    if not settings.full_pdf_protection:
        blockers.append("full PDF protection must remain enabled")
    if settings.automatic_phase_advance:
        blockers.append("automatic phase advance must remain disabled")
    return Phase3C4Preflight(
        target_block_count=len(units), cached_block_count=cached,
        estimated_new_content_calls=new_calls,
        maximum_content_calls=settings.phase3c4_maximum_content_calls,
        maximum_model_list_calls=settings.phase3c4_maximum_model_list_calls,
        total_source_characters=source_chars, total_context_characters=context_chars,
        estimated_input_tokens=input_tokens, estimated_output_tokens=output_tokens,
        maximum_output_tokens=max_output,
        estimated_cost_lower_cny=round(lower, 6), estimated_cost_upper_cny=round(upper, 6),
        maximum_cash_cost_cny=settings.phase3c4_maximum_cash_cost_cny,
        api_key_set=key_set, blockers=blockers, ready_for_real_call=not blockers,
    )


def _ledger(settings: ProjectSettings, root: Path) -> tuple[Path, dict[str, Any]]:
    path = _data_root(settings, root) / "call_ledger.json"
    if path.is_file():
        return path, load_json(path)
    return path, {"schema_version": "1.0", "phase": "3C+4", "content_calls_started": 0, "entries": []}


def run_phase3c4_translation(
    settings: ProjectSettings, *, allow_api: bool, confirmed: bool,
    root: Path | None = None,
    key_status_resolver: Callable[[ProjectSettings, Path], tuple[bool, str | None]] = translation_api_key_status,
    key_loader: Callable[[ProjectSettings, Path], tuple[str, str]] = load_translation_api_key,
    provider_factory: Callable[..., Any] = DeepSeekOpenAICompatibleProvider,
) -> Phase3C4BatchResult:
    root = (root or project_root()).resolve()
    preflight = phase3c4_preflight(settings, root=root, key_status_resolver=key_status_resolver)
    units = build_phase3c4_units(settings, root=root)
    prompt = resolve_project_path(settings.phase3c4_prompt_path, root=root).read_text(encoding="utf-8")
    profile = _profile_text(settings, root)
    system_prompt = prompt + "\n\n语言风格配置（仅作约束）：\n" + profile
    prompt_hash, profile_hash = sha256_text(prompt), sha256_text(profile)
    results: list[Phase3C4TranslationResult] = []
    pending: list[tuple[Phase3C4TranslationUnit, str, dict[str, Path]]] = []
    for unit in units:
        fingerprint = phase3c4_request_fingerprint(unit, settings, prompt_sha256=prompt_hash, profile_sha256=profile_hash)
        paths = _call_paths(settings, root, fingerprint)
        cached = _load_cached(paths["cache"], fingerprint)
        if cached:
            results.append(cached)
        else:
            pending.append((unit, fingerprint, paths))
    cache_hits = len(results)
    if not pending or not allow_api:
        return Phase3C4BatchResult(results=results, api_calls=0, cache_hits=cache_hits, failed=0, preflight=preflight)
    if not confirmed:
        raise PermissionError("Phase 3C+4 real calls require explicit confirmation")
    if not preflight.ready_for_real_call:
        raise RuntimeError("Phase 3C+4 preflight failed: " + "; ".join(preflight.blockers))
    api_key, _ = key_loader(settings, root)
    provider = provider_factory(api_key=api_key, base_url=settings.translation_base_url, timeout_seconds=settings.translation_request_timeout_seconds)
    ledger_path, ledger = _ledger(settings, root)
    calls = 0
    for unit, fingerprint, paths in pending:
        if int(ledger.get("content_calls_started", 0)) >= settings.phase3c4_maximum_content_calls:
            raise RuntimeError("Phase 3C+4 persistent 24-call limit reached")
        request = {
            "schema_version": "1.0", "phase": "3C+4", "request_fingerprint": fingerprint,
            "provider": settings.translation_provider, "base_url": settings.translation_base_url,
            "model": settings.translation_model, "thinking_mode": settings.translation_thinking_mode,
            "prompt_version": settings.phase3c4_prompt_version, "prompt_sha256": prompt_hash,
            "language_profile_sha256": profile_hash, "payload": unit.provider_payload(),
        }
        atomic_write_json(paths["request"], request)
        ledger["content_calls_started"] = int(ledger.get("content_calls_started", 0)) + 1
        ledger.setdefault("entries", []).append({"fingerprint": fingerprint, "target_block_id": unit.target_block_id, "started_at": _now().isoformat(), "status": "started"})
        atomic_write_json(ledger_path, ledger)
        calls += 1
        try:
            response = provider.translate_one(
                model=settings.translation_model, system_prompt=system_prompt,
                user_payload=unit.provider_payload(), max_output_tokens=settings.translation_max_output_tokens,
                temperature=settings.translation_temperature, thinking_mode=settings.translation_thinking_mode,
            )
            atomic_write_json(paths["raw"], response.raw_response)
            if response.usage is not None:
                atomic_write_json(paths["usage"], response.usage)
            normalized = normalize_phase3c4_translation(
                content=response.content, unit=unit, settings=settings, fingerprint=fingerprint,
                prompt_sha256=prompt_hash, profile_sha256=profile_hash, raw_path=paths["raw"],
                usage=response.usage, request_id=response.request_id,
            )
            atomic_write_json(paths["normalized"], normalized)
            atomic_write_json(paths["cache"], normalized)
            ledger["entries"][-1]["status"] = "completed"
            ledger["entries"][-1]["completed_at"] = _now().isoformat()
            atomic_write_json(ledger_path, ledger)
            results.append(normalized)
        except Exception as error:
            ledger["entries"][-1]["status"] = "failed"
            ledger["entries"][-1]["error_type"] = type(error).__name__
            atomic_write_json(ledger_path, ledger)
            raise
    order = {unit.target_block_id: index for index, unit in enumerate(units)}
    results.sort(key=lambda item: order[item.target_block_id])
    return Phase3C4BatchResult(results=results, api_calls=calls, cache_hits=cache_hits, failed=0, preflight=preflight)


def build_phase3c4_bilingual_document(
    settings: ProjectSettings, results: list[Phase3C4TranslationResult], *, root: Path | None = None,
) -> Phase3C4BilingualDocument:
    root = (root or project_root()).resolve()
    source_path, source = _source_document(settings, root)
    units = build_phase3c4_units(settings, root=root)
    expected = [unit.target_block_id for unit in units]
    counts = Counter(item.target_block_id for item in results)
    duplicates = sorted(item for item, count in counts.items() if count > 1)
    if duplicates:
        raise ValueError(f"duplicate translation results: {duplicates}")
    actual = set(counts)
    missing = [item for item in expected if item not in actual]
    if missing:
        raise ValueError(f"missing translation results: {missing}")
    extra = sorted(actual - set(expected))
    if extra:
        raise ValueError(f"extra translation results: {extra}")
    by_result = {item.target_block_id: item for item in results}
    by_source = {item.logical_block_id: item for item in source.entries}
    blocks: list[Phase3C4BilingualBlock] = []
    modified = False
    for block_id in expected:
        source_block = by_source[block_id]
        result = by_result[block_id]
        if result.source_text_sha256 != sha256_text(source_block.source_text):
            modified = True
            raise ValueError(f"source text hash changed for {block_id}")
        blocks.append(Phase3C4BilingualBlock(
            block_id=block_id, block_type=source_block.block_type,
            chapter_id=source_block.chapter_id, section_id=source_block.section_id,
            source_pages=source_block.source_pages, source_fragment_ids=source_block.source_fragment_ids,
            source_block_ids=source_block.source_block_ids, source_text=source_block.source_text,
            completeness_status=source_block.completeness_status,
            translation=result.translation, untranslated_source_terms=result.untranslated_source_terms,
            warnings=result.warnings, raw_response_reference=result.raw_response_path,
            cache_fingerprint=result.request_fingerprint,
            usage_reference=str(_call_paths(settings, root, result.request_fingerprint)["usage"])
            if _call_paths(settings, root, result.request_fingerprint)["usage"].is_file() else None,
        ))
    audit = Phase3C4AlignmentAudit(
        source_block_count=len(expected), translated_block_count=len(blocks),
        missing_translation_ids=[], duplicate_translation_ids=[], extra_translation_ids=[],
        source_text_modified=modified,
        untranslated_source_term_count=sum(len(item.untranslated_source_terms) for item in blocks),
        strict_passed=True,
    )
    document = Phase3C4BilingualDocument(
        document_id=source.document_id, source_document_path=str(source_path),
        source_document_sha256=sha256_file(source_path), page_count=12, logical_blocks=blocks,
        chapters=[Phase3C4StructuralEntry(
            block_id=item.block_id, block_type=item.block_type,
            source_text=item.source_text, translation=item.translation,
        ) for item in blocks if item.block_type == "chapter_title"],
        sections=[Phase3C4StructuralEntry(
            block_id=item.block_id, block_type=item.block_type,
            source_text=item.source_text, translation=item.translation,
        ) for item in blocks if item.block_type == "section_title"],
        audit=audit, created_at=_now(),
    )
    atomic_write_json(resolve_project_path(settings.phase3c4_master_path, root=root), document)
    return document


def _markdown(document: Phase3C4BilingualDocument, *, bilingual: bool) -> str:
    lines: list[str] = []
    for block in document.logical_blocks:
        prefix = "# " if block.block_type == "chapter_title" else "## " if block.block_type == "section_title" else ""
        lines.append(prefix + block.source_text.strip())
        if bilingual:
            lines.extend(["", block.translation.strip()])
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _style_doc(document: Document) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, before, after in (("Title", 18, 18, 10), ("Heading 1", 16, 18, 10), ("Heading 2", 13, 12, 6)):
        style = document.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.color.rgb = RGBColor(45, 58, 76)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)


def _docx(document_data: Phase3C4BilingualDocument, path: Path, *, bilingual: bool, canonical_hash: str) -> None:
    document = Document()
    _style_doc(document)
    document.core_properties.comments = f"Generated from bilingual_document_sample12_zh-Hans_v1.json sha256={canonical_hash}"
    for block in document_data.logical_blocks:
        if block.block_type == "chapter_title":
            paragraph = document.add_paragraph(block.source_text, style="Heading 1")
        elif block.block_type == "section_title":
            paragraph = document.add_paragraph(block.source_text, style="Heading 2")
        else:
            paragraph = document.add_paragraph(block.source_text)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if "\n" in block.source_text
                else WD_ALIGN_PARAGRAPH.JUSTIFY
            )
        if bilingual:
            translation = document.add_paragraph(block.translation)
            translation.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            translation.paragraph_format.space_after = Pt(12)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def export_phase3c4_candidates(
    settings: ProjectSettings, *, master_path: str | Path, root: Path | None = None,
) -> Phase3C4ExportResult:
    root = (root or project_root()).resolve()
    canonical = resolve_project_path(master_path, root=root)
    data = Phase3C4BilingualDocument.model_validate(load_json(canonical))
    if not data.audit.strict_passed:
        raise ValueError("Candidate export requires a passing bilingual alignment audit")
    destination = resolve_project_path(settings.phase3c4_candidate_directory, root=root)
    source_md = destination / "source_english_sample12.md"
    source_docx = destination / "source_english_sample12.docx"
    bilingual_md = destination / "bilingual_zh-Hans_sample12.md"
    bilingual_docx = destination / "bilingual_zh-Hans_sample12.docx"
    canonical_hash = sha256_file(canonical)
    atomic_write_text(source_md, _markdown(data, bilingual=False))
    atomic_write_text(bilingual_md, _markdown(data, bilingual=True))
    _docx(data, source_docx, bilingual=False, canonical_hash=canonical_hash)
    _docx(data, bilingual_docx, bilingual=True, canonical_hash=canonical_hash)
    return Phase3C4ExportResult(
        source_json_sha256=canonical_hash, source_markdown_path=str(source_md),
        source_docx_path=str(source_docx), bilingual_markdown_path=str(bilingual_md),
        bilingual_docx_path=str(bilingual_docx),
    )


def detect_libreoffice(settings: ProjectSettings, *, root: Path | None = None) -> Path | None:
    root = (root or project_root()).resolve()
    for item in settings.phase3c4_soffice_candidates:
        path = resolve_project_path(item, root=root)
        if path.is_file():
            return path
    return None


def run_libreoffice_smoke(
    settings: ProjectSettings, *, source_docx: str | Path, bilingual_docx: str | Path,
    executable: str | Path | None = None,
    command_runner: Callable[..., Any] = subprocess.run,
    root: Path | None = None,
) -> LibreOfficeSmokeResult:
    root = (root or project_root()).resolve()
    exe = Path(executable) if executable else detect_libreoffice(settings, root=root)
    if exe is None or not exe.is_file():
        return LibreOfficeSmokeResult(
            libreoffice_found=False, executable_path=None, completed=False, fatal_error=False,
            source_pdf_path=None, bilingual_pdf_path=None,
            source_pdf_page_count=None, bilingual_pdf_page_count=None,
            warnings=["LibreOffice was not found; DOCX files remain available."],
        )
    output = resolve_project_path(settings.phase3c4_rendered_directory, root=root)
    output.mkdir(parents=True, exist_ok=True)
    profile = resolve_project_path(settings.phase3c4_libreoffice_profile_directory, root=root)
    profile.mkdir(parents=True, exist_ok=True)
    profile_url = "file:///" + quote(str(profile).replace("\\", "/"), safe="/:/-")
    warnings: list[str] = []
    converted: list[Path] = []
    try:
        for docx in (Path(source_docx), Path(bilingual_docx)):
            command = [str(exe), "--headless", f"-env:UserInstallation={profile_url}", "--convert-to", "pdf", "--outdir", str(output), str(docx)]
            completed = command_runner(command, capture_output=True, text=True, timeout=180, check=False)
            if completed.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed for {docx.name}")
            pdf = output / f"{docx.stem}.pdf"
            if not pdf.is_file() or pdf.stat().st_size == 0:
                raise RuntimeError(f"LibreOffice did not produce {pdf.name}")
            converted.append(pdf)
        counts = []
        for pdf_path in converted:
            with fitz.open(pdf_path) as pdf:
                if pdf.page_count < 1:
                    raise RuntimeError(f"Rendered PDF is empty: {pdf_path.name}")
                counts.append(pdf.page_count)
        return LibreOfficeSmokeResult(
            libreoffice_found=True, executable_path=str(exe), completed=True, fatal_error=False,
            source_pdf_path=str(converted[0]), bilingual_pdf_path=str(converted[1]),
            source_pdf_page_count=counts[0], bilingual_pdf_page_count=counts[1], warnings=warnings,
        )
    except Exception as error:
        return LibreOfficeSmokeResult(
            libreoffice_found=True, executable_path=str(exe), completed=False, fatal_error=True,
            source_pdf_path=str(converted[0]) if converted else None,
            bilingual_pdf_path=str(converted[1]) if len(converted) > 1 else None,
            source_pdf_page_count=None, bilingual_pdf_page_count=None,
            warnings=[str(error)],
        )
