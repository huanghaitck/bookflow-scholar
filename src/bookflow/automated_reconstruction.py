"""Offline Phase 2B.2 automated, traceable, fail-closed reconstruction."""

from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel

from .io_utils import atomic_write_json, atomic_write_jsonl, load_json, sha256_file, stable_hash
from .page_pipeline import build_context
from .paths import ProjectSettings, project_root, resolve_project_path
from .phase2a1 import VisionNormalizedPageV11
from .phase2b_calls import load_all_page_results, load_latest_boundaries
from .phase2b2_schemas import (
    AutomatedBoundary,
    AutomatedLogicalBlock,
    AutomatedPageRecord,
    BilingualDocument,
    BilingualEntry,
    CarryRecord,
    LogicalReconstructionAudit,
    SourceCoverageAudit,
    SourceFragment,
    TranslationAlignmentAudit,
    TranslationContextV2,
)
from .schemas import PageRecord


CONTENT_TYPES = {"body", "chapter_title", "section_title", "footnote", "caption"}
BODY_TYPE = "body"


class AutomatedRunResult(BaseModel):
    page_records_path: str
    boundary_records_path: str
    logical_blocks_path: str
    translation_context_path: str
    source_audit_path: str
    logical_audit_path: str
    translation_audit_path: str
    master_document_path: str
    diagnostic_markdown_path: str | None
    diagnostic_word_path: str | None
    pages: int
    boundaries: int
    resolved_boundaries: int
    unresolved_boundaries: int
    logical_blocks: int
    cross_page_blocks: int
    translation_ready_true: int
    translation_ready_false: int
    fragments: int
    source_audit_passed: bool
    logical_audit_passed: bool
    strict_export_ready: bool
    api_calls: int = 0
    deepseek_calls: int = 0
    translation_calls: int = 0


class ExportResult(BaseModel):
    mode: str
    markdown_path: str | None
    word_path: str | None
    blocked: bool
    blockers: list[str]


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _sample_only(pdf_path: str | Path, settings: ProjectSettings, root: Path) -> Path:
    source = resolve_project_path(pdf_path, root=root)
    sample = resolve_project_path(settings.sample_pdf, root=root)
    protected = resolve_project_path(settings.source_pdf, root=root)
    if source != sample or source == protected:
        raise PermissionError("Phase 2B.2 only accepts the configured 11-page sample")
    with fitz.open(source) as document:
        if document.page_count != 11:
            raise ValueError(f"Phase 2B.2 requires the actual 11-page sample, found {document.page_count}")
    return source


def _record(context: Any, page: int) -> PageRecord:
    path = Path(context.record_directory) / f"page_{page:04d}.json"
    return PageRecord.model_validate(load_json(path))


def _latest_legacy_path(settings: ProjectSettings, root: Path, page: int) -> Path:
    directory = (
        resolve_project_path(settings.vision_normalized_v11_directory, root=root)
        / re.sub(r"[^A-Za-z0-9_.-]+", "_", settings.vision_provider)
        / re.sub(r"[^A-Za-z0-9_.-]+", "_", settings.vision_model)
        / f"page_{page:04d}"
    )
    candidates: list[tuple[datetime, Path]] = []
    for path in directory.glob("*.json") if directory.is_dir() else []:
        try:
            item = VisionNormalizedPageV11.model_validate(load_json(path))
        except Exception:
            continue
        candidates.append((item.normalized_at, path))
    if not candidates:
        raise FileNotFoundError(f"No cached normalized visual result for page {page}")
    return max(candidates, key=lambda item: item[0])[1]


def _clean_for_compare(text: str) -> str:
    value = text.casefold().replace("-­", "")
    value = re.sub(r"-\s+", "", value)
    value = re.sub(r"[^a-z0-9]+", "", value)
    return value


def _similarity(first: str, second: str) -> float:
    left, right = _clean_for_compare(first), _clean_for_compare(second)
    if not left and not right:
        return 1.0
    if not left or not right:
        return 0.0
    # Character-level book text contains many repeated letters; disabling autojunk
    # prevents long pages from being falsely scored near zero.
    return SequenceMatcher(None, left, right, autojunk=False).ratio()


def _strip_closing(text: str) -> str:
    return text.rstrip().rstrip('"\'”’)]}')


def _ends_sentence(text: str) -> bool:
    value = _strip_closing(text)
    return bool(value) and value[-1] in ".!?…"


def _starts_lower(text: str) -> bool:
    match = re.search(r"[A-Za-z]", text)
    return bool(match and match.group(0).islower())


def _starts_quote(text: str) -> bool:
    return text.lstrip().startswith(('"', "'", "“", "‘"))


def _within_page_continuation(left: str, right: str) -> bool:
    if _ends_sentence(left):
        return False
    ending = _strip_closing(left).rstrip()
    return bool(
        ending
        and (
            ending[-1] in ",;:—–-"
            or _starts_lower(right)
            or _starts_quote(right)
        )
    )


def _edge_relation(previous_text: str, next_text: str | None, chapter_break: bool = False) -> str:
    if next_text is None:
        return "unresolved" if not _ends_sentence(previous_text) else "no_join"
    if chapter_break:
        return "no_join"
    if _ends_sentence(previous_text):
        return "no_join"
    if _starts_lower(next_text) or next_text.lstrip().startswith((",", ";", ":")):
        return "continue"
    return "unresolved"


def _fragment_id(document_id: str, page: int, block_id: str, text: str) -> str:
    return "fragment_" + stable_hash(
        {"document_id": document_id, "page": page, "block_id": block_id, "text": text}
    )[:20]


def build_automated_pages(
    pdf_path: str | Path,
    settings: ProjectSettings,
    *,
    root: Path | None = None,
) -> list[AutomatedPageRecord]:
    root = (root or project_root()).resolve()
    source = _sample_only(pdf_path, settings, root)
    legacy = load_all_page_results(settings, root)
    if len(legacy) != 11:
        raise RuntimeError("All eleven cached visual results are required")
    context = build_context(
        source,
        settings,
        pages=list(range(1, 12)),
        dpi=settings.render_dpi,
        color_mode=settings.render_color_mode,
        image_format=settings.render_format,
        root=root,
    )
    raw_body: dict[int, list[Any]] = {
        page: [block for block in sorted(legacy[page].blocks, key=lambda item: item.order) if block.block_type == BODY_TYPE]
        for page in range(1, 12)
    }
    external: dict[tuple[int, int], str] = {}
    for page in range(1, 11):
        previous = raw_body[page][-1].text if raw_body[page] else ""
        following = raw_body[page + 1][0].text if raw_body[page + 1] else ""
        chapter = any(block.block_type == "chapter_title" for block in legacy[page + 1].blocks)
        external[(page, page + 1)] = _edge_relation(previous, following, chapter)
    external[(11, 12)] = _edge_relation(raw_body[11][-1].text if raw_body[11] else "", None)

    pages: list[AutomatedPageRecord] = []
    with fitz.open(source) as document:
        for page in range(1, 12):
            item = legacy[page]
            manifest_record = _record(context, page)
            selected = [block for block in sorted(item.blocks, key=lambda value: value.order) if block.block_type in CONTENT_TYPES]
            body = [block for block in selected if block.block_type == BODY_TYPE]
            first_body_id = body[0].block_id if body else None
            last_body_id = body[-1].block_id if body else None
            fragments: list[SourceFragment] = []
            previous_body: Any | None = None
            for block in selected:
                is_body = block.block_type == BODY_TYPE
                body_index = body.index(block) if is_body else -1
                next_body = body[body_index + 1] if is_body and body_index + 1 < len(body) else None
                starts_external = bool(
                    is_body
                    and block.block_id == first_body_id
                    and page > 1
                    and external.get((page - 1, page)) == "continue"
                )
                starts_internal = bool(
                    is_body
                    and previous_body is not None
                    and _within_page_continuation(previous_body.text, block.text)
                )
                ends_internal = bool(
                    is_body and next_body is not None and _within_page_continuation(block.text, next_body.text)
                )
                ends_external = bool(
                    is_body
                    and block.block_id == last_body_id
                    and external.get((page, page + 1)) in {"continue", "unresolved"}
                )
                fragment = SourceFragment(
                    fragment_id=_fragment_id(item.document_id, page, block.block_id, block.text),
                    text=block.text,
                    source_page=page,
                    source_block_ids=[f"p{page:04d}:{block.block_id}"],
                    block_type=block.block_type,
                    order=block.order,
                    starts_mid_sentence=starts_external or starts_internal,
                    ends_mid_sentence=ends_external or ends_internal,
                    starts_mid_paragraph=starts_external or starts_internal,
                    ends_mid_paragraph=ends_external or ends_internal,
                    visible_trailing_hyphen=block.text.rstrip().endswith("-"),
                    uncertainty=list(item.uncertain_characters) if block.uncertain else [],
                )
                fragments.append(fragment)
                if is_body:
                    previous_body = block
            by_block = {fragment.source_block_ids[0].split(":", 1)[1]: fragment for fragment in fragments}
            head = by_block.get(first_body_id) if first_body_id and page > 1 and external.get((page - 1, page)) == "continue" else None
            tail = by_block.get(last_body_id) if last_body_id and external.get((page, page + 1)) in {"continue", "unresolved"} else None
            titles: list[str] = []
            for value in [item.title] + [block.text for block in selected if block.block_type in {"chapter_title", "section_title"}]:
                if value and value not in titles:
                    titles.append(value)
            visible_parts: list[str] = []
            for value in [item.running_header, item.title] + [block.text for block in selected] + [item.footer, item.page_number_text]:
                if value and value not in visible_parts:
                    visible_parts.append(value)
            full_visible = "\n\n".join(visible_parts)
            text_layer = document.load_page(page - 1).get_text("text")
            ratio = _similarity(full_visible, text_layer)
            all_blocks_recorded = len(fragments) == len(selected) and all(fragment.text for fragment in fragments)
            coverage = "complete" if all_blocks_recorded and ratio >= settings.automated_text_layer_coverage_threshold else "partial"
            pages.append(
                AutomatedPageRecord(
                    schema_version=settings.automated_page_schema_version,
                    document_id=item.document_id,
                    pdf_page=page,
                    printed_page=item.printed_page,
                    page_type=item.page_type,
                    full_visible_text=full_visible,
                    complete_blocks=[
                        fragment.fragment_id
                        for fragment in fragments
                        if not any(
                            [
                                fragment.starts_mid_sentence,
                                fragment.ends_mid_sentence,
                                fragment.starts_mid_paragraph,
                                fragment.ends_mid_paragraph,
                            ]
                        )
                    ],
                    head_fragment=head,
                    tail_fragment=tail,
                    content_fragments=fragments,
                    running_header=item.running_header,
                    footer=item.footer,
                    page_number_text=item.page_number_text,
                    titles=titles,
                    image_sha256=manifest_record.image_sha256,
                    text_layer_text=text_layer,
                    text_layer_similarity=ratio,
                    transcription_status=item.status,
                    source_coverage_status=coverage,
                    legacy_normalized_path=str(_latest_legacy_path(settings, root, page).resolve()),
                    created_at=_now(),
                )
            )
    output = resolve_project_path(settings.automated_page_directory, root=root)
    atomic_write_jsonl(output / "sample_11_pages.pages.jsonl", pages)
    for page in pages:
        atomic_write_json(output / "pages" / f"page_{page.pdf_page:04d}.json", page)
    return pages


def _legacy_model_evidence(settings: ProjectSettings, root: Path) -> dict[str, Any]:
    effective = load_latest_boundaries(settings, root, "pair")
    effective.update(load_latest_boundaries(settings, root, "triple"))
    return effective


def build_automated_boundaries(
    pages: list[AutomatedPageRecord],
    settings: ProjectSettings,
    *,
    root: Path,
    include_open_boundary: bool = True,
    legacy_evidence: dict[str, Any] | None = None,
    artifact_stem: str = "sample_11_pages",
    boundary_page_pairs: list[tuple[int, int]] | None = None,
) -> list[AutomatedBoundary]:
    legacy = _legacy_model_evidence(settings, root) if legacy_evidence is None else legacy_evidence
    by_page = {page.pdf_page: page for page in pages}
    if sorted(by_page) != list(range(1, len(pages) + 1)):
        raise ValueError("Automated page records must be continuous and one-based")
    results: list[AutomatedBoundary] = []
    if boundary_page_pairs is None:
        boundary_count = len(pages) if include_open_boundary else max(0, len(pages) - 1)
        pairs = [(previous, previous + 1) for previous in range(1, boundary_count + 1)]
    else:
        pairs = list(boundary_page_pairs)
    for previous, next_page in pairs:
        left_page = by_page[previous]
        left = next(
            fragment for fragment in reversed(left_page.content_fragments) if fragment.block_type == BODY_TYPE
        )
        right_page = by_page.get(next_page)
        right = (
            next(fragment for fragment in right_page.content_fragments if fragment.block_type == BODY_TYPE)
            if right_page and any(
                fragment.block_type == BODY_TYPE for fragment in right_page.content_fragments
            )
            else None
        )
        boundary_id = f"boundary_p{previous:04d}_p{next_page:04d}"
        chapter = bool(
            right_page
            and any(fragment.block_type == "chapter_title" for fragment in right_page.content_fragments)
        )
        missing_transcription = bool(
            "missing_visual_transcription" in left.uncertainty
            or (right is not None and "missing_visual_transcription" in right.uncertainty)
        )
        relation = "unresolved" if missing_transcription else _edge_relation(
            left.text, right.text if right else None, chapter
        )
        support: list[str] = []
        conflicts: list[str] = []
        method = "local_text_layer_and_deterministic_rules"
        if chapter:
            support.append("The next page contains an explicit chapter_title block.")
        if missing_transcription:
            conflicts.append("A required adjacent page has no completed visual transcription.")
        if _ends_sentence(left.text):
            support.append("The previous visible fragment ends with sentence-final punctuation.")
        else:
            support.append("The previous visible fragment does not end with sentence-final punctuation.")
        if right and _starts_lower(right.text):
            support.append("The next visible fragment begins with a lowercase word.")
        old = None if missing_transcription else legacy.get(boundary_id)
        validated_pair = None
        if old is not None:
            source_inputs = [left.fragment_id, right.fragment_id if right else "missing_page", old.boundary_id]
            if relation == "continue":
                if old.sentence_continuation is True and old.paragraph_continuation is True:
                    support.append("Saved VLM pair/triple evidence also reports sentence and paragraph continuation.")
                else:
                    conflicts.append("Saved VLM boundary evidence disagrees with local sentence/paragraph continuation.")
            elif relation == "no_join" and (old.sentence_continuation is True or old.paragraph_continuation is True):
                conflicts.append("Saved VLM boundary evidence proposes continuation despite a local structural stop.")
        else:
            source_inputs = [left.fragment_id, right.fragment_id if right else "missing_page"]
        if old is not None:
            from .main_text_edition import validated_pair_resolution

            left_tokens = re.findall(r"[A-Za-z]+", left.text)
            right_tokens = re.findall(r"[A-Za-z]+", right.text if right else "")

            validated_pair = validated_pair_resolution(
                model_status=old.model_review_status,
                structural_break=old.structural_break,
                join_operation=old.join_operation,
                hyphen_type=old.hyphen_type,
                word_continuation=old.word_continuation,
                sentence_continuation=old.sentence_continuation,
                paragraph_continuation=old.paragraph_continuation,
                visible_trailing_hyphen=left.visible_trailing_hyphen,
                left_token=left_tokens[-1] if left_tokens else "",
                right_token=right_tokens[0] if right_tokens else "",
            )
        if relation == "no_join":
            structural = "chapter_break" if chapter else "paragraph_break"
            join = "no_join"
            word = sentence = paragraph = False
            status = "resolved_primary"
            reason = "Visible punctuation and structure safely prohibit a cross-page join."
        elif relation == "continue" and not left.visible_trailing_hyphen:
            structural = "none"
            join = "insert_space"
            word = False
            sentence = paragraph = True
            status = "resolved_primary"
            reason = "Both visible edge tokens are complete alphabetic words; continuation requires one space, never zero spaces."
            if old is not None:
                if old.word_continuation is True or old.join_operation == "concatenate_without_space":
                    conflicts.append("Saved VLM word/join fields would remove a required visible word boundary.")
        elif relation == "continue" and left.visible_trailing_hyphen:
            structural = "none"
            sentence = paragraph = True
            if validated_pair is not None:
                structural = str(validated_pair["structural_break"])
                join = str(validated_pair["join_operation"])
                word = bool(validated_pair["word_continuation"])
                sentence = bool(validated_pair["sentence_continuation"])
                paragraph = bool(validated_pair["paragraph_continuation"])
                status = str(validated_pair["auto_resolution_status"])
                support.append("Python validated the visible cross-page hyphen using pair evidence.")
                reason = "The visible layout hyphen is removed deterministically while preserving the model response as evidence."
            elif old is not None and old.hyphen_type == "line_break_hyphen":
                join = "remove_layout_hyphen"
                word = True
                status = "resolved_pair" if len(old.review_window) == 2 else "resolved_triple"
                support.append("The visible hyphen and saved adjacent-page evidence agree that it is a layout break.")
                reason = "A visibly present layout hyphen is removed after adjacent-page corroboration."
            elif old is not None and old.hyphen_type == "lexical_hyphen":
                join = "preserve_lexical_hyphen"
                word = True
                status = "resolved_pair" if len(old.review_window) == 2 else "resolved_triple"
                support.append("The visible hyphen and saved adjacent-page evidence agree that it is lexical.")
                reason = "A visibly present lexical hyphen is preserved after adjacent-page corroboration."
            else:
                join = "unresolved"
                word = None
                status = "unresolved"
                conflicts.append("A visible trailing hyphen cannot be classified safely from the available automatic evidence.")
                reason = "Fail closed until a text boundary adjudicator can classify the visible hyphen."
        elif relation == "unresolved" and validated_pair is not None:
            structural = str(validated_pair["structural_break"])
            join = str(validated_pair["join_operation"])
            word = bool(validated_pair["word_continuation"])
            sentence = bool(validated_pair["sentence_continuation"])
            paragraph = bool(validated_pair["paragraph_continuation"])
            status = str(validated_pair["auto_resolution_status"])
            support.append("Saved adjacent-page evidence was validated by deterministic token and structure rules.")
            reason = "Python validated the pair observation without allowing model word-continuation fields to control spacing."
        elif (
            relation == "unresolved"
            and old is not None
            and old.model_review_status == "completed"
            and old.structural_break == "none"
            and old.sentence_continuation is True
            and old.paragraph_continuation is True
            and old.word_continuation is False
            and old.join_operation == "concatenate_with_space"
            and not left.visible_trailing_hyphen
        ):
            structural = "none"
            join = "insert_space"
            word = False
            sentence = paragraph = True
            status = "resolved_pair" if len(old.review_window) == 2 else "resolved_triple"
            support.append("Adjacent-page visual evidence resolves an uppercase continuation with two complete words.")
            reason = "The primary rule was uncertain; validated adjacent-page evidence requires one visible word boundary."
        elif (
            relation == "unresolved"
            and old is not None
            and old.model_review_status == "completed"
            and old.structural_break != "none"
            and old.join_operation == "no_join"
            and old.sentence_continuation is False
            and old.paragraph_continuation is False
        ):
            structural = old.structural_break
            join = "no_join"
            word = sentence = paragraph = False
            status = "resolved_pair" if len(old.review_window) == 2 else "resolved_triple"
            support.append("Adjacent-page visual evidence identifies an explicit structural break.")
            reason = "The primary rule was uncertain; adjacent-page structure safely prohibits joining."
        else:
            structural = "unknown"
            join = "unresolved"
            word = sentence = paragraph = None
            status = "unresolved"
            reason = "The visible edge lacks enough deterministic evidence for a safe join or break."
        results.append(
            AutomatedBoundary(
                schema_version=settings.automated_boundary_schema_version,
                boundary_id=boundary_id,
                document_id=left_page.document_id,
                previous_page=previous,
                next_page=next_page,
                next_page_available=right is not None,
                previous_fragment_id=left.fragment_id,
                next_fragment_id=right.fragment_id if right else None,
                previous_tail_text=left.text,
                next_head_text=right.text if right else "",
                word_continuation=word,
                sentence_continuation=sentence,
                paragraph_continuation=paragraph,
                structural_break=structural,
                join_operation=join,
                visible_trailing_hyphen=left.visible_trailing_hyphen,
                resolution_method=method,
                supporting_evidence=support,
                conflicting_evidence=conflicts,
                resolution_reason=reason,
                auto_resolution_status=status,
                source_inputs=source_inputs,
                text_adjudicator_called=False,
                translation_called=False,
                created_at=_now(),
            )
        )
    output = resolve_project_path(settings.automated_boundary_directory, root=root)
    atomic_write_jsonl(output / f"{artifact_stem}.boundaries.jsonl", results)
    for boundary in results:
        atomic_write_json(output / "boundaries" / f"{boundary.boundary_id}.json", boundary)
    return results


class _UnionFind:
    def __init__(self, values: Iterable[str]) -> None:
        self.parent = {value: value for value in values}

    def find(self, value: str) -> str:
        while self.parent[value] != value:
            self.parent[value] = self.parent[self.parent[value]]
            value = self.parent[value]
        return value

    def union(self, left: str, right: str) -> None:
        first, second = self.find(left), self.find(right)
        if first != second:
            self.parent[second] = first


def _apply_join(left: str, right: str, operation: str) -> str:
    if operation == "insert_space":
        return left.rstrip() + " " + right.lstrip()
    if operation == "remove_layout_hyphen":
        value = left.rstrip()
        if not value.endswith("-"):
            raise ValueError("layout hyphen removal requires a visible trailing hyphen")
        return value[:-1] + right.lstrip()
    if operation == "preserve_lexical_hyphen":
        return left.rstrip() + right.lstrip()
    raise ValueError(f"Unsupported join operation: {operation}")


def build_automated_logical_blocks(
    pages: list[AutomatedPageRecord],
    boundaries: list[AutomatedBoundary],
    settings: ProjectSettings,
    *,
    root: Path,
    artifact_stem: str = "sample_11_pages",
) -> tuple[list[AutomatedLogicalBlock], list[TranslationContextV2]]:
    fragments = [fragment for page in pages for fragment in page.content_fragments]
    order = {fragment.fragment_id: index for index, fragment in enumerate(fragments)}
    by_id = {fragment.fragment_id: fragment for fragment in fragments}
    union = _UnionFind(by_id)
    join_operations: dict[tuple[str, str], str] = {}
    by_page_body: dict[int, list[SourceFragment]] = {
        page.pdf_page: [fragment for fragment in page.content_fragments if fragment.block_type == BODY_TYPE]
        for page in pages
    }
    for page, body in by_page_body.items():
        for left, right in zip(body, body[1:]):
            if _within_page_continuation(left.text, right.text):
                union.union(left.fragment_id, right.fragment_id)
                join_operations[(left.fragment_id, right.fragment_id)] = "insert_space"
    boundary_by_id = {item.boundary_id: item for item in boundaries}
    incoming_by_fragment = {
        item.next_fragment_id: item for item in boundaries if item.next_fragment_id
    }
    outgoing_by_fragment = {item.previous_fragment_id: item for item in boundaries}
    for boundary in boundaries:
        if (
            boundary.auto_resolution_status != "unresolved"
            and boundary.paragraph_continuation is True
            and boundary.next_fragment_id
            and boundary.join_operation in {"insert_space", "remove_layout_hyphen", "preserve_lexical_hyphen"}
        ):
            union.union(boundary.previous_fragment_id, boundary.next_fragment_id)
            join_operations[(boundary.previous_fragment_id, boundary.next_fragment_id)] = boundary.join_operation
    groups: dict[str, list[str]] = defaultdict(list)
    for fragment in fragments:
        groups[union.find(fragment.fragment_id)].append(fragment.fragment_id)
    ordered_groups = sorted(groups.values(), key=lambda values: min(order[value] for value in values))
    logical: list[AutomatedLogicalBlock] = []
    chapter_id: str | None = None
    section_id: str | None = None
    for group in ordered_groups:
        group.sort(key=lambda value: order[value])
        members = [by_id[value] for value in group]
        text = members[0].text
        for left, right in zip(members, members[1:]):
            operation = join_operations[(left.fragment_id, right.fragment_id)]
            text = _apply_join(text, right.text, operation)
        unresolved: list[str] = []
        first, last = members[0], members[-1]
        from .main_text_edition import boundary_leaves_fragment_unresolved

        group_set = set(group)
        if first.block_type == BODY_TYPE and first.starts_mid_paragraph:
            incoming = incoming_by_fragment.get(first.fragment_id)
            if boundary_leaves_fragment_unresolved(incoming, group_set, incoming=True):
                unresolved.append(incoming.boundary_id if incoming else f"boundary_p{first.source_page - 1:04d}_p{first.source_page:04d}")
        if last.block_type == BODY_TYPE and last.ends_mid_paragraph:
            outgoing = outgoing_by_fragment.get(last.fragment_id)
            if boundary_leaves_fragment_unresolved(outgoing, group_set, incoming=False):
                unresolved.append(outgoing.boundary_id if outgoing else f"boundary_p{last.source_page:04d}_p{last.source_page + 1:04d}")
        source_pages = sorted({item.source_page for item in members})
        source_block_ids = [block_id for member in members for block_id in member.source_block_ids]
        coverage = bool(
            all(member.text.strip() and member.source_block_ids for member in members)
            and len(group) == len(set(group))
        )
        complete = not unresolved
        logical_id = "logical2_" + stable_hash(
            {"fragments": group, "text": text, "schema": settings.automated_logical_schema_version}
        )[:20]
        if first.block_type == "chapter_title":
            chapter_id = logical_id
            section_id = None
        elif first.block_type == "section_title":
            section_id = logical_id
        ready = bool(complete and coverage and text.strip())
        logical.append(
            AutomatedLogicalBlock(
                schema_version=settings.automated_logical_schema_version,
                logical_block_id=logical_id,
                document_id=pages[0].document_id,
                source_pages=source_pages,
                source_fragment_ids=group,
                source_block_ids=source_block_ids,
                source_text=text,
                block_type=first.block_type,
                chapter_id=chapter_id,
                section_id=section_id,
                cross_page=len(source_pages) > 1,
                sentence_complete=complete,
                paragraph_complete=complete,
                coverage_complete=coverage,
                unresolved_boundaries=sorted(set(unresolved)),
                header_footer_page_number_clean=all(
                    member.block_type not in {"header", "footer", "page_number"} for member in members
                ),
                translation_ready=ready,
                created_at=_now(),
            )
        )
    contexts: list[TranslationContextV2] = []
    body = [item for item in logical if item.block_type == BODY_TYPE]
    for index, item in enumerate(body):
        before = body[index - 1] if index else None
        after = body[index + 1] if index + 1 < len(body) else None
        same_before = before if before and before.chapter_id == item.chapter_id and before.paragraph_complete else None
        same_after = after if after and after.chapter_id == item.chapter_id and after.paragraph_complete else None
        contexts.append(
            TranslationContextV2(
                target_logical_block_id=item.logical_block_id,
                source_text=item.source_text,
                context_before_text=same_before.source_text if same_before else "",
                context_after_text=same_after.source_text if same_after else "",
                chapter_context=(
                    next((candidate.source_text for candidate in logical if candidate.logical_block_id == item.chapter_id), None)
                    if item.chapter_id
                    else None
                ),
                translate_target_only=True,
                context_complete=bool(same_before or same_after) or item.translation_ready,
                translation_ready=item.translation_ready,
            )
        )
    logical_root = resolve_project_path(settings.automated_logical_directory, root=root)
    context_root = resolve_project_path(settings.automated_context_directory, root=root)
    atomic_write_jsonl(logical_root / f"{artifact_stem}.logical_blocks.jsonl", logical)
    atomic_write_jsonl(context_root / f"{artifact_stem}.translation_context.jsonl", contexts)
    return logical, contexts


def build_carry_records(
    pages: list[AutomatedPageRecord], boundaries: list[AutomatedBoundary], settings: ProjectSettings
) -> list[CarryRecord]:
    by_boundary = {item.previous_page: item for item in boundaries}
    records: list[CarryRecord] = []
    start = 1
    total = len(pages)
    while start <= total:
        end = min(total, start + settings.automated_batch_size - 1)
        if start > 1:
            overlap_page = start
            fragment = next(
                item for item in pages[overlap_page - 1].content_fragments if item.block_type == BODY_TYPE
            )
            records.append(
                CarryRecord(
                    direction="carry_in",
                    batch_start=start,
                    batch_end=end,
                    fragment_id=fragment.fragment_id,
                    source_page=fragment.source_page,
                    text=fragment.text,
                    completeness_status="incomplete_start" if fragment.starts_mid_paragraph else "complete",
                    expected_next_page=None,
                )
            )
        boundary = by_boundary.get(end)
        if boundary and (boundary.paragraph_continuation is True or boundary.auto_resolution_status == "unresolved"):
            fragment = next(
                item for item in reversed(pages[end - 1].content_fragments) if item.block_type == BODY_TYPE
            )
            records.append(
                CarryRecord(
                    direction="carry_out",
                    batch_start=start,
                    batch_end=end,
                    fragment_id=fragment.fragment_id,
                    source_page=end,
                    text=fragment.text,
                    completeness_status=(
                        "unresolved" if boundary.auto_resolution_status == "unresolved" else "incomplete_end"
                    ),
                    expected_next_page=end + 1,
                )
            )
        if end == total:
            break
        start = end - settings.automated_batch_overlap + 1
    return records


def run_audits(
    pages: list[AutomatedPageRecord],
    boundaries: list[AutomatedBoundary],
    logical: list[AutomatedLogicalBlock],
    settings: ProjectSettings,
    *,
    root: Path,
    expected_page_count: int = 11,
) -> tuple[SourceCoverageAudit, LogicalReconstructionAudit, TranslationAlignmentAudit]:
    page_numbers = [item.pdf_page for item in pages]
    duplicates = sorted(number for number, count in Counter(page_numbers).items() if count > 1)
    expected_pages = set(range(1, expected_page_count + 1))
    actual_pages = set(page_numbers)
    missing = sorted(expected_pages - actual_pages)
    discontinuous = sorted(
        page for page in page_numbers if page > 1 and page - 1 not in actual_pages
    )
    source_audit = SourceCoverageAudit(
        total_pages=expected_page_count,
        processed_pages=len(page_numbers),
        missing_pages=missing,
        duplicate_pages=duplicates,
        discontinuous_pages=discontinuous,
        page_hashes_valid=all(bool(page.image_sha256) for page in pages),
        transcription_missing_pages=[
            page.pdf_page for page in pages
            if not page.content_fragments or page.transcription_status.startswith("failed")
        ],
        partial_coverage_pages=[page.pdf_page for page in pages if page.source_coverage_status != "complete"],
        all_visible_source_recorded=all(page.source_coverage_status == "complete" for page in pages),
        passed=False,
    )
    source_audit.passed = not any(
        [
            source_audit.missing_pages,
            source_audit.duplicate_pages,
            source_audit.discontinuous_pages,
            source_audit.transcription_missing_pages,
            source_audit.partial_coverage_pages,
        ]
    ) and source_audit.page_hashes_valid
    expected_fragments = [fragment.fragment_id for page in pages for fragment in page.content_fragments]
    referenced = [fragment_id for block in logical for fragment_id in block.source_fragment_ids]
    counts = Counter(referenced)
    unreferenced = sorted(set(expected_fragments) - set(referenced))
    duplicate_fragments = sorted(fragment_id for fragment_id, count in counts.items() if count != 1)
    unsupported_no_space = [
        item.boundary_id
        for item in boundaries
        if item.join_operation in {"remove_layout_hyphen", "preserve_lexical_hyphen"}
        and not item.visible_trailing_hyphen
    ]
    chapter_violations = [
        item.boundary_id
        for item in boundaries
        if item.structural_break == "chapter_break" and item.join_operation != "no_join"
    ]
    unresolved = [
        item.boundary_id for item in boundaries if item.auto_resolution_status == "unresolved"
    ]
    internal_unresolved = [
        item.boundary_id
        for item in boundaries
        if item.auto_resolution_status == "unresolved" and item.next_page_available
    ]
    external_open = [
        item.boundary_id
        for item in boundaries
        if item.auto_resolution_status == "unresolved" and not item.next_page_available
    ]
    untraceable = [
        item.logical_block_id
        for item in logical
        if not item.source_fragment_ids or not item.source_block_ids or not item.source_pages
    ]
    logical_audit = LogicalReconstructionAudit(
        expected_fragment_count=len(expected_fragments),
        referenced_fragment_count=len(referenced),
        unreferenced_fragment_ids=unreferenced,
        duplicate_fragment_ids=duplicate_fragments,
        unresolved_boundary_ids=unresolved,
        internal_unresolved_boundary_ids=internal_unresolved,
        external_open_boundary_ids=external_open,
        internal_boundaries_passed=not internal_unresolved,
        unsupported_no_space_boundaries=unsupported_no_space,
        untraceable_logical_block_ids=untraceable,
        chapter_break_join_violations=chapter_violations,
        passed=not any([unreferenced, duplicate_fragments, unresolved, unsupported_no_space, untraceable, chapter_violations]),
    )
    ready_ids = [item.logical_block_id for item in logical if item.translation_ready]
    translation_audit = TranslationAlignmentAudit(
        status="not_run",
        ready_source_blocks=len(ready_ids),
        translated_blocks=0,
        missing_translation_ids=ready_ids,
        duplicate_translation_ids=[],
        context_leak_ids=[],
        incomplete_translated_ids=[],
        passed=False,
    )
    audit_root = resolve_project_path(settings.automated_audit_directory, root=root)
    atomic_write_json(audit_root / "source_coverage_audit.json", source_audit)
    atomic_write_json(audit_root / "logical_reconstruction_audit.json", logical_audit)
    atomic_write_json(audit_root / "translation_alignment_audit.json", translation_audit)
    return source_audit, logical_audit, translation_audit


def build_master_document(
    pdf_path: Path,
    pages: list[AutomatedPageRecord],
    logical: list[AutomatedLogicalBlock],
    carry: list[CarryRecord],
    audits: tuple[SourceCoverageAudit, LogicalReconstructionAudit, TranslationAlignmentAudit],
    settings: ProjectSettings,
    *,
    root: Path,
) -> BilingualDocument:
    source_audit, logical_audit, translation_audit = audits
    blockers: list[str] = []
    if not source_audit.passed:
        blockers.append("source_coverage_audit_failed")
    if not logical_audit.passed:
        blockers.append("logical_reconstruction_audit_failed")
    if not translation_audit.passed:
        blockers.append("translation_alignment_audit_not_passed")
    if any(not item.sentence_complete or not item.paragraph_complete for item in logical):
        blockers.append("incomplete_source_blocks_exist")
    document = BilingualDocument(
        schema_version="1.0",
        document_id=pages[0].document_id,
        source_pdf=str(pdf_path.resolve()),
        source_pdf_sha256=sha256_file(pdf_path),
        entries=[
            BilingualEntry(
                logical_block_id=item.logical_block_id,
                source_pages=item.source_pages,
                source_text=item.source_text,
                chinese_text=None,
                translation_status="not_translated",
                translation_ready=item.translation_ready,
                unresolved_boundaries=item.unresolved_boundaries,
            )
            for item in logical
        ],
        source_coverage_audit=source_audit,
        logical_reconstruction_audit=logical_audit,
        translation_alignment_audit=translation_audit,
        carry_records=carry,
        strict_export_ready=not blockers,
        strict_blockers=blockers,
        api_calls=0,
        deepseek_calls=0,
        translation_calls=0,
        created_at=_now(),
    )
    atomic_write_json(resolve_project_path(settings.automated_master_path, root=root), document)
    return document


def _diagnostic_sections(document: BilingualDocument) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    for entry in document.entries:
        markers: list[str] = []
        if entry.unresolved_boundaries:
            markers.append("[UNRESOLVED_BOUNDARY] " + ", ".join(entry.unresolved_boundaries))
        if not entry.translation_ready:
            markers.append("[INCOMPLETE_SOURCE]")
        markers.append("[TRANSLATION_NOT_AVAILABLE]")
        pages = f"[原书页码范围: {entry.source_pages[0]}-{entry.source_pages[-1]}]"
        sections.append((pages, markers + [entry.source_text, "[中文译文尚未生成]"]))
    return sections


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float = 11,
    color: str = "000000",
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _configure_diagnostic_document(word: Document, mode: str) -> None:
    """Apply the compact_reference_guide preset with a diagnostic masthead override."""

    section = word.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = word.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, "2E74B5"),
        ("Heading 2", 13, 14, 7, "2E74B5"),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(header.add_run("Bilingual Book Workflow | Diagnostic Draft"), size=8.5, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(footer.add_run("MACHINE DIAGNOSTIC - NOT FINAL"), size=8.5, color="9B1C1C", bold=True)
    title = word.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_run_font(
        title.add_run("英汉双语文档" if mode == "strict" else "英汉双语诊断草稿"),
        size=23,
        color="0B2545",
        bold=True,
    )
    subtitle = word.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _set_run_font(
        subtitle.add_run("正式输出" if mode == "strict" else "非final | 自动重建与失败关闭检查"),
        size=11,
        color="6B7280",
        italic=mode != "strict",
    )


def export_from_master(
    master_path: str | Path,
    output_directory: str | Path,
    *,
    mode: str = "strict",
) -> ExportResult:
    document = BilingualDocument.model_validate(load_json(master_path))
    output = Path(output_directory).resolve()
    if mode not in {"strict", "permissive"}:
        raise ValueError("mode must be strict or permissive")
    if mode == "strict" and not document.strict_export_ready:
        return ExportResult(
            mode=mode,
            markdown_path=None,
            word_path=None,
            blocked=True,
            blockers=document.strict_blockers,
        )
    prefix = "bilingual" if mode == "strict" else "bilingual_diagnostic"
    markdown_path = output / f"{prefix}.md"
    word_path = output / f"{prefix}.docx"
    sections = _diagnostic_sections(document) if mode == "permissive" else [
        (
            f"[原书页码范围: {entry.source_pages[0]}-{entry.source_pages[-1]}]",
            [entry.source_text, entry.chinese_text or ""],
        )
        for entry in document.entries
    ]
    markdown_lines = ["# 英汉双语文档" if mode == "strict" else "# 英汉双语诊断草稿（非final）", ""]
    word = Document()
    _configure_diagnostic_document(word, mode)
    for heading, paragraphs in sections:
        markdown_lines.extend([heading, ""])
        heading_paragraph = word.add_paragraph()
        heading_paragraph.paragraph_format.space_before = Pt(10)
        heading_paragraph.paragraph_format.space_after = Pt(4)
        heading_paragraph.paragraph_format.keep_with_next = True
        _set_run_font(heading_paragraph.add_run(heading), size=9.5, color="2E74B5", bold=True)
        for paragraph in paragraphs:
            markdown_lines.extend([paragraph, ""])
            item = word.add_paragraph()
            if paragraph.startswith("[UNRESOLVED_BOUNDARY]") or paragraph == "[INCOMPLETE_SOURCE]":
                _set_run_font(item.add_run(paragraph), size=10, color="9B1C1C", bold=True)
            elif paragraph.startswith("[TRANSLATION_NOT_AVAILABLE]") or paragraph == "[中文译文尚未生成]":
                _set_run_font(item.add_run(paragraph), size=10, color="6B7280", italic=True)
            else:
                _set_run_font(item.add_run(paragraph), size=11, color="000000")
    markdown_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_path.write_text("\n".join(markdown_lines).rstrip() + "\n", encoding="utf-8")
    word_path.parent.mkdir(parents=True, exist_ok=True)
    word.save(word_path)
    return ExportResult(
        mode=mode,
        markdown_path=str(markdown_path),
        word_path=str(word_path),
        blocked=False,
        blockers=[],
    )


def run_automated_reconstruction(
    pdf_path: str | Path,
    settings: ProjectSettings,
    *,
    root: Path | None = None,
    create_diagnostic: bool = True,
) -> AutomatedRunResult:
    root = (root or project_root()).resolve()
    source = _sample_only(pdf_path, settings, root)
    pages = build_automated_pages(source, settings, root=root)
    boundaries = build_automated_boundaries(pages, settings, root=root)
    logical, contexts = build_automated_logical_blocks(pages, boundaries, settings, root=root)
    carry = build_carry_records(pages, boundaries, settings)
    audits = run_audits(pages, boundaries, logical, settings, root=root)
    master = build_master_document(source, pages, logical, carry, audits, settings, root=root)
    diagnostic = (
        export_from_master(
            resolve_project_path(settings.automated_master_path, root=root),
            resolve_project_path(settings.automated_export_directory, root=root),
            mode="permissive",
        )
        if create_diagnostic
        else None
    )
    return AutomatedRunResult(
        page_records_path=str(
            (resolve_project_path(settings.automated_page_directory, root=root) / "sample_11_pages.pages.jsonl").resolve()
        ),
        boundary_records_path=str(
            (resolve_project_path(settings.automated_boundary_directory, root=root) / "sample_11_pages.boundaries.jsonl").resolve()
        ),
        logical_blocks_path=str(
            (resolve_project_path(settings.automated_logical_directory, root=root) / "sample_11_pages.logical_blocks.jsonl").resolve()
        ),
        translation_context_path=str(
            (resolve_project_path(settings.automated_context_directory, root=root) / "sample_11_pages.translation_context.jsonl").resolve()
        ),
        source_audit_path=str(
            (resolve_project_path(settings.automated_audit_directory, root=root) / "source_coverage_audit.json").resolve()
        ),
        logical_audit_path=str(
            (resolve_project_path(settings.automated_audit_directory, root=root) / "logical_reconstruction_audit.json").resolve()
        ),
        translation_audit_path=str(
            (resolve_project_path(settings.automated_audit_directory, root=root) / "translation_alignment_audit.json").resolve()
        ),
        master_document_path=str(resolve_project_path(settings.automated_master_path, root=root)),
        diagnostic_markdown_path=diagnostic.markdown_path if diagnostic else None,
        diagnostic_word_path=diagnostic.word_path if diagnostic else None,
        pages=len(pages),
        boundaries=len(boundaries),
        resolved_boundaries=sum(item.auto_resolution_status != "unresolved" for item in boundaries),
        unresolved_boundaries=sum(item.auto_resolution_status == "unresolved" for item in boundaries),
        logical_blocks=len(logical),
        cross_page_blocks=sum(item.cross_page for item in logical),
        translation_ready_true=sum(item.translation_ready for item in logical),
        translation_ready_false=sum(not item.translation_ready for item in logical),
        fragments=sum(len(page.content_fragments) for page in pages),
        source_audit_passed=audits[0].passed,
        logical_audit_passed=audits[1].passed,
        strict_export_ready=master.strict_export_ready,
        api_calls=0,
        deepseek_calls=0,
        translation_calls=0,
    )
