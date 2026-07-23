"""Language-neutral, install-safe workspace lifecycle for new books."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
import unicodedata
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .io_utils import atomic_write_json, atomic_write_jsonl, atomic_write_text, sha256_file
from .publication_notes import (
    NOTE_PLACEHOLDER_RE,
    annotate_note_references,
    classify_note_blocks,
    find_note_reference_labels,
    note_id,
    normalize_note_label,
    split_note_entries,
)

SUPPORTED_LANGUAGES = ("zh-Hans", "en", "fr", "de", "ja", "es")
OUTPUT_ROLES = ("source", "target", "bilingual")
OUTPUT_ROLE_CHOICES = (*OUTPUT_ROLES, "all")
LAYOUT_MODES = ("text", "structure", "publication")
BILINGUAL_LAYOUTS = ("stacked", "parallel-columns")
WORKSPACE_MANIFEST = "bookflow_workspace.json"


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def validate_language(value: str, *, allow_auto: bool = False) -> str:
    allowed = set(SUPPORTED_LANGUAGES) | ({"auto"} if allow_auto else set())
    if value not in allowed:
        raise ValueError(f"unsupported language {value!r}; expected one of {sorted(allowed)}")
    return value


def detect_language(text: str) -> str:
    sample = text[:50000]
    if any("\u3040" <= c <= "\u30ff" for c in sample):
        return "ja"
    cjk = sum("\u3400" <= c <= "\u9fff" for c in sample)
    if cjk > max(20, len(sample) // 20):
        return "zh-Hans"
    lowered = f" {sample.lower()} "
    scores = {
        "fr": sum(lowered.count(x) for x in (" le ", " la ", " les ", " des ", " une ", " et ", " qu'", "é")),
        "de": sum(lowered.count(x) for x in (" der ", " die ", " das ", " und ", " nicht ", " ist ", "ß")),
        "es": sum(lowered.count(x) for x in (" el ", " los ", " las ", " una ", " y ", " que ", "ñ", "¿")),
        "en": sum(lowered.count(x) for x in (" the ", " and ", " of ", " to ", " is ", " in ")),
    }
    return max(scores, key=scores.get) if max(scores.values()) else "en"


def _load(workspace: Path) -> dict[str, Any]:
    path = workspace.resolve() / WORKSPACE_MANIFEST
    if not path.is_file():
        raise FileNotFoundError(f"workspace manifest not found: {path}")
    value = json.loads(path.read_text("utf-8"))
    validate_language(value["source_language"])
    validate_language(value["target_language"])
    return value


def _save(workspace: Path, manifest: dict[str, Any]) -> None:
    manifest["updated_at"] = _now()
    atomic_write_json(workspace / WORKSPACE_MANIFEST, manifest)


def create_workspace(
    workspace: Path, source_pdf: Path, source_language: str, target_language: str,
    *, output_directory: Path | None = None, profile: str = "reading", copy_source: bool = True,
    layout_mode: str = "text", bilingual_layout: str = "stacked",
    output_role: str = "all",
    metadata: dict[str, str] | None = None,
    workspace_id: str | None = None,
) -> dict[str, Any]:
    workspace = workspace.resolve()
    source_pdf = source_pdf.resolve()
    if workspace.exists() and any(workspace.iterdir()):
        raise FileExistsError(f"workspace is not empty: {workspace}")
    validate_language(source_language, allow_auto=True)
    validate_language(target_language)
    if layout_mode not in LAYOUT_MODES:
        raise ValueError(f"invalid layout mode {layout_mode!r}")
    if bilingual_layout not in BILINGUAL_LAYOUTS:
        raise ValueError(f"invalid bilingual layout {bilingual_layout!r}")
    if output_role not in OUTPUT_ROLE_CHOICES:
        raise ValueError(f"invalid output role {output_role!r}")
    if not source_pdf.is_file():
        raise FileNotFoundError(source_pdf)
    workspace.mkdir(parents=True, exist_ok=True)
    for name in ("input", "data", "cache", "checkpoints", "logs", "output", "manual_review"):
        (workspace / name).mkdir()
    stored = workspace / "input" / f"source{source_pdf.suffix.lower()}"
    if copy_source:
        shutil.copy2(source_pdf, stored)
    else:
        stored = source_pdf
    pdf_sha = sha256_file(stored)
    manifest = {
        "schema_version": "bookflow-workspace-1.0", "workspace_id": workspace_id or f"book-{pdf_sha[:16]}",
        "source_pdf": str(stored), "source_pdf_sha256": pdf_sha,
        "source_filename": source_pdf.name,
        "source_language_requested": source_language, "source_language": source_language if source_language != "auto" else "en",
        "source_language_detection": "pending" if source_language == "auto" else "explicit",
        "target_language": target_language, "language_pair": None, "profile": profile,
        "layout_mode": layout_mode, "bilingual_layout": bilingual_layout,
        "output_role": output_role,
        "output_directory": str((output_directory or workspace / "output").resolve()),
        "created_at": _now(), "updated_at": _now(), "stage": "created", "provider_calls": 0,
    }
    if metadata:
        allowed_metadata = {"book_title", "author", "volume", "year", "source_institution", "rights_notice"}
        manifest.update({key: str(value).strip() for key, value in metadata.items()
                         if key in allowed_metadata and str(value).strip()})
    if source_language != "auto":
        manifest["language_pair"] = f"{source_language}-{target_language}"
    _save(workspace, manifest)
    return manifest


def update_workspace_metadata(workspace: Path, metadata: dict[str, str]) -> dict[str, Any]:
    allowed = {"book_title", "author", "volume", "year", "source_institution", "rights_notice"}
    unknown = sorted(set(metadata) - allowed)
    if unknown:
        raise ValueError(f"unsupported metadata fields: {unknown}")
    workspace = workspace.resolve(); manifest = _load(workspace)
    manifest.update({key: str(value).strip() for key, value in metadata.items() if str(value).strip()})
    _save(workspace, manifest)
    return {key: manifest.get(key) for key in sorted(allowed)}


def _chunks(text: str, limit: int = 3500, *, preserve_structure: bool = False) -> list[str]:
    if preserve_structure:
        lines = text.replace("\r", "").strip().split("\n")
        result: list[str] = []
        current: list[str] = []
        current_size = 0
        for line in lines:
            if current and current_size + len(line) + 1 > limit:
                result.append("\n".join(current)); current = []; current_size = 0
            current.append(line); current_size += len(line) + 1
        if current:
            result.append("\n".join(current))
        return result
    paragraphs = [x.strip() for x in text.replace("\r", "").split("\n") if x.strip()]
    result: list[str] = []
    current = ""
    for paragraph in paragraphs:
        pieces = [paragraph[i:i + limit] for i in range(0, len(paragraph), limit)] or [""]
        for piece in pieces:
            if current and len(current) + len(piece) + 1 > limit:
                result.append(current); current = piece
            else:
                current = f"{current}\n{piece}".strip()
    if current:
        result.append(current)
    return result


def _page_marker_token(physical_page: int, label: str) -> str:
    return f"[[BOOKFLOW_PAGE_BREAK:{physical_page}:{label}]]"


def _continues_across_page(left: str, right: str) -> bool:
    """Conservatively identify a prose continuation at a physical page edge."""
    left = left.rstrip()
    right = right.lstrip()
    if not left or not right:
        return False
    if left.endswith(("-", "\u2010", "\u2011", "\u2012", "\u2013")):
        return True
    if left[-1] in ".!?。！？；;：:":
        return False
    if right[0].islower() or right[0].isdigit():
        return True
    return left[-1] in ",，、(" or left[-1].isalnum()


def _attach_page_markers_and_translation_groups(
    units: list[dict[str, Any]], layout_elements: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Keep layout objects independent while translating page-edge prose together."""
    by_id = {str(unit["translation_unit_id"]): unit for unit in units}
    body_by_page: dict[int, list[dict[str, Any]]] = {}
    marker_labels: dict[int, tuple[str, str]] = {}
    for element in layout_elements:
        page_no = int(element["source_page"])
        if element["element_type"] == "pagination":
            marker_labels.setdefault(
                page_no, _original_page_marker_label(element.get("source_text"), page_no),
            )
            continue
        if element["element_type"] != "body":
            continue
        unit = by_id.get(str(element.get("translation_unit_id") or ""))
        if unit is not None:
            body_by_page.setdefault(page_no, []).append(unit)
    for values in body_by_page.values():
        values.sort(key=lambda unit: tuple(unit.get("reading_order") or ()))

    pages = sorted({int(element["source_page"]) for element in layout_elements})
    for page_no in pages:
        label, marker_source = marker_labels.get(
            page_no, _original_page_marker_label(None, page_no),
        )
        marker = {"physical_page": page_no, "printed_page": label,
                  "marker_source": marker_source}
        if page_no == pages[0] and body_by_page.get(page_no):
            body_by_page[page_no][0].setdefault("page_markers_before", []).append(marker)
            continue
        previous_pages = [value for value in pages if value < page_no and body_by_page.get(value)]
        if previous_pages:
            body_by_page[previous_pages[-1]][-1].setdefault("page_markers_after", []).append(marker)
        elif body_by_page.get(page_no):
            body_by_page[page_no][0].setdefault("page_markers_before", []).append(marker)

    edges: list[tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]] = []
    for left_page, right_page in zip(pages, pages[1:]):
        if right_page != left_page + 1:
            continue
        left_values = body_by_page.get(left_page) or []
        right_values = body_by_page.get(right_page) or []
        if not left_values or not right_values:
            continue
        left_block = str(left_values[-1].get("source_block_id") or "")
        right_block = str(right_values[0].get("source_block_id") or "")
        left_units = [unit for unit in left_values if unit.get("source_block_id") == left_block]
        right_units = [unit for unit in right_values if unit.get("source_block_id") == right_block]
        left_text = "\n".join(str(unit["source_text"]) for unit in left_units)
        right_text = "\n".join(str(unit["source_text"]) for unit in right_units)
        if not _continues_across_page(left_text, right_text):
            continue
        label, marker_source = marker_labels.get(
            right_page, _original_page_marker_label(None, right_page),
        )
        edges.append((left_units, right_units, {
            "physical_page": right_page, "printed_page": label,
            "marker_source": marker_source,
        }))

    components: list[dict[str, Any]] = []
    for left_units, right_units, marker in edges:
        member_ids = {str(unit["translation_unit_id"]) for unit in left_units + right_units}
        touching = [item for item in components if member_ids & set(item["unit_ids"])]
        if touching:
            component = touching[0]
            for extra in touching[1:]:
                component["unit_ids"].extend(extra["unit_ids"])
                component["breaks"].extend(extra["breaks"])
                components.remove(extra)
            component["unit_ids"].extend(member_ids)
            component["breaks"].append(marker)
        else:
            components.append({"unit_ids": list(member_ids), "breaks": [marker]})

    order = {str(unit["translation_unit_id"]): index for index, unit in enumerate(units)}
    groups: list[dict[str, Any]] = []
    for component in components:
        member_ids = sorted(set(component["unit_ids"]), key=order.__getitem__)
        members = [by_id[uid] for uid in member_ids]
        breaks_by_page = {int(item["physical_page"]): item for item in component["breaks"]}
        source_parts: list[str] = []
        markers: list[dict[str, Any]] = []
        for index, unit in enumerate(members):
            if index:
                page_no = int(unit["source_page"])
                marker = breaks_by_page.get(page_no)
                if marker:
                    token = _page_marker_token(page_no, str(marker["printed_page"]))
                    source_parts.append(token)
                    markers.append({**marker, "token": token, "kind": "page",
                                    "after_unit_id": member_ids[index - 1]})
                else:
                    token = f"[[BOOKFLOW_UNIT_BREAK:{index}]]"
                    source_parts.append(token)
                    markers.append({"token": token, "kind": "unit",
                                    "after_unit_id": member_ids[index - 1]})
            source_parts.append(str(unit["source_text"]))
        source_text = "\n".join(source_parts)
        fingerprint = "|".join(member_ids + [item["token"] for item in markers])
        group_id = "tug_" + hashlib.sha256(fingerprint.encode("utf-8")).hexdigest()[:24]
        for unit in members:
            unit["translation_group_id"] = group_id
        groups.append({
            "translation_group_id": group_id,
            "unit_ids": member_ids,
            "source_text": source_text,
            "source_text_sha256": hashlib.sha256(source_text.encode("utf-8")).hexdigest(),
            "markers": markers,
            "source_pages": sorted({int(unit["source_page"]) for unit in members}),
        })
    return groups


def inspect_workspace(workspace: Path, *, page_text_overrides: dict[int, str] | None = None,
                      page_quality_records: list[dict[str, Any]] | None = None,
                      page_route_records: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    workspace = workspace.resolve(); manifest = _load(workspace)
    pdf_path = Path(manifest["source_pdf"])
    if sha256_file(pdf_path) != manifest["source_pdf_sha256"]:
        raise ValueError("source PDF SHA-256 mismatch")
    pdf = fitz.open(pdf_path)
    if pdf.needs_pass:
        raise ValueError("encrypted PDF requires a password")
    metadata_title = str((pdf.metadata or {}).get("title") or "").strip()
    if metadata_title: manifest["book_title"] = metadata_title
    pages: list[dict[str, Any]] = []; units: list[dict[str, Any]] = []; all_text: list[str] = []
    quality_by_page = {index + 1: item for index, item in enumerate(page_quality_records or [])}
    route_by_page = {index + 1: item for index, item in enumerate(page_route_records or [])}
    for page_index, page in enumerate(pdf):
        physical_page = page_index + 1
        text = (page_text_overrides or {}).get(physical_page, page.get_text("text"))
        all_text.append(text)
        images = len(page.get_images(full=True)); rotation = int(page.rotation)
        quality = quality_by_page.get(physical_page)
        route = route_by_page.get(physical_page)
        pages.append({"physical_page": physical_page, "width": page.rect.width, "height": page.rect.height,
                      "rotation": rotation, "text_characters": len(text), "image_count": images,
                      "textless": not bool(text.strip()),
                      "text_quality_score": quality.get("quality_score") if quality else None,
                      "text_quality_passed": quality.get("passed") if quality else None,
                      "extraction_route": route.get("route") if route else "pymupdf",
                      "review_required": bool(route and route.get("status") != "accepted"),
                      "high_risk": rotation != 0 or not text.strip() or images > 4 or bool(route and route.get("status") != "accepted")})
        for sequence, chunk in enumerate(_chunks(text), 1):
            stable = f"{manifest['source_pdf_sha256']}|{page_index + 1}|{sequence}|{chunk}"
            uid = "tu_" + hashlib.sha256(stable.encode("utf-8")).hexdigest()[:24]
            units.append({"translation_unit_id": uid, "source_object_id": f"page-{page_index + 1:04d}-text-{sequence:04d}",
                          "source_page": page_index + 1, "source_text": chunk,
                          "source_text_sha256": hashlib.sha256(chunk.encode("utf-8")).hexdigest(),
                          "source_language": manifest["source_language"], "target_language": manifest["target_language"],
                          "status": "pending"})
    pdf.close()
    detected = detect_language("\n".join(all_text))
    if manifest["source_language_requested"] == "auto":
        manifest["source_language"] = detected; manifest["source_language_detection"] = "automatic"
    else:
        manifest["detected_language"] = detected
    manifest["language_pair"] = f"{manifest['source_language']}-{manifest['target_language']}"
    manifest.update(stage="inspected", page_count=len(pages), unit_count=len(units))
    atomic_write_jsonl(workspace / "data/pages.jsonl", pages)
    atomic_write_jsonl(workspace / "data/translation_units.jsonl", units)
    report = {"source_pdf_sha256": manifest["source_pdf_sha256"], "pages": len(pages), "units": len(units),
              "source_language": manifest["source_language"], "detected_language": detected,
              "textless_pages": [p["physical_page"] for p in pages if p["textless"]],
              "high_risk_pages": [p["physical_page"] for p in pages if p["high_risk"]]}
    atomic_write_json(workspace / "data/inspection_report.json", report); _save(workspace, manifest)
    return report


def _normalized_rect(rect: fitz.Rect, page: fitz.Page) -> list[float]:
    return [float(rect.x0 / page.rect.width), float(rect.y0 / page.rect.height),
            float(rect.x1 / page.rect.width), float(rect.y1 / page.rect.height)]


def _rect_overlap_ratio(left: list[float], right: list[float]) -> float:
    x0, y0 = max(left[0], right[0]), max(left[1], right[1])
    x1, y1 = min(left[2], right[2]), min(left[3], right[3])
    intersection = max(0.0, x1 - x0) * max(0.0, y1 - y0)
    area = max((left[2] - left[0]) * (left[3] - left[1]), 1e-9)
    return intersection / area


def _geometric_text_blocks(page: fitz.Page) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    payload = page.get_text("dict")
    for raw_block in payload.get("blocks", []):
        if raw_block.get("type") != 0:
            continue
        lines: list[str] = []
        sizes: list[float] = []
        for line in raw_block.get("lines", []):
            spans = line.get("spans", [])
            value = "".join(str(span.get("text") or "") for span in spans).strip()
            if value:
                lines.append(value)
            sizes.extend(float(span.get("size") or 0) for span in spans if span.get("size"))
        text = "\n".join(lines).strip()
        if not text:
            continue
        rect = fitz.Rect(raw_block.get("bbox", (0, 0, 0, 0))) & page.rect
        if rect.is_empty:
            continue
        blocks.append({"text": text, "bbox": _normalized_rect(rect, page),
                       "font_size": sum(sizes) / len(sizes) if sizes else 0.0})
    return blocks


def _merge_wrapped_text_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Join adjacent same-style line blocks while retaining paragraph gaps."""
    merged: list[dict[str, Any]] = []
    for current in blocks:
        if not merged:
            merged.append(dict(current))
            continue
        previous = merged[-1]
        previous_box, current_box = previous["bbox"], current["bbox"]
        previous_height = max(0.001, float(previous_box[3]) - float(previous_box[1]))
        vertical_gap = float(current_box[1]) - float(previous_box[3])
        font_limit = max(1.0, max(float(previous.get("font_size", 0)), float(current.get("font_size", 0))) * 0.2)
        same_style = abs(float(previous.get("font_size", 0)) - float(current.get("font_size", 0))) <= font_limit
        previous_center = (float(previous_box[0]) + float(previous_box[2])) / 2.0
        current_center = (float(current_box[0]) + float(current_box[2])) / 2.0
        aligned = (
            abs(float(previous_box[0]) - float(current_box[0])) <= 0.025
            or abs(previous_center - current_center) <= 0.05
        )
        wrapped_line = (
            -0.002 <= vertical_gap <= max(0.012, previous_height * 0.8)
            and aligned
            and same_style
            and not previous.get("ocr_element_type")
            and not current.get("ocr_element_type")
        )
        if not wrapped_line:
            merged.append(dict(current))
            continue
        previous["text"] = f"{str(previous['text']).rstrip()} {str(current['text']).lstrip()}"
        previous["bbox"] = [
            min(float(previous_box[0]), float(current_box[0])),
            min(float(previous_box[1]), float(current_box[1])),
            max(float(previous_box[2]), float(current_box[2])),
            max(float(previous_box[3]), float(current_box[3])),
        ]
        previous["font_size"] = (
            float(previous.get("font_size", 0)) + float(current.get("font_size", 0))
        ) / 2.0
    return merged


def _markdown_cells(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _normalize_ocr_markdown_table(paragraph: str) -> str:
    """Make OCR Markdown tables rectangular, flattening a simple spanning header."""
    lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
    if len(lines) < 2 or not all(line.startswith("|") and line.endswith("|") for line in lines):
        return paragraph
    rows = [_markdown_cells(line) for line in lines]
    separator = lambda row: bool(row) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)
    content_rows = [row for row in rows if not separator(row)]
    width = max((len(row) for row in content_rows), default=0)
    if width < 2:
        return paragraph

    header = rows[0]
    remove_subheader = False
    if len(header) < width and len(rows) >= 4 and separator(rows[1]):
        subheader = rows[2]
        expansion = width - len(header) + 1
        candidates = [
            index for index in range(0, len(subheader) - expansion + 1)
            if all(subheader[offset].strip() for offset in range(index, index + expansion))
            and index < len(header) and header[index].strip()
        ]
        if candidates:
            index = candidates[0]
            parent = header[index]
            header = (
                header[:index]
                + [f"{parent} {subheader[offset]}".strip() for offset in range(index, index + expansion)]
                + header[index + 1:]
            )
            rows[0] = header
            rows.pop(2)
            remove_subheader = True

    normalized: list[str] = []
    for row_index, row in enumerate(rows):
        if separator(row):
            cells = ["---"] * width
        else:
            cells = (row + [""] * width)[:width]
        normalized.append("| " + " | ".join(cells) + " |")
    return "\n".join(normalized) if remove_subheader or any(len(row) != width for row in rows) else paragraph


def _ocr_fallback_blocks(text: str, classification: dict[str, Any]) -> list[dict[str, Any]]:
    """Anchor page-level OCR text without treating it as text inside every visual.

    OCR providers intentionally return reading-order text, not fabricated word
    coordinates.  Structural regions still let us keep headers, raster-table
    text, notes, prose, and pagination in separate translation units.
    """
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n", text) if item.strip()]
    semantic = [item for item in classification.get("semantic_regions", []) if isinstance(item, dict)]
    visuals = [item for item in classification.get("visual_regions", []) if isinstance(item, dict)]
    header_region = next((item for item in semantic if item.get("type") == "header"), None)
    footnote_regions = [item for item in semantic if item.get("type") == "footnote"]
    table_region = next((item for item in visuals if item.get("type") == "table"), None)
    if header_region is None and table_region and isinstance(table_region.get("caption_bbox"), list):
        header_region = {"type": "header", "bbox": table_region["caption_bbox"]}
    blocks: list[dict[str, Any]] = []
    body_index = 0
    for index, paragraph in enumerate(paragraphs):
        lines = paragraph.splitlines()
        is_table = (sum(line.count("|") >= 2 for line in lines) >= 2
                    or sum(line.count("\t") >= 2 for line in lines) >= 2)
        if is_table and sum(line.count("|") >= 2 for line in lines) >= 2:
            paragraph = _normalize_ocr_markdown_table(paragraph)
        is_note = bool(re.match(r"^(?:[*\u2020\u2021]|\(?\d{1,3}[.)])\s*", paragraph))
        is_pagination = bool(re.fullmatch(r"\s*(?:\d{1,4}|[ivxlcdm]{1,8})[.!]?\s*", paragraph, re.I))
        if index == 0 and header_region and isinstance(header_region.get("bbox"), list):
            element_type, bbox, source = "header", header_region["bbox"], "ocr_header"
        elif is_table and table_region and isinstance(table_region.get("bbox"), list):
            element_type, bbox, source = "body", table_region["bbox"], "ocr_table"
        elif is_note:
            region = footnote_regions[min(len(blocks), len(footnote_regions) - 1)] if footnote_regions else None
            bbox = region["bbox"] if region and isinstance(region.get("bbox"), list) else [0.08, 0.72, 0.92, 0.82]
            element_type, source = "footnote", "ocr_footnote"
        elif is_pagination:
            element_type, bbox, source = "pagination", [0.46, 0.94, 0.54, 0.98], "ocr_pagination"
        else:
            top = min(0.84, 0.60 + body_index * 0.1)
            element_type, bbox, source = "body", [0.08, top, 0.92, min(0.92, top + 0.09)], "ocr_body"
            body_index += 1
        blocks.append({"text": paragraph, "bbox": [float(value) for value in bbox],
                       "font_size": 0.0, "source": source, "ocr_element_type": element_type})
    return blocks


def rebuild_structured_translation_units(
    workspace: Path, *, page_text_overrides: dict[int, str] | None = None,
) -> dict[str, Any]:
    """Create independently translatable, geometry-anchored publication units.

    Body, headers, footers, footnotes, and figure captions intentionally have
    different unit types. Text inside a retained visual object is excluded so
    map labels and raster OCR cannot leak into body translation. Pagination is
    layout metadata and is not sent to a language provider.
    """
    workspace = workspace.resolve(); manifest = _load(workspace)
    classification_path = workspace / "data/page_classification.jsonl"
    if not classification_path.is_file():
        raise RuntimeError("structured translation units require page classification")
    classifications = {int(item["physical_page"]): item for item in _jsonl(classification_path)}
    document = fitz.open(Path(manifest["source_pdf"]))
    page_blocks: dict[int, list[dict[str, Any]]] = {}
    edge_occurrences: dict[tuple[str, str], set[int]] = {}
    try:
        for page_index, page in enumerate(document):
            page_no = page_index + 1
            blocks = _merge_wrapped_text_blocks(_geometric_text_blocks(page))
            if not blocks and (page_text_overrides or {}).get(page_no, "").strip():
                blocks = _ocr_fallback_blocks(
                    str((page_text_overrides or {})[page_no]).strip(), classifications.get(page_no, {}),
                )
            page_blocks[page_no] = blocks
            for block in blocks:
                edge = "header" if block["bbox"][1] <= 0.13 else "footer" if block["bbox"][3] >= 0.87 else ""
                key = _running_artifact_key(block["text"])
                if edge and key and len(block["text"]) <= 140:
                    edge_occurrences.setdefault((edge, key), set()).add(page_no)

        repeated_edges = {item for item, pages in edge_occurrences.items() if len(pages) >= 2}
        note_blocks = classify_note_blocks(page_blocks, classifications)
        block_types: dict[tuple[int, int], str] = {}
        page_number_pattern = re.compile(r"^\s*(?:\d{1,4}|[ivxlcdm]{1,8})[.!]?\s*$", re.I)

        # Pass one assigns a semantic type before any translation unit is made.
        # This keeps note discovery independent from provider output and permits
        # body references to resolve to note sections that occur later.
        for page_no, blocks in page_blocks.items():
            classification = classifications.get(page_no, {})
            visual_regions = [item for item in classification.get("visual_regions", []) if isinstance(item, dict)]
            semantic_regions = [item for item in classification.get("semantic_regions", []) if isinstance(item, dict)]
            sizes = sorted(block["font_size"] for block in blocks if block["font_size"])
            median_size = sizes[len(sizes) // 2] if sizes else 0.0
            for block_index, block in enumerate(blocks, 1):
                bbox = block["bbox"]; text = block["text"]
                caption_match = next((item for item in visual_regions
                                      if isinstance(item.get("caption_bbox"), list)
                                      and _rect_overlap_ratio(bbox, item["caption_bbox"]) >= 0.35), None)
                visual_match = next((item for item in visual_regions if isinstance(item.get("bbox"), list)
                                     and _rect_overlap_ratio(bbox, item["bbox"]) >= 0.40), None)
                semantic_match = next((item for item in semantic_regions if isinstance(item.get("bbox"), list)
                                       and _rect_overlap_ratio(bbox, item["bbox"]) >= 0.35), None)
                edge = "header" if bbox[1] <= 0.13 else "footer" if bbox[3] >= 0.87 else ""
                key = _running_artifact_key(text)
                if block.get("ocr_element_type"):
                    element_type = str(block["ocr_element_type"])
                elif page_number_pattern.fullmatch(text) and edge:
                    element_type = "pagination"
                elif caption_match is not None:
                    element_type = "figure_caption"
                elif visual_match is not None:
                    element_type = "visual_text"
                elif semantic_match is not None:
                    element_type = str(semantic_match["type"])
                elif edge and (edge, key) in repeated_edges:
                    element_type = edge
                elif (classification.get("page_class") == "body" and bbox[1] >= 0.68
                      and median_size and block["font_size"] <= median_size * 0.86
                      and re.match(r"^(?:[*\u2020\u2021]|\(?\d{1,3}[.)])\s*", text)):
                    element_type = "footnote"
                else:
                    element_type = "body"
                if (page_no, block_index) in note_blocks and element_type not in {"visual_text", "pagination"}:
                    element_type = str(note_blocks[(page_no, block_index)]["element_type"])
                block_types[(page_no, block_index)] = element_type

        note_plans: dict[tuple[int, int], list[dict[str, Any]]] = {}
        note_catalog: dict[str, list[dict[str, Any]]] = {}
        last_note_by_section: dict[str, str] = {}
        last_numeric_label_by_section: dict[str, int] = {}
        last_plan_by_section: dict[str, dict[str, Any]] = {}
        note_numbering_anomalies: list[dict[str, Any]] = []
        for page_no, blocks in page_blocks.items():
            for block_index, block in enumerate(blocks, 1):
                element_type = block_types[(page_no, block_index)]
                if element_type not in {"footnote", "chapter_endnote", "document_endnote"}:
                    continue
                note_info = note_blocks.get((page_no, block_index), {})
                scope = str(note_info.get("note_scope") or element_type)
                section_id = str(note_info.get("note_section_id") or f"page-{page_no:04d}-footnotes")
                entries = split_note_entries(block["text"])
                if element_type == "footnote" and entries and entries[0][0] is None:
                    bare = re.match(r"^\s*([*\u2020\u2021]|\d{1,4})[.)]?\s+(.+)$", block["text"], re.S)
                    if bare:
                        entries = [(bare.group(1), bare.group(2).strip())]
                plans: list[dict[str, Any]] = []
                for entry_index, (raw_label, text) in enumerate(entries, 1):
                    label = normalize_note_label(raw_label, last_numeric_label_by_section.get(section_id))
                    numeric_label = int(label) if label and label.isdigit() else None
                    previous_label = last_numeric_label_by_section.get(section_id)
                    if (numeric_label is not None and previous_label is not None
                            and numeric_label > previous_label + 25):
                        previous_plan = plans[-1] if plans else last_plan_by_section.get(section_id)
                        if previous_plan is not None:
                            previous_plan["text"] = previous_plan["text"].rstrip() + f"\n{label}. {text}"
                            note_numbering_anomalies.append({"source_page": page_no, "label": label,
                                                             "previous_label": str(previous_label),
                                                             "action": "retained_inside_previous_note"})
                            continue
                    current_note_id = note_id(scope, section_id, label) if label else None
                    continuation_of = None if current_note_id else last_note_by_section.get(section_id)
                    if current_note_id:
                        last_note_by_section[section_id] = current_note_id
                        if numeric_label is not None:
                            last_numeric_label_by_section[section_id] = numeric_label
                        candidate = {"label": label, "note_id": current_note_id, "scope": scope,
                                     "section_id": section_id, "source_page": page_no}
                        note_catalog.setdefault(str(label), []).append(candidate)
                    plan = {"label": label, "text": text, "note_id": current_note_id,
                            "continuation_of_note_id": continuation_of, "scope": scope,
                            "section_id": section_id, "entry_index": entry_index}
                    plans.append(plan); last_plan_by_section[section_id] = plan
                note_plans[(page_no, block_index)] = plans

        def notes_for_page(page_no: int) -> dict[str, str]:
            resolved: dict[str, str] = {}
            for label, candidates in note_catalog.items():
                same_page = [item for item in candidates if item["scope"] == "footnote"
                             and int(item["source_page"]) == page_no]
                future_chapter = [item for item in candidates if item["scope"] == "chapter_endnote"
                                  and int(item["source_page"]) >= page_no]
                document_notes = [item for item in candidates if item["scope"] == "document_endnote"]
                choices = same_page or sorted(future_chapter, key=lambda item: int(item["source_page"]))
                choices = choices or sorted(document_notes, key=lambda item: abs(int(item["source_page"]) - page_no))
                if len({item["note_id"] for item in choices}) == 1:
                    resolved[label] = choices[0]["note_id"]
            return resolved

        units: list[dict[str, Any]] = []
        layout_elements: list[dict[str, Any]] = []
        skipped_visual_text = 0
        element_counts: Counter[str] = Counter()
        note_graph_notes: list[dict[str, Any]] = []
        note_graph_references: list[dict[str, Any]] = []
        unresolved_references: list[dict[str, Any]] = []
        occurrence_counts: dict[str, int] = {}
        for page_no, blocks in page_blocks.items():
            classification = classifications.get(page_no, {})
            visual_regions = [item for item in classification.get("visual_regions", []) if isinstance(item, dict)]
            page_elements: list[dict[str, Any]] = []
            for visual_index, visual in enumerate(visual_regions, 1):
                bbox = visual.get("bbox")
                if isinstance(bbox, list) and len(bbox) == 4:
                    page_elements.append({"element_type": "visual", "source_page": page_no,
                                          "bbox": [float(value) for value in bbox],
                                          "visual_index": visual_index, "object_type": visual.get("type")})
            known_notes = notes_for_page(page_no)
            for block_index, block in enumerate(blocks, 1):
                bbox = block["bbox"]; element_type = block_types[(page_no, block_index)]
                if element_type == "visual_text":
                    skipped_visual_text += 1
                    continue
                if element_type == "pagination":
                    page_elements.append({"element_type": "pagination", "source_page": page_no,
                                          "bbox": bbox, "source_text": block["text"]})
                    continue
                plans = note_plans.get((page_no, block_index))
                source_entries = plans or [{"label": None, "text": block["text"], "note_id": None,
                                            "continuation_of_note_id": None, "scope": None,
                                            "section_id": None, "entry_index": 1}]
                for entry in source_entries:
                    source_text = str(entry["text"])
                    note_links: list[dict[str, str]] = []
                    if element_type == "body":
                        candidates = find_note_reference_labels(source_text)
                        source_text, note_links = annotate_note_references(
                            source_text, known_notes, page_no=page_no, occurrence_counts=occurrence_counts)
                        resolved_labels = {item["label"] for item in note_links}
                        unresolved_references.extend(
                            {"source_page": page_no, "label": label, "source_object_block": block_index}
                            for label in candidates if label not in resolved_labels
                        )
                    preserve_structure = block.get("source") == "ocr_table"
                    for chunk_index, chunk in enumerate(
                        _chunks(source_text, preserve_structure=preserve_structure), 1,
                    ):
                        stable = (f"{manifest['source_pdf_sha256']}|{page_no}|{element_type}|"
                                  f"{','.join(f'{value:.6f}' for value in bbox)}|{entry['entry_index']}|"
                                  f"{chunk_index}|{chunk}")
                        uid = "tu_" + hashlib.sha256(stable.encode("utf-8")).hexdigest()[:24]
                        source_object_id = (f"page-{page_no:04d}-{element_type}-{block_index:04d}-"
                                             f"{entry['entry_index']:02d}-{chunk_index:02d}")
                        source_block_id = (f"page-{page_no:04d}-{element_type}-{block_index:04d}-"
                                           f"{entry['entry_index']:02d}")
                        unit = {"translation_unit_id": uid, "source_object_id": source_object_id,
                                "source_block_id": source_block_id,
                                "source_page": page_no, "source_text": chunk,
                                "source_text_sha256": hashlib.sha256(chunk.encode("utf-8")).hexdigest(),
                                "source_language": manifest["source_language"],
                                "target_language": manifest["target_language"], "status": "pending",
                                "element_type": element_type, "bbox": bbox,
                                "reading_order": [round(bbox[1], 6), round(bbox[0], 6), block_index,
                                                  entry["entry_index"], chunk_index]}
                        if preserve_structure:
                            unit["preserve_structure"] = True
                        if note_links:
                            chunk_links = [item for item in note_links if item["placeholder"] in chunk]
                            if chunk_links:
                                unit["note_links"] = chunk_links
                                note_graph_references.extend({**item, "source_page": page_no,
                                                              "translation_unit_id": uid}
                                                             for item in chunk_links)
                        if entry["note_id"] and chunk_index == 1:
                            unit.update(note_id=entry["note_id"], note_label=entry["label"],
                                        note_scope=entry["scope"], note_section_id=entry["section_id"])
                            note_graph_notes.append({"note_id": entry["note_id"], "label": entry["label"],
                                                     "scope": entry["scope"], "section_id": entry["section_id"],
                                                     "source_page": page_no, "translation_unit_id": uid})
                        elif entry["continuation_of_note_id"] or (entry["note_id"] and chunk_index > 1):
                            unit["continuation_of_note_id"] = (entry["continuation_of_note_id"]
                                                               or entry["note_id"])
                            unit["note_scope"] = entry["scope"]
                            unit["note_section_id"] = entry["section_id"]
                        units.append(unit)
                        page_element = {"element_type": element_type, "source_page": page_no,
                                        "bbox": bbox, "translation_unit_id": uid,
                                        "source_object_id": source_object_id}
                        page_elements.append(page_element)
                        element_counts[element_type] += 1
            type_rank = {"header": 0, "body": 1, "visual": 1, "figure_caption": 1,
                         "note_heading": 1, "chapter_endnote": 1, "document_endnote": 1,
                         "footnote": 2, "footer": 3, "pagination": 4}
            page_elements.sort(key=lambda item: (type_rank.get(item["element_type"], 1),
                                                  item["bbox"][1], item["bbox"][0]))
            for sequence, element in enumerate(page_elements, 1):
                element["page_sequence"] = sequence
                layout_elements.append(element)
    finally:
        document.close()
    translation_groups = _attach_page_markers_and_translation_groups(units, layout_elements)
    atomic_write_jsonl(workspace / "data/translation_units.jsonl", units)
    atomic_write_jsonl(workspace / "data/page_layout_elements.jsonl", layout_elements)
    atomic_write_jsonl(workspace / "data/translation_groups.jsonl", translation_groups)
    manifest.update(stage="inspected", unit_count=len(units),
                    translation_group_count=len(translation_groups),
                    translation_unit_schema="anchored-publication-v3")
    _save(workspace, manifest)
    note_graph = {"schema_version": "publication-note-graph-v1", "notes": note_graph_notes,
                  "references": note_graph_references, "unresolved_references": unresolved_references,
                  "numbering_anomalies": note_numbering_anomalies,
                  "note_count": len(note_graph_notes), "reference_count": len(note_graph_references),
                  "unresolved_reference_count": len(unresolved_references)}
    atomic_write_json(workspace / "data/note_graph.json", note_graph)
    report = {"schema_version": "anchored-publication-v3", "units": len(units),
              "translation_groups": len(translation_groups),
              "element_counts": dict(sorted(element_counts.items())),
              "visual_text_blocks_excluded": skipped_visual_text,
              "layout_elements": len(layout_elements),
              "note_count": len(note_graph_notes), "note_reference_count": len(note_graph_references),
              "unresolved_note_reference_count": len(unresolved_references),
              "note_numbering_anomaly_count": len(note_numbering_anomalies)}
    atomic_write_json(workspace / "data/structured_translation_report.json", report)
    return report


def _jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(x) for x in path.read_text("utf-8").splitlines() if x.strip()]


def plan_workspace(workspace: Path) -> dict[str, Any]:
    workspace = workspace.resolve(); manifest = _load(workspace)
    units = _jsonl(workspace / "data/translation_units.jsonl")
    cache_dir = workspace / "cache" / manifest["language_pair"]
    pending = [u for u in units if not (cache_dir / f"{u['translation_unit_id']}.json").is_file()]
    report = {"workspace_id": manifest["workspace_id"], "language_pair": manifest["language_pair"],
              "unit_count": len(units), "pending": len(pending), "cached": len(units) - len(pending),
              "roles": list(OUTPUT_ROLES), "api_calls": 0}
    atomic_write_json(workspace / "data/translation_plan.json", report)
    return report


def _provider_call_count(workspace: Path) -> int:
    path = workspace / "logs/provider_calls.jsonl"
    if not path.is_file(): return 0
    total = 0
    for item in _jsonl(path):
        if "transport_calls" in item: total += int(item["transport_calls"])
        elif item.get("provider") == "mock": total += 1
        else: total += len(item.get("unit_ids", []))
    return total


def translate_workspace(workspace: Path, provider: Any, *, provider_name: str = "mock", model: str = "mock-v1",
                        batch_size: int = 8, max_units: int | None = None,
                        control: Callable[[], None] | None = None) -> dict[str, Any]:
    workspace = workspace.resolve(); manifest = _load(workspace); units = _jsonl(workspace / "data/translation_units.jsonl")
    overlay_path = workspace / "manual_review/imported_objects.json"
    if overlay_path.is_file():
        payload = json.loads(overlay_path.read_text("utf-8")); records = payload.get("objects", payload)
        overlays = {item["object_id"]: item for item in records if isinstance(item, dict) and item.get("object_id")}
        for unit in units:
            correction = overlays.get(unit["source_object_id"], {}).get("source_text")
            if correction is not None:
                unit["source_text"] = str(correction)
                unit["source_text_sha256"] = hashlib.sha256(str(correction).encode("utf-8")).hexdigest()
    cache_dir = workspace / "cache" / manifest["language_pair"]; cache_dir.mkdir(parents=True, exist_ok=True)
    pending = [u for u in units if not (cache_dir / f"{u['translation_unit_id']}.json").is_file()]
    if max_units is not None: pending = pending[:max_units]
    pending_ids = {str(unit["translation_unit_id"]) for unit in pending}
    units_by_id = {str(unit["translation_unit_id"]): unit for unit in units}
    groups_path = workspace / "data/translation_groups.jsonl"
    groups = _jsonl(groups_path) if groups_path.is_file() else []
    grouped_pending: set[str] = set()
    work_items: list[dict[str, Any]] = []
    for group in groups:
        member_ids = [str(value) for value in group.get("unit_ids", [])]
        if not pending_ids.intersection(member_ids):
            continue
        members = [units_by_id[value] for value in member_ids]
        grouped_pending.update(member_ids)
        markers = [dict(item) for item in group.get("markers", [])]
        marker_by_after = {str(item["after_unit_id"]): str(item["token"]) for item in markers}
        source_parts: list[str] = []
        for member in members:
            source_parts.append(str(member["source_text"]))
            token = marker_by_after.get(str(member["translation_unit_id"]))
            if token:
                source_parts.append(token)
        grouped_source = "\n".join(source_parts)
        work_items.append({
            **members[0],
            "translation_unit_id": str(group["translation_group_id"]),
            "source_object_id": str(group["translation_group_id"]),
            "source_text": grouped_source,
            "source_text_sha256": hashlib.sha256(grouped_source.encode("utf-8")).hexdigest(),
            "protected_terms": [str(item["token"]) for item in markers],
            "_translation_group": {"unit_ids": member_ids, "markers": markers},
        })
    work_items.extend(unit for unit in pending
                      if str(unit["translation_unit_id"]) not in grouped_pending)
    manifest["provider_calls"] = _provider_call_count(workspace)
    calls = 0; completed = 0
    for start in range(0, len(work_items), batch_size):
        if control:
            control()
        batch = work_items[start:start + batch_size]
        results = provider.translate_batch(batch)
        if control:
            control()
        transport_calls = 1 if provider_name == "mock" else len(batch)
        calls += transport_calls
        by_id = {x["translation_unit_id"]: x for x in results}
        if set(by_id) != {u["translation_unit_id"] for u in batch}:
            raise ValueError("provider result IDs do not match requested units")
        completed_ids: list[str] = []
        for work in batch:
            value = by_id[work["translation_unit_id"]]
            group = work.get("_translation_group")
            if group:
                translated = str(value["translated_text"])
                segments: list[str] = []
                remainder = translated
                for marker in group["markers"]:
                    token = str(marker["token"])
                    if token not in remainder:
                        raise ValueError(
                            f"provider removed protected page/unit boundary marker: {token}"
                        )
                    before, remainder = remainder.split(token, 1)
                    segments.append(before.strip())
                segments.append(remainder.strip())
                if len(segments) != len(group["unit_ids"]):
                    raise ValueError("translation group segment count mismatch")
                pairs = zip(group["unit_ids"], segments)
            else:
                pairs = [(str(work["translation_unit_id"]), str(value["translated_text"]))]
            for unit_id, translated_text in pairs:
                unit = units_by_id[unit_id]
                atomic_write_json(cache_dir / f"{unit_id}.json", {
                    "translation_unit_id": unit_id, "source_object_id": unit["source_object_id"],
                    "source_text_sha256": unit["source_text_sha256"],
                    "source_language": manifest["source_language"],
                    "target_language": manifest["target_language"],
                    "translated_text": translated_text,
                    "provider": provider_name, "model": model, "created_at": _now(),
                    "translation_group_id": work["translation_unit_id"] if group else None,
                })
                if unit_id in pending_ids:
                    completed += 1
                completed_ids.append(unit_id)
        with (workspace / "logs/provider_calls.jsonl").open("a", encoding="utf-8") as handle:
            handle.write(json.dumps({"timestamp": _now(), "provider": provider_name, "model": model,
                                     "unit_ids": completed_ids,
                                     "dispatch_ids": [u["translation_unit_id"] for u in batch],
                                     "transport_calls": transport_calls, "status": "completed"}) + "\n")
        manifest["provider_calls"] = int(manifest.get("provider_calls", 0)) + transport_calls
        manifest["stage"] = "translating"; _save(workspace, manifest)
        atomic_write_json(workspace / "checkpoints" / f"translation-{manifest['language_pair']}.json",
                          {"status": "in_progress", "completed": len(units) - len(pending) + completed,
                           "total": len(units), "provider_calls": manifest["provider_calls"], "updated_at": _now()})
    final = plan_workspace(workspace)
    status = "completed" if final["pending"] == 0 else "in_progress"
    atomic_write_json(workspace / "checkpoints" / f"translation-{manifest['language_pair']}.json",
                      {"status": status, "completed": len(units) - final["pending"], "total": len(units),
                       "provider_calls": manifest["provider_calls"], "updated_at": _now()})
    manifest["stage"] = "translated" if status == "completed" else "translating"; _save(workspace, manifest)
    return {"translated": completed, "api_calls": calls, "pending": final["pending"], "status": status}


def status_workspace(workspace: Path) -> dict[str, Any]:
    workspace = workspace.resolve(); manifest = _load(workspace); plan = plan_workspace(workspace)
    checkpoint = workspace / "checkpoints" / f"translation-{manifest['language_pair']}.json"
    calls = _provider_call_count(workspace)
    if calls != manifest.get("provider_calls", 0): manifest["provider_calls"] = calls; _save(workspace, manifest)
    return {**plan, "stage": manifest["stage"], "provider_calls": calls,
            "checkpoint": json.loads(checkpoint.read_text("utf-8")) if checkpoint.is_file() else None}


def control_workspace(workspace: Path, action: str) -> dict[str, Any]:
    """Record a durable, non-destructive pause/cancel transition."""
    if action not in {"pause", "cancel"}:
        raise ValueError("action must be pause or cancel")
    workspace = workspace.resolve()
    manifest = _load(workspace)
    previous = manifest["stage"]
    manifest["stage_before_control"] = previous
    manifest["stage"] = "paused" if action == "pause" else "cancelled"
    manifest["next_action"] = "resume" if action == "pause" else "inspect_or_resume"
    _save(workspace, manifest)
    atomic_write_json(workspace / "checkpoints" / "workspace-control.json", {
        "action": action, "previous_stage": previous, "stage": manifest["stage"],
        "updated_at": manifest["updated_at"],
    })
    return {"workspace": str(workspace), "previous_stage": previous,
            "stage": manifest["stage"], "cache_preserved": True}


def _content(workspace: Path, role: str) -> tuple[dict[str, Any], list[tuple[str, str, int]]]:
    if role not in OUTPUT_ROLES: raise ValueError(f"invalid output role {role!r}")
    manifest = _load(workspace); units = _jsonl(workspace / "data/translation_units.jsonl")
    overlay_path = workspace / "manual_review/imported_objects.json"
    overlays: dict[str, dict[str, Any]] = {}
    if overlay_path.is_file():
        payload = json.loads(overlay_path.read_text("utf-8")); records = payload.get("objects", payload)
        overlays = {item["object_id"]: item for item in records if isinstance(item, dict) and item.get("object_id")}
    cache_dir = workspace / "cache" / manifest["language_pair"]; values = []
    for unit in units:
        cache = cache_dir / f"{unit['translation_unit_id']}.json"
        target = json.loads(cache.read_text("utf-8"))["translated_text"] if cache.is_file() else ""
        overlay = overlays.get(unit["source_object_id"], {})
        source = overlay.get("source_text", unit["source_text"]); target = overlay.get("translated_text", target)
        if role != "source" and not target: raise RuntimeError(f"missing translation: {unit['translation_unit_id']}")
        values.append((source, target, int(unit["source_page"])))
    return manifest, values


def _publication_visual_assets(workspace: Path, manifest: dict[str, Any]) -> dict[int, list[Path]]:
    """Extract bounded visual objects without ever using a full source page.

    A PDF page commonly contains a full-page raster scan plus separately placed
    photographs or maps.  Treating the mere presence of an image object as a
    publication figure used to embed the whole source page in every edition.
    Only bounded image regions are valid publication assets here; page
    facsimiles remain source-viewer evidence and never enter the finished book.
    """
    classification_path = workspace / "data/page_classification.jsonl"
    if not classification_path.is_file():
        return {}
    classifications = {int(item["physical_page"]): item for item in _jsonl(classification_path)}
    output = workspace / "data/publication_assets"; output.mkdir(parents=True, exist_ok=True)
    document = fitz.open(Path(manifest["source_pdf"]))
    assets: dict[int, list[Path]] = {}
    objects: list[dict[str, Any]] = []
    try:
        for page_index, page in enumerate(document):
            page_number = page_index + 1
            classification = classifications.get(page_number, {})
            regions = classification.get("visual_regions")
            if not isinstance(regions, list):
                continue
            page_area = max(float(page.rect.get_area()), 1.0)
            accepted_regions: list[tuple[str, float, fitz.Rect, list[float] | None]] = []
            for region in regions:
                if not isinstance(region, dict) or region.get("type") not in {"photo", "map", "illustration", "diagram", "table"}:
                    continue
                if region.get("publication_role", "body_figure") != "body_figure":
                    continue
                try:
                    normalized = [float(value) for value in region.get("bbox", [])]
                    confidence = float(region.get("confidence", 0))
                except (TypeError, ValueError):
                    continue
                if (len(normalized) != 4 or confidence < 0.75
                        or not all(0 <= value <= 1 for value in normalized)
                        or normalized[2] <= normalized[0] or normalized[3] <= normalized[1]):
                    continue
                rect = fitz.Rect(normalized[0] * page.rect.width, normalized[1] * page.rect.height,
                                 normalized[2] * page.rect.width, normalized[3] * page.rect.height) & page.rect
                caption_bbox = region.get("caption_bbox")
                if isinstance(caption_bbox, list) and len(caption_bbox) == 4:
                    try:
                        caption = fitz.Rect(float(caption_bbox[0]) * page.rect.width,
                                            float(caption_bbox[1]) * page.rect.height,
                                            float(caption_bbox[2]) * page.rect.width,
                                            float(caption_bbox[3]) * page.rect.height) & page.rect
                    except (TypeError, ValueError):
                        caption_bbox = None
                area_ratio = float(rect.get_area() / page_area)
                if rect.is_empty or area_ratio >= 0.80 or rect.width < 24 or rect.height < 24:
                    continue
                accepted_regions.append((str(region["type"]), confidence, rect,
                                         caption_bbox if isinstance(caption_bbox, list) else None))
            for object_index, (object_type, confidence, rect, caption_bbox) in enumerate(accepted_regions, start=1):
                path = output / f"source-page-{page_number:04d}-{object_type}-{object_index:03d}.png"
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=rect, alpha=False)
                pixmap.save(path)
                assets.setdefault(page_number, []).append(path)
                objects.append({
                    "object_id": f"source-page-{page_number:04d}-object-{object_index:03d}",
                    "source_page": page_number,
                    "object_type": object_type,
                    "confidence": confidence,
                    "source": "vision_layout",
                    "representation_mode": "cropped_visual_object",
                    "bbox": [round(value, 3) for value in rect],
                    "normalized_bbox": [round(rect.x0 / page.rect.width, 6),
                                        round(rect.y0 / page.rect.height, 6),
                                        round(rect.x1 / page.rect.width, 6),
                                        round(rect.y1 / page.rect.height, 6)],
                    "caption_bbox": caption_bbox,
                    "path": str(path),
                    "sha256": sha256_file(path),
                })
    finally:
        document.close()
    atomic_write_json(workspace / "data/publication_visual_assets.json", {
        "schema_version": "bookflow-publication-visual-assets-v2",
        "source_pdf_sha256": manifest["source_pdf_sha256"],
        "full_page_facsimile_count": 0,
        "objects": objects,
    })
    return assets


def reflow_text(text: str, language: str, *, preserve_structure: bool = False) -> str:
    """Remove OCR visual line breaks while retaining explicit paragraph boundaries."""
    value = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    paragraphs = re.split(r"\n\s*\n", value)
    latin = language in {"en", "fr", "de", "es"}
    result: list[str] = []
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if preserve_structure:
            result.append(paragraph)
            continue
        if latin:
            paragraph = re.sub(r"(?<=\w)-\s*\n\s*(?=[\w\u00c0-\u024f])", "", paragraph)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in paragraph.split("\n") if line.strip()]
        groups: list[list[str]] = []; locked: list[bool] = []
        for line in lines:
            letters = [char for char in line if char.isalpha()]
            uppercase = bool(letters) and sum(char.isupper() for char in letters) / len(letters) >= 0.8
            semantic_boundary = (len(line) <= 100 and uppercase) or bool(re.match(
                r"^(chapter|chapitre|kapitel|cap[ií]tulo|appendix|annexe|table|contents|index)\b", line, re.I))
            semantic_boundary = semantic_boundary or bool(re.match(
                r"^第[一二三四五六七八九十百零〇0-9]+章(?:\s|$)", line))
            list_item = bool(re.match(r"^(?:[-*•]|\d+[.)]|[IVXLCDM]+[.)])\s+", line))
            if semantic_boundary or list_item:
                groups.append([line]); locked.append(True)
            elif groups and not locked[-1]:
                groups[-1].append(line)
            else:
                groups.append([line]); locked.append(False)
        separator = "" if language in {"zh-Hans", "ja"} else " "
        result.extend(separator.join(group).strip() for group in groups if group)
    return "\n\n".join(result)


def _running_artifact_key(line: str) -> str:
    value = re.sub(r"\s+", " ", line).strip().casefold()
    value = re.sub(r"^[\W_]+|[\W_]+$", "", value)
    return value


def _discover_running_artifacts(texts: list[str]) -> set[str]:
    """Find repeated short lines confined to page edges in the current book."""
    pages_by_key: dict[str, set[int]] = {}
    for page_index, text in enumerate(texts):
        lines = [line.strip() for line in text.replace("\r", "").split("\n") if line.strip()]
        edge_lines = lines[:3] + lines[-3:]
        for line in edge_lines:
            key = _running_artifact_key(line)
            words = key.split()
            if not key or len(line) > 60 or len(words) > 8 or _is_chapter_heading(line):
                continue
            compact = re.sub(r"\s+", "", key)
            contains_cjk = bool(re.search(r"[\u3400-\u9fff\u3040-\u30ff]", compact))
            if len(words) == 1 and not (len(compact) >= 6 or (contains_cjk and len(compact) >= 2)):
                continue
            if re.fullmatch(r"(?:\d{1,4}|[ivxlcdm]{1,6})", key, re.I):
                continue
            pages_by_key.setdefault(key, set()).add(page_index)
    return {key for key, pages in pages_by_key.items()
            if len(pages) >= (2 if len(key) >= 12 or len(key.split()) >= 3 else 3)}


def _strip_running_artifacts(text: str, signatures: set[str]) -> tuple[str, int, int]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    nonempty = [index for index, line in enumerate(lines) if line.strip()]
    edge = set(nonempty[:3] + nonempty[-3:])
    artifact_indexes = {index for index in edge if _running_artifact_key(lines[index]) in signatures}
    page_marker_indexes: set[int] = set()
    marker_pattern = re.compile(r"^\s*(?:\d{1,4}|[ivxlcdm]{1,6})[.!]?\s*$", re.I)
    for index in edge:
        if not marker_pattern.fullmatch(lines[index]):
            continue
        nearest = min(artifact_indexes, key=lambda item: abs(item - index), default=None)
        if nearest is not None and abs(nearest - index) <= 2:
            page_marker_indexes.add(index)
    removed = artifact_indexes | page_marker_indexes
    cleaned = "\n".join(line for index, line in enumerate(lines) if index not in removed)
    return cleaned, len(artifact_indexes), len(page_marker_indexes)


def _publication_clean_values(values: list[tuple[str, str, int]],
                              classifications: list[dict[str, Any]] | None = None
                              ) -> tuple[list[tuple[str, str, int]], dict[str, Any]]:
    eligible_pages = {int(item["physical_page"]) for item in classifications or []
                      if item.get("page_class") in {"body", "toc", "index", "appendix"}}
    discovery_pages = {int(item["physical_page"]) for item in classifications or []
                       if int(item["physical_page"]) in eligible_pages and not item.get("chapter_boundary")}
    if not classifications:
        eligible_pages = {page for _, _, page in values}
        discovery_pages = eligible_pages
    source_signatures = _discover_running_artifacts(
        [source for source, _, page in values if page in discovery_pages])
    target_signatures = _discover_running_artifacts(
        [target for _, target, page in values if target and page in discovery_pages])
    cleaned: list[tuple[str, str, int]] = []
    source_artifacts = source_markers = target_artifacts = target_markers = 0
    for source, target, page in values:
        if page not in eligible_pages:
            cleaned.append((source, target, page))
            continue
        source, artifacts, markers = _strip_running_artifacts(source, source_signatures)
        source_artifacts += artifacts; source_markers += markers
        if target:
            target, artifacts, markers = _strip_running_artifacts(target, target_signatures)
            target_artifacts += artifacts; target_markers += markers
        cleaned.append((source, target, page))
    return cleaned, {
        "source_signatures": sorted(source_signatures), "target_signatures": sorted(target_signatures),
        "source_artifact_lines_removed": source_artifacts, "source_page_markers_removed": source_markers,
        "target_artifact_lines_removed": target_artifacts, "target_page_markers_removed": target_markers,
    }


def reading_reflow_diagnostics(values: list[tuple[str, str]], source_language: str,
                               target_language: str) -> dict[str, Any]:
    before = [text for pair in values for text in pair if text]
    after = [reflow_text(source, source_language) for source, _ in values]
    after.extend(reflow_text(target, target_language) for _, target in values if target)
    return {
        "input_hard_line_breaks": sum(text.count("\n") for text in before),
        "output_hard_line_breaks": sum(len(re.findall(r"(?<!\n)\n(?!\n)", text)) for text in after),
        "suspected_line_end_hyphens_before": sum(len(re.findall(r"\w-\s*\n\s*\w", text)) for text in before),
        "suspected_line_end_hyphens_after": sum(len(re.findall(r"\w-[ \t]*\n(?!\n)[ \t]*\w", text)) for text in after),
    }


def _is_chapter_heading(text: str) -> bool:
    return bool(re.match(
        r"^(?:(?:chapter|chapitre|kapitel|cap[ií]tulo|appendix|annexe)\s+(?:[IVXLCDM]+|\d+)\b|"
        r"第[一二三四五六七八九十百零〇0-9]+章)", text, re.I))


def _add_structured_paragraph(document: Document, text: str, *, heading: bool = False) -> None:
    document.add_heading(text, level=1) if heading else document.add_paragraph(text)


def _commit_generated_file(temporary: Path, destination: Path) -> Path:
    """Commit to the stable name, or preserve a new version when that file is open."""
    try:
        temporary.replace(destination)
        return destination
    except PermissionError:
        stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        fallback = destination.with_name(f"{destination.stem}-{stamp}{destination.suffix}")
        temporary.replace(fallback)
        return fallback


def _publication_font(language: str) -> tuple[Path, str]:
    fonts = Path(r"C:\Windows\Fonts")
    choices = {
        "zh-Hans": [("NotoSerifSC-VF.ttf", "Noto Serif SC"), ("msyh.ttc", "Microsoft YaHei")],
        "ja": [("NotoSerifJP-VF.ttf", "Noto Serif JP"), ("yumin.ttf", "Yu Mincho")],
    }.get(language, [("NotoSerif-Regular.ttf", "Noto Serif"), ("times.ttf", "Times New Roman")])
    for filename, family in choices:
        path = fonts / filename
        if path.is_file():
            return path, family
    raise RuntimeError(f"no embeddable publication font available for {language}")


def _unit_render_parts(unit: dict[str, Any], role: str, manifest: dict[str, Any]) -> list[tuple[str, str]]:
    preserve = bool(unit.get("preserve_structure"))
    source = unicodedata.normalize(
        "NFKC",
        reflow_text(str(unit.get("source_text") or ""), manifest["source_language"],
                    preserve_structure=preserve),
    )
    target = unicodedata.normalize(
        "NFKC",
        reflow_text(str(unit.get("translated_text") or ""), manifest["target_language"],
                    preserve_structure=preserve),
    )
    before = [
        (f"【{item['printed_page']}】", "zh-Hans")
        for item in unit.get("page_markers_before", [])
    ]
    after = [
        (f"【{item['printed_page']}】", "zh-Hans")
        for item in unit.get("page_markers_after", [])
    ]
    if role == "source":
        return [*before, (source, manifest["source_language"]), *after]
    if role == "target":
        return [*before, (target, manifest["target_language"]), *after]
    return [*before, (source, manifest["source_language"]), *after,
            (target, manifest["target_language"])]


def _primary_content_part(parts: list[tuple[str, str]]) -> int:
    for index, (text, _) in enumerate(parts):
        if not re.fullmatch(r"【[^】]+】", text.strip()):
            return index
    return 0


def _original_page_marker_label(value: Any, physical_page: int) -> tuple[str, str]:
    """Prefer a detected printed label and fall back to the source PDF page number."""
    raw = str(value or "").strip().strip("【】").strip()
    match = re.fullmatch(r"(\d{1,4}|[ivxlcdm]{1,8})[.!]?", raw, re.I)
    if match:
        return match.group(1), "detected_printed_page"
    return str(physical_page), "physical_page_fallback"


def _anchored_publication_pages(workspace: Path, role: str, manifest: dict[str, Any],
                                 visual_assets: dict[int, list[Path]]) -> dict[int, list[dict[str, Any]]]:
    units = {item["translation_unit_id"]: item for item in _jsonl(workspace / "data/translation_units.jsonl")}
    overlays: dict[str, dict[str, Any]] = {}
    overlay_path = workspace / "manual_review/imported_objects.json"
    if overlay_path.is_file():
        payload = json.loads(overlay_path.read_text("utf-8"))
        records = payload.get("objects", payload)
        overlays = {
            item["object_id"]: item
            for item in records
            if isinstance(item, dict) and item.get("object_id")
        }
    cache_dir = workspace / "cache" / manifest["language_pair"]
    for uid, unit in units.items():
        cache = cache_dir / f"{uid}.json"
        unit["translated_text"] = json.loads(cache.read_text("utf-8"))["translated_text"] if cache.is_file() else ""
        overlay = overlays.get(unit["source_object_id"], {})
        if "source_text" in overlay:
            unit["source_text"] = str(overlay["source_text"])
        if "translated_text" in overlay:
            unit["translated_text"] = str(overlay["translated_text"])
        if role != "source" and not unit["translated_text"]:
            raise RuntimeError(f"missing translation: {uid}")
    pages: dict[int, list[dict[str, Any]]] = {}
    marker_labels: dict[int, tuple[str, str]] = {}
    embedded_marker_pages = {
        int(marker["physical_page"])
        for unit in units.values()
        for key in ("page_markers_before", "page_markers_after")
        for marker in unit.get(key, [])
    }
    for raw in _jsonl(workspace / "data/page_layout_elements.jsonl"):
        element = dict(raw); page_no = int(element["source_page"])
        if element["element_type"] == "pagination":
            marker_labels.setdefault(
                page_no, _original_page_marker_label(element.get("source_text"), page_no),
            )
            continue
        if element["element_type"] == "visual":
            visual_index = int(element.get("visual_index", 0))
            paths = visual_assets.get(page_no, [])
            if 1 <= visual_index <= len(paths):
                element["path"] = paths[visual_index - 1]
            else:
                continue
        else:
            unit = units.get(element.get("translation_unit_id"))
            if unit is None:
                continue
            element["parts"] = _unit_render_parts(unit, role, manifest)
            for key in ("note_links", "note_id", "note_label", "note_scope", "note_section_id",
                        "continuation_of_note_id"):
                if key in unit:
                    element[key] = unit[key]
        pages.setdefault(page_no, []).append(element)
    classification_path = workspace / "data/page_classification.jsonl"
    classified_pages = {
        int(item["physical_page"]) for item in _jsonl(classification_path)
        if item.get("physical_page") is not None
    } if classification_path.is_file() else set()
    for page_no in sorted(set(pages) | classified_pages):
        if page_no in embedded_marker_pages:
            continue
        label, marker_source = marker_labels.get(
            page_no, _original_page_marker_label(None, page_no),
        )
        pages.setdefault(page_no, []).append({
            "element_type": "original_page_marker",
            "source_page": page_no,
            "page_sequence": -1,
            "bbox": [0.44, 0.0, 0.56, 0.035],
            "parts": [(f"【{label}】", "zh-Hans")],
            "marker_label": label,
            "marker_source": marker_source,
        })
    references_by_note: dict[str, list[dict[str, str]]] = {}
    for elements in pages.values():
        for element in elements:
            for link in element.get("note_links", []):
                references_by_note.setdefault(link["note_id"], []).append(link)
    for elements in pages.values():
        for element in elements:
            if element.get("note_id") in references_by_note:
                element["note_backlinks"] = references_by_note[element["note_id"]]
    for elements in pages.values():
        elements.sort(key=lambda item: int(item.get("page_sequence", 0)))
    return pages


def _docx_bookmark_name(value: str) -> str:
    return ("bf_" + re.sub(r"[^0-9A-Za-z_]", "_", value))[:40]


def _docx_bookmark(paragraph: Any, value: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), _docx_bookmark_name(value))
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(start); paragraph._p.append(end)


def _docx_hyperlink(paragraph: Any, text: str, target: str, *, language: str, size: float,
                    italic: bool = False) -> None:
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("w:anchor"), _docx_bookmark_name(target))
    run = paragraph.add_run(text)
    _, family = _publication_font(language)
    run.font.name = family; run.font.size = Pt(size); run.font.italic = italic
    run.font.underline = True; run.font.color.rgb = None
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), family)
    paragraph._p.remove(run._element); hyperlink.append(run._element); paragraph._p.append(hyperlink)


def _set_docx_text(paragraph: Any, parts: list[tuple[str, str]], *, size: float = 10.5,
                   italic: bool = False, note_links: list[dict[str, str]] | None = None,
                   bookmark_counter: list[int] | None = None) -> None:
    links = {item["placeholder"]: item for item in note_links or []}
    anchor_part = _primary_content_part(parts)
    for position, (text, language) in enumerate(parts):
        if position:
            paragraph.add_run("\n")
        cursor = 0
        for match in NOTE_PLACEHOLDER_RE.finditer(text):
            if match.start() > cursor:
                run = paragraph.add_run(text[cursor:match.start()])
                _, family = _publication_font(language)
                run.font.name = family; run.font.size = Pt(size); run.font.italic = italic
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), family)
            link = links.get(match.group(0))
            if link:
                if position == anchor_part and bookmark_counter is not None:
                    _docx_bookmark(paragraph, link["reference_id"], bookmark_counter[0])
                    bookmark_counter[0] += 1
                _docx_hyperlink(paragraph, link["label"], link["note_id"],
                                language=language, size=size, italic=italic)
            else:
                paragraph.add_run(match.group(1))
            cursor = match.end()
        if cursor < len(text):
            run = paragraph.add_run(text[cursor:])
            _, family = _publication_font(language)
            run.font.name = family; run.font.size = Pt(size); run.font.italic = italic
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), family)


def _markdown_linked_text(text: str, item: dict[str, Any], *, anchor_references: bool) -> str:
    links = {link["placeholder"]: link for link in item.get("note_links", [])}

    def replace(match: re.Match[str]) -> str:
        link = links.get(match.group(0))
        if not link:
            return match.group(1)
        anchor = f'<a id="{link["reference_id"]}"></a>' if anchor_references else ""
        return f'{anchor}<a href="#{link["note_id"]}">{link["label"]}</a>'

    return NOTE_PLACEHOLDER_RE.sub(replace, text)


def _measure_wrapped_lines(text: str, font: fitz.Font, fontsize: float, width: float) -> list[str]:
    output: list[str] = []
    for paragraph in text.replace("\r", "").split("\n"):
        if not paragraph:
            output.append(""); continue
        tokens = re.findall(r"[\u3400-\u9fff\u3040-\u30ff]|\s+|[^\s\u3400-\u9fff\u3040-\u30ff]+", paragraph)
        line = ""
        for token in tokens:
            candidate = (line + token).lstrip() if not line else line + token
            if line and font.text_length(candidate, fontsize=fontsize) > width:
                output.append(line.rstrip()); line = token.lstrip()
                while line and font.text_length(line, fontsize=fontsize) > width:
                    split_at = max(1, int(len(line) * width / max(font.text_length(line, fontsize=fontsize), 1)))
                    output.append(line[:split_at]); line = line[split_at:]
            else:
                line = candidate
        output.append(line.rstrip())
    return output or [""]


def _render_anchored_publication(workspace: Path, role: str, formats: tuple[str, ...],
                                 pdf_renderer: str, renderer_config: Path | None,
                                 manifest: dict[str, Any], out: Path, stem: str,
                                 values: list[tuple[str, str, int]]) -> dict[str, Any]:
    visual_assets = _publication_visual_assets(workspace, manifest)
    pages = _anchored_publication_pages(workspace, role, manifest, visual_assets)
    outputs: dict[str, Any] = {}
    if "md" in formats:
        path = out / f"{stem}.md"; page_markdown: list[str] = []

        def markdown_item(item: dict[str, Any]) -> str:
            rendered_parts: list[str] = []
            parts = item.get("parts", [])
            anchor_part = _primary_content_part(parts)
            for position, (text, _) in enumerate(parts):
                rendered = _markdown_linked_text(
                    text, item, anchor_references=position == anchor_part,
                )
                if item.get("note_id"):
                    prefix = (
                        f'<a id="{item["note_id"]}"></a>' if position == anchor_part else ""
                    )
                    rendered = prefix + f'{item.get("note_label")}. ' + rendered
                rendered_parts.append(rendered)
            rendered = "\n\n".join(rendered_parts)
            if item.get("note_backlinks"):
                returns = " ".join(
                    f'<a href="#{link["reference_id"]}">back {index}</a>'
                    for index, link in enumerate(item["note_backlinks"], 1)
                )
                rendered += " " + returns
            return rendered

        for page_no, elements in sorted(pages.items()):
            header = [markdown_item(item) for item in elements if item["element_type"] == "header"]
            footer = [markdown_item(item) for item in elements if item["element_type"] == "footer"]
            footnotes = [markdown_item(item) for item in elements if item["element_type"] == "footnote"]
            body: list[str] = []
            for item in elements:
                if item["element_type"] in {"header", "footer", "footnote"}:
                    continue
                if item["element_type"] == "visual":
                    body.append(f"![{item.get('object_type', 'visual')} from source page {page_no}]({item['path'].as_posix()})")
                else:
                    rendered = markdown_item(item)
                    if rendered:
                        body.append(("## " if item["element_type"] == "note_heading" else "") + rendered)
            chunks = [f"<!-- source-page: {page_no} -->"]
            if header: chunks.append("<header>" + "<br>".join(header) + "</header>")
            chunks.extend(body)
            if footnotes: chunks.append("---\n\n" + "\n\n".join(f"<small>{text}</small>" for text in footnotes))
            if footer: chunks.append("<footer>" + "<br>".join(footer) + "</footer>")
            page_markdown.append("\n\n".join(chunks))
        atomic_write_text(path, "\n\n<div class=\"source-page-break\"></div>\n\n".join(page_markdown) + "\n")
        outputs["md"] = str(path)

    docx_path: Path | None = None
    if "docx" in formats or ("pdf" in formats and pdf_renderer == "office"):
        path = out / f"{stem}.docx"; doc = Document()
        title = manifest.get("book_title") or Path(manifest.get("source_filename", "Book")).stem or "Book"
        doc.core_properties.title = title
        bookmark_counter = [1]

        def set_docx_item(paragraph: Any, item: dict[str, Any], *, size: float,
                          italic: bool = False) -> None:
            if item.get("note_id"):
                _docx_bookmark(paragraph, item["note_id"], bookmark_counter[0])
                bookmark_counter[0] += 1
                paragraph.add_run(f'{item.get("note_label")}. ')
            _set_docx_text(paragraph, item.get("parts", []), size=size, italic=italic,
                           note_links=item.get("note_links"), bookmark_counter=bookmark_counter)
            for index, link in enumerate(item.get("note_backlinks", []), 1):
                paragraph.add_run(" ")
                language = item.get("parts", [("", manifest["source_language"])])[0][1]
                _docx_hyperlink(paragraph, f"back {index}", link["reference_id"],
                                language=language, size=size, italic=italic)

        for page_position, (page_no, elements) in enumerate(sorted(pages.items())):
            section = doc.sections[0] if page_position == 0 else doc.add_section(WD_SECTION.NEW_PAGE)
            section.left_margin = Inches(0.82); section.right_margin = Inches(0.82)
            section.top_margin = Inches(0.72); section.bottom_margin = Inches(0.72)
            section.header.is_linked_to_previous = False; section.footer.is_linked_to_previous = False
            header_items = [item for item in elements if item["element_type"] == "header"]
            footer_items = [item for item in elements if item["element_type"] == "footer"]
            footnote_items = [item for item in elements if item["element_type"] == "footnote"]
            for position, item in enumerate(header_items):
                paragraph = section.header.paragraphs[0] if position == 0 else section.header.add_paragraph()
                set_docx_item(paragraph, item, size=8.5)
            if footnote_items:
                footnote_paragraph = section.footer.paragraphs[0]
                footnote_paragraph.add_run("____________________________\n")
                for position, item in enumerate(footnote_items):
                    if position:
                        footnote_paragraph.add_run("\n")
                    set_docx_item(footnote_paragraph, item, size=8.0)
            for position, item in enumerate(footer_items):
                footer_paragraph = (section.footer.add_paragraph() if footnote_items or position
                                    else section.footer.paragraphs[0])
                set_docx_item(footer_paragraph, item, size=8.5)
            for item in elements:
                kind = item["element_type"]
                if kind in {"header", "footer", "footnote"}:
                    continue
                if kind == "visual":
                    doc.add_picture(str(item["path"]), width=Inches(5.9))
                else:
                    paragraph = doc.add_paragraph()
                    if kind == "note_heading":
                        paragraph.style = doc.styles["Heading 2"]
                    if kind == "original_page_marker":
                        paragraph.alignment = 1
                    set_docx_item(paragraph, item, size=8.5 if kind == "original_page_marker"
                                  else 9.5 if kind == "figure_caption" else 10.5,
                                  italic=kind == "figure_caption")
        temporary_docx = path.with_suffix(".docx.tmp"); doc.save(temporary_docx)
        docx_path = _commit_generated_file(temporary_docx, path)
        if "docx" in formats: outputs["docx"] = str(docx_path)

    if "pdf" in formats:
        path = out / f"{stem}.pdf"
        if pdf_renderer == "office":
            from .renderer_backend import convert_docx_to_pdf, load_renderer_config
            if docx_path is None: raise RuntimeError("office PDF rendering requires a generated DOCX")
            conversion = convert_docx_to_pdf(load_renderer_config(renderer_config), docx_path, path)
            outputs["pdf"] = conversion if conversion["status"] != "generated" else conversion["path"]
        elif pdf_renderer == "native_pdf":
            pdf = fitz.open(); font_cache: dict[str, tuple[Path, fitz.Font, str]] = {}
            pdf_destinations: dict[str, tuple[int, fitz.Point]] = {}
            pdf_pending_links: list[tuple[int, fitz.Rect, str]] = []
            for language in {manifest["source_language"], manifest["target_language"], "zh-Hans"}:
                font_path, _ = _publication_font(language)
                font_cache[language] = (font_path, fitz.Font(fontfile=str(font_path)),
                                        "bf_" + language.replace("-", "_"))

            def new_page(header_items: list[dict[str, Any]], footer_items: list[dict[str, Any]]) -> tuple[fitz.Page, float]:
                page = pdf.new_page(width=595, height=842)
                for font_path, _, alias in font_cache.values():
                    page.insert_font(fontname=alias, fontfile=str(font_path))
                for items, y in ((header_items, 30.0), (footer_items, 817.0)):
                    x = 48.0
                    for item in items:
                        for text, language in item["parts"]:
                            _, font, alias = font_cache[language]
                            for line in _measure_wrapped_lines(text, font, 7.5, 499):
                                page.insert_text((x, y), line, fontsize=7.5, fontname=alias, color=(0.3, 0.3, 0.3))
                                y += 9
                return page, 58.0

            def marked_text(text: str, item: dict[str, Any]) -> str:
                links = {link["placeholder"]: link for link in item.get("note_links", [])}
                return NOTE_PLACEHOLDER_RE.sub(
                    lambda match: "\x01" + str(links.get(match.group(0), {}).get("label", match.group(1))) + "\x02",
                    text,
                )

            def insert_linked_line(page: fitz.Page, line: str, language: str, y: float,
                                   fontsize: float, item: dict[str, Any], link_position: list[int],
                                   *, anchor_references: bool) -> None:
                _, font, alias = font_cache[language]
                x = 48.0; cursor = 0; links = item.get("note_links", [])
                for match in re.finditer(r"\x01([^\x02]*)\x02", line):
                    prefix = line[cursor:match.start()]
                    if prefix:
                        page.insert_text((x, y), prefix, fontsize=fontsize, fontname=alias)
                        x += font.text_length(prefix, fontsize=fontsize)
                    label = match.group(1)
                    link = links[link_position[0]] if link_position[0] < len(links) else None
                    if link:
                        if anchor_references:
                            pdf_destinations.setdefault(link["reference_id"], (page.number, fitz.Point(x, y)))
                        label_width = max(font.text_length(label, fontsize=fontsize), fontsize * 0.55)
                        pdf_pending_links.append((page.number,
                                                  fitz.Rect(x, y - fontsize, x + label_width, y + 2),
                                                  link["note_id"]))
                        link_position[0] += 1
                    page.insert_text((x, y), label, fontsize=fontsize, fontname=alias, color=(0.0, 0.2, 0.65))
                    x += font.text_length(label, fontsize=fontsize); cursor = match.end()
                suffix = line[cursor:]
                if suffix:
                    page.insert_text((x, y), suffix, fontsize=fontsize, fontname=alias)

            def insert_pdf_backlinks(page: fitz.Page, item: dict[str, Any], language: str,
                                     y: float, fontsize: float) -> None:
                _, font, alias = font_cache[language]; x = 48.0
                for index, link in enumerate(item.get("note_backlinks", []), 1):
                    label = f"back {index}"
                    width = max(font.text_length(label, fontsize=fontsize), fontsize)
                    page.insert_text((x, y), label, fontsize=fontsize, fontname=alias, color=(0.0, 0.2, 0.65))
                    pdf_pending_links.append((page.number, fitz.Rect(x, y - fontsize, x + width, y + 2),
                                              link["reference_id"]))
                    x += width + fontsize

            for _, elements in sorted(pages.items()):
                header_items = [item for item in elements if item["element_type"] == "header"]
                footer_items = [item for item in elements if item["element_type"] == "footer"]
                footnote_items = [item for item in elements if item["element_type"] == "footnote"]
                footnote_lines: list[dict[str, Any]] = []
                for item in footnote_items:
                    anchor_part = _primary_content_part(item["parts"])
                    for part_position, (text, language) in enumerate(item["parts"]):
                        _, font, _ = font_cache[language]
                        value = marked_text(text, item)
                        if item.get("note_id"):
                            value = f'{item.get("note_label")}. ' + value
                        cursor = [0]
                        for line_position, line in enumerate(_measure_wrapped_lines(value, font, 7.5, 499)):
                            footnote_lines.append({"line": line, "language": language, "item": item,
                                                   "link_position": cursor, "anchor": part_position == anchor_part,
                                                   "note_destination": bool(item.get("note_id") and part_position == anchor_part
                                                                            and line_position == 0)})
                    if item.get("note_backlinks"):
                        footnote_lines.append({"backlinks": True, "language": item["parts"][0][1], "item": item})
                footnote_top = max(650.0, 794.0 - len(footnote_lines) * 9.0 - (12.0 if footnote_lines else 0.0))
                page, y = new_page(header_items, footer_items)
                for item in elements:
                    kind = item["element_type"]
                    if kind in {"header", "footer", "footnote"}:
                        continue
                    if kind == "visual":
                        image = fitz.Pixmap(str(item["path"])); available = footnote_top - y - 12
                        scale = min(499 / image.width, max(80.0, available) / image.height, 1.0)
                        image_width, image_height = image.width * scale, image.height * scale
                        if y + image_height > footnote_top:
                            page, y = new_page(header_items, footer_items)
                        page.insert_image(fitz.Rect(48, y, 48 + image_width, y + image_height),
                                          filename=str(item["path"]), keep_proportion=True)
                        y += image_height + 10
                        continue
                    fontsize = 8.2 if kind == "figure_caption" else 9.5
                    content_bottom = footnote_top - (24.0 if item.get("note_backlinks") else 0.0)
                    anchor_part = _primary_content_part(item["parts"])
                    for part_position, (text, language) in enumerate(item["parts"]):
                        _, font, alias = font_cache[language]
                        value = marked_text(text, item)
                        if item.get("note_id"):
                            value = f'{item.get("note_label")}. ' + value
                        link_position = [0]
                        for line_position, line in enumerate(_measure_wrapped_lines(value, font, fontsize, 499)):
                            if y + 12 > content_bottom:
                                page, y = new_page(header_items, footer_items)
                            if item.get("note_id") and part_position == anchor_part and line_position == 0:
                                pdf_destinations[item["note_id"]] = (page.number, fitz.Point(48, y))
                            if line:
                                insert_linked_line(page, line, language, y, fontsize, item, link_position,
                                                   anchor_references=part_position == anchor_part)
                            y += fontsize * 1.38
                        y += 4
                    if item.get("note_backlinks"):
                        if y + 12 > footnote_top:
                            page, y = new_page(header_items, footer_items)
                        insert_pdf_backlinks(page, item, item["parts"][0][1], y, 8.2); y += 12
                if footnote_lines:
                    page.draw_line((48, footnote_top), (230, footnote_top), color=(0.35, 0.35, 0.35), width=0.5)
                    fy = footnote_top + 11
                    for row in footnote_lines:
                        if row.get("backlinks"):
                            insert_pdf_backlinks(page, row["item"], row["language"], fy, 7.0)
                            fy += 9
                            continue
                        if row.get("note_destination"):
                            pdf_destinations[row["item"]["note_id"]] = (page.number, fitz.Point(48, fy))
                        insert_linked_line(page, row["line"], row["language"], fy, 7.5, row["item"],
                                           row["link_position"], anchor_references=row["anchor"])
                        fy += 9
            for source_page, rect, target in pdf_pending_links:
                destination = pdf_destinations.get(target)
                if destination is None:
                    continue
                destination_page, point = destination
                pdf[source_page].insert_link({"kind": fitz.LINK_GOTO, "from": rect,
                                              "page": destination_page, "to": point})
            temporary_pdf = path.with_suffix(".pdf.tmp"); pdf.save(temporary_pdf, garbage=4, deflate=True); pdf.close()
            outputs["pdf"] = str(_commit_generated_file(temporary_pdf, path))
        else:
            raise ValueError(f"unknown PDF renderer {pdf_renderer!r}")

    manifest_outputs = {key: value if isinstance(value, dict) else
                        {"status": "generated", "path": value, "sha256": sha256_file(value)}
                        for key, value in outputs.items()}
    report = {"role": role, "source_language": manifest["source_language"],
              "target_language": manifest["target_language"], "language_pair": manifest["language_pair"],
              "layout_mode": "publication", "anchored_layout": True, "pdf_renderer": pdf_renderer,
              "visual_object_count": sum(len(paths) for paths in visual_assets.values()),
              "full_page_facsimile_count": 0,
              "font_files": sorted({str(_publication_font(language)[0]) for language in
                                    {manifest["source_language"], manifest["target_language"], "zh-Hans"}}),
              "outputs": manifest_outputs,
              "reflow": reading_reflow_diagnostics([(source, target) for source, target, _ in values],
                                                   manifest["source_language"], manifest["target_language"])}
    atomic_write_json(out / "render_manifest.json", report)
    return report


def _safe_edition_filename(value: str, *, fallback: str = "Book") -> str:
    cleaned = re.sub(r'[\x00-\x1f<>:"/\\|?*]+', " ", str(value))
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    if not cleaned:
        cleaned = fallback
    if cleaned.upper() in {
        "CON", "PRN", "AUX", "NUL",
        *(f"COM{index}" for index in range(1, 10)),
        *(f"LPT{index}" for index in range(1, 10)),
    }:
        cleaned = f"_{cleaned}"
    return cleaned[:120].rstrip(" .") or fallback


def _language_filename_suffix(language: str) -> str:
    normalized = str(language or "und").strip().replace("_", "-").lower()
    return (normalized.split("-", 1)[0] or "und")


def _translated_book_title(workspace: Path, manifest: dict[str, Any]) -> str | None:
    explicit = str(manifest.get("translated_book_title") or "").strip()
    if explicit:
        return explicit
    units_path = workspace / "data/translation_units.jsonl"
    if not units_path.is_file():
        return None
    source_title = str(
        manifest.get("book_title")
        or Path(str(manifest.get("source_filename") or "Book")).stem
    ).strip()
    normalized_title = re.sub(r"\W+", "", source_title, flags=re.UNICODE).casefold()
    cache_dir = workspace / "cache" / str(manifest.get("language_pair") or "")
    candidates: list[tuple[int, int, dict[str, Any]]] = []
    for position, unit in enumerate(_jsonl(units_path)):
        source_text = str(unit.get("source_text") or "").strip()
        normalized_source = re.sub(r"\W+", "", source_text, flags=re.UNICODE).casefold()
        element_type = str(unit.get("element_type") or "").casefold()
        if normalized_title and normalized_source == normalized_title:
            rank = 0
        elif any(token in element_type for token in ("title", "heading")) and len(source_text) <= 240:
            rank = 1
        else:
            continue
        candidates.append((rank, position, unit))
    for _, _, unit in sorted(candidates, key=lambda item: (item[0], item[1])):
        cache_path = cache_dir / f"{unit['translation_unit_id']}.json"
        if not cache_path.is_file():
            continue
        translated = str(json.loads(cache_path.read_text("utf-8")).get("translated_text") or "").strip()
        translated = re.sub(r"\[\[BOOKFLOW_[A-Z_]+:[^\]]+\]\]", "", translated).strip()
        if translated:
            return translated.splitlines()[0].strip()
    return None


def edition_output_stem(workspace: Path, role: str) -> str:
    manifest = _load(workspace.resolve())
    source_stem = Path(str(manifest.get("source_filename") or "Book")).stem or "Book"
    source_code = _language_filename_suffix(str(manifest.get("source_language") or "und"))
    target_code = _language_filename_suffix(str(manifest.get("target_language") or "und"))
    if role == "source":
        return _safe_edition_filename(f"{source_stem}_{source_code}")
    target_title = _translated_book_title(workspace.resolve(), manifest) or source_stem
    if role == "target":
        return _safe_edition_filename(f"{target_title}_{target_code}")
    if role == "bilingual":
        bilingual_labels = {
            "zh": "双语版",
            "en": "bilingual",
            "fr": "bilingue",
            "de": "zweisprachig",
            "ja": "バイリンガル",
            "es": "bilingüe",
        }
        label = bilingual_labels.get(target_code, "bilingual")
        return _safe_edition_filename(f"{target_title} {label}")
    raise ValueError(f"invalid output role {role!r}")


def render_workspace(workspace: Path, role: str, formats: tuple[str, ...] = ("md", "docx", "pdf"),
                     *, layout_mode: str | None = None, bilingual_layout: str | None = None,
                     pdf_renderer: str = "native_pdf", renderer_config: Path | None = None) -> dict[str, Any]:
    workspace = workspace.resolve(); manifest, values = _content(workspace, role)
    layout_mode = layout_mode or manifest.get("layout_mode", "text")
    bilingual_layout = bilingual_layout or manifest.get("bilingual_layout", "stacked")
    if layout_mode not in LAYOUT_MODES: raise ValueError(f"invalid layout mode {layout_mode!r}")
    if bilingual_layout not in BILINGUAL_LAYOUTS: raise ValueError(f"invalid bilingual layout {bilingual_layout!r}")
    publication_path = workspace / "data/publication_reconstruction.json"
    if layout_mode == "publication" and not publication_path.is_file():
        raise RuntimeError("publication mode requires completed structure analysis")
    out = Path(manifest["output_directory"]) / role; out.mkdir(parents=True, exist_ok=True)
    stem = edition_output_stem(workspace, role)
    if layout_mode == "publication" and (workspace / "data/page_layout_elements.jsonl").is_file():
        return _render_anchored_publication(workspace, role, formats, pdf_renderer, renderer_config,
                                            manifest, out, stem, values)
    blocks: list[tuple[str, str, int]] = []
    preserve = layout_mode == "structure"
    classification_records = _jsonl(workspace / "data/page_classification.jsonl") if publication_path.is_file() else []
    cleanup = {"source_signatures": [], "target_signatures": [], "source_artifact_lines_removed": 0,
               "source_page_markers_removed": 0, "target_artifact_lines_removed": 0,
               "target_page_markers_removed": 0}
    render_values = values
    if layout_mode == "publication":
        render_values, cleanup = _publication_clean_values(values, classification_records)
    normalized_values = [(reflow_text(source, manifest["source_language"], preserve_structure=preserve),
                           reflow_text(target, manifest["target_language"], preserve_structure=preserve), page)
                         for source, target, page in render_values]
    for source, target, page in normalized_values:
        blocks.extend([(source, manifest["source_language"], page)] if role == "source" else
                      [(target, manifest["target_language"], page)] if role == "target" else
                      [(source, manifest["source_language"], page), (target, manifest["target_language"], page)])
    plain_blocks = [text for text, _, _ in blocks]
    visual_assets = _publication_visual_assets(workspace, manifest) if layout_mode == "publication" else {}
    heading_pages = {int(item["physical_page"]) for item in classification_records if item.get("chapter_boundary")}

    def visual_marker(index: int, page: int) -> str:
        if page not in visual_assets or (index + 1 < len(normalized_values) and normalized_values[index + 1][2] == page):
            return ""
        return "".join(
            f"\n\n![Cropped source visual page {page} object {object_index}]({path.as_posix()})"
            for object_index, path in enumerate(visual_assets[page], start=1)
        )
    outputs: dict[str, Any] = {}
    if "md" in formats:
        path = out / f"{stem}.md"
        if role == "bilingual":
            md_blocks = [f"### Source\n\n{source}\n\n### Target\n\n{target}{visual_marker(index, page)}"
                         for index, (source, target, page) in enumerate(normalized_values)]
        else:
            md_blocks = [(source if role == "source" else target) + visual_marker(index, page)
                         for index, (source, target, page) in enumerate(normalized_values)]
        atomic_write_text(path, "\n\n".join(md_blocks) + "\n"); outputs["md"] = str(path)
    docx_path: Path | None = None
    if "docx" in formats or ("pdf" in formats and pdf_renderer == "office"):
        path = out / f"{stem}.docx"; doc = Document()
        title = manifest.get("book_title") or Path(manifest.get("source_filename", "Book")).stem or "Book"
        doc.core_properties.title = title
        doc.core_properties.author = manifest.get("author", "")
        doc.core_properties.subject = f"{role} edition; {manifest['source_language']} to {manifest['target_language']}"
        doc.add_heading(title, 0)
        section = doc.sections[0]; section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)
        if role == "bilingual" and bilingual_layout == "parallel-columns":
            for index, (source, target, page) in enumerate(normalized_values):
                table = doc.add_table(rows=1, cols=2); table.autofit = False
                usable = section.page_width - section.left_margin - section.right_margin
                table.columns[0].width = usable // 2; table.columns[1].width = usable // 2
                table.cell(0, 0).text = source; table.cell(0, 1).text = target
                doc.add_paragraph()
                if visual_marker(index, page):
                    for visual_path in visual_assets[page]:
                        doc.add_picture(str(visual_path), width=Inches(6.4))
        elif role == "bilingual":
            for index, (source, target, page) in enumerate(normalized_values):
                source_paragraphs = source.split("\n\n"); target_paragraphs = target.split("\n\n")
                source_heading = next((position for position, paragraph in enumerate(source_paragraphs)
                                       if page in heading_pages and _is_chapter_heading(paragraph)), None)
                for position, paragraph in enumerate(source_paragraphs):
                    _add_structured_paragraph(doc, paragraph, heading=position == source_heading)
                target_heading = next((position for position, paragraph in enumerate(target_paragraphs)
                                       if page in heading_pages and _is_chapter_heading(paragraph)), None)
                if target_heading is None and source_heading is not None:
                    target_heading = 0
                for position, paragraph in enumerate(target_paragraphs):
                    _add_structured_paragraph(doc, paragraph, heading=position == target_heading)
                if visual_marker(index, page):
                    for visual_path in visual_assets[page]:
                        doc.add_picture(str(visual_path), width=Inches(6.4))
        else:
            for index, (source, target, page) in enumerate(normalized_values):
                source_paragraphs = source.split("\n\n"); rendered_paragraphs = (source if role == "source" else target).split("\n\n")
                source_heading = next((position for position, paragraph in enumerate(source_paragraphs)
                                       if page in heading_pages and _is_chapter_heading(paragraph)), None)
                rendered_heading = source_heading
                if role != "source":
                    rendered_heading = next((position for position, paragraph in enumerate(rendered_paragraphs)
                                             if page in heading_pages and _is_chapter_heading(paragraph)), None)
                    if rendered_heading is None and source_heading is not None:
                        rendered_heading = 0
                for position, paragraph in enumerate(rendered_paragraphs):
                    _add_structured_paragraph(doc, paragraph, heading=position == rendered_heading)
                if visual_marker(index, page):
                    for visual_path in visual_assets[page]:
                        doc.add_picture(str(visual_path), width=Inches(6.4))
        temporary_docx = path.with_suffix(".docx.tmp")
        doc.save(temporary_docx)
        docx_path = _commit_generated_file(temporary_docx, path)
        if "docx" in formats: outputs["docx"] = str(docx_path)
    if "pdf" in formats:
        path = out / f"{stem}.pdf"
        if pdf_renderer == "office":
            from .renderer_backend import convert_docx_to_pdf, load_renderer_config
            if docx_path is None:
                raise RuntimeError("office PDF rendering requires a generated DOCX")
            conversion = convert_docx_to_pdf(load_renderer_config(renderer_config), docx_path, path)
            outputs["pdf"] = conversion if conversion["status"] != "generated" else conversion["path"]
        elif pdf_renderer == "native_pdf":
            pdf = fitz.open(); page = pdf.new_page(); y = 50
            inserted_visual_pages: set[int] = set()
            for block_index, (text, block_language, block_page) in enumerate(blocks):
                font = "japan" if block_language == "ja" else "china-s" if block_language == "zh-Hans" else "helv"
                width = 42 if block_language in {"zh-Hans", "ja"} else 88
                visual_lines: list[str] = []
                for raw_line in text.replace("\r", "").split("\n"):
                    visual_lines.extend(raw_line[i:i + width] for i in range(0, len(raw_line), width))
                    if not raw_line: visual_lines.append("")
                for line in visual_lines or [""]:
                    if y > 790: page = pdf.new_page(); y = 50
                    if line: page.insert_text((45, y), line, fontsize=9, fontname=font)
                    y += 13
                y += 6
                last_block_for_page = block_index + 1 == len(blocks) or blocks[block_index + 1][2] != block_page
                if last_block_for_page and block_page in visual_assets and block_page not in inserted_visual_pages:
                    for visual_path in visual_assets[block_page]:
                        image = fitz.Pixmap(str(visual_path))
                        available_width = page.rect.width - 90
                        maximum_height = page.rect.height - 100
                        scale = min(available_width / image.width, maximum_height / image.height)
                        image_width = image.width * scale
                        image_height = image.height * scale
                        if y + image_height > page.rect.height - 45:
                            page = pdf.new_page(); y = 50
                        image_rect = fitz.Rect(45, y, 45 + image_width, y + image_height)
                        page.insert_image(image_rect, filename=str(visual_path), keep_proportion=True)
                        y = image_rect.y1 + 18
                    inserted_visual_pages.add(block_page)
            temporary_pdf = path.with_suffix(".pdf.tmp"); pdf.save(temporary_pdf); pdf.close()
            outputs["pdf"] = str(_commit_generated_file(temporary_pdf, path))
        else:
            raise ValueError(f"unknown PDF renderer {pdf_renderer!r}")
    manifest_outputs = {}
    for key, value in outputs.items():
        manifest_outputs[key] = value if isinstance(value, dict) else {
            "status": "generated", "path": value, "sha256": sha256_file(value)}
    report = {"role": role, "source_language": manifest["source_language"], "target_language": manifest["target_language"],
              "language_pair": manifest["language_pair"], "layout_mode": layout_mode,
              "bilingual_layout": bilingual_layout, "page_body_width_inches": 6.67,
              "publication_reconstruction": str(publication_path) if publication_path.is_file() else "not_generated",
              "publication_status": ("generated_with_review_pending" if layout_mode == "publication" and manifest.get("review_pending")
                                     else "generated" if layout_mode == "publication" else "not_requested"),
              "pdf_renderer": pdf_renderer,
               "visual_object_count": sum(len(paths) for paths in visual_assets.values()),
               "full_page_facsimile_count": 0,
               "visual_objects": [{"source_page": page, "path": str(path)}
                                  for page, paths in visual_assets.items() for path in paths],
               "publication_cleanup": cleanup,
               "reflow": reading_reflow_diagnostics([(source, target) for source, target, _ in values],
                                                     manifest["source_language"], manifest["target_language"]),
              "outputs": manifest_outputs}
    atomic_write_json(out / "render_manifest.json", report); return report


def validate_workspace(workspace: Path) -> dict[str, Any]:
    workspace = workspace.resolve(); manifest = _load(workspace); plan = plan_workspace(workspace)
    errors = []
    if manifest["source_language"] not in SUPPORTED_LANGUAGES or manifest["target_language"] not in SUPPORTED_LANGUAGES: errors.append("unsupported language")
    if manifest["source_language"] == manifest["target_language"]: errors.append("source and target languages must differ")
    if plan["pending"]: errors.append(f"{plan['pending']} translations pending")
    requested_roles = OUTPUT_ROLES if manifest.get("output_role", "all") == "all" else (manifest["output_role"],)
    for role in requested_roles:
        r = Path(manifest["output_directory"]) / role / "render_manifest.json"
        if not r.is_file(): errors.append(f"missing {role} render manifest")
        else:
            render = json.loads(r.read_text("utf-8"))
            if render.get("layout_mode") == "text" and render.get("reflow", {}).get("suspected_line_end_hyphens_after"):
                errors.append(f"{role} contains suspected line-end hyphenation")
            docx_path = render.get("outputs", {}).get("docx", {}).get("path")
            if docx_path and Path(docx_path).is_file():
                import zipfile
                with zipfile.ZipFile(docx_path) as archive:
                    xml = archive.read("word/document.xml")
                if render.get("layout_mode") == "text" and b"<w:br" in xml:
                    errors.append(f"{role} contains hard DOCX line breaks")
    report = {"valid": not errors, "errors": errors, "workspace_id": manifest["workspace_id"], "language_pair": manifest["language_pair"]}
    atomic_write_json(workspace / "data/validation_report.json", report); return report


def build_workspace(workspace: Path, formats: tuple[str, ...] = ("md", "docx", "pdf"),
                    *, layout_mode: str | None = None, bilingual_layout: str | None = None,
                    pdf_renderer: str = "native_pdf", renderer_config: Path | None = None) -> dict[str, Any]:
    workspace = workspace.resolve(); manifest = _load(workspace)
    roles = OUTPUT_ROLES if manifest.get("output_role", "all") == "all" else (manifest["output_role"],)
    rendered = {role: render_workspace(workspace, role, formats,
        layout_mode=layout_mode, bilingual_layout=bilingual_layout, pdf_renderer=pdf_renderer,
        renderer_config=renderer_config) for role in roles}
    validation = validate_workspace(workspace)
    if not validation["valid"]: raise RuntimeError("workspace validation failed: " + "; ".join(validation["errors"]))
    manifest = _load(workspace); manifest["stage"] = "built"; _save(workspace, manifest)
    blocked_renderer = any(item["outputs"].get("pdf", {}).get("status") == "blocked_by_renderer" for item in rendered.values())
    result = {"workspace_id": manifest["workspace_id"], "language_pair": manifest["language_pair"], "roles": rendered,
              "status": "blocked_by_renderer" if blocked_renderer else "generated",
              "validation": validation, "built_at": _now()}
    atomic_write_json(workspace / "output/build_manifest.json", result); return result
