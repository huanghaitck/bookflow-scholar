"""Fail-closed Phase 3A selection, preflight, cache, and translation flow."""

from __future__ import annotations

import json
import math
import os
import re
import time
from contextlib import contextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Iterator

import yaml

from .io_utils import (
    atomic_write_json,
    atomic_write_text,
    load_json,
    sha256_text,
    stable_hash,
)
from .paths import ProjectSettings, project_root, resolve_project_path
from .phase2b2_schemas import AutomatedLogicalBlock
from .secret_store import load_translation_api_key, translation_api_key_status
from .translation_provider import DeepSeekOpenAICompatibleProvider
from .translation_schemas import (
    NormalizedTranslationResult,
    TranslationBatchResult,
    TranslationCandidate,
    TranslationModelPayload,
    TranslationPreflightReport,
    TranslationRequestPayload,
)


BODY_TYPE = "body"
TRANSLATABLE_TYPES = {
    "book_title",
    "subtitle",
    "chapter_title",
    "section_title",
    "subsection_title",
    "body",
    "footnote",
    "caption",
    "table_title",
    "table_cell",
    "epigraph",
    "other_translatable",
}
STRUCTURAL_TITLE_TYPES = {
    "book_title",
    "subtitle",
    "chapter_title",
    "section_title",
    "subsection_title",
    "table_title",
}
SELECTION_ORDER = [
    "structural_title",
    "ordinary_single_page",
    "cross_page",
    "rhetorical_long_form",
    "proper_names_or_historical_voice",
]
HISTORICAL_MARKERS = (
    "jap ",
    "chinaman",
    "manchu",
    "sikh",
    "annamese",
    "tonkinese",
    "bluejacket",
    "native",
    "coolie",
)


def _jsonl_models(path: Path, model_type):
    return [
        model_type.model_validate(json.loads(line))
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


def _logical_path(settings: ProjectSettings, root: Path) -> Path:
    name = Path(settings.sample_pdf).stem + ".logical_blocks.jsonl"
    path = resolve_project_path(settings.automated_logical_directory, root=root) / name
    if not path.is_file():
        raise FileNotFoundError(f"Automated logical blocks not found: {path}")
    return path


def _load_logical(settings: ProjectSettings, root: Path) -> list[AutomatedLogicalBlock]:
    return _jsonl_models(_logical_path(settings, root), AutomatedLogicalBlock)


def _context_neighbors(
    blocks: list[AutomatedLogicalBlock], target: AutomatedLogicalBlock
) -> tuple[AutomatedLogicalBlock | None, AutomatedLogicalBlock | None]:
    body = [item for item in blocks if item.block_type == BODY_TYPE]
    index = next(index for index, item in enumerate(body) if item.logical_block_id == target.logical_block_id)
    before = body[index - 1] if index else None
    after = body[index + 1] if index + 1 < len(body) else None
    before = before if before and before.chapter_id == target.chapter_id and before.translation_ready else None
    after = after if after and after.chapter_id == target.chapter_id and after.translation_ready else None
    return before, after


def _title_context(
    blocks: list[AutomatedLogicalBlock],
    target: AutomatedLogicalBlock,
    title_id: str | None,
) -> tuple[str | None, str | None]:
    if not title_id or title_id == target.logical_block_id:
        return None, None
    title = next(
        (
            item
            for item in blocks
            if item.logical_block_id == title_id
            and item.block_type in STRUCTURAL_TITLE_TYPES
        ),
        None,
    )
    if title is None:
        return None, None
    return title.logical_block_id, title.source_text


def build_translation_units(
    settings: ProjectSettings, *, root: Path | None = None
) -> list[TranslationRequestPayload]:
    """Build one target-only unit for every complete translatable logical block."""

    root = (root or project_root()).resolve()
    blocks = _load_logical(settings, root)
    return [
        build_translation_request(item.logical_block_id, settings, root=root)
        for item in blocks
        if item.block_type in TRANSLATABLE_TYPES
        and item.translation_ready
        and not item.unresolved_boundaries
        and bool(item.source_text.strip())
    ]


def _proper_name_score(text: str) -> int:
    tokens = re.findall(r"\b[A-Z][a-z]+(?:-[A-Z][a-z]+)?\b", text)
    return len(tokens)


def _rhetorical_score(text: str) -> int:
    return len(text) + 120 * sum(text.count(mark) for mark in ("!", ";", ":", "—"))


def _historical_score(text: str) -> int:
    lowered = text.lower()
    return sum(lowered.count(marker) for marker in HISTORICAL_MARKERS)


def select_representative_blocks(
    settings: ProjectSettings, *, root: Path | None = None
) -> list[TranslationCandidate]:
    """Select one title plus four non-trivial body blocks for five API calls."""

    root = (root or project_root()).resolve()
    blocks = _load_logical(settings, root)
    eligible_body = [
        item
        for item in blocks
        if item.block_type == BODY_TYPE
        and item.translation_ready
        and not item.unresolved_boundaries
        and bool(item.source_text.strip())
    ]
    eligible_titles = [
        item
        for item in blocks
        if item.block_type in STRUCTURAL_TITLE_TYPES
        and item.translation_ready
        and not item.unresolved_boundaries
        and bool(item.source_text.strip())
    ]
    if len(eligible_body) < 4 or not eligible_titles:
        raise RuntimeError("A complete title and four complete body blocks are required")
    used: set[str] = set()

    def take(kind: str, pool: list[AutomatedLogicalBlock], key, reason: str, reverse: bool = True):
        available = [item for item in pool if item.logical_block_id not in used]
        if not available:
            raise RuntimeError(f"No candidate is available for {kind}")
        selected = sorted(available, key=lambda item: (key(item), item.logical_block_id), reverse=reverse)[0]
        used.add(selected.logical_block_id)
        before, after = (
            _context_neighbors(blocks, selected)
            if selected.block_type == BODY_TYPE
            else (None, None)
        )
        return TranslationCandidate(
            logical_block_id=selected.logical_block_id,
            block_type=selected.block_type,
            source_pages=selected.source_pages,
            cross_page=selected.cross_page,
            source_text_character_count=len(selected.source_text),
            selection_type=kind,
            selection_reason=reason,
            context_before_block_id=before.logical_block_id if before else None,
            context_after_block_id=after.logical_block_id if after else None,
            translation_ready=True,
            unresolved_boundaries=[],
        )

    title = take(
        "structural_title",
        eligible_titles,
        lambda item: (
            item.block_type == "section_title",
            len(item.source_text),
        ),
        "正式章节或节标题，单独验证结构标题能够被翻译且不与运行页眉混淆。",
    )
    special = take(
        "proper_names_or_historical_voice",
        eligible_body,
        lambda item: (
            _historical_score(item.source_text),
            _proper_name_score(item.source_text),
            len(item.source_text),
        ),
        "含较多专名、旧式拼写、时代称谓或历史用语，用于结构化术语测试。",
    )
    rhetorical = take(
        "rhetorical_long_form",
        eligible_body,
        lambda item: _rhetorical_score(item.source_text),
        "长句、并列结构、感叹或修辞密集，用于测试历史游记语气。",
    )
    cross_pool = [item for item in eligible_body if item.cross_page]
    cross = take(
        "cross_page",
        cross_pool,
        lambda item: -len(item.source_text),
        "跨两页或多页且已完成自动逻辑重建。",
    )
    ordinary_pool = [
        item for item in eligible_body
        if not item.cross_page and 250 <= len(item.source_text) <= 800
    ]
    ordinary = take(
        "ordinary_single_page",
        ordinary_pool,
        lambda item: -abs(len(item.source_text) - 450),
        "普通单页完整正文段落，长度适中且没有未解决边界。",
    )
    by_type = {
        item.selection_type: item
        for item in [title, special, rhetorical, cross, ordinary]
    }
    return [by_type[kind] for kind in SELECTION_ORDER]


def build_translation_request(
    logical_block_id: str,
    settings: ProjectSettings,
    *,
    root: Path | None = None,
) -> TranslationRequestPayload:
    root = (root or project_root()).resolve()
    blocks = _load_logical(settings, root)
    target = next((item for item in blocks if item.logical_block_id == logical_block_id), None)
    if target is None:
        raise KeyError(f"Unknown logical block: {logical_block_id}")
    if (
        target.block_type not in TRANSLATABLE_TYPES
        or not target.translation_ready
        or target.unresolved_boundaries
    ):
        raise ValueError(
            "Only translation_ready translatable blocks without unresolved boundaries may be translated"
        )
    before, after = (
        _context_neighbors(blocks, target)
        if target.block_type == BODY_TYPE
        else (None, None)
    )
    chapter_title_block_id, chapter_title_source = _title_context(
        blocks, target, target.chapter_id
    )
    section_title_block_id, section_title_source = _title_context(
        blocks, target, target.section_id
    )
    profile_path = resolve_project_path(settings.translation_language_profile_path, root=root)
    profile = yaml.safe_load(profile_path.read_text(encoding="utf-8"))
    if not isinstance(profile, dict):
        raise ValueError("Translation language profile must be a YAML mapping")
    return TranslationRequestPayload(
        translation_unit_id=f"translation_{target.logical_block_id}",
        target_block_id=target.logical_block_id,
        block_type=target.block_type,
        source_text=target.source_text,
        chapter_id=target.chapter_id,
        section_id=target.section_id,
        chapter_title_block_id=chapter_title_block_id,
        section_title_block_id=section_title_block_id,
        chapter_title_context_source=chapter_title_source,
        chapter_title_context_translation=None,
        section_title_context_source=section_title_source,
        section_title_context_translation=None,
        context_before_block_ids=[before.logical_block_id] if before else [],
        context_after_block_ids=[after.logical_block_id] if after else [],
        context_before_text=before.source_text if before else None,
        context_after_text=after.source_text if after else None,
        source_pages=target.source_pages,
        source_language=settings.translation_source_language,
        target_language=settings.translation_target_language,
        translate_target_only=True,
        glossary=[],
        translation_profile=profile,
    )


def translation_preflight(
    settings: ProjectSettings,
    *,
    root: Path | None = None,
    key_status_resolver: Callable[[ProjectSettings, Path], tuple[bool, str | None]] = translation_api_key_status,
) -> TranslationPreflightReport:
    """Perform an offline-only selection, safety, and conservative cost preflight."""

    root = (root or project_root()).resolve()
    candidates = select_representative_blocks(settings, root=root)
    payloads = [build_translation_request(item.logical_block_id, settings, root=root) for item in candidates]
    prompt_path = resolve_project_path(settings.translation_prompt_path, root=root)
    profile_path = resolve_project_path(settings.translation_language_profile_path, root=root)
    prompt = prompt_path.read_text(encoding="utf-8")
    profile = profile_path.read_text(encoding="utf-8")
    source_chars = sum(len(item.source_text) for item in payloads)
    context_chars = sum(
        len(item.context_before_text or "")
        + len(item.context_after_text or "")
        + len(item.chapter_title_context_source or "")
        + len(item.chapter_title_context_translation or "")
        + len(item.section_title_context_source or "")
        + len(item.section_title_context_translation or "")
        for item in payloads
    )
    request_chars = sum(
        len(prompt) + len(json.dumps(item.model_dump(mode="json"), ensure_ascii=False))
        for item in payloads
    ) + len(profile)
    estimated_input_tokens = math.ceil(request_chars / 4)
    maximum_output_tokens_total = len(payloads) * settings.translation_max_output_tokens
    lower = (
        estimated_input_tokens
        * settings.translation_input_cache_hit_price_cny_per_million_tokens
        / 1_000_000
    )
    upper = (
        estimated_input_tokens
        * settings.translation_input_cache_miss_price_cny_per_million_tokens
        + maximum_output_tokens_total
        * settings.translation_output_price_cny_per_million_tokens
    ) / 1_000_000
    key_set, _ = key_status_resolver(settings, root)
    blockers: list[str] = []
    if len(candidates) != 5:
        blockers.append("Exactly five representative logical blocks are required")
    if not key_set:
        blockers.append(f"API key is not configured in {settings.translation_api_key_env}")
    if not settings.full_pdf_protection:
        blockers.append("Full PDF protection must remain enabled")
    if settings.automatic_phase_advance:
        blockers.append("Automatic phase advance must remain disabled")
    if settings.translation_automatic_retry:
        blockers.append("Automatic retry must remain disabled")
    if settings.translation_maximum_real_calls != 5:
        blockers.append("Maximum content calls must equal five")
    if upper > settings.translation_maximum_cash_cost_cny:
        blockers.append("Conservative cash-cost upper bound exceeds the Phase 3A hard limit")
    return TranslationPreflightReport(
        provider=settings.translation_provider,
        model=settings.translation_model,
        base_url=settings.translation_base_url,
        api_key_env=settings.translation_api_key_env,
        api_key_set=key_set,
        candidates=candidates,
        target_block_count=len(candidates),
        total_source_characters=source_chars,
        total_context_characters=context_chars,
        estimated_input_tokens=estimated_input_tokens,
        maximum_output_tokens_per_call=settings.translation_max_output_tokens,
        maximum_output_tokens_total=maximum_output_tokens_total,
        maximum_model_list_calls=settings.translation_maximum_model_list_calls,
        maximum_content_calls=settings.translation_maximum_real_calls,
        input_cache_miss_price_cny_per_million_tokens=settings.translation_input_cache_miss_price_cny_per_million_tokens,
        output_price_cny_per_million_tokens=settings.translation_output_price_cny_per_million_tokens,
        estimated_cost_lower_cny=round(lower, 6),
        estimated_cost_upper_cny=round(upper, 6),
        maximum_cash_cost_cny=settings.translation_maximum_cash_cost_cny,
        pricing_reference_url=settings.translation_pricing_reference_url,
        pricing_checked_date=settings.translation_pricing_checked_date,
        usage_fields_expected=["prompt_tokens", "completion_tokens", "total_tokens"],
        actual_charge_note="API usage does not prove the actual account deduction; verify in the DeepSeek platform.",
        blockers=blockers,
        ready_for_real_call=not blockers,
        api_called=False,
    )


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _safe_component(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("_.") or "value"


def _translation_paths(
    settings: ProjectSettings, root: Path, fingerprint: str
) -> dict[str, Path]:
    provider = _safe_component(settings.translation_provider)
    model = _safe_component(settings.translation_model)
    language = _safe_component(settings.translation_target_language)
    request_root = resolve_project_path(settings.translation_request_directory, root=root)
    raw_root = resolve_project_path(settings.translation_raw_directory, root=root)
    normalized_root = resolve_project_path(settings.translation_normalized_directory, root=root)
    cache_root = resolve_project_path(settings.translation_cache_directory, root=root)
    usage_root = resolve_project_path(settings.translation_usage_directory, root=root)
    return {
        "request": request_root / "records" / f"{fingerprint}.json",
        "ledger": request_root / "phase3a_call_ledger.json",
        "lock": request_root / "phase3a_call_ledger.lock",
        "model_check": request_root / "phase3a_model_check.json",
        "raw": raw_root / provider / model / f"{fingerprint}.json",
        "normalized": normalized_root / language / f"{fingerprint}.json",
        "failure": normalized_root / language / f"{fingerprint}.failure.json",
        "cache": cache_root / f"{fingerprint}.json",
        "usage": usage_root / f"{fingerprint}.json",
    }


def _load_call_ledger(path: Path, maximum: int) -> dict[str, Any]:
    if not path.is_file():
        return {
            "schema_version": "1.0",
            "phase": "3A",
            "maximum_content_calls": maximum,
            "content_calls_started": 0,
            "attempts": [],
        }
    ledger = load_json(path)
    if not isinstance(ledger, dict):
        raise RuntimeError("Phase 3A call ledger is invalid")
    return ledger


@contextmanager
def _exclusive_call_lock(path: Path) -> Iterator[None]:
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        descriptor = os.open(path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
    except FileExistsError as exc:
        raise RuntimeError("Another Phase 3A call may be in progress") from exc
    try:
        os.write(descriptor, f"pid={os.getpid()}\n".encode("ascii"))
        os.close(descriptor)
        yield
    finally:
        if path.exists():
            path.unlink()


def _request_fingerprint(
    payload: TranslationRequestPayload,
    settings: ProjectSettings,
    *,
    prompt_sha256: str,
    profile_sha256: str,
) -> str:
    context = {
        "before": payload.context_before_text,
        "after": payload.context_after_text,
        "chapter_source": payload.chapter_title_context_source,
        "chapter_translation": payload.chapter_title_context_translation,
        "section_source": payload.section_title_context_source,
        "section_translation": payload.section_title_context_translation,
    }
    return stable_hash(
        {
            "phase": "3A",
            "provider": settings.translation_provider,
            "model": settings.translation_model,
            "base_url": settings.translation_base_url.rstrip("/"),
            "thinking_mode": settings.translation_thinking_mode,
            "block_type": payload.block_type,
            "source_text_sha256": sha256_text(payload.source_text),
            "context_sha256": stable_hash(context),
            "prompt_version": settings.translation_prompt_version,
            "prompt_sha256": prompt_sha256,
            "language_profile_version": settings.translation_language_profile_version,
            "language_profile_sha256": profile_sha256,
            "source_language": settings.translation_source_language,
            "target_language": settings.translation_target_language,
            "response_format": {"type": "json_object"},
            "max_tokens": settings.translation_max_output_tokens,
            "temperature": settings.translation_temperature,
            "automatic_retry": False,
        }
    )


def _compact_whitespace(value: str) -> str:
    return " ".join(value.split())


def _local_integrity_warnings(
    source_text: str, translation: str
) -> list[str]:
    warnings: list[str] = []
    for number in sorted(set(re.findall(r"\b\d+(?:[.,]\d+)?\b", source_text))):
        if number not in translation:
            warnings.append(f"possible_missing_number:{number}")
    if re.search(r"\b(?:not|never|no|without|cannot|could not|would not)\b", source_text, re.I):
        if not any(marker in translation for marker in ("不", "未", "无", "没有", "非", "勿")):
            warnings.append("possible_missing_negation")
    return warnings


def _normalize_translation(
    *,
    content: str,
    payload: TranslationRequestPayload,
    settings: ProjectSettings,
    fingerprint: str,
    prompt_sha256: str,
    profile_sha256: str,
    raw_path: Path,
    usage: dict[str, Any] | None,
    request_id: str | None,
) -> NormalizedTranslationResult:
    if not content.strip():
        raise ValueError("Provider returned empty JSON content")
    parsed = json.loads(content.strip())
    if not isinstance(parsed, dict):
        raise ValueError("Translation response must be one JSON object")
    model_payload = TranslationModelPayload.model_validate(parsed)
    if model_payload.target_block_id != payload.target_block_id:
        raise ValueError("Returned target_block_id does not match the request")
    if model_payload.block_type != payload.block_type:
        raise ValueError("Returned block_type does not match the request")
    translation_compact = _compact_whitespace(model_payload.translation)
    source_compact = _compact_whitespace(payload.source_text)
    if source_compact and source_compact in translation_compact:
        raise ValueError("Translation repeats the complete English source_text")
    for label, context in (
        ("context_before_text", payload.context_before_text),
        ("context_after_text", payload.context_after_text),
        ("chapter_title_context_source", payload.chapter_title_context_source),
        ("chapter_title_context_translation", payload.chapter_title_context_translation),
        ("section_title_context_source", payload.section_title_context_source),
        ("section_title_context_translation", payload.section_title_context_translation),
    ):
        compact = _compact_whitespace(context or "")
        if compact and compact in translation_compact:
            raise ValueError(f"Translation includes complete {label}")
    extra_ids = set(re.findall(r"logical2_[A-Za-z0-9]+", model_payload.translation))
    if extra_ids - {payload.target_block_id}:
        raise ValueError("Translation contains an unrelated logical block ID")
    minimum_length = (
        1
        if payload.block_type in STRUCTURAL_TITLE_TYPES
        else max(12, int(len(payload.source_text) * 0.12))
    )
    if len(model_payload.translation.strip()) < minimum_length:
        raise ValueError("Translation is too short and may be an obvious summary")
    warnings = list(model_payload.warnings)
    warnings.extend(_local_integrity_warnings(payload.source_text, model_payload.translation))
    context_sha = stable_hash(
        {
            "before": payload.context_before_text,
            "after": payload.context_after_text,
            "chapter_source": payload.chapter_title_context_source,
            "chapter_translation": payload.chapter_title_context_translation,
            "section_source": payload.section_title_context_source,
            "section_translation": payload.section_title_context_translation,
        }
    )
    return NormalizedTranslationResult(
        schema_version=settings.translation_schema_version,
        translation_unit_id=f"translation_{payload.target_block_id}",
        target_block_id=payload.target_block_id,
        block_type=payload.block_type,
        source_text_sha256=sha256_text(payload.source_text),
        context_sha256=context_sha,
        prompt_version=settings.translation_prompt_version,
        prompt_sha256=prompt_sha256,
        language_profile_version=settings.translation_language_profile_version,
        language_profile_sha256=profile_sha256,
        provider=settings.translation_provider,
        model=settings.translation_model,
        thinking_mode="disabled",
        target_language=settings.translation_target_language,
        request_fingerprint=fingerprint,
        raw_response_path=str(raw_path.resolve()),
        translation=model_payload.translation,
        uncertain_terms=model_payload.uncertain_terms,
        historical_terms=model_payload.historical_terms,
        warnings=warnings,
        usage=usage,
        request_id=request_id,
        api_called=True,
        cache_hit=False,
        status="translated",
        created_at=_now(),
    )


def _write_diagnostic_docx(path: Path, derived: dict[str, Any]) -> None:
    """Create a non-final Word diagnostic from the same normalized data as Markdown."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Phase 3A translation diagnostic | NOT FINAL"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string("666666")
    footer = section.footer.paragraphs[0]
    footer.text = "Generated from phase3a_translation_sample.json"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("777777")

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("Phase 3A 翻译诊断样本（非 final）")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor.from_string("1F4D78")

    for item in derived["translations"]:
        pages = item["source_pages"]
        meta = document.add_paragraph()
        meta.paragraph_format.space_before = Pt(8)
        meta.paragraph_format.space_after = Pt(3)
        meta_run = meta.add_run(
            f"[原书页码范围: {pages[0]}-{pages[-1]} | {item['block_type']}]"
        )
        meta_run.italic = True
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor.from_string("666666")
        if item["block_type"] == "chapter_title":
            english_style = "Heading 1"
        elif item["block_type"] in {"section_title", "subsection_title"}:
            english_style = "Heading 2"
        else:
            english_style = None
        english = document.add_paragraph(style=english_style)
        english.add_run(item["source_text"])
        chinese = document.add_paragraph()
        chinese_run = chinese.add_run(item["translation"])
        chinese_fonts = (
            chinese_run._element.get_or_add_rPr().get_or_add_rFonts()
        )
        chinese_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        chinese.paragraph_format.space_after = Pt(10)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_derived_outputs(
    settings: ProjectSettings,
    root: Path,
    candidates: list[TranslationCandidate],
    payloads: dict[str, TranslationRequestPayload],
    results: list[NormalizedTranslationResult],
    *,
    model_list_calls: int,
    content_calls: int,
) -> tuple[str, str, str]:
    master_path = resolve_project_path(settings.automated_master_path, root=root)
    master = load_json(master_path)
    by_id = {item.target_block_id: item for item in results}
    translations: list[dict[str, Any]] = []
    for candidate in candidates:
        result = by_id[candidate.logical_block_id]
        payload = payloads[candidate.logical_block_id]
        translations.append(
            {
                "target_block_id": result.target_block_id,
                "block_type": result.block_type,
                "source_pages": payload.source_pages,
                "cross_page": candidate.cross_page,
                "selection_type": candidate.selection_type,
                "selection_reason": candidate.selection_reason,
                "source_text": payload.source_text,
                "translation": result.translation,
                "translation_status": result.status,
                "uncertain_terms": [item.model_dump(mode="json") for item in result.uncertain_terms],
                "historical_terms": [item.model_dump(mode="json") for item in result.historical_terms],
                "warnings": result.warnings,
                "usage": result.usage,
                "provider": result.provider,
                "model": result.model,
                "thinking_mode": result.thinking_mode,
                "request_fingerprint": result.request_fingerprint,
            }
        )
    request_root = resolve_project_path(settings.translation_request_directory, root=root)
    ledger_path = request_root / "phase3a_call_ledger.json"
    model_check_path = request_root / "phase3a_model_check.json"
    ledger = load_json(ledger_path) if ledger_path.is_file() else {}
    model_check = load_json(model_check_path) if model_check_path.is_file() else {}
    cumulative_content_calls = (
        int(ledger.get("content_calls_started", 0)) if isinstance(ledger, dict) else 0
    )
    cumulative_model_list_calls = (
        int(model_check.get("model_list_call_count", 0))
        if isinstance(model_check, dict)
        else 0
    )
    derived = {
        "schema_version": "1.1",
        "phase": "3A",
        "document_id": master.get("document_id"),
        "source_pdf_sha256": master.get("source_pdf_sha256"),
        "source_document_path": str(master_path.resolve()),
        "source_document_modified": False,
        "model_list_calls": cumulative_model_list_calls,
        "content_translation_calls": cumulative_content_calls,
        "current_run_model_list_calls": model_list_calls,
        "current_run_content_translation_calls": content_calls,
        "automatic_retries": 0,
        "translations": translations,
        "final_document": False,
        "created_at": _now().isoformat(),
    }
    derived_path = resolve_project_path(settings.translation_derived_document_path, root=root)
    atomic_write_json(derived_path, derived)
    lines = ["# Phase 3A 翻译诊断样本（非final）", ""]
    for item in translations:
        pages = item["source_pages"]
        lines.extend(
            [
                f"[原书页码范围: {pages[0]}-{pages[-1]}]",
                "",
                item["source_text"],
                "",
                item["translation"],
                "",
            ]
        )
    markdown_path = resolve_project_path(settings.translation_diagnostic_markdown_path, root=root)
    atomic_write_text(markdown_path, "\n".join(lines).rstrip() + "\n")
    docx_path = resolve_project_path(settings.translation_diagnostic_docx_path, root=root)
    _write_diagnostic_docx(docx_path, derived)
    return (
        str(derived_path.resolve()),
        str(markdown_path.resolve()),
        str(docx_path.resolve()),
    )


def run_translation_sample(
    settings: ProjectSettings,
    *,
    root: Path | None = None,
    allow_api: bool = False,
    confirm_five_calls: bool = False,
    key_status_resolver: Callable[[ProjectSettings, Path], tuple[bool, str | None]] = translation_api_key_status,
    key_loader: Callable[[ProjectSettings, Path], tuple[str, str]] = load_translation_api_key,
    provider_factory: Callable[..., Any] = DeepSeekOpenAICompatibleProvider,
) -> TranslationBatchResult:
    """Run zero or at most five translation calls; successful fingerprints are cached."""

    root = (root or project_root()).resolve()
    preflight = translation_preflight(
        settings, root=root, key_status_resolver=key_status_resolver
    )
    selected_ids = [item.logical_block_id for item in preflight.candidates]
    if not allow_api:
        return TranslationBatchResult(
            selected_block_ids=selected_ids,
            results=[],
            api_calls=0,
            cache_hits=0,
            failed=0,
        )
    if not confirm_five_calls:
        raise PermissionError("--confirm-five-calls is required for Phase 3A")
    if not preflight.ready_for_real_call:
        raise RuntimeError("Preflight failed: " + "; ".join(preflight.blockers))

    prompt_path = resolve_project_path(settings.translation_prompt_path, root=root)
    profile_path = resolve_project_path(settings.translation_language_profile_path, root=root)
    prompt = prompt_path.read_text(encoding="utf-8")
    prompt_sha = sha256_text(prompt)
    profile_sha = sha256_text(profile_path.read_text(encoding="utf-8"))
    payloads = {
        item.logical_block_id: build_translation_request(item.logical_block_id, settings, root=root)
        for item in preflight.candidates
    }
    fingerprints = {
        block_id: _request_fingerprint(
            payload,
            settings,
            prompt_sha256=prompt_sha,
            profile_sha256=profile_sha,
        )
        for block_id, payload in payloads.items()
    }
    first_paths = _translation_paths(settings, root, next(iter(fingerprints.values())))
    model_check_path = first_paths["model_check"]
    successful_cache: dict[str, NormalizedTranslationResult] = {}
    for block_id, fingerprint in fingerprints.items():
        paths = _translation_paths(settings, root, fingerprint)
        if not paths["cache"].is_file():
            continue
        cache = load_json(paths["cache"])
        if (
            isinstance(cache, dict)
            and cache.get("request_fingerprint") == fingerprint
            and cache.get("status") == "translated"
            and paths["raw"].is_file()
            and paths["normalized"].is_file()
            and paths["usage"].is_file()
        ):
            stored = NormalizedTranslationResult.model_validate(load_json(paths["normalized"]))
            successful_cache[block_id] = stored.model_copy(
                update={"api_called": False, "cache_hit": True}
            )
    external_work_needed = len(successful_cache) != len(selected_ids)
    model_check = load_json(model_check_path) if model_check_path.is_file() else None
    model_check_valid = bool(
        isinstance(model_check, dict)
        and model_check.get("provider") == settings.translation_provider
        and model_check.get("model") == settings.translation_model
        and model_check.get("base_url") == settings.translation_base_url
        and model_check.get("checked") is True
    )
    model_list_calls = 0
    model_available: bool | None = None
    provider = None
    if model_check_valid:
        model_available = bool(model_check.get("model_available"))
    if external_work_needed and not model_check_valid:
        api_key, _ = key_loader(settings, root)
        provider = provider_factory(
            api_key=api_key,
            base_url=settings.translation_base_url,
            timeout_seconds=settings.translation_request_timeout_seconds,
        )
        model_ids = provider.list_model_ids()
        model_list_calls = 1
        model_available = settings.translation_model in model_ids
        atomic_write_json(
            model_check_path,
            {
                "schema_version": "1.0",
                "provider": settings.translation_provider,
                "model": settings.translation_model,
                "base_url": settings.translation_base_url,
                "checked": True,
                "model_available": model_available,
                "model_list_call_count": 1,
                "unrelated_model_ids_recorded": False,
                "checked_at": _now().isoformat(),
            },
        )
    if model_available is False:
        return TranslationBatchResult(
            selected_block_ids=selected_ids,
            results=list(successful_cache.values()),
            api_calls=0,
            cache_hits=len(successful_cache),
            failed=0,
            model_list_calls=model_list_calls,
            model_available=False,
        )
    if not external_work_needed:
        results = [successful_cache[block_id] for block_id in selected_ids]
        derived_path, markdown_path, docx_path = _write_derived_outputs(
            settings,
            root,
            preflight.candidates,
            payloads,
            results,
            model_list_calls=0,
            content_calls=0,
        )
        return TranslationBatchResult(
            selected_block_ids=selected_ids,
            results=results,
            api_calls=0,
            cache_hits=5,
            failed=0,
            model_list_calls=0,
            model_available=True,
            derived_document_path=derived_path,
            diagnostic_markdown_path=markdown_path,
            diagnostic_docx_path=docx_path,
            strict_export_ready=True,
        )
    if provider is None:
        api_key, _ = key_loader(settings, root)
        provider = provider_factory(
            api_key=api_key,
            base_url=settings.translation_base_url,
            timeout_seconds=settings.translation_request_timeout_seconds,
        )

    results_by_id = dict(successful_cache)
    api_calls = 0
    failed = 0
    with _exclusive_call_lock(first_paths["lock"]):
        ledger = _load_call_ledger(first_paths["ledger"], settings.translation_maximum_real_calls)
        for candidate in preflight.candidates:
            block_id = candidate.logical_block_id
            if block_id in results_by_id:
                continue
            if int(ledger.get("content_calls_started", 0)) >= settings.translation_maximum_real_calls:
                raise RuntimeError("Persistent Phase 3A content-call limit is exhausted")
            payload = payloads[block_id]
            fingerprint = fingerprints[block_id]
            paths = _translation_paths(settings, root, fingerprint)
            started_at = _now()
            request_record = {
                "schema_version": "1.0",
                "phase": "3A",
                "status": "in_flight",
                "target_block_id": block_id,
                "translation_unit_id": payload.translation_unit_id,
                "block_type": payload.block_type,
                "source_pages": payload.source_pages,
                "request_fingerprint": fingerprint,
                "provider": settings.translation_provider,
                "model": settings.translation_model,
                "base_url": settings.translation_base_url,
                "api_key_env": settings.translation_api_key_env,
                "api_key_recorded": False,
                "source_text_sha256": sha256_text(payload.source_text),
                "context_sha256": stable_hash(
                    {
                        "before": payload.context_before_text,
                        "after": payload.context_after_text,
                        "chapter_source": payload.chapter_title_context_source,
                        "chapter_translation": payload.chapter_title_context_translation,
                        "section_source": payload.section_title_context_source,
                        "section_translation": payload.section_title_context_translation,
                    }
                ),
                "prompt_version": settings.translation_prompt_version,
                "prompt_sha256": prompt_sha,
                "language_profile_version": settings.translation_language_profile_version,
                "language_profile_sha256": profile_sha,
                "thinking_mode": "disabled",
                "response_format": {"type": "json_object"},
                "automatic_retry": False,
                "retries": 0,
                "started_at": started_at.isoformat(),
            }
            atomic_write_json(paths["request"], request_record)
            ledger["content_calls_started"] = int(ledger.get("content_calls_started", 0)) + 1
            ledger.setdefault("attempts", []).append(
                {
                    "request_fingerprint": fingerprint,
                    "target_block_id": block_id,
                    "status": "in_flight",
                    "started_at": started_at.isoformat(),
                }
            )
            atomic_write_json(paths["ledger"], ledger)
            call_start = time.perf_counter()
            api_calls += 1
            try:
                response = provider.translate_one(
                    model=settings.translation_model,
                    system_prompt=prompt,
                    user_payload=payload.model_dump(mode="json"),
                    max_output_tokens=settings.translation_max_output_tokens,
                    temperature=settings.translation_temperature,
                    thinking_mode=settings.translation_thinking_mode,
                )
                atomic_write_json(
                    paths["raw"],
                    {
                        "record_type": "raw_translation_provider_response",
                        "request_fingerprint": fingerprint,
                        "target_block_id": block_id,
                        "block_type": payload.block_type,
                        "provider": settings.translation_provider,
                        "model": settings.translation_model,
                        "api_called": True,
                        "received_at": _now().isoformat(),
                        "response": response.raw_response,
                    },
                )
                normalized = _normalize_translation(
                    content=response.content,
                    payload=payload,
                    settings=settings,
                    fingerprint=fingerprint,
                    prompt_sha256=prompt_sha,
                    profile_sha256=profile_sha,
                    raw_path=paths["raw"],
                    usage=response.usage,
                    request_id=response.request_id,
                )
                atomic_write_json(paths["normalized"], normalized)
                atomic_write_json(
                    paths["usage"],
                    {
                        "request_fingerprint": fingerprint,
                        "target_block_id": block_id,
                        "block_type": payload.block_type,
                        "request_id": response.request_id,
                        "usage": response.usage,
                        "actual_cash_charge_cny": None,
                        "actual_charge_returned_by_api": False,
                        "note": "Usage does not establish the actual account deduction.",
                    },
                )
                atomic_write_json(
                    paths["cache"],
                    {
                        "request_fingerprint": fingerprint,
                        "target_block_id": block_id,
                        "block_type": payload.block_type,
                        "status": "translated",
                        "raw_response_path": str(paths["raw"].resolve()),
                        "normalized_output_path": str(paths["normalized"].resolve()),
                        "usage_path": str(paths["usage"].resolve()),
                    },
                )
                request_record.update(
                    {
                        "status": "completed",
                        "request_id": response.request_id,
                        "elapsed_seconds": round(time.perf_counter() - call_start, 3),
                        "ended_at": _now().isoformat(),
                    }
                )
                ledger["attempts"][-1].update(
                    {"status": "completed", "request_id": response.request_id, "ended_at": _now().isoformat()}
                )
                atomic_write_json(paths["request"], request_record)
                atomic_write_json(paths["ledger"], ledger)
                results_by_id[block_id] = normalized
            except Exception as exc:
                failed += 1
                if not paths["raw"].is_file():
                    atomic_write_json(
                        paths["raw"],
                        {
                            "record_type": "translation_provider_or_validation_error",
                            "request_fingerprint": fingerprint,
                            "target_block_id": block_id,
                            "block_type": payload.block_type,
                            "provider": settings.translation_provider,
                            "model": settings.translation_model,
                            "api_called": True,
                            "error_type": type(exc).__name__,
                            "error_message": str(exc)[:1000],
                            "automatic_retry": False,
                            "retries": 0,
                            "recorded_at": _now().isoformat(),
                        },
                    )
                atomic_write_json(
                    paths["failure"],
                    {
                        "schema_version": settings.translation_schema_version,
                        "target_block_id": block_id,
                        "block_type": payload.block_type,
                        "request_fingerprint": fingerprint,
                        "status": "translation_failed",
                        "raw_response_path": str(paths["raw"].resolve()),
                        "error_type": type(exc).__name__,
                        "error_message": str(exc)[:1000],
                        "automatic_retry": False,
                        "retries": 0,
                        "created_at": _now().isoformat(),
                    },
                )
                request_record.update(
                    {
                        "status": "failed",
                        "error_type": type(exc).__name__,
                        "error_message": str(exc)[:1000],
                        "elapsed_seconds": round(time.perf_counter() - call_start, 3),
                        "ended_at": _now().isoformat(),
                    }
                )
                ledger["attempts"][-1].update(
                    {"status": "failed", "error_type": type(exc).__name__, "ended_at": _now().isoformat()}
                )
                atomic_write_json(paths["request"], request_record)
                atomic_write_json(paths["ledger"], ledger)
                break

    ordered_results = [
        results_by_id[block_id]
        for block_id in selected_ids
        if block_id in results_by_id
    ]
    derived_path = None
    markdown_path = None
    docx_path = None
    if len(ordered_results) == 5 and failed == 0:
        derived_path, markdown_path, docx_path = _write_derived_outputs(
            settings,
            root,
            preflight.candidates,
            payloads,
            ordered_results,
            model_list_calls=model_list_calls,
            content_calls=api_calls,
        )
    return TranslationBatchResult(
        selected_block_ids=selected_ids,
        results=ordered_results,
        api_calls=api_calls,
        cache_hits=len(successful_cache),
        failed=failed,
        retries=0,
        model_list_calls=model_list_calls,
        model_available=True,
        derived_document_path=derived_path,
        diagnostic_markdown_path=markdown_path,
        diagnostic_docx_path=docx_path,
        strict_export_ready=(len(ordered_results) == 5 and failed == 0),
    )
