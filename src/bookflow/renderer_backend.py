"""Configurable office renderer discovery, validation, and isolated conversion."""

from __future__ import annotations

import os
import ctypes
import shutil
import subprocess
import tempfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

import fitz
import yaml
from docx import Document


@dataclass(frozen=True)
class OfficeRendererConfig:
    provider_id: str = "libreoffice"
    executable: str | None = None
    discovery: str = "auto"
    timeout_seconds: int = 180
    headless: bool = True
    temp_profile_root: str | None = None


def load_renderer_config(path: Path | None = None) -> OfficeRendererConfig:
    if path is None:
        return OfficeRendererConfig()
    data = yaml.safe_load(path.read_text("utf-8")) or {}
    raw = (data.get("renderers") or {}).get("office") or {}
    if not isinstance(raw, dict):
        raise ValueError("renderers.office must be a mapping")
    return OfficeRendererConfig(
        provider_id=str(raw.get("provider_id", "libreoffice")), executable=raw.get("executable"),
        discovery=str(raw.get("discovery", "auto")), timeout_seconds=int(raw.get("timeout_seconds", 180)),
        headless=bool(raw.get("headless", True)), temp_profile_root=raw.get("temp_profile_root"),
    )


def _runtime_profile_path() -> Path:
    root = os.getenv("APPDATA") or os.getenv("LOCALAPPDATA")
    return Path(root) / "Bookflow" / "runtime.yaml" if root else Path("runtime.yaml")


def _windows_file_version(path: Path) -> str | None:
    if os.name != "nt":
        return None
    try:
        class FixedInfo(ctypes.Structure):
            _fields_ = [(name, ctypes.c_uint32) for name in (
                "dwSignature", "dwStrucVersion", "dwFileVersionMS", "dwFileVersionLS",
                "dwProductVersionMS", "dwProductVersionLS", "dwFileFlagsMask", "dwFileFlags",
                "dwFileOS", "dwFileType", "dwFileSubtype", "dwFileDateMS", "dwFileDateLS")]
        handle = ctypes.c_uint32()
        size = ctypes.windll.version.GetFileVersionInfoSizeW(str(path), ctypes.byref(handle))
        if not size: return None
        buffer = ctypes.create_string_buffer(size)
        if not ctypes.windll.version.GetFileVersionInfoW(str(path), 0, size, buffer): return None
        pointer = ctypes.c_void_p(); length = ctypes.c_uint()
        if not ctypes.windll.version.VerQueryValueW(buffer, "\\", ctypes.byref(pointer), ctypes.byref(length)): return None
        info = ctypes.cast(pointer, ctypes.POINTER(FixedInfo)).contents
        return ".".join(str(x) for x in (info.dwFileVersionMS >> 16, info.dwFileVersionMS & 0xFFFF,
                                           info.dwFileVersionLS >> 16, info.dwFileVersionLS & 0xFFFF))
    except (AttributeError, OSError, ValueError):
        return None


def _runtime_executable() -> str | None:
    path = _runtime_profile_path()
    if not path.is_file():
        return None
    data = yaml.safe_load(path.read_text("utf-8")) or {}
    value = (data.get("renderers") or {}).get("office", {}).get("executable")
    return str(value) if value else None


def _registry_candidates() -> list[str]:
    if os.name != "nt":
        return []
    try:
        import winreg
    except ImportError:
        return []
    values = []
    for hive in (winreg.HKEY_CURRENT_USER, winreg.HKEY_LOCAL_MACHINE):
        for key_name in (r"SOFTWARE\LibreOffice\UNO\InstallPath", r"SOFTWARE\WOW6432Node\LibreOffice\UNO\InstallPath"):
            try:
                with winreg.OpenKey(hive, key_name) as key:
                    directory = winreg.QueryValue(key, None)
                    values.append(str(Path(directory) / "soffice.exe"))
            except OSError:
                continue
    return values


def discovery_candidates(config: OfficeRendererConfig) -> list[dict[str, str]]:
    values: list[tuple[str, str | None]] = [
        ("explicit_config", config.executable),
        ("environment", os.getenv("BOOKFLOW_LIBREOFFICE_PATH")),
    ]
    if config.discovery == "manual":
        return [{"source": source, "path": str(Path(value).expanduser())} for source, value in values if value]
    values.append(("runtime_profile", _runtime_executable()))
    values.extend(("windows_registry", value) for value in _registry_candidates())
    values.extend(("path", shutil.which(name)) for name in ("soffice", "libreoffice"))
    for variable in ("ProgramFiles", "ProgramFiles(x86)"):
        root = os.getenv(variable)
        if root:
            values.append(("common_directory", str(Path(root) / "LibreOffice" / "program" / "soffice.exe")))
    seen = set(); result = []
    for source, value in values:
        if not value:
            continue
        normalized = str(Path(value).expanduser())
        key = normalized.casefold()
        if key not in seen:
            result.append({"source": source, "path": normalized}); seen.add(key)
    return result


def detect_renderer(config: OfficeRendererConfig) -> dict[str, Any]:
    checked = []
    first_existing: tuple[Path, dict[str, Any]] | None = None
    for candidate in discovery_candidates(config):
        path = Path(candidate["path"])
        exists = path.is_file()
        checked.append({**candidate, "exists": exists})
        if exists:
            validation = validate_renderer(path, timeout_seconds=min(config.timeout_seconds, 3))
            if first_existing is None:
                first_existing = (path, validation)
            if validation["valid"]:
                return {"status": "available", "provider_id": config.provider_id, "executable": str(path.resolve()),
                        "version": validation["version"], "checked": checked, "validation": validation}
    if first_existing is not None:
        path, validation = first_existing
        return {"status": "detected_unverified", "provider_id": config.provider_id,
                "executable": str(path.resolve()), "version": validation.get("version"),
                "checked": checked, "validation": validation,
                "reason": "executable exists but version probe did not complete; run renderers test"}
    return {"status": "renderer_unavailable", "provider_id": config.provider_id, "executable": None,
            "checked": checked, "reason": "no runnable LibreOffice executable discovered",
            "next_action": "bookflow renderers detect --config <renderer.yaml> or set BOOKFLOW_LIBREOFFICE_PATH"}


def validate_renderer(executable: Path, *, timeout_seconds: int = 30) -> dict[str, Any]:
    try:
        with tempfile.TemporaryDirectory(prefix="bookflow-office-detect-") as temporary:
            profile = Path(temporary) / "profile"; profile.mkdir()
            run = subprocess.run([str(executable), "--headless", f"-env:UserInstallation={profile.as_uri()}", "--version"],
                                 capture_output=True, text=False, timeout=timeout_seconds, shell=False)
    except subprocess.TimeoutExpired as exc:
        # Some Windows LibreOffice builds attach the version request to an
        # existing process and never close stdout.  The signed executable's
        # version resource is a deterministic local fallback; conversion is
        # still exercised independently by ``renderers test``.
        version = _windows_file_version(executable)
        return {"valid": bool(version), "error": None if version else type(exc).__name__,
                "probe_error": type(exc).__name__, "version": version,
                "validation_method": "windows_file_version" if version else "version_probe"}
    except OSError as exc:
        return {"valid": False, "error": type(exc).__name__, "version": None}
    raw = run.stdout or run.stderr or b""
    text = raw.decode("utf-8", errors="replace").strip()
    version = text.splitlines()[0] if text else _windows_file_version(executable)
    return {"valid": run.returncode == 0 and bool(version), "returncode": run.returncode,
            "version": version, "error": None if run.returncode == 0 else "nonzero_exit",
            "validation_method": "version_command" if text else "windows_file_version"}


def test_renderer(config: OfficeRendererConfig) -> dict[str, Any]:
    detected = detect_renderer(config)
    if detected["status"] not in {"available", "detected_unverified"}:
        return {**detected, "test_conversion": "blocked_by_renderer"}
    base = Path(config.temp_profile_root).resolve() if config.temp_profile_root else None
    try:
        with tempfile.TemporaryDirectory(prefix="bookflow-office-", dir=str(base) if base else None) as temporary:
            root = Path(temporary); source = root / "Bookflow 渲染测试 路径 日本語.docx"; output_dir = root / "output"; output_dir.mkdir()
            document = Document(); document.add_heading("Bookflow renderer test", 0); document.add_paragraph("Unicode: é è ç 中文 日本語")
            document.save(source)
            profile = root / "profile"; profile.mkdir()
            command = [detected["executable"], "--headless", f"-env:UserInstallation={profile.as_uri()}",
                       "--convert-to", "pdf", "--outdir", str(output_dir), str(source)]
            run = subprocess.run(command, capture_output=True, text=False, timeout=config.timeout_seconds, shell=False)
            output = output_dir / f"{source.stem}.pdf"
            if run.returncode or not output.is_file() or output.stat().st_size < 100:
                return {**detected, "test_conversion": "failed", "returncode": run.returncode,
                        "output_exists": output.is_file(),
                        "error": (run.stderr or run.stdout or b"")[-1000:].decode("utf-8", errors="replace")}
            pdf = fitz.open(output); pages = pdf.page_count; pdf.close()
            return {**detected, "status": "available" if pages > 0 else "detected_unverified",
                    "test_conversion": "passed" if pages > 0 else "failed",
                    "returncode": run.returncode, "pdf_pages": pages, "output_size": output.stat().st_size,
                    "isolated_profile": True, "shell": False}
    except (OSError, subprocess.TimeoutExpired, ValueError) as exc:
        return {**detected, "test_conversion": "failed", "error": type(exc).__name__}


def convert_docx_to_pdf(config: OfficeRendererConfig, source: Path, output: Path) -> dict[str, Any]:
    """Convert with an isolated profile and atomically commit a verified PDF."""
    detected = detect_renderer(config)
    if detected["status"] not in {"available", "detected_unverified"}:
        return {"status": "blocked_by_renderer", "renderer": detected}
    base = Path(config.temp_profile_root).resolve() if config.temp_profile_root else None
    try:
        with tempfile.TemporaryDirectory(prefix="bookflow-office-job-", dir=str(base) if base else None) as temporary:
            root = Path(temporary); profile = root / "profile"; profile.mkdir(); converted = root / "converted"; converted.mkdir()
            command = [detected["executable"], "--headless", f"-env:UserInstallation={profile.as_uri()}",
                       "--convert-to", "pdf", "--outdir", str(converted), str(source.resolve())]
            run = subprocess.run(command, capture_output=True, text=False, timeout=config.timeout_seconds, shell=False)
            candidate = converted / f"{source.stem}.pdf"
            if run.returncode or not candidate.is_file() or candidate.stat().st_size < 100:
                return {"status": "blocked_by_renderer", "renderer": detected, "returncode": run.returncode,
                        "output_exists": candidate.is_file(), "reason": "conversion did not create a valid-size PDF"}
            document = fitz.open(candidate); pages = document.page_count; document.close()
            if pages < 1:
                return {"status": "blocked_by_renderer", "renderer": detected, "reason": "converted PDF has no pages"}
            output.parent.mkdir(parents=True, exist_ok=True)
            temporary_output = output.with_suffix(output.suffix + ".tmp")
            shutil.copy2(candidate, temporary_output)
            try:
                temporary_output.replace(output)
                committed = output
            except PermissionError:
                from datetime import datetime, timezone
                stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
                committed = output.with_name(f"{output.stem}-{stamp}{output.suffix}")
                temporary_output.replace(committed)
            return {"status": "generated", "renderer": detected["provider_id"], "path": str(committed),
                    "pages": pages, "size_bytes": committed.stat().st_size, "isolated_profile": True}
    except (OSError, subprocess.TimeoutExpired, ValueError) as exc:
        return {"status": "blocked_by_renderer", "renderer": detected, "reason": type(exc).__name__}


def renderer_list(config: OfficeRendererConfig) -> dict[str, Any]:
    return {"office": asdict(config), "native_pdf": {"provider_id": "pymupdf", "status": "available",
            "capabilities": ["text_pdf"], "external_executable": False}}
