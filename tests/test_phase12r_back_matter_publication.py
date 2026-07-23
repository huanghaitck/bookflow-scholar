from __future__ import annotations

import hashlib
import json
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from PIL import Image

from bookflow.back_matter_publication import (
    LAYOUT_OBJECT_TYPES,
    build_render_plan_manifest,
    render_layout_docx,
    render_layout_markdown,
    validate_layout_model,
    validate_quality_gates,
)
from bookflow.production import calculate_phase12r_readiness


ROOT = Path(__file__).resolve().parents[1]
PHASE12R = ROOT / "data/fullbook/back_matter/phase12r"


def _json(path: Path) -> dict:
    return json.loads(path.read_text("utf-8"))


def _sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_page_normalization_and_rotated_landscape_regression() -> None:
    manifest = _json(PHASE12R / "normalized/normalization_manifest.json")
    pages = {item["physical_page"]: item for item in manifest["pages"]}
    assert set(pages) == {21, 22, *range(381, 409)}
    assert all(item["orientation_degrees_clockwise"] == (90 if page == 387 else 0) for page, item in pages.items())
    assert all(item["operation_order"] == ["rotate", "crop", "deskew"] for item in pages.values())
    assert all(len(item["raw_to_normalized_transform"]) == 3 for item in pages.values())
    with Image.open(PHASE12R / "normalized/page_0387.normalized.png") as page:
        assert page.width > page.height
    region = PHASE12R / "normalized/regions/appendix_a_rotated_table_p0387.png"
    with Image.open(region) as image:
        assert image.width > image.height


def test_illustration_list_is_complete_and_uses_original_printed_locators() -> None:
    model = _json(PHASE12R / "illustration_list_reading_order_v1.json")
    entries = model["entries"]
    assert len(entries) == 63
    assert [entry["reading_order"] for entry in entries] == list(range(1, 64))
    assert sum(entry["physical_page"] == 21 for entry in entries) == 31
    assert sum(entry["physical_page"] == 22 for entry in entries) == 32
    assert len(model["groups"][0]["entry_ids"]) == 61
    assert len(model["groups"][1]["entry_ids"]) == 2
    assert entries[0]["printed_locator"] == "Frontispiece"
    assert entries[0]["locator_kind"] == "frontispiece"
    assert all(entry["source_region_id"] for entry in entries)
    assert all(str(entry["printed_locator"]) != str(entry["physical_page"]) for entry in entries)
    assert {entry["latin_name"] for entry in entries if entry["latin_name"]} >= {
        "Budorcas bedfordi",
        "Capreolus bedfordi",
        "Cervus kansuensis",
        "Gazella gutturosa",
        "Gazella przewalskii",
        "Nemorhaedus argyrochaetes",
    }


def test_layout_object_contract_and_publication_quality_gates() -> None:
    layout = _json(PHASE12R / "layout_object_model_v1.json")
    normalization = _json(PHASE12R / "normalized/normalization_manifest.json")
    illustrations = _json(PHASE12R / "illustration_list_reading_order_v1.json")
    validation = validate_layout_model(layout)
    assert validation["valid"], validation
    assert {item["object_type"] for item in layout["objects"]} <= LAYOUT_OBJECT_TYPES
    gates = validate_quality_gates(
        normalization=normalization,
        layout_model=layout,
        illustration_model=illustrations,
    )
    assert gates["valid"], gates
    rotated = next(item for item in layout["objects"] if item["object_id"] == "lo_appendix_a_rotated_table_0387")
    assert rotated["rows"] == []
    assert rotated["numeric_status"] == "unresolved"
    assert rotated["provenance"]["legacy_ocr_rows"] == "rejected_as_formal_content"


def test_render_profiles_and_cross_format_counts() -> None:
    layout = _json(PHASE12R / "layout_object_model_v1.json")
    plan = build_render_plan_manifest(layout)
    reading = plan["profiles"]["reading"]
    evidence = plan["profiles"]["evidence"]
    assert reading["technical_label_count"] == 0
    assert reading["full_page_facsimile_count"] == 0
    assert evidence["technical_label_count"] > 0
    assert len({json.dumps(value, sort_keys=True) for value in reading["format_object_counts"].values()}) == 1
    assert len({json.dumps(value, sort_keys=True) for value in evidence["format_object_counts"].values()}) == 1


def test_native_table_landscape_index_and_dot_leaders(tmp_path: Path) -> None:
    illustrations = {
        "groups": [{"group_id": "illustrations", "entry_ids": ["ill_1"]}],
        "entries": [{"entry_id": "ill_1", "source_text": "Plate", "printed_locator": "18"}],
    }
    table = {
        "object_id": "table_1",
        "object_type": "rotated_table",
        "source_pages": [1],
        "bbox": [0, 0, 1, 1],
        "orientation": 90,
        "reading_order": 2,
        "confidence": 1,
        "provenance": {},
        "render_policy": {"reading": "native", "evidence": "native"},
        "schema_candidate": {"column_count": 2, "columns": ["Name", "Value"]},
        "rows": [{"cells": ["A", "1"]}],
        "publication": {"wide": True, "overflow_policy": "landscape_section"},
    }
    illustration = {
        "object_id": "illustrations",
        "object_type": "illustration_list",
        "source_pages": [1],
        "bbox": [0, 0, 1, 1],
        "orientation": 0,
        "reading_order": 1,
        "confidence": 1,
        "provenance": {},
        "render_policy": {"reading": "native", "evidence": "native"},
    }
    index = deepcopy(illustration)
    index.update(object_id="index", object_type="index", reading_order=3)
    model = {"objects": [illustration, table, index]}
    markdown, md_manifest = render_layout_markdown(model, illustrations, profile="reading")
    assert "Plate ............ 18" in markdown
    assert "| Name | Value |" in markdown
    destination = tmp_path / "back_matter.docx"
    docx_manifest = render_layout_docx(model, illustrations, destination, profile="reading")
    document = Document(destination)
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 1).text == "1"
    with zipfile.ZipFile(destination) as archive:
        xml = archive.read("word/document.xml")
    assert b"w:leader=\"dot\"" in xml or b"w:leader=\"middleDot\"" in xml
    assert b"w:orient=\"landscape\"" in xml
    assert b"w:num=\"2\"" in xml
    assert md_manifest["object_counts"] == docx_manifest["object_counts"]


def test_index_artifacts_are_not_formal_nodes() -> None:
    model = _json(ROOT / "data/fullbook/back_matter/index_reading_order_v1.json")
    forbidden = ("digitized by", "univ calif", "printed by hazell")
    assert all(node["term"].strip().upper() != "INDEX" for node in model["nodes"])
    assert all(not node["term"].strip().isdigit() for node in model["nodes"])
    assert all(not any(value in node["source_display_text"].lower() for value in forbidden) for node in model["nodes"])


def test_phase12r_release_gate_accepts_formal_provider_results_and_translation_overlay() -> None:
    en = calculate_phase12r_readiness(ROOT, "en")
    zh = calculate_phase12r_readiness(ROOT, "zh-Hans")
    assert en["ready"] and not en["blockers"]
    assert zh["ready"] and not zh["blockers"]


def test_vision_plan_is_offline_and_uses_only_deterministic_regions() -> None:
    plan = _json(ROOT / "reports/PHASE12R_VISION_EXTRACTION_PLAN.json")
    regions = {
        item["region_id"]: item["region_asset_ref"]
        for item in _json(PHASE12R / "normalized/normalization_manifest.json")["regions"]
    }
    assert plan["execution_status"] == "planned_offline_only"
    assert plan["real_api_allowed"] is False
    assert plan["api_calls"] == 0 and plan["api_tokens"] == 0
    assert len(plan["requests"]) == 30
    assert all(request["region_id"] in regions for request in plan["requests"])
    assert all(regions[request["region_id"]].endswith(request["asset_ref"]) for request in plan["requests"])


def test_first_book_frozen_hashes_are_unchanged() -> None:
    expected = {
        "data/fullbook/canonical/canonical_book_document_v1.json": "16c1c9ba4d60d1c2a4124433291a1a56bf499384215c720f6988e6e183c01326",
        "data/fullbook/main_text/boundaries/main_text.boundaries.jsonl": "b08c4bab8506f6d85cfd5e48b54ec801bd1868e10e6fb1375779011c08faf5a1",
        "data/fullbook/main_text/source_document_main_text_v1.json": "f18ad3eefa24eec1241dbe69a8baa0ecd8512a998a4796b5636c8f980cc01c8a",
        "data/fullbook/main_text/bilingual_document_main_text_zh-Hans_v1.json": "d00f42fecbd0f8410019a2b9cfb42eeba0a2dc6f97188fdcb00cce8acf4bdc8b",
    }
    assert {path: _sha(ROOT / path) for path in expected} == expected


def test_candidate_release_manifests_remain_unchanged() -> None:
    candidates = {
        "output/fullbook/big-game-en-reading-release-20260717T155541Z877193/render_manifest.json": "f83046ef7fbbe0a2f399819494faaaf1c8270cf9887d2bd9c07c2a28e8790a08",
        "output/fullbook/big-game-zh-Hans-reading-release-20260717T155700Z478010/render_manifest.json": "b40de82df2d0b5dc0d89b8615e785d18390cf9b3bea3f31dceaccfe3f45056bb",
        "output/fullbook/big-game-bilingual-reading-release-20260717T160109Z973759/render_manifest.json": "dd4159921085e494725460ca8232e43e8a361f8f184d73df7a83465ea2ae8c8e",
    }
    for path, markdown_sha in candidates.items():
        manifest = _json(ROOT / path)
        assert manifest["outputs"]["markdown"]["sha256"] == markdown_sha
        assert manifest["api_calls"] == 0 and manifest["api_tokens"] == 0
