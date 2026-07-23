from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from typer.testing import CliRunner

from bookflow.cli import app
from bookflow.production import CANONICAL_SHA, RenderInputs, TranslationResolver, build, calculate_release_eligibility, enforce_release_gate, render_docx, render_markdown


ROOT = Path(__file__).resolve().parents[1]


def test_inputs_and_resolver_use_frozen_layers() -> None:
    inputs = RenderInputs.load(ROOT)
    assert len(inputs.canonical["logical_units"]) == 971
    assert inputs.canonical["metadata"]["total_chapters"] == 30
    assert inputs.status_counts == {"reused_frozen": 971, "validated": 764, "preserve_source": 334, "blocked_by_source_quality": 7}
    first = inputs.canonical["logical_units"][0]
    text, status = TranslationResolver(inputs).resolve(first["logical_block_id"], first["source_text"], "zh-Hans", "preview")
    assert text and status == "reused_frozen"


def test_markdown_fullbook_and_release_gate(tmp_path: Path) -> None:
    inputs = RenderInputs.load(ROOT)
    en = render_markdown(inputs, tmp_path / "en", "en", "release")
    zh = render_markdown(inputs, tmp_path / "zh", "zh-Hans", "preview")
    bi = render_markdown(inputs, tmp_path / "bi", "bilingual", "preview")
    for result in (en, zh, bi):
        text = Path(result["path"]).read_text("utf-8")
        assert result["source_logical_units"] == 971 and result["chapters"] == 30
        assert "THE WHITE-MANED SEROW" in text
        assert "ACCOUNT OF" in text
        assert "ESTIMATE OF EXPENSES" in text
        assert "TABLE OF DISTANCES AND STAGES" in text
        assert "## Index" in text
        assert "Digitized by Microsoft" not in text
        assert "Univ Calif" not in text
        assert "CHAPTER XXVI" in text and "238" in text
    assert Path(en["path"]).read_text("utf-8").count("## CHAPTER I\n") == 1
    release = render_markdown(inputs, tmp_path / "release", "zh-Hans", "release")
    assert release["index_element_count"] == 298
    assert release["appendix_element_count"] == 462


def test_docx_structural_gate(tmp_path: Path) -> None:
    result = render_docx(RenderInputs.load(ROOT), tmp_path, "en", "release")
    doc = Document(result["path"])
    names = {s.name for s in doc.styles}
    assert result["chapters"] == 30 and result["images"] > 0
    assert {"BodyEnglish", "BodyChinese", "ChapterTitle", "AppendixHeading", "IndexEntry"} <= names
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "THE WHITE-MANED SEROW" in text and "ACCOUNT OF" in text
    assert "Digitized by Microsoft" not in text and "Univ Calif" not in text


def test_chinese_and_bilingual_reading_headers(tmp_path: Path) -> None:
    inputs = RenderInputs.load(ROOT)
    zh = render_docx(inputs, tmp_path / "zh", "zh-Hans", "preview")
    bi = render_docx(inputs, tmp_path / "bi", "bilingual", "preview")
    ztext = "\n".join(p.text for p in Document(zh["path"]).paragraphs)
    btext = "\n".join(p.text for p in Document(bi["path"]).paragraphs)
    assert "中国中西部的大型猎物" in ztext and "第一章" in ztext and "红神的召唤" in ztext
    assert "CHAPTER I / 第一章" in btext and bi["bilingual_pairs"] > 650


def test_build_is_immutable_and_manifested(tmp_path: Path) -> None:
    result = build(ROOT, "en", "release", ("md",), output_root=tmp_path)
    manifest = json.loads(Path(result["manifest_path"]).read_text("utf-8"))
    assert manifest["canonical_sha256"] == CANONICAL_SHA
    assert manifest["api_calls"] == manifest["api_tokens"] == 0
    assert manifest["validation"] == "passed"


def test_unified_cli_help_dry_run_and_release_gate() -> None:
    runner = CliRunner()
    assert runner.invoke(app, ["--help"]).exit_code == 0
    assert runner.invoke(app, ["translate", "--dry-run"]).exit_code == 0
    eligibility = calculate_release_eligibility(ROOT)
    enforce_release_gate(eligibility, "zh-Hans", "release")
    enforce_release_gate(eligibility, "bilingual", "release")
    assert eligibility["eligible"] and eligibility["pending_translatable"] == 0
