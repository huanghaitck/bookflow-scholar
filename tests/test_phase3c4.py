from __future__ import annotations

import json
import subprocess
from pathlib import Path

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from bookflow.io_utils import load_json, sha256_file, sha256_text
from bookflow.paths import load_settings, project_root
from bookflow.phase3c4 import (
    Phase3C4ModelPayload,
    build_phase3c4_bilingual_document,
    build_phase3c4_units,
    detect_libreoffice,
    export_phase3c4_candidates,
    is_phase3c4_translatable_type,
    normalize_phase3c4_translation,
    phase3c4_preflight,
    phase3c4_request_fingerprint,
    run_libreoffice_smoke,
    run_phase3c4_translation,
)
from bookflow.translation_provider import TranslationProviderResponse


ROOT = project_root()


def _settings(tmp_path: Path):
    return load_settings().model_copy(
        update={
            "phase3c4_data_directory": str(tmp_path / "phase3c4"),
            "phase3c4_master_path": str(tmp_path / "bilingual_document_sample12_zh-Hans_v1.json"),
            "phase3c4_candidate_directory": str(tmp_path / "candidate"),
            "phase3c4_rendered_directory": str(tmp_path / "rendered"),
            "phase3c4_libreoffice_profile_directory": str(tmp_path / "lo-profile"),
            "phase3c4_soffice_candidates": [str(tmp_path / "missing" / "soffice.exe")],
        }
    )


class FakeTranslationProvider:
    instances: list["FakeTranslationProvider"] = []

    def __init__(self, **_: object) -> None:
        self.calls: list[dict[str, object]] = []
        self.__class__.instances.append(self)

    def translate_one(self, **kwargs: object) -> TranslationProviderResponse:
        payload = dict(kwargs["user_payload"])
        self.calls.append(payload)
        number = len(self.calls)
        terms = ["Tonkinese"] if "Tonkinese" in str(payload["source_text"]) else []
        repeat = max(3, len(str(payload["source_text"])) // 35)
        result = {
            "target_block_id": payload["target_block_id"],
            "block_type": payload["block_type"],
            "translation": (f"这是第{number}个完整中文译文，忠实保留原文信息和叙述语气。" * repeat)
            + (" Tonkinese" if terms else ""),
            "untranslated_source_terms": terms,
            "warnings": [],
        }
        content = json.dumps(result, ensure_ascii=False)
        usage = {
            "prompt_tokens": 1000,
            "prompt_tokens_details": {"cached_tokens": 100},
            "completion_tokens": 200,
            "total_tokens": 1200,
        }
        return TranslationProviderResponse(
            raw_response={
                "id": f"mock-{number}",
                "model": "deepseek-v4-pro",
                "choices": [{"message": {"content": content}}],
                "usage": usage,
            },
            content=content,
            request_id=f"mock-{number}",
            usage=usage,
            response_model="deepseek-v4-pro",
        )


@pytest.fixture(autouse=True)
def _clear_fake_instances():
    FakeTranslationProvider.instances.clear()


@pytest.fixture()
def mock_batch(tmp_path: Path):
    settings = _settings(tmp_path)
    source_before = sha256_file(ROOT / "data/source_document_sample12_v1.json")
    phase3a_before = sha256_file(ROOT / "data/phase3a_translation_sample.json")
    batch = run_phase3c4_translation(
        settings,
        allow_api=True,
        confirmed=True,
        root=ROOT,
        key_status_resolver=lambda *_: (True, "DEEPSEEK_API_KEY"),
        key_loader=lambda *_: ("fake-key", "DEEPSEEK_API_KEY"),
        provider_factory=FakeTranslationProvider,
    )
    assert sha256_file(ROOT / "data/source_document_sample12_v1.json") == source_before
    assert sha256_file(ROOT / "data/phase3a_translation_sample.json") == phase3a_before
    return settings, batch


def test_new_prompt_version_and_fingerprint_are_phase3c4_specific(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    units = build_phase3c4_units(settings, root=ROOT)
    fingerprint = phase3c4_request_fingerprint(
        units[0], settings, prompt_sha256="a" * 64, profile_sha256="b" * 64
    )
    assert settings.phase3c4_prompt_version == "translation_en_zh_v2.0"
    assert settings.phase3c4_prompt_version != settings.translation_prompt_version
    assert len(fingerprint) == 64


def test_builds_exactly_24_complete_translation_units(tmp_path: Path) -> None:
    units = build_phase3c4_units(_settings(tmp_path), root=ROOT)
    assert len(units) == 24
    assert all(unit.translation_ready for unit in units)
    assert all(unit.source_text.strip() for unit in units)
    assert all(unit.translate_target_only for unit in units)


def test_titles_are_independent_units_and_structural_noise_is_excluded(tmp_path: Path) -> None:
    units = build_phase3c4_units(_settings(tmp_path), root=ROOT)
    title_units = [item for item in units if item.block_type in {"chapter_title", "section_title"}]
    assert len(title_units) == 4
    assert is_phase3c4_translatable_type("chapter_title") is True
    assert is_phase3c4_translatable_type("section_title") is True
    assert is_phase3c4_translatable_type("running_header") is False
    assert is_phase3c4_translatable_type("page_number") is False


def test_body_request_separates_target_from_read_only_context(tmp_path: Path) -> None:
    unit = next(item for item in build_phase3c4_units(_settings(tmp_path), root=ROOT) if item.block_type == "body")
    assert unit.source_text
    assert unit.context_after_text or unit.context_before_text
    assert unit.source_text not in {unit.context_before_text, unit.context_after_text}
    assert unit.chapter_title_context is not None
    assert unit.section_title_context is not None


def test_preserved_english_term_is_valid_model_output(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    unit = build_phase3c4_units(settings, root=ROOT)[0]
    payload = Phase3C4ModelPayload.model_validate(
        {
            "target_block_id": unit.target_block_id,
            "block_type": unit.block_type,
            "translation": "中文译文保留 Tonkinese。",
            "untranslated_source_terms": ["Tonkinese"],
            "warnings": [],
        }
    )
    result = normalize_phase3c4_translation(
        content=payload.model_dump_json(),
        unit=unit,
        settings=settings,
        fingerprint="f" * 64,
        prompt_sha256="a" * 64,
        profile_sha256="b" * 64,
        raw_path=tmp_path / "raw.json",
        usage=None,
        request_id="mock",
    )
    assert result.untranslated_source_terms == ["Tonkinese"]
    assert result.status == "translated"


def test_listed_source_term_missing_from_translation_is_added_without_failing(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    unit = build_phase3c4_units(settings, root=ROOT)[0]
    content = json.dumps(
        {
            "target_block_id": unit.target_block_id,
            "block_type": unit.block_type,
            "translation": "第一章",
            "untranslated_source_terms": ["Kamo- gawa"],
            "warnings": [],
        },
        ensure_ascii=False,
    )
    result = normalize_phase3c4_translation(
        content=content,
        unit=unit,
        settings=settings,
        fingerprint="f" * 64,
        prompt_sha256="a" * 64,
        profile_sha256="b" * 64,
        raw_path=tmp_path / "raw.json",
        usage=None,
        request_id="mock",
    )
    assert "Kamo- gawa" in result.translation
    assert result.untranslated_source_terms == ["Kamo- gawa"]
    assert any("deterministically appended" in warning for warning in result.warnings)


def test_context_cannot_be_copied_into_translation(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    unit = next(item for item in build_phase3c4_units(settings, root=ROOT) if item.context_after_text)
    content = json.dumps(
        {
            "target_block_id": unit.target_block_id,
            "block_type": unit.block_type,
            "translation": unit.context_after_text,
            "untranslated_source_terms": [],
            "warnings": [],
        },
        ensure_ascii=False,
    )
    with pytest.raises(ValueError, match="context_after"):
        normalize_phase3c4_translation(
            content=content,
            unit=unit,
            settings=settings,
            fingerprint="f" * 64,
            prompt_sha256="a" * 64,
            profile_sha256="b" * 64,
            raw_path=tmp_path / "raw.json",
            usage=None,
            request_id="mock",
        )


def test_preflight_reports_24_calls_and_stays_below_two_cny(tmp_path: Path) -> None:
    report = phase3c4_preflight(
        _settings(tmp_path), root=ROOT, key_status_resolver=lambda *_: (True, "DEEPSEEK_API_KEY")
    )
    assert report.target_block_count == 24
    assert report.estimated_new_content_calls == 24
    assert report.maximum_content_calls == 24
    assert report.total_source_characters == 17629
    assert report.total_context_characters > 0
    assert report.estimated_cost_upper_cny < 2.00
    assert report.ready_for_real_call is True
    assert report.api_called is False


def test_mock_run_translates_each_block_once_and_second_run_is_cached(mock_batch) -> None:
    settings, first = mock_batch
    assert first.api_calls == 24
    assert first.cache_hits == 0
    assert first.failed == 0
    assert len(first.results) == 24
    second = run_phase3c4_translation(
        settings,
        allow_api=True,
        confirmed=True,
        root=ROOT,
        key_status_resolver=lambda *_: (True, "DEEPSEEK_API_KEY"),
        key_loader=lambda *_: ("fake-key", "DEEPSEEK_API_KEY"),
        provider_factory=FakeTranslationProvider,
    )
    assert second.api_calls == 0
    assert second.cache_hits == 24
    assert second.retries == 0


def test_phase3a_artifacts_are_not_overwritten_by_mock_run(mock_batch) -> None:
    _, _ = mock_batch
    assert sha256_file(ROOT / "data/phase3a_translation_sample.json") == "1b5bb632c4b5223b8d64a5f67b9fd48e52569e864b4d5af0ad77e6faee7b6738"


def test_bilingual_master_detects_missing_duplicate_and_extra_results(mock_batch) -> None:
    settings, batch = mock_batch
    with pytest.raises(ValueError, match="missing"):
        build_phase3c4_bilingual_document(settings, batch.results[:-1], root=ROOT)
    with pytest.raises(ValueError, match="duplicate"):
        build_phase3c4_bilingual_document(settings, batch.results + [batch.results[0]], root=ROOT)
    extra = batch.results[0].model_copy(update={"target_block_id": "logical_extra"})
    with pytest.raises(ValueError, match="extra"):
        build_phase3c4_bilingual_document(settings, batch.results + [extra], root=ROOT)


def test_bilingual_master_aligns_all_ids_and_does_not_build_terminology_table(mock_batch) -> None:
    settings, batch = mock_batch
    document = build_phase3c4_bilingual_document(settings, batch.results, root=ROOT)
    assert len(document.logical_blocks) == 24
    assert document.audit.missing_translation_ids == []
    assert document.audit.duplicate_translation_ids == []
    assert document.audit.extra_translation_ids == []
    assert document.audit.source_text_modified is False
    encoded = document.model_dump_json()
    assert "terminology_table" not in encoded
    assert "glossary" not in encoded
    assert "doing King" in encoded
    assert "doingKing" not in encoded


def test_four_candidate_outputs_come_from_one_bilingual_json(mock_batch) -> None:
    settings, batch = mock_batch
    document = build_phase3c4_bilingual_document(settings, batch.results, root=ROOT)
    result = export_phase3c4_candidates(settings, master_path=settings.phase3c4_master_path, root=ROOT)
    paths = [
        Path(result.source_markdown_path),
        Path(result.source_docx_path),
        Path(result.bilingual_markdown_path),
        Path(result.bilingual_docx_path),
    ]
    assert all(path.is_file() for path in paths)
    source_md = paths[0].read_text(encoding="utf-8")
    bilingual_md = paths[2].read_text(encoding="utf-8")
    source_word = "\n".join(p.text for p in Document(paths[1]).paragraphs)
    bilingual_word = "\n".join(p.text for p in Document(paths[3]).paragraphs)
    for block in document.logical_blocks:
        assert block.source_text in source_md
        assert block.source_text in source_word
        assert block.source_text in bilingual_md
        assert block.translation in bilingual_md
        assert block.source_text in bilingual_word
        assert block.translation in bilingual_word
    assert result.source_json_sha256 == sha256_file(settings.phase3c4_master_path)


def test_source_lines_are_not_stretched_by_justification(mock_batch) -> None:
    settings, batch = mock_batch
    document = build_phase3c4_bilingual_document(settings, batch.results, root=ROOT)
    result = export_phase3c4_candidates(settings, master_path=settings.phase3c4_master_path, root=ROOT)
    multiline = next(block for block in document.logical_blocks if "\n" in block.source_text)
    paragraph = next(
        item for item in Document(result.source_docx_path).paragraphs
        if item.text == multiline.source_text
    )
    assert paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY


def test_candidates_hide_internal_ids_and_api_metadata(mock_batch) -> None:
    settings, batch = mock_batch
    build_phase3c4_bilingual_document(settings, batch.results, root=ROOT)
    result = export_phase3c4_candidates(settings, master_path=settings.phase3c4_master_path, root=ROOT)
    visible = Path(result.bilingual_markdown_path).read_text(encoding="utf-8")
    assert "logical2_" not in visible
    assert "source_fragment" not in visible
    assert "request_fingerprint" not in visible
    assert "prompt_tokens" not in visible


def test_libreoffice_missing_is_safe_degradation(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    assert detect_libreoffice(settings, root=ROOT) is None
    result = run_libreoffice_smoke(
        settings,
        source_docx=tmp_path / "source.docx",
        bilingual_docx=tmp_path / "bilingual.docx",
        root=ROOT,
    )
    assert result.libreoffice_found is False
    assert result.completed is False
    assert result.fatal_error is False


def test_libreoffice_success_writes_two_readable_pdfs(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source_docx = tmp_path / "source.docx"
    bilingual_docx = tmp_path / "bilingual.docx"
    Document().save(source_docx)
    Document().save(bilingual_docx)
    executable = tmp_path / "soffice.exe"
    executable.parent.mkdir(parents=True, exist_ok=True)
    executable.write_bytes(b"mock")

    def fake_runner(command, **_: object):
        outdir = Path(command[command.index("--outdir") + 1])
        docx = Path(command[-1])
        outdir.mkdir(parents=True, exist_ok=True)
        pdf = fitz.open()
        pdf.new_page().insert_text((72, 72), docx.stem)
        pdf.save(outdir / f"{docx.stem}.pdf")
        pdf.close()
        return subprocess.CompletedProcess(command, 0, "converted", "")

    result = run_libreoffice_smoke(
        settings,
        source_docx=source_docx,
        bilingual_docx=bilingual_docx,
        executable=executable,
        command_runner=fake_runner,
        root=ROOT,
    )
    assert result.completed is True
    assert result.source_pdf_page_count == 1
    assert result.bilingual_pdf_page_count == 1
    assert Path(result.source_pdf_path).is_file()
    assert Path(result.bilingual_pdf_path).is_file()


def test_mock_phase3c4_never_calls_glm_or_real_network(mock_batch) -> None:
    _, batch = mock_batch
    assert batch.glm_calls == 0
    assert batch.model_list_calls == 0
    assert batch.retries == 0
    assert len(FakeTranslationProvider.instances) == 1
